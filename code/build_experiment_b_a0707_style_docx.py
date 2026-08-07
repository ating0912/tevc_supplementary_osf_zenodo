from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r".")
SUMMARY_DIR = ROOT / "p0_lite_outputs" / "experiment_b_configuration_summary_20260713"
FIG_DIR = SUMMARY_DIR / "figures"
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "陳羿婷_實驗B數據報告_0713.docx"

METHOD_LABELS = {
    "MetaDesigned_ECMADE_MOO": "Meta-designed",
    "BayesianConfig_ECMADE_MOO": "Bayesian",
    "RandomConfig_ECMADE_MOO": "Random",
    "HandCrafted_ECMADE_MOO": "Hand-crafted",
}

METRIC_DIRECTIONS = {
    "HV": "max",
    "IGD": "min",
    "PF_Overlap": "max",
    "PF_Drift": "min",
    "Diversity": "max",
    "Runtime": "min",
    "RankScore": "min",
}


def set_ea(run, east_asia="Microsoft JhengHei") -> None:
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for style_name in ("List Paragraph", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.0


def style_table(table, widths: list[float]) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx < len(widths):
                cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if row_idx == 0 or (col_idx > 0 and len(cell.text) < 18):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_ea(run)
                    run.font.size = Pt(10)
    for cell in table.rows[0].cells:
        set_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph()


def add_para(doc: Document, text: str, *, justify=True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else None
    r = p.add_run(text)
    set_ea(r)
    r.font.size = Pt(12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_ea(r)
    r.font.size = Pt(12)


def fmt(value, digits=4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def fmt_p(value) -> str:
    value = float(value)
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}"


def best_baseline(overall: pd.DataFrame, metric: str) -> str:
    baselines = overall[overall["method"] != "MetaDesigned_ECMADE_MOO"].copy()
    col = f"mean_{metric}" if metric != "RankScore" else "mean_RankScore"
    if METRIC_DIRECTIONS[metric] == "max":
        return str(baselines.sort_values(col, ascending=False).iloc[0]["method"])
    return str(baselines.sort_values(col, ascending=True).iloc[0]["method"])


def result_symbol(row: pd.Series) -> str:
    if bool(row["significant_0_05"]) and float(row["median_improvement"]) > 0:
        return "+"
    if float(row["median_improvement"]) < 0 and int(row["losses"]) > int(row["wins"]):
        return "-"
    return "="


def method_rows(overall: pd.DataFrame, metrics: list[str]) -> list[list[str]]:
    rows = []
    for _, row in overall.iterrows():
        values = [METHOD_LABELS[row["method"]]]
        for metric in metrics:
            col = f"mean_{metric}" if metric != "RankScore" else "mean_RankScore"
            values.append(fmt(row[col], 4 if metric != "Runtime" else 3))
        rows.append(values)
    return rows


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    overall = pd.read_csv(SUMMARY_DIR / "overall_configuration_comparison.csv")
    ranked = pd.read_csv(SUMMARY_DIR / "combined_instance_method_metrics_ranked.csv")
    stats = pd.read_csv(SUMMARY_DIR / "statistical_tests_meta_vs_baselines.csv")
    friedman = pd.read_csv(SUMMARY_DIR / "friedman_tests_all_methods.csv")
    theta_usage = pd.read_csv(SUMMARY_DIR / "theta_usage_by_method.csv")

    overall = overall.sort_values(["overall_RankScore", "mean_RankScore", "method"]).reset_index(drop=True)

    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Experiment B 數據報告")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor.from_string("0B2545")
    set_ea(tr)

    add_para(
        doc,
        "Experiment B 以 Experiment A 的 baseline comparison 報告格式，整理 ECMADE-MOO 在不同 configuration strategy 下的結果。此實驗不比較完全不同的演算法核心，而是比較同一 ECMADE-MOO search engine 在不同 theta configuration selection strategy 下，是否能在 unseen portfolio instances 上取得較佳泛化表現。",
    )
    add_bullet(doc, "比較方法包含 Hand-crafted、Random configuration、Bayesian configuration 與 Meta-designed ECMADE-MOO。")
    add_bullet(doc, "所有方法使用相同 32 個 unseen test instances，每個 instance 執行 30 independent runs。")
    add_bullet(doc, "Random、Bayesian 與 Meta-designed 皆從相同 L24 theta configuration set 選擇參數組合。")
    add_bullet(doc, "Bayesian configuration 先在 validation instances 上搜尋單一 global theta，再將該 theta 套用至全部 unseen test instances。")
    add_bullet(doc, "後處理以 common reference front 計算 HV、IGD、PF overlap、PF drift、diversity、runtime 與 RankScore。")
    add_bullet(doc, "統計檢定包含 Friedman overall test，以及 Meta-designed vs baselines 的 one-sided Wilcoxon signed-rank test with Holm correction。")

    doc.add_heading("公平性固定設定", level=1)
    add_table(
        doc,
        ["設定", "固定值 / 說明"],
        [
            ["底層演算法", "ECMADE-MOO"],
            ["比較對象", "Hand-crafted、Random configuration、Bayesian configuration、Meta-designed"],
            ["Problem class", "Constrained portfolio optimization unseen test instances"],
            ["Instance 數", "32"],
            ["Independent runs", "每個 method × instance 30 runs"],
            ["Population size", "N = 100"],
            ["Termination budget", "maxFE = 10000"],
            ["Configuration set", "L24 theta configuration set"],
            ["Hand-crafted", "使用固定人工設定 configuration"],
            ["Random", "從 L24 theta set 隨機選擇 configuration"],
            ["Bayesian", "在 validation split 上以 discrete Gaussian-process surrogate + expected improvement 搜尋單一 global theta，final test 全部 instances 使用同一 theta"],
            ["Meta-designed", "根據 instance features 選擇 theta configuration"],
            ["Reference front", "同一 instance 內所有方法與 runs 的 raw Pareto fronts union 建立 common reference front"],
            ["Output metrics", "HV、IGD、PF overlap、PF drift、diversity、runtime、RankScore"],
        ],
        [1.55, 4.97],
    )

    doc.add_heading("指標計算方式", level=1)
    add_table(
        doc,
        ["指標", "計算方式與解讀"],
        [
            ["HV", "以 common reference front 與 normalized objective values 計算 hypervolume，越高越好。"],
            ["IGD", "以 common reference front 為參考，計算 reference points 到 run PF 的平均最近距離，越低越好。"],
            ["PF Overlap", "對 common reference front 的每個點，檢查其到 run PF 的最近距離是否小於 tolerance；越高代表 run PF 對共同參考前緣的覆蓋度越好，並非 run-to-run overlap。"],
            ["PF Drift", "衡量 normalized PF centroid across runs 的偏移距離，越低代表 front 越穩定。"],
            ["Diversity", "衡量 final/archive front 的分布廣度，越高代表解集覆蓋較廣。"],
            ["Runtime", "直接讀取每次 run 的 runtime 並取平均，越低越好。"],
            ["RankScore", "在同一 instance 內，對六個 metrics 分別排序後取平均，越低越好。"],
            ["Overall RankScore", "先計算各方法的 mean metric，再對六個 mean metrics 排名並取平均，越低越好。"],
            ["First-place", "計算 OverallInstanceRank == 1 的 instances 數；並列第一不計入 strict first-place。"],
        ],
        [1.15, 5.37],
    )

    doc.add_heading("Result", level=1)

    doc.add_heading("Solution Quality", level=2)
    add_para(doc, "Meta-designed ECMADE-MOO 在 solution quality 指標上取得最高 mean HV 與最低 mean IGD，顯示 instance-aware theta selection 能改善 unseen instances 上的 Pareto-front 解品質與收斂表現。")
    add_table(
        doc,
        ["Method", "HV↑", "IGD↓"],
        method_rows(overall, ["HV", "IGD"]),
        [2.25, 2.1, 2.1],
    )

    doc.add_heading("PF Stability", level=2)
    add_para(doc, "PF stability 方面，Meta-designed 同時取得最高 PF Overlap 與最低 PF Drift，表示其產生的 front 與 common reference front 的重疊程度較高，且 across-run front drift 較小。")
    add_table(
        doc,
        ["Method", "PF Overlap", "PF Drift"],
        method_rows(overall, ["PF_Overlap", "PF_Drift"]),
        [2.25, 2.1, 2.1],
    )

    doc.add_heading("Configuration Strategy Ranking", level=2)
    add_para(doc, "為對應 Experiment A 的 overall method comparison，本節列出 mean RankScore、Overall RankScore 與 first-place instances。Overall RankScore 越低表示跨指標整體排名越好。")
    add_table(
        doc,
        ["Method", "Mean RankScore↓", "Mean Instance Rank↓", "Overall RankScore↓", "First-place"],
        [
            [
                METHOD_LABELS[row["method"]],
                fmt(row["mean_RankScore"], 3),
                fmt(row["mean_InstanceRank"], 3),
                fmt(row["overall_RankScore"], 3),
                f"{int(row['first_place_instances'])}/32",
            ]
            for _, row in overall.iterrows()
        ],
        [1.75, 1.25, 1.25, 1.25, 1.02],
    )

    doc.add_heading("Search Behavior", level=2)
    meta_usage = theta_usage[theta_usage["method"] == "MetaDesigned_ECMADE_MOO"].copy()
    add_para(doc, "Search behavior 在本實驗中主要由 diversity 與 theta usage 觀察。Random configuration 的 mean diversity 最高，但 Meta-designed 在 diversity 排名第 2，且其 theta usage 顯示 meta-learner 並非固定選單一參數，而是依 instance 分派多個 theta configuration。")
    add_table(
        doc,
        ["Method", "Diversity", "Dominant theta / usage"],
        [
            [
                METHOD_LABELS[row["method"]],
                fmt(row["mean_Diversity"], 4),
                ", ".join(
                    theta_usage[theta_usage["method"] == row["method"]]
                    .sort_values("instances", ascending=False)
                    .head(2)
                    .apply(lambda r: f"{r['theta_id']} ({int(r['instances'])})", axis=1)
                    .tolist()
                ),
            ]
            for _, row in overall.iterrows()
        ],
        [1.65, 1.25, 3.62],
    )

    doc.add_heading("Computational Cost", level=2)
    add_para(doc, "Runtime 結果顯示 Bayesian configuration 在 final test 上最快，Meta-designed 並非最低 runtime；因此 Experiment B 的主要貢獻應描述為品質與穩定性提升，而非速度最佳。")
    add_table(
        doc,
        ["Method", "Runtime↓"],
        method_rows(overall, ["Runtime"]),
        [3.25, 3.25],
    )

    doc.add_heading("視覺化分析", level=2)
    figures = [
        ("Figure B1", "Overall RankScore", FIG_DIR / "overall_rank_score.png"),
        ("Figure B2", "Per-instance RankScore boxplot", FIG_DIR / "per_instance_rankscore_boxplot.png"),
        ("Figure B3", "Meta-designed median improvement by metric", FIG_DIR / "meta_median_improvement_by_metric.png"),
        ("Figure B4", "Meta-designed theta usage", FIG_DIR / "meta_theta_usage.png"),
    ]
    for label, caption, path in figures:
        if path.exists():
            p = doc.add_paragraph()
            r = p.add_run(f"{label}. {caption}")
            r.bold = True
            set_ea(r)
            doc.add_picture(str(path), width=Inches(5.85))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("統計檢定摘要", level=1)
    add_para(doc, "本報告採用與 Experiment A 相同的統計呈現邏輯：先報 Friedman overall test，確認四種 strategy 在同一 metric 上是否存在整體差異；再報 Meta-designed ECMADE-MOO 相對最佳 baseline 的 Wilcoxon signed-rank post-hoc test，並加入 Holm correction。")
    add_para(doc, "Wilcoxon post-hoc test 採 one-sided alternative，檢定方向固定為 Meta-designed 是否優於 baseline。對 Runtime 而言，差值定義為 baseline runtime - Meta-designed runtime，因此正值代表 Meta-designed 較快；若 Meta-designed 較慢，單尾 p-value 會接近 1。Cliff's δ 則保留方向性，負值代表 baseline 較佔優勢。")

    doc.add_heading("Friedman 整體差異檢定", level=2)
    add_table(
        doc,
        ["Metric", "Statistic", "p-value"],
        [
            [row["metric"], fmt(row["friedman_chi_square"], 3), fmt_p(row["p_value"])]
            for _, row in friedman.iterrows()
        ],
        [2.05, 2.1, 2.35],
    )

    doc.add_heading("Wilcoxon signed-rank test", level=2)
    posthoc_rows = []
    for metric in ["HV", "IGD", "PF_Overlap", "PF_Drift", "Diversity", "Runtime", "RankScore"]:
        baseline = best_baseline(overall, metric)
        test_row = stats[(stats["metric"] == metric) & (stats["baseline"] == baseline)].iloc[0]
        fr = friedman[friedman["metric"] == metric].iloc[0]
        posthoc_rows.append(
            [
                metric,
                fmt_p(fr["p_value"]),
                "Yes" if fr["p_value"] < 0.05 else "No",
                METHOD_LABELS[baseline],
                fmt_p(test_row["holm_p_value"]),
                fmt(test_row["cliffs_delta"], 3),
                result_symbol(test_row),
            ]
        )
    add_table(
        doc,
        ["Metric", "Friedman p", "Sig.", "Meta-designed vs Best Baseline", "Holm adj. p", "Cliff's δ", "Result"],
        posthoc_rows,
        [0.85, 1.0, 0.55, 2.1, 0.95, 0.65, 0.42],
    )
    add_para(doc, "Result 欄位中，+ 表示 Meta-designed 顯著優於該 baseline；= 表示未達顯著差異；- 表示最佳 baseline 在該指標方向上較佔優勢。")

    doc.add_heading("Experiment A → Experiment B", level=1)
    add_table(
        doc,
        ["Experiment A Finding", "Interpretation", "Experiment B Design"],
        [
            ["ECMADE-MOO 在 OR-Library 上具有高 HV 與低 IGD", "ECMADE-MOO 可作為具競爭力的 search engine", "Experiment B 保留 ECMADE-MOO 作為底層演算法，只比較 configuration selection strategy"],
            ["Synthetic instances 上固定 configuration 的表現較受 instance distribution 影響", "固定手調參數存在 problem-dependent limitation", "導入 instance features，讓 Meta-designed strategy 依問題特性選 theta"],
            ["高 solution quality 不必然伴隨最高 repeated-run stability", "需要同時檢查 PF overlap、PF drift 與 diversity", "Experiment B 同時報告 quality、stability、search behavior 與 runtime"],
            ["Experiment A 的 baseline comparison 提供演算法層級比較", "下一步需檢查 ECMADE-MOO 內部參數選擇是否可改善泛化", "Experiment B 將比較單位從 algorithm baseline 轉為 configuration strategy baseline"],
            ["RankScore 可作為輔助總結，但不能取代分構面解讀", "不同 metric 反映不同研究問題", "Experiment B 同時提供 Overall RankScore 與分構面表格，避免只用單一數字下結論"],
        ],
        [1.85, 2.2, 2.47],
    )

    doc.add_heading("結論", level=1)
    add_para(doc, "Experiment B 的結果顯示，Meta-designed ECMADE-MOO 在四種 configuration strategy 中取得最低 Overall RankScore，並在 HV、IGD、PF Overlap 與 PF Drift 上皆取得最佳 overall rank。")
    add_para(doc, "相較於 Hand-crafted 與 Random baselines，Meta-designed strategy 說明 theta configuration selection 不應只依賴固定人工設定或隨機抽樣；instance-aware selection 能更穩定地改善 unseen instances 上的 front quality 與 stability。")
    add_para(doc, "相較於 Bayesian configuration，Meta-designed 在整體 RankScore 與多數品質/穩定性指標上仍具優勢，但 runtime 並非最佳，因此論文敘事應聚焦於 generalization 與 front quality/stability trade-off，而非宣稱速度最佳。")
    add_para(doc, "整體而言，Experiment B 可作為 Experiment A 之後的延伸：當 Experiment A 指出 fixed ECMADE-MOO configuration 在 heterogeneous instances 上仍有侷限時，Experiment B 進一步證明 meta-designed configuration selection 能改善 ECMADE-MOO 的 unseen-instance generalization。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Experiment B data report")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)
    set_ea(fr)

    doc.save(OUT_PATH)
    print(f"DOCX={OUT_PATH}")


if __name__ == "__main__":
    main()

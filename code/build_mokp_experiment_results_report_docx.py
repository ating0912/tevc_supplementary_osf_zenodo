from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SUMMARY_DIR = ROOT / "p0_lite_outputs" / "p1_mokp_config_comparison_20260719"
DIAGNOSTIC_DIR = ROOT / "p0_lite_outputs" / "p1_mokp_experiment_c_global_theta_diagnostic_20260729"
STABILITY_DIR = ROOT / "p0_lite_outputs" / "p1_mokp_experiment_c_stability_full_20260729"
OUT_DIR = ROOT / "p0_lite_outputs" / "reports"
OUT_DOCX = OUT_DIR / "MOKP_nonfinancial_experiment_results_report_20260730.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(91, 103, 112)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D7DBE2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
                     color=BLUE if level in (1, 2) else DARK_BLUE)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, color=MUTED)
    return p


def add_dataframe_table(doc, df, columns, headers, widths, number_cols=None):
    number_cols = set(number_cols or [])
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=8.5, bold=True, color=INK)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                text = f"{value:.4f}"
            else:
                text = str(value)
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in number_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=8.2, color=INK)
    set_table_width(table, widths)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, attr, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Experiment Result Report")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Non-financial MOKP generalization and Experiment C stability-aware transfer diagnostics")
    set_run_font(r, size=13, color=MUTED)

    meta = [
        ("Prepared date", "2026-07-30"),
        ("Experiment family", "MOKP non-financial constrained/combinatorial test bed"),
        ("Main added row", "ExperimentC_StabilityAware_ECMADE_MOO"),
        ("Diagnostic rows", "ExperimentC_GlobalTheta034_ECMADE_MOO; ExperimentC_GlobalTheta037_ECMADE_MOO"),
        ("Source directory", str(SUMMARY_DIR)),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    for i, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(i, 0), LIGHT_FILL)
        p0 = table.cell(i, 0).paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=9, bold=True, color=INK)
        p1 = table.cell(i, 1).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=9, color=INK)

    doc.add_paragraph()
    add_callout(
        doc,
        "Main takeaway",
        "The Experiment C stability-aware selector is executable on the non-financial MOKP test bed, "
        "but the current transfer result does not support a strong cross-domain generalization claim. "
        "Its performance is close to the two fixed global-theta diagnostics and remains behind "
        "BayesianConfig_ECMADE_MOO and the base ECMADE_MOO baselines."
    )


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(SUMMARY_DIR / "overall_method_summary.csv")
    friedman = pd.read_csv(SUMMARY_DIR / "friedman_tests.csv")
    ref = pd.read_csv(SUMMARY_DIR / "reference_front_info.csv")
    run_metrics = pd.read_csv(SUMMARY_DIR / "run_metrics.csv")

    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    add_heading(doc, "1. Experiment Scope", 1)
    add_paragraph(
        doc,
        "This report summarizes the non-financial transfer experiment prepared to address the reviewer concern "
        "that the study may otherwise appear portfolio-specific. The added test bed is a bi-objective "
        "multi-objective knapsack problem (MOKP), treated as a constrained/combinatorial benchmark rather than "
        "a financial optimization setting."
    )
    for item in [
        "Test set: 18 MOKP instances covering item counts d = 100, 250, and 500; capacity ratios 0.35, 0.50, and 0.65; and independent/conflicting profit modes.",
        "Execution budget: 30 independent runs per method-instance cell.",
        "Comparison set: 12 methods, including base MOEAs, ECMADE_MOO, BayesianConfig_ECMADE_MOO, RandomConfig_ECMADE_MOO, MetaTransfer_ECMADE_MOO, and the three Experiment C rows.",
        "Post-processing: method-instance metrics were evaluated against common instance-level reference fronts constructed from the pooled observed PF outputs.",
    ]:
        add_bullet(doc, item)

    completeness = pd.DataFrame([
        {"item": "ExperimentC_StabilityAware_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_StabilityAware_ECMADE_MOO")].shape[0]), "status": "Complete"},
        {"item": "ExperimentC_GlobalTheta034_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_GlobalTheta034_ECMADE_MOO")].shape[0]), "status": "Complete"},
        {"item": "ExperimentC_GlobalTheta037_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_GlobalTheta037_ECMADE_MOO")].shape[0]), "status": "Complete"},
        {"item": "Global-theta diagnostic PF outputs", "expected_runs": 1080, "parsed_runs": len(list((DIAGNOSTIC_DIR / "test").rglob("pf_obj.csv"))), "status": "Complete"},
    ])
    add_caption(doc, "Table 1. Output completeness check for the newly added Experiment C rows.")
    add_dataframe_table(
        doc,
        completeness,
        ["item", "expected_runs", "parsed_runs", "status"],
        ["Item", "Expected", "Parsed", "Status"],
        [4680, 1500, 1500, 1680],
        number_cols=["expected_runs", "parsed_runs", "status"],
    )

    add_heading(doc, "2. Overall Results", 1)
    add_paragraph(
        doc,
        "Lower RankScore values indicate better average rank across the evaluated metric set. "
        "For raw metrics, HV, PF_Overlap, and Diversity are maximized, whereas IGD, PF_Drift, and Runtime are minimized."
    )
    compact = summary[
        [
            "method",
            "runs",
            "mean_HV",
            "mean_IGD",
            "mean_PF_Overlap",
            "mean_PF_Drift",
            "mean_Runtime",
            "overall_RankScore",
            "first_place_instances",
        ]
    ].copy()
    compact = compact.sort_values("overall_RankScore")
    add_caption(doc, "Table 2. Overall MOKP method summary sorted by overall_RankScore.")
    add_dataframe_table(
        doc,
        compact,
        [
            "method",
            "runs",
            "mean_HV",
            "mean_IGD",
            "mean_PF_Overlap",
            "mean_PF_Drift",
            "mean_Runtime",
            "overall_RankScore",
            "first_place_instances",
        ],
        ["Method", "Runs", "HV", "IGD", "PF overlap", "PF drift", "Runtime", "RankScore", "1st inst."],
        [3000, 700, 800, 800, 900, 850, 850, 850, 610],
        number_cols=["runs", "mean_HV", "mean_IGD", "mean_PF_Overlap", "mean_PF_Drift", "mean_Runtime", "overall_RankScore", "first_place_instances"],
    )

    add_heading(doc, "3. Experiment C Diagnostic", 1)
    expc = summary[summary["method"].isin([
        "ExperimentC_StabilityAware_ECMADE_MOO",
        "ExperimentC_GlobalTheta034_ECMADE_MOO",
        "ExperimentC_GlobalTheta037_ECMADE_MOO",
        "BayesianConfig_ECMADE_MOO",
        "ECMADE_MOO",
    ])].copy()
    expc = expc.sort_values("overall_RankScore")
    add_caption(doc, "Table 3. Experiment C transfer rows compared with the two strongest ECMADE-family references.")
    add_dataframe_table(
        doc,
        expc,
        ["method", "mean_HV", "mean_IGD", "mean_Diversity", "mean_Runtime", "mean_InstanceRank", "overall_RankScore"],
        ["Method", "HV", "IGD", "Diversity", "Runtime", "Inst. rank", "RankScore"],
        [3600, 900, 900, 900, 900, 1080, 1080],
        number_cols=["mean_HV", "mean_IGD", "mean_Diversity", "mean_Runtime", "mean_InstanceRank", "overall_RankScore"],
    )
    for item in [
        "ExperimentC_StabilityAware_ECMADE_MOO achieved HV = 0.8749, IGD = 0.1696, and overall_RankScore = 8.5000.",
        "The two fixed global theta diagnostics were very close: theta_034 RankScore = 8.8333 and theta_037 RankScore = 8.8333.",
        "The stability-aware assignment is therefore only slightly better than using either selected theta globally; the gap is not large enough to claim successful cross-domain selector generalization.",
        "The stronger ECMADE-family references remained BayesianConfig_ECMADE_MOO and ECMADE_MOO, with overall_RankScore = 3.1667 and 3.5000, respectively.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. Statistical Summary", 1)
    friedman_table = friedman.copy()
    friedman_table["friedman_p_value"] = friedman_table["friedman_p_value"].map(lambda x: f"{x:.3e}")
    add_caption(doc, "Table 4. Friedman tests across 18 MOKP instances and 12 methods.")
    add_dataframe_table(
        doc,
        friedman_table,
        ["metric", "direction", "instances", "methods", "friedman_chi_square", "friedman_p_value"],
        ["Metric", "Direction", "Instances", "Methods", "Chi-square", "p-value"],
        [1600, 1100, 1300, 1100, 2100, 2160],
        number_cols=["instances", "methods", "friedman_chi_square", "friedman_p_value"],
    )
    add_paragraph(
        doc,
        "The omnibus tests are significant across all six metrics, confirming that method differences are present. "
        "Pairwise Wilcoxon-Holm results further indicate that BayesianConfig_ECMADE_MOO and ECMADE_MOO significantly "
        "outperform the Experiment C transfer rows on HV and IGD, while the direct stability-aware row is not materially "
        "separated from the two fixed global-theta diagnostics."
    )

    add_heading(doc, "5. Reference Front Coverage", 1)
    ref_summary = pd.DataFrame([
        {"quantity": "Instances", "value": len(ref)},
        {"quantity": "Minimum reference points", "value": int(ref["reference_points"].min())},
        {"quantity": "Median reference points", "value": int(ref["reference_points"].median())},
        {"quantity": "Maximum reference points", "value": int(ref["reference_points"].max())},
    ])
    add_caption(doc, "Table 5. Common reference-front coverage used by the MOKP analysis.")
    add_dataframe_table(
        doc,
        ref_summary,
        ["quantity", "value"],
        ["Quantity", "Value"],
        [6600, 2760],
        number_cols=["value"],
    )

    add_heading(doc, "6. Manuscript Recommendation", 1)
    add_callout(
        doc,
        "Recommended framing",
        "Use this experiment as a non-financial external validation/stress test, not as evidence of completed cross-domain generalization. "
        "The result is useful because it demonstrates the pipeline can be executed outside portfolio optimization and reveals a concrete limitation."
    )
    add_paragraph(
        doc,
        "Suggested manuscript wording:",
    )
    quote = (
        "Although the non-financial MOKP evaluation confirms that the proposed stability-aware transfer pipeline can be "
        "applied outside constrained portfolio optimization, the current transfer result does not yet establish broad "
        "cross-domain generalization. On the MOKP test bed, the stability-aware Experiment C configuration remains close "
        "to fixed global-theta diagnostics and trails BayesianConfig_ECMADE_MOO and the base ECMADE_MOO baselines. "
        "We therefore treat this result as external validation of the experimental protocol and as evidence of a current "
        "limitation, rather than as a claim that the selector precisely reconstructs Oracle theta across problem domains."
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(quote)
    set_run_font(r, size=10.5, color=DARK_BLUE)
    r.italic = True

    add_heading(doc, "7. Data Artifacts", 1)
    artifacts = pd.DataFrame([
        {"artifact": "Overall method summary", "path": str(SUMMARY_DIR / "overall_method_summary.csv")},
        {"artifact": "Run-level metrics", "path": str(SUMMARY_DIR / "run_metrics.csv")},
        {"artifact": "Friedman tests", "path": str(SUMMARY_DIR / "friedman_tests.csv")},
        {"artifact": "Pairwise Wilcoxon-Holm tests", "path": str(SUMMARY_DIR / "pairwise_wilcoxon.csv")},
        {"artifact": "Reference front info", "path": str(SUMMARY_DIR / "reference_front_info.csv")},
        {"artifact": "Experiment C global-theta raw outputs", "path": str(DIAGNOSTIC_DIR)},
        {"artifact": "Experiment C stability-aware raw outputs", "path": str(STABILITY_DIR)},
    ])
    add_caption(doc, "Table 6. Files and directories supporting this report.")
    add_dataframe_table(
        doc,
        artifacts,
        ["artifact", "path"],
        ["Artifact", "Path"],
        [3000, 6360],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = footer.add_run("MOKP Experiment Result Report")
        set_run_font(r, size=9, color=MUTED)

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_report())

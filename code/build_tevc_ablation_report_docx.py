from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE_DIR = ROOT / "docx_outputs"
OUT_PATH = SOURCE_DIR / "TEVC_消融實驗數據報告_20260720.docx"


def find_source_markdown() -> Path:
    matches = sorted(SOURCE_DIR.glob("TEVC_*20260720.md"))
    if not matches:
        raise FileNotFoundError("Could not find TEVC markdown report in docx_outputs")
    return matches[0]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def mark_repeating_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, after=6, line_spacing=1.10, color=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_markdown_inline(paragraph, text: str, size=11, bold_default=False, color=None) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, bold=True, color=color or "1F4D78")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold_default, color=color)


def table_widths(headers: list[str]) -> list[float]:
    n = len(headers)
    if n == 6 and headers[0] == "目標":
        return [1.55, 0.95, 0.95, 1.10, 1.10, 0.85]
    if n == 7 and headers[0] == "selector":
        return [1.85, 0.70, 0.70, 0.85, 0.90, 0.75, 0.75]
    if n == 6 and headers[0] == "factor":
        return [1.35, 0.90, 0.95, 0.85, 0.85, 1.60]
    return [6.5 / max(n, 1)] * n


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = rows[0]
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, table_widths(headers))
    mark_repeating_header(table.rows[0])
    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_inline(p, value, size=8.5 if len(headers) >= 6 else 10, bold_default=(r_idx == 0))
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
    doc.add_paragraph()


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def build_docx(markdown: str) -> Document:
    doc = Document()
    apply_document_styles(doc)
    lines = markdown.splitlines()

    # Use the H1 as a polished title paragraph instead of Word's built-in Title style.
    first_title_done = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_markdown_table(lines, i)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and not first_title_done:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_markdown_inline(p, text, size=20, bold_default=True, color="0B2545")
                first_title_done = True
            else:
                p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
                add_markdown_inline(p, text, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold_default=True)
            i += 1
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            add_markdown_inline(p, numbered_match.group(2), size=11)
            style_paragraph(p, size=11, after=4, line_spacing=1.167)
            i += 1
            continue

        p = doc.add_paragraph()
        add_markdown_inline(p, stripped, size=11)
        style_paragraph(p, size=11, after=6, line_spacing=1.10)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("TEVC 消融實驗數據報告")
    set_run_font(run, size=9, color="666666")
    return doc


def main() -> None:
    source = find_source_markdown()
    markdown = source.read_text(encoding="utf-8")
    doc = build_docx(markdown)
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()

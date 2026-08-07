from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "TEVC_P0_P1_紅字補數據對照表_20260724.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: object, bold: bool = False, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("" if pd.isna(text) else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_table(doc: Document, title: str, columns: list[str], rows: list[list[object]], font_size: int = 8) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[idx], col, bold=True, size=font_size)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()


def add_df_table(
    doc: Document,
    title: str,
    frame: pd.DataFrame,
    columns: list[str] | None = None,
    rename: dict[str, str] | None = None,
    max_rows: int | None = None,
    font_size: int = 8,
) -> None:
    df = frame.copy()
    if columns is not None:
        df = df[columns]
    if max_rows is not None:
        df = df.head(max_rows)
    if rename:
        df = df.rename(columns=rename)
    rows = df.fillna("NA").values.tolist()
    add_table(doc, title, list(df.columns), rows, font_size=font_size)


def add_note(doc: Document, text: str, color: str = "FFF2CC") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_text(cell, text, size=9)
    doc.add_paragraph()


def add_source(doc: Document, path: str) -> None:
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    run = p.add_run(f"資料檔：{path}")
    run.font.size = Pt(8)
    run.font.name = "Consolas"


def df_csv(path: str) -> pd.DataFrame:
    return pd.read_csv(ROOT / path, encoding="utf-8-sig")


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft JhengHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TEVC P0/P1 紅字待補數據對照表")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph("用途：說明每一段紅字待補處應放入哪些已產生的 P0/P1 數據、表格與建議文字。")
    doc.add_paragraph("來源 Markdown：TEVC_P0_P1_補稿表格文字_20260723.md")
    doc.add_paragraph("產出日期：2026-07-24")

    add_note(
        doc,
        "使用方式：先看第 1 張總表，找到母稿中的紅字錨點；再依「要補的數據/表格」與「建議補入文字」貼到對應段落。"
        "P0 多數補在 Method/Experiment C/SI；P1 多數補在 Generalization、Real-market validation、Discussion/Conclusion。",
    )

    mapping_rows = [
        [
            "【尚待加強｜theta 表】",
            "P0",
            "Table S1. 24/L24 theta configuration table；包含 S、operator、migration、elite ratio、stagnation threshold、archive/constraint handling、Experiment C test count。",
            "Method / Experimental design / Supplementary Table S1",
            "p0_lite_outputs/theta_configuration_paper_table_20260723/theta_configuration_table_for_paper.csv",
            "說明 selector 實際選的是具可解釋演算法意義的 theta configuration，而不是黑箱代號。",
        ],
        [
            "【尚待加強｜統計檢定】",
            "P0",
            "Table S2. Experiment C Friedman tests；pairwise Wilcoxon + Holm correction。",
            "Results / Experiment C / Statistical validation",
            "p0_lite_outputs/tevc_p0_statistical_tests_20260718/unified_friedman_tests.csv；unified_pairwise_wilcoxon.csv",
            "寫成 primary endpoint 與支援指標上的 Holm-corrected paired evidence，避免只用未校正 sign-test。",
        ],
        [
            "【尚待加強｜feature importance】",
            "P0",
            "Table S4/S5. Top feature importances；selector validation summary。",
            "Results / Selector interpretability",
            "p0_lite_outputs/experiment_c_stability_selector_training/feature_importance.csv；validation_selector_summary.csv",
            "強調 selector 使用 instance features 與 algorithm-design features；但 top1/top3 hit rate = 0，不可寫精準重建 Oracle。",
        ],
        [
            "【尚待加強｜label protocol】",
            "P0",
            "Training rows = 3216；validation rows = 696；theta candidates = 24；label-generation optimizer runs 與 runtime。",
            "Method / Label generation protocol / Reproducibility",
            "p0_lite_outputs/experiment_c_stability_selector_training/training_config.json；p0_lite_outputs/tevc_cost_runtime_summary_20260723/",
            "說明 label rows 如何產生、training/validation/test split，以及 test instances 不參與 label selection。",
        ],
        [
            "【尚待加強｜統一命名】",
            "P0",
            "Table S3. RankScore / C_LabelScore / C_ThetaRank 定義、ties、normalization scope。",
            "Method / Evaluation metrics",
            "TEVC_P0_P1_補稿表格文字_20260723.md 的 P0-3；相關 scored/ranked CSV",
            "統一用 RankScore 與 C_LabelScore；C_LabelScore = -0.2 rank_HV - 0.2 rank_IGD - 0.3 rank_PF_Overlap - 0.3 rank_PF_Drift。",
        ],
        [
            "【尚待加強｜real-market validation】與 4.1",
            "P1",
            "Table S8 rolling-window validation；Table S9 CVaR95；Table S10 10/20/50 bps cost sensitivity。",
            "Results / Real-market validation / Supplementary risk table",
            "p0_lite_outputs/p1_rolling_window_market_validation_20260719/summary/；cvar_sensitivity/",
            "寫成 robustness / limitation evidence。ECMADE_MOO CVaR95 loss 排第 1，但 after-cost return 不是最佳。",
        ],
        [
            "【尚待加強｜real-market validation 執行】",
            "P1",
            "33 rolling windows；annual net return、Sharpe、Sortino、MDD、turnover、runtime、CVaR95、10/20/50 bps sensitivity。",
            "4.1 Real-market validation 補強設計",
            "summary/method_overall_summary.csv；cvar_sensitivity/method_cvar_overall_summary.csv；method_transaction_cost_sensitivity.csv",
            "補入最小可行版本：rolling-window 設計 + out-of-sample 指標 + 交易成本敏感度。",
        ],
        [
            "【尚待加強｜非金融泛化】",
            "P1",
            "Table S7. MOKP non-financial constrained/combinatorial test bed。",
            "Results / Generalization beyond portfolio optimization",
            "p0_lite_outputs/p1_mokp_analysis_independent_20260719/overall_method_summary.csv",
            "18 MOKP instances，每方法 540 runs；ECMADE_MOO overall RankScore = 2.000，first-place = 13/18。",
        ],
        [
            "【尚待加強｜結論語氣】",
            "P0/P1",
            "整合 P0/P1 結果限制：selector 非 Oracle reconstruction；real-market return 非全面勝出。",
            "Discussion / Conclusion",
            "TEVC_P0_P1_補稿表格文字_20260723.md 的 P1-5",
            "寫成：MOKP 泛化成立，real-market downside risk 較穩；return dominance remains problem-dependent。",
        ],
        [
            "【尚待加強｜可重現性 package】",
            "P0/P1",
            "code、README、environment.yml、MATLAB/PlatEMO version、CPU/GPU/OS、run logs、tables、figures、raw PF csv。",
            "Supplementary package / Reproducibility statement",
            "p0_lite_outputs/；各 summary/README 檔",
            "這不是新實驗數據，需整理 supplementary package 與路徑說明。",
        ],
        [
            "【尚待加強｜數學式】",
            "非數據",
            "objective functions、risk definition、return estimation、weight bounds、repair operator、constraint handling 公式。",
            "Method / Problem formulation",
            "需人工補公式，不由 P0/P1 數據替代",
            "這裡不是貼結果表，而是補完整問題定義與約束處理公式。",
        ],
        [
            "【尚待加強｜演算法細節】",
            "非數據",
            "Algorithm 1 ECMADE-MOO；Algorithm 2 theta selection；Algorithm 3 stability-aware label generation。",
            "Method / Algorithm",
            "可參照 TEVC_P0_P1_補稿表格文字_20260723.md 文字與 theta 表",
            "這裡需補 pseudo-code 與 archive/migration/stagnation/repair 流程。",
        ],
        [
            "【Reference QC 待辦】與 P2 文獻",
            "P2/非數據",
            "補 DOI、卷期頁碼、IEEE TEVC bibliography style；補 2024-2026 最新文獻。",
            "References / Related work",
            "需另做文獻 QC",
            "這裡不放實驗數據，放最新文獻與正式出版資訊。",
        ],
    ]
    add_table(
        doc,
        "1. 紅字位置與應補數據對照總表",
        ["紅字錨點", "分類", "要補的數據/表格", "建議放置位置", "資料檔/來源", "建議寫法"],
        mapping_rows,
        font_size=7,
    )

    doc.add_heading("2. P0 可直接放入的表格與文字", level=1)
    theta = df_csv("p0_lite_outputs/theta_configuration_paper_table_20260723/theta_configuration_table_for_paper.csv")
    add_df_table(
        doc,
        "Table S1. 24/L24 theta configuration table",
        theta,
        columns=[
            "L24_row",
            "paper_theta_id",
            "source_theta_id",
            "subpops",
            "operator",
            "migration",
            "elite_ratio",
            "stagnation_threshold",
            "archive_strategy",
            "constraint_handling",
            "experiment_c_test_assignment_count",
        ],
        font_size=7,
    )
    add_source(doc, "p0_lite_outputs/theta_configuration_paper_table_20260723/theta_configuration_table_for_paper.csv")
    add_note(
        doc,
        "應補在紅字：【尚待加強｜theta 表】、表格待辦列「補完整 theta configuration table 與 theta encoding」。"
        "正文可寫：theta jointly specifies subpopulations, operator mode, migration mode, elite ratio, and stagnation threshold。",
    )

    friedman = df_csv("p0_lite_outputs/tevc_p0_statistical_tests_20260718/unified_friedman_tests.csv")
    friedman_c = friedman[friedman["scope"] == "Experiment_C_StabilityComparison"]
    add_df_table(
        doc,
        "Table S2. Experiment C Friedman tests",
        friedman_c,
        columns=[
            "metric",
            "direction",
            "instances",
            "methods",
            "friedman_chi_square",
            "friedman_p_value",
            "nemenyi_cd_alpha_0_05",
        ],
        rename={
            "metric": "Metric",
            "direction": "Direction",
            "instances": "Instances",
            "methods": "Methods",
            "friedman_chi_square": "Friedman chi-square",
            "friedman_p_value": "p-value",
            "nemenyi_cd_alpha_0_05": "Nemenyi CD",
        },
        font_size=8,
    )
    add_source(doc, "p0_lite_outputs/tevc_p0_statistical_tests_20260718/unified_friedman_tests.csv")
    add_note(
        doc,
        "應補在紅字：【尚待加強｜統計檢定】、表格待辦列「補 Experiment C 的 Holm correction 或 primary endpoint paired test」。"
        "Pairwise Holm 校正表請引用 unified_pairwise_wilcoxon.csv。",
    )

    feature = df_csv("p0_lite_outputs/experiment_c_stability_selector_training/feature_importance.csv")
    add_df_table(
        doc,
        "Table S4. Top feature importances",
        feature,
        columns=["feature", "importance"],
        rename={"feature": "Feature", "importance": "Importance"},
        max_rows=10,
        font_size=8,
    )
    selector = df_csv("p0_lite_outputs/experiment_c_stability_selector_training/validation_selector_summary.csv")
    add_df_table(
        doc,
        "Table S5. Experiment C selector validation",
        selector,
        font_size=7,
    )
    add_source(doc, "p0_lite_outputs/experiment_c_stability_selector_training/feature_importance.csv；validation_selector_summary.csv")
    add_note(
        doc,
        "應補在紅字：【尚待加強｜feature importance】與【尚待加強｜label protocol】。"
        "注意 validation top1/top3 hit rate = 0，結論要寫成 stability-aware trade-off selector。",
    )

    score_rows = [
        [
            "RankScore",
            "Experiment A/B general ranking",
            "lower is better",
            "arithmetic mean of metric ranks",
            "average rank",
            "within each instance-K group",
        ],
        [
            "C_LabelScore",
            "Experiment C stability label",
            "higher is better",
            "-0.2 rank_HV - 0.2 rank_IGD - 0.3 rank_PF_Overlap - 0.3 rank_PF_Drift",
            "average rank",
            "within each instance-K group",
        ],
        [
            "C_ThetaRank",
            "Experiment C theta order",
            "lower is better",
            "rank of C_LabelScore in descending order",
            "first/average according to table generation",
            "within each instance-K group",
        ],
    ]
    add_table(
        doc,
        "Table S3. Unified score naming and formula",
        ["Score name", "Scope", "Direction", "Formula / rule", "Tie handling", "Normalization scope"],
        score_rows,
        font_size=8,
    )
    add_note(
        doc,
        "應補在紅字：【尚待加強｜統一命名】。建議修稿後避免再混用 OverallRankScore / J-score，正文固定用 RankScore 與 C_LabelScore。",
    )

    cost = df_csv("p0_lite_outputs/tevc_cost_runtime_summary_20260723/label_generation_runtime_summary.csv")
    add_df_table(
        doc,
        "Table S6. Offline label-generation cost",
        cost,
        columns=[
            "cost_component",
            "label_rows",
            "runs_per_label",
            "estimated_optimizer_runs",
            "total_runtime_hours",
            "mean_runtime_sec_per_optimizer_run",
        ],
        rename={
            "cost_component": "Component",
            "label_rows": "Label rows",
            "runs_per_label": "Runs per label",
            "estimated_optimizer_runs": "Estimated optimizer runs",
            "total_runtime_hours": "Total runtime hours",
            "mean_runtime_sec_per_optimizer_run": "Mean sec/run",
        },
        font_size=8,
    )
    add_source(doc, "p0_lite_outputs/tevc_cost_runtime_summary_20260723/label_generation_runtime_summary.csv")
    add_note(
        doc,
        "應補在紅字：表格待辦列「補 configuration cost 與 meta-training cost」、【尚待加強｜可重現性 package】。"
        "文字要分開 offline label-generation cost 與 online final-optimization runtime。",
    )

    doc.add_heading("3. P1 可直接放入的表格與文字", level=1)
    mokp = df_csv("p0_lite_outputs/p1_mokp_analysis_independent_20260719/overall_method_summary.csv")
    add_df_table(
        doc,
        "Table S7. Non-financial MOKP test-bed results",
        mokp,
        columns=[
            "method",
            "instances",
            "runs",
            "mean_HV",
            "mean_IGD",
            "mean_PF_Overlap",
            "mean_PF_Drift",
            "mean_Runtime",
            "overall_RankScore",
            "first_place_instances",
        ],
        rename={
            "method": "Method",
            "instances": "Instances",
            "runs": "Runs",
            "mean_HV": "mean HV",
            "mean_IGD": "mean IGD",
            "mean_PF_Overlap": "mean PF_Overlap",
            "mean_PF_Drift": "mean PF_Drift",
            "mean_Runtime": "mean Runtime",
            "overall_RankScore": "Overall RankScore",
            "first_place_instances": "First-place instances",
        },
        font_size=7,
    )
    add_source(doc, "p0_lite_outputs/p1_mokp_analysis_independent_20260719/overall_method_summary.csv")
    add_note(
        doc,
        "應補在紅字：【尚待加強｜非金融泛化】。建議寫：18 MOKP instances、每方法 540 runs；ECMADE_MOO overall RankScore = 2.000，13/18 first-place。",
    )

    rolling = df_csv("p0_lite_outputs/p1_rolling_window_market_validation_20260719/summary/method_overall_summary.csv")
    add_df_table(
        doc,
        "Table S8. Real-market rolling-window validation",
        rolling,
        columns=[
            "method",
            "windows",
            "mean_annual_net_return",
            "mean_sharpe",
            "mean_sortino",
            "mean_max_drawdown",
            "mean_rebalance_turnover",
            "mean_runtime_sec",
            "mean_RankScore",
            "first_place_windows",
        ],
        rename={
            "method": "Method",
            "windows": "Windows",
            "mean_annual_net_return": "Annual net return",
            "mean_sharpe": "Sharpe",
            "mean_sortino": "Sortino",
            "mean_max_drawdown": "MDD",
            "mean_rebalance_turnover": "Turnover",
            "mean_runtime_sec": "Runtime sec",
            "mean_RankScore": "RankScore",
            "first_place_windows": "First-place windows",
        },
        font_size=7,
    )
    add_source(doc, "p0_lite_outputs/p1_rolling_window_market_validation_20260719/summary/method_overall_summary.csv")
    add_note(
        doc,
        "應補在紅字：【尚待加強｜real-market validation】、4.1、【尚待加強｜real-market validation 執行】。"
        "請寫成 robustness / limitation evidence；ECMADE_MOO 不是 annual return 最佳。",
    )

    cvar = df_csv("p0_lite_outputs/p1_rolling_window_market_validation_20260719/cvar_sensitivity/method_cvar_overall_summary.csv")
    add_df_table(
        doc,
        "Table S9. CVaR(95%) downside-risk summary",
        cvar,
        columns=["method", "windows", "mean_cvar95_return", "mean_cvar95_loss", "rank_cvar95_loss"],
        rename={
            "method": "Method",
            "windows": "Windows",
            "mean_cvar95_return": "mean CVaR95 return",
            "mean_cvar95_loss": "mean CVaR95 loss",
            "rank_cvar95_loss": "CVaR rank",
        },
        font_size=8,
    )
    add_source(doc, "p0_lite_outputs/p1_rolling_window_market_validation_20260719/cvar_sensitivity/method_cvar_overall_summary.csv")
    add_note(
        doc,
        "應補在 4.1 real-market validation 附近。重點：ECMADE_MOO mean CVaR95 loss 最低，顯示 downside-risk stability。",
    )

    sens = df_csv("p0_lite_outputs/p1_rolling_window_market_validation_20260719/cvar_sensitivity/method_transaction_cost_sensitivity.csv")
    sens_small = sens[sens["method"].isin(["MOEAD", "ECMADE_MOO"])].copy()
    add_df_table(
        doc,
        "Table S10. Transaction-cost sensitivity: key rows",
        sens_small,
        columns=[
            "cost_scenario",
            "method",
            "mean_scenario_annual_net_return",
            "mean_turnover",
            "rank_annual_net_return",
        ],
        rename={
            "cost_scenario": "Cost",
            "method": "Method",
            "mean_scenario_annual_net_return": "Annual net return",
            "mean_turnover": "Turnover",
            "rank_annual_net_return": "Rank",
        },
        font_size=8,
    )
    add_source(doc, "p0_lite_outputs/p1_rolling_window_market_validation_20260719/cvar_sensitivity/method_transaction_cost_sensitivity.csv")
    add_note(
        doc,
        "應補在 real-market validation 執行/敏感度段落。重點：10/20/50 bps 下 MOEAD annual net return 最高，ECMADE_MOO 排第 5；請保守解讀。",
    )

    doc.add_heading("4. 可直接貼入母稿的結論語氣", level=1)
    conclusion_rows = [
        [
            "P0 結論",
            "selector 具可解釋性並使用 instance/problem/algorithm-design features；但 top1/top3 hit rate 為 0，因此不可宣稱精準重建 Oracle。",
        ],
        [
            "P1 結論",
            "MOKP 顯示非金融 constrained/combinatorial 泛化能力；real-market 顯示 CVaR95 downside-risk 較穩，但 after-cost return 不是最佳。",
        ],
        [
            "總結語氣",
            "The proposed design generalizes beyond portfolio-only benchmarks and improves stability-oriented behavior, while real-market return dominance remains problem-dependent.",
        ],
    ]
    add_table(doc, "建議結論用語", ["段落", "建議文字"], conclusion_rows, font_size=9)

    doc.add_heading("5. 原始輸出檔索引", level=1)
    index_rows = [
        ["P0 統計檢定", "p0_lite_outputs/tevc_p0_statistical_tests_20260718/"],
        ["P0 feature importance / selector", "p0_lite_outputs/experiment_c_stability_selector_training/"],
        ["P0 theta table", "p0_lite_outputs/theta_configuration_paper_table_20260723/"],
        ["P0/P1 cost runtime", "p0_lite_outputs/tevc_cost_runtime_summary_20260723/"],
        ["P1 MOKP", "p0_lite_outputs/p1_mokp_analysis_independent_20260719/"],
        ["P1 rolling market", "p0_lite_outputs/p1_rolling_window_market_validation_20260719/summary/"],
        ["P1 CVaR and cost sensitivity", "p0_lite_outputs/p1_rolling_window_market_validation_20260719/cvar_sensitivity/"],
    ]
    add_table(doc, "資料檔索引", ["內容", "路徑"], index_rows, font_size=8)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

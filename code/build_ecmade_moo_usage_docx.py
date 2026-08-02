from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "ECMADE-MOO_程式參數與使用說明.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B7C9DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    dxa_widths = [int(w * 1440) for w in widths_in]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa_widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in dxa_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(dxa_widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_text(paragraph, text: str, bold=False, italic=False, color=None, size=None, font=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return run


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.style = "Code"
        add_text(p, line, font="Consolas", size=9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_widths(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], HEADER_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, text, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, text)
            set_cell_margins(cells[i])
    set_table_widths(table, widths)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    code = styles.add_style("Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.line_spacing = 1.05

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(footer, "ECMADE-MOO 使用說明", color="666666", size=9)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(p, "ECMADE-MOO 程式參數與使用說明", bold=True, color=DARK_BLUE, size=22)
    p = doc.add_paragraph()
    add_text(
        p,
        "適用於多目標投資組合最佳化、OR-Library 資料、ZDT、DTLZ 與 UF benchmark 的 Python 執行指南。",
        color="555555",
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    p = doc.add_paragraph()
    add_text(p, "程式位置：", bold=True)
    add_text(p, str(ROOT / "ecmade_moo.py"), font="Consolas", size=9)

    doc.add_heading("1. 程式用途", level=1)
    doc.add_paragraph(
        "ecmade_moo.py 實作 ECMADE-MOO：保留 ECMADE 的多子群協同演化、自適應 F/CR 參數控制與資訊交換機制，並以 NSGA-II 的 Pareto dominance、Fast Non-dominated Sorting、Crowding Distance 與 Elitism 完成多目標環境選擇。"
    )
    doc.add_paragraph(
        "在投資組合問題中，程式以風險最小化與負報酬最小化作為兩個目標；在 benchmark 問題中，程式直接輸出近似 Pareto Front 並可計算 IGD。"
    )

    doc.add_heading("2. 基本執行方式", level=1)
    add_code_block(
        doc,
        [
            "python ecmade_moo.py --problem ZDT1 --pop-size 100 --max-fe 10000",
            "python ecmade_moo.py --problem DTLZ2 --dimension 12 --objectives 3",
            "python ecmade_moo.py --problem UF10 --max-fe 30000",
            "python ecmade_moo.py --problems ZDT1 ZDT2 DTLZ1 DTLZ2 UF1 UF2",
        ],
    )

    doc.add_heading("3. 主要參數", level=1)
    add_table(
        doc,
        ["參數", "預設值", "說明"],
        [
            ["--problem", "ZDT1", "單一問題名稱。支援 ZDT1-4、ZDT6、DTLZ1-7、UF1-10、ORLIB、PORTFOLIO。"],
            ["--problems", "無", "批次執行多個 benchmark，例如 --problems ZDT1 DTLZ2 UF1。若設定此參數，會覆蓋 --problem。"],
            ["--pop-size", "100", "族群大小 N。環境選擇後每一代保留 N 個個體。"],
            ["--max-fe", "10000", "最大函數評估次數 MaxFE。初始族群評估也會計入。"],
            ["--subpops", "3", "子群數量 M。預設對應三種 DE 搜尋策略。"],
            ["--seed", "2026", "隨機種子，用於重現實驗結果。"],
            ["--dimension", "依問題預設", "決策變數維度 D。ZDT/UF 通常使用 30；DTLZ 可依需求調整。"],
            ["--objectives", "3 for DTLZ", "DTLZ 的目標數 M。ZDT/UF 目標數由問題定義決定。"],
            ["--out-dir", "ecmade_moo_outputs", "輸出資料夾。所有 population、Pareto front、history 與 metrics 會寫入此處。"],
            ["--pf-samples", "10000", "reference Pareto Front 取樣數，用於 benchmark IGD 計算。"],
        ],
        [1.35, 1.05, 4.1],
    )

    doc.add_heading("4. 演算法內部參數", level=1)
    add_table(
        doc,
        ["名稱", "預設值", "用途"],
        [
            ["archive_size", "20", "每個個體各自保存最近 H 筆成功 F/CR，形成 RSF 與 RSCR。"],
            ["theta", "1/13", "自適應更新係數；F 使用帶時間權重的 Lehmer mean，CR 使用帶時間權重的 arithmetic mean。"],
            ["stagnation_threshold", "50", "自適應資訊交換門檻。當 Pareto archive 連續超過 50 代沒有新增非支配解時，才觸發菁英複製與子群資訊交換。"],
            ["exploitation_alpha", "0.8", "論文公式 (16) 的 exploitation 權重 α，用於降低 best-guided 搜尋過早收斂的風險。"],
            ["initial_mu_f", "(0.9, 0.8, 0.8)", "各子群 F 的初始平均值。"],
            ["initial_mu_cr", "(0.9, 0.5, 0.5)", "各子群 CR 的初始平均值。"],
        ],
        [1.7, 1.35, 3.45],
    )
    doc.add_paragraph(
        "上述內部參數位於 ECMADEMOOConfig，可直接在 ecmade_moo.py 中修改；一般實驗優先調整 --pop-size、--max-fe、--subpops 與 --seed。"
    )
    doc.add_paragraph(
        "資訊交換採 ECMADE 論文精神的停滯觸發式設計：ECMADE-MOO 以 Pareto archive 是否新增非支配解判斷搜尋是否有進展；若連續停滯超過門檻，才將 Pareto 精英複製到各子群並取代較差個體。"
    )
    doc.add_paragraph(
        "Mutation 所需的 r1、r2、r3、r4、r5 皆從目標個體所屬子群內抽樣，索引互異且排除 target；因此每個子群至少需要 6 個個體，建議 pop-size 不小於 6 × subpops。"
    )
    add_table(
        doc,
        ["子群", "搜尋角色", "論文對應策略", "目前程式公式"],
        [
            ["p1", "Exploration", "DE/rand/2，公式 (15)", "v = x_r1 + F(x_r2 - x_r3) + F(x_r4 - x_r5)"],
            ["p2", "Exploitation", "best-guided DE/rand·best/2 變形，公式 (16)", "v = α·x_best + F(x_r1 - x_r2) + F(x_r3 - x_r4)"],
            ["p3", "Balance", "DE/rand/1 + DE/current-to-best/1，公式 (17)", "v = (1 - ω)·DE/rand/1 + ω·DE/current-to-best/1，ω = G/MG"],
        ],
        [0.75, 1.15, 2.05, 2.55],
    )

    doc.add_heading("5. 支援問題與預設維度", level=1)
    add_table(
        doc,
        ["類型", "支援名稱", "預設維度與目標數", "備註"],
        [
            ["ZDT", "ZDT1, ZDT2, ZDT3, ZDT4, ZDT6", "ZDT1-3: D=30, M=2；ZDT4/ZDT6: D=10, M=2", "ZDT4 的第 2 到 D 個變數範圍為 [-5, 5]。"],
            ["DTLZ", "DTLZ1-DTLZ7", "預設 M=3；DTLZ1: D=7；DTLZ7: D=22；其他: D=M+9", "可用 --dimension 與 --objectives 調整。"],
            ["UF", "UF1-UF10", "預設 D=30；UF1-7: M=2；UF8-10: M=3", "變數上下界依 CEC2009 UF 定義設定。"],
            ["ORLIB", "ORLIB", "依檔案資產數決定 D；M=2", "讀取 OR-Library portfolio 格式，輸出風險與負報酬。"],
            ["PORTFOLIO", "PORTFOLIO", "依 returns CSV 欄數決定 D；M=2", "CSV 每欄一個資產，每列一期報酬率。"],
        ],
        [1.0, 1.65, 2.1, 1.75],
    )

    doc.add_heading("6. 投資組合輸入格式", level=1)
    doc.add_heading("6.1 Returns CSV", level=2)
    doc.add_paragraph("每一欄代表一個資產，每一列代表一期報酬率。程式會自動計算平均報酬向量與共變異數矩陣。")
    add_code_block(
        doc,
        [
            "python ecmade_moo.py --problem PORTFOLIO --returns-csv returns.csv",
            "",
            "CSV 範例：",
            "0.010,0.005,0.002",
            "0.011,0.004,0.003",
            "0.009,0.006,0.001",
        ],
    )
    doc.add_heading("6.2 OR-Library", level=2)
    doc.add_paragraph(
        "ORLIB 模式會嘗試讀取常見 OR-Library portfolio 資料排列，包括資產數、平均報酬，以及共變異數矩陣或三欄式 covariance pair。若你的 OR-Library 檔案有特殊格式，可先確認前幾列是否包含資產數與報酬資料。"
    )
    add_code_block(doc, ["python ecmade_moo.py --problem ORLIB --orlib-path port1.txt"])

    doc.add_heading("7. 輸出檔案", level=1)
    add_table(
        doc,
        ["檔案", "內容"],
        [
            ["*_population_variables.csv", "最後一代完整族群的決策變數。投資組合問題中每列為一組資產權重。"],
            ["*_population_objectives.csv", "最後一代完整族群的目標值。"],
            ["*_pareto_variables.csv", "最後一代非支配解的決策變數。"],
            ["*_pareto_objectives.csv", "最後一代非支配解的目標值，即近似 Pareto Front。"],
            ["*_history.csv", "每一代的 generation、evaluations、nondominated 數量。"],
            ["*_reference_front.csv", "benchmark 問題的 reference Pareto Front 取樣點。投資組合問題不會產生此檔。"],
            ["*_metrics.txt", "評估次數、Pareto front 大小；若有 reference front，另包含 IGD。"],
        ],
        [2.35, 4.15],
    )

    doc.add_heading("8. 投資組合目標函數解讀", level=1)
    add_table(
        doc,
        ["目標", "公式", "最佳化方向"],
        [
            ["f1", "w^T Sigma w", "最小化投資組合風險或變異數。"],
            ["f2", "- w^T mu", "最小化負報酬，等價於最大化預期報酬。"],
        ],
        [0.9, 2.1, 3.5],
    )
    doc.add_paragraph("程式會使用 simplex repair 將投資組合權重限制在 [0, 1]，並使所有權重總和為 1。")

    doc.add_heading("9. 建議實驗設定", level=1)
    add_table(
        doc,
        ["情境", "建議設定", "說明"],
        [
            ["快速測試", "--pop-size 30 --max-fe 1000", "確認程式與輸入資料格式是否正確。"],
            ["一般 benchmark", "--pop-size 100 --max-fe 10000", "可和常見 NSGA-II 設定進行初步比較。"],
            ["較困難 UF/DTLZ", "--pop-size 100 --max-fe 30000 或更高", "UF10、DTLZ3 等問題通常需要更多評估次數。"],
            ["正式統計實驗", "多個 --seed 分別執行", "建議至少 20 到 30 次獨立執行，再統計 IGD 平均與標準差。"],
            ["投資組合問題", "--pop-size 100 --max-fe 10000 起", "若資產數很多，可提高 max-fe 或 pop-size。"],
        ],
        [1.45, 2.1, 2.95],
    )

    doc.add_heading("10. 常見問題", level=1)
    doc.add_heading("沒有產生 IGD", level=2)
    doc.add_paragraph("投資組合資料通常沒有已知 reference Pareto Front，因此 metrics 只會記錄 evaluations 與 pareto_size。")
    doc.add_heading("Pareto 解太少", level=2)
    doc.add_paragraph("可增加 --max-fe、--pop-size，或嘗試不同 --seed。若是投資組合資料，資產報酬與風險高度相近時，非支配解數量也可能較少。")
    doc.add_heading("OR-Library 讀取錯誤", level=2)
    doc.add_paragraph("請確認檔案中是否包含資產數、每個資產的平均報酬，以及完整共變異數資料。若格式非常特殊，建議先轉成 returns CSV 或自行整理成平均報酬與共變異數後再接入程式。")

    doc.add_heading("11. 對應程式區塊", level=1)
    add_table(
        doc,
        ["功能", "程式位置"],
        [
            ["ECMADE-MOO 主流程", "class ECMADEMOO"],
            ["NSGA-II 環境選擇", "nondominated_sort、crowding_distance、environmental_selection_indices"],
            ["ZDT 問題", "build_zdt"],
            ["DTLZ 問題", "build_dtlz"],
            ["UF 問題", "build_uf、evaluate_uf_2obj、evaluate_uf_3obj"],
            ["OR-Library 讀取", "build_orlibrary_problem"],
            ["Returns CSV 讀取", "build_returns_csv_problem"],
            ["輸出檔案", "save_outputs"],
        ],
        [2.0, 4.5],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)

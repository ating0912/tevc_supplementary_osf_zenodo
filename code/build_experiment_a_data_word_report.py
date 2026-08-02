import os
import math
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = r"C:\Users\yiting\Documents\Playground"
SYN_DIR = os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio", "experiment_A_report_20260705_094925")
OR_DIR = os.path.join(ROOT, "p0_lite_outputs", "orlib_constrained_portfolio", "experiment_A_orlib_report_20260705_235837")
OUT_DIR = os.path.join(ROOT, "docx_outputs")
OUT_DOCX = os.path.join(OUT_DIR, "Experiment_A_數據來源與計算流程說明_統計檢定修正版.docx")
POSTHOC_CSV = os.path.join(OUT_DIR, "Experiment_A_posthoc_ECMADE_vs_baselines_Holm_A12.csv")
COMPACT_STATS_CSV = os.path.join(OUT_DIR, "Experiment_A_compact_statistical_tests.csv")

STAT_METRICS = [
    ("HV", "max"),
    ("IGD", "min"),
    ("PF_Overlap", "max"),
    ("EAF_Band_Width", "min"),
    ("PF_Drift", "min"),
    ("Diversity", "max"),
    ("Runtime", "min"),
]

def rgb(hex_value):
    hex_value = hex_value.strip("#")
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


def set_east_asia_font(run, font_name="Microsoft JhengHei"):
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    if run._element.rPr.rFonts is None:
        r_fonts = OxmlElement("w:rFonts")
        run._element.rPr.append(r_fonts)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, widths):
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    set_east_asia_font(run)
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    style_table(table, widths)
    doc.add_paragraph()
    return table


def add_path_paragraph(doc, label, path):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(label + ": ")
    run.bold = True
    p.add_run(path)


def fmt(value, digits=4):
    try:
        value = float(value)
    except Exception:
        return str(value)
    return f"{value:.{digits}f}"


def fmt_p(value):
    value = float(value)
    if value == 0:
        return "<1e-300"
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}"


def normal_cdf(x):
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def wilcoxon_approx(x, y, direction):
    diffs = np.asarray(x, dtype=float) - np.asarray(y, dtype=float)
    diffs = diffs[~np.isnan(diffs)]
    diffs = diffs[np.abs(diffs) > 1e-12]
    n = len(diffs)
    if n == 0:
        return math.nan, math.nan
    signed = diffs if direction == "max" else -diffs
    absdiff = np.abs(signed)
    order = np.argsort(absdiff)
    ranks = np.empty(n)
    i = 0
    while i < n:
        j = i
        while j + 1 < n and absdiff[order[j + 1]] == absdiff[order[i]]:
            j += 1
        rank = (i + 1 + j + 1) / 2.0
        ranks[order[i : j + 1]] = rank
        i = j + 1
    wpos = ranks[signed > 0].sum()
    mean = n * (n + 1) / 4.0
    var = n * (n + 1) * (2 * n + 1) / 24.0
    if var <= 0:
        return float(wpos), math.nan
    z = (wpos - mean - 0.5 * (1 if wpos > mean else -1)) / math.sqrt(var)
    p = 2 * min(normal_cdf(z), 1 - normal_cdf(z))
    return float(wpos), float(max(min(p, 1.0), 0.0))


def holm_adjust(p_values):
    m = len(p_values)
    indexed = sorted(enumerate(p_values), key=lambda item: (math.inf if pd.isna(item[1]) else item[1]))
    adjusted = [math.nan] * m
    running = 0.0
    for rank, (idx, p) in enumerate(indexed, start=1):
        if pd.isna(p):
            adjusted[idx] = math.nan
            continue
        adj = min(1.0, p * (m - rank + 1))
        running = max(running, adj)
        adjusted[idx] = running
    return adjusted


def vargha_delaney_a12(x, y, direction):
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    x = x[~np.isnan(x)]
    y = y[~np.isnan(y)]
    if direction == "min":
        x = -x
        y = -y
    if len(x) == 0 or len(y) == 0:
        return math.nan
    greater = 0.0
    ties = 0.0
    for xv in x:
        greater += np.sum(xv > y)
        ties += np.sum(xv == y)
    return float((greater + 0.5 * ties) / (len(x) * len(y)))


def load_overall(path):
    df = pd.read_csv(os.path.join(path, "overall_method_summary.csv"))
    if "method" not in df.columns:
        df = df.rename(columns={df.columns[0]: "method"})
    return df.sort_values("RankScore")


def method_rows(df):
    rows = []
    for rank, (_, r) in enumerate(df.iterrows(), 1):
        rows.append([
            rank,
            r["method"],
            fmt(r["RankScore"], 3),
            fmt(r["mean_HV"]),
            fmt(r["mean_IGD"]),
            fmt(r["mean_PF_Overlap"]),
            fmt(r["mean_EAF_Band_Width"]),
            fmt(r["mean_Runtime"], 3),
            fmt(r["mean_Feasible_Rate"]),
        ])
    return rows


def quality_rows(dataset, df):
    rows = []
    for _, r in df.sort_values(["mean_HV", "mean_IGD"], ascending=[False, True]).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_HV"]), fmt(r["mean_IGD"])])
    return rows


def stability_rows(dataset, df):
    rows = []
    for _, r in df.sort_values("mean_PF_Overlap", ascending=False).iterrows():
        rows.append([
            dataset,
            r["method"],
            fmt(r["mean_PF_Overlap"]),
            fmt(r["mean_EAF_Band_Width"]),
            fmt(r["mean_PF_Drift"]),
        ])
    return rows


def behavior_rows(dataset, df):
    rows = []
    for _, r in df.sort_values("mean_Diversity", ascending=False).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_Diversity"]), fmt(r["mean_Spacing"])])
    return rows


def cost_rows(dataset, df):
    rows = []
    for _, r in df.sort_values("mean_Runtime", ascending=True).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_Runtime"], 3)])
    return rows


def friedman_rows(path):
    stats = pd.read_csv(os.path.join(path, "statistical_tests.csv"))
    rows = []
    for _, r in stats[stats["test"].eq("Friedman")].iterrows():
        rows.append([
            r["metric"],
            fmt(r["statistic"], 3),
            fmt_p(r["p_value"]),
            "Yes" if float(r["p_value"]) < 0.05 else "No",
        ])
    return rows


def ecmade_pairwise_highlights(path):
    stats = pd.read_csv(os.path.join(path, "statistical_tests.csv"))
    keep = stats[
        stats["test"].str.contains("Wilcoxon", na=False)
        & stats["comparison"].str.contains("ECMADE_MOO vs SPEA2", na=False)
        & stats["metric"].isin(["HV", "IGD", "PF_Overlap", "EAF_Band_Width", "Runtime"])
    ]
    rows = []
    for _, r in keep.iterrows():
        rows.append([r["metric"], r["comparison"], fmt(r["statistic"], 3), fmt_p(r["p_value"])])
    return rows


def build_posthoc_tables():
    datasets = [
        ("Synthetic", SYN_DIR),
        ("OR-Library", OR_DIR),
    ]
    posthoc_rows = []
    compact_rows = []
    for dataset, path in datasets:
        inst = pd.read_csv(os.path.join(path, "instance_method_metrics.csv"))
        overall = load_overall(path)
        stats = pd.read_csv(os.path.join(path, "statistical_tests.csv"))
        friedman = stats[stats["test"].eq("Friedman")].set_index("metric")
        for metric, direction in STAT_METRICS:
            wide = inst.pivot_table(index="instance", columns="method", values=metric, aggfunc="mean")
            baselines = [m for m in wide.columns if m != "ECMADE_MOO"]
            baseline_means = overall[overall["method"].isin(baselines)].set_index("method")[f"mean_{metric}"]
            best_baseline = baseline_means.idxmax() if direction == "max" else baseline_means.idxmin()
            metric_rows = []
            for baseline in baselines:
                paired = wide[["ECMADE_MOO", baseline]].dropna()
                stat, p = wilcoxon_approx(paired["ECMADE_MOO"].to_numpy(), paired[baseline].to_numpy(), direction)
                a12 = vargha_delaney_a12(paired["ECMADE_MOO"].to_numpy(), paired[baseline].to_numpy(), direction)
                e_mean = overall.loc[overall["method"].eq("ECMADE_MOO"), f"mean_{metric}"].iloc[0]
                b_mean = overall.loc[overall["method"].eq(baseline), f"mean_{metric}"].iloc[0]
                metric_rows.append(
                    {
                        "Dataset": dataset,
                        "Metric": metric,
                        "Direction": "higher is better" if direction == "max" else "lower is better",
                        "Comparison": f"ECMADE_MOO vs {baseline}",
                        "Baseline": baseline,
                        "Best baseline for metric": "Yes" if baseline == best_baseline else "No",
                        "ECMADE mean": e_mean,
                        "Baseline mean": b_mean,
                        "Wilcoxon statistic": stat,
                        "raw p-value": p,
                        "A12": a12,
                    }
                )
            adjusted = holm_adjust([r["raw p-value"] for r in metric_rows])
            for r, adj in zip(metric_rows, adjusted):
                r["Holm adjusted p-value"] = adj
            posthoc_rows.extend(metric_rows)
            best_row = next(r for r in metric_rows if r["Baseline"] == best_baseline)
            f_p = friedman.loc[metric, "p_value"] if metric in friedman.index else math.nan
            compact_rows.append(
                {
                    "Dataset": dataset,
                    "Metric": metric,
                    "Friedman p-value": f_p,
                    "Significant": "Yes" if float(f_p) < 0.05 else "No",
                    "ECMADE-MOO vs Best Baseline": best_baseline,
                    "Holm adjusted p-value": best_row["Holm adjusted p-value"],
                    "Effect Size A12": best_row["A12"],
                }
            )
    posthoc = pd.DataFrame(posthoc_rows)
    compact = pd.DataFrame(compact_rows)
    posthoc.to_csv(POSTHOC_CSV, index=False, encoding="utf-8-sig")
    compact.to_csv(COMPACT_STATS_CSV, index=False, encoding="utf-8-sig")
    return compact, posthoc


def compact_stats_rows(compact):
    rows = []
    for _, r in compact.iterrows():
        rows.append(
            [
                r["Dataset"],
                r["Metric"],
                fmt_p(r["Friedman p-value"]),
                r["Significant"],
                r["ECMADE-MOO vs Best Baseline"],
                fmt_p(r["Holm adjusted p-value"]),
                fmt(r["Effect Size A12"], 3),
            ]
        )
    return rows


def file_size_mb(path):
    return os.path.getsize(path) / (1024 * 1024)


def build_doc():
    os.makedirs(OUT_DIR, exist_ok=True)
    syn_overall = load_overall(SYN_DIR)
    or_overall = load_overall(OR_DIR)
    compact_stats, _ = build_posthoc_tables()
    syn_runs = len(pd.read_csv(os.path.join(SYN_DIR, "run_metrics.csv"), usecols=["method"]))
    or_runs = len(pd.read_csv(os.path.join(OR_DIR, "run_metrics.csv"), usecols=["method"]))
    syn_missing = len(pd.read_csv(os.path.join(SYN_DIR, "missing_outputs.csv")))
    or_missing = len(pd.read_csv(os.path.join(OR_DIR, "missing_outputs.csv")))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Experiment A 數據來源與計算流程說明")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = rgb("0B2545")
    set_east_asia_font(title_run)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("Synthetic constrained portfolio + OR-Library portfolio instances")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = rgb("555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. 文件目的與完成狀態", level=1)
    doc.add_paragraph(
        "本文件整理 Experiment A 的數據來源、實驗輸出如何產生、後處理指標如何計算，以及所有主要檔案的位置。"
        "Experiment A 已包含 synthetic constrained portfolio instances 與正式 OR-Library portfolio instances 兩組資料。"
    )
    add_table(
        doc,
        ["資料集", "完整 runs", "缺漏輸出", "狀態", "報告目錄"],
        [
            ["Synthetic", syn_runs, syn_missing, "Complete" if syn_missing == 0 else "Incomplete", SYN_DIR],
            ["OR-Library", or_runs, or_missing, "Complete" if or_missing == 0 else "Incomplete", OR_DIR],
        ],
        [1200, 1100, 1000, 1100, 4960],
    )

    doc.add_heading("2. 數據從哪裡來", level=1)
    doc.add_heading("2.1 Synthetic constrained portfolio", level=2)
    doc.add_paragraph(
        "Synthetic 資料由 workspace 內的 synthetic constrained portfolio manifest 定義。每筆 manifest row 指向一個合成投資組合資料檔，"
        "包含資產數、K、相關結構、報酬分布、風險結構、replicate 與 seed。實驗 runner 根據 manifest 逐一呼叫 PlatEMO 演算法並保存每次 run 的 population、PF、runtime、feasible rate 與 generation snapshots。"
    )
    add_path_paragraph(doc, "Synthetic manifest", os.path.join(ROOT, "data", "synthetic_constrained_portfolio", "manifest.csv"))
    add_path_paragraph(doc, "Synthetic run output root", os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio"))
    add_path_paragraph(doc, "Synthetic final report folder", SYN_DIR)

    doc.add_heading("2.2 OR-Library portfolio instances", level=2)
    doc.add_paragraph(
        "OR-Library 資料只使用正式 port1.txt 到 port5.txt。這五個原始檔已複製到 workspace，並由 OR-Library manifest 展開為 "
        "5 個 port x 4 個固定 K levels 的 20 個測試 instance。固定 K levels 為 K = 5, 10, 20, 30。"
    )
    add_path_paragraph(doc, "OR-Library source folder", r"C:\Users\yiting\Desktop\NCHU\lab\TEVC\OR-Library")
    add_path_paragraph(doc, "Workspace OR-Library data folder", os.path.join(ROOT, "data", "orlib_constrained_portfolio"))
    add_path_paragraph(doc, "OR-Library manifest", os.path.join(ROOT, "data", "orlib_constrained_portfolio", "manifest.csv"))
    add_path_paragraph(doc, "OR-Library final report folder", OR_DIR)

    or_manifest = pd.read_csv(os.path.join(ROOT, "data", "orlib_constrained_portfolio", "manifest.csv"))
    port_summary = or_manifest.groupby(["source_file", "assets"])["K"].apply(lambda x: ", ".join(map(str, sorted(x)))).reset_index()
    add_table(
        doc,
        ["OR-Library file", "Assets", "K levels", "Workspace path"],
        [
            [
                r["source_file"],
                int(r["assets"]),
                r["K"],
                os.path.join(ROOT, "data", "orlib_constrained_portfolio", "instances", r["source_file"]),
            ]
            for _, r in port_summary.iterrows()
        ],
        [1500, 900, 1300, 5660],
    )

    doc.add_heading("3. 數據如何技術產出", level=1)
    steps = [
        ("Manifest 定義 instance", "Synthetic 與 OR-Library 都先由 manifest.csv 定義 instance、split、assets、K、資料檔 path 與資料結構標籤。"),
        ("Runner 執行演算法", "SyntheticRunner/ORLibraryRunner 依 manifest 逐列呼叫演算法，方法包含 NSGAII、SPEA2、MOEAD、GDE3、ECMADE_MOO、A_MPMO。"),
        ("固定公平設定", "所有方法使用相同 function evaluations、population size、random seed block 與 PortfolioORLIB repair-based constraint handling。正式設定為 N=100、maxFE=10000、每個 method x instance 30 independent runs。"),
        ("保存 run-level output", "每次 run 輸出 pf_obj.csv、pf_dec.csv、population_obj.csv、population_dec.csv、runtime.csv、feasible_rate.csv、generation snapshots 等檔案。"),
        ("後處理計算指標", "報告程式從 final PF、generation PF、runtime、feasible rate、constraint metrics 重新計算 HV、IGD、PF overlap、EAF band width、PF drift、diversity、spacing、runtime 與 feasibility。"),
        ("統計檢定與圖表", "報告程式產生 Friedman tests、ECMADE-MOO pairwise Wilcoxon approximations，以及 dashboard、PF overlay、heatmap、EAF、runtime、stability-diversity figures。"),
    ]
    add_table(doc, ["步驟", "說明"], steps, [1900, 7460])

    doc.add_heading("4. 指標如何計算", level=1)
    metric_rows = [
        ["HV", "每個 instance 以所有方法與 runs 的 final PF union 建立 empirical reference front，再對 normalized PF 計算 2D hypervolume。越高越好。"],
        ["IGD", "以 empirical reference front 為參考，計算 reference points 到 run PF 的平均最近距離。越低越好。"],
        ["PF overlap", "normalized reference front 中，距離 run PF 小於 tolerance 的比例；主 tolerance = 0.01，另有 0.005/0.02 sensitivity。越高越好。"],
        ["EAF band width", "以 repeated-run attainment curves 在 201-point grid 上的 band width 衡量穩定性。越低代表越穩定。"],
        ["PF drift", "normalized PF centroid across runs 的平均漂移距離。越低越穩定。"],
        ["Diversity / spacing", "由 final PF objective range 與 nearest-neighbor spacing 推導，反映解集分散與均勻程度。"],
        ["Runtime", "直接讀取每次 run 的 runtime.csv。越低越好。"],
        ["Feasible rate / violation", "由 saved decision vectors 檢查 cardinality、sum-to-one、bounds；constraint_metrics.csv 保存 mean/max violation。"],
    ]
    add_table(doc, ["指標", "計算方式與解讀"], metric_rows, [1600, 7760])

    doc.add_heading("5. Overall Performance Comparison", level=1)
    doc.add_paragraph(
        "本文件不再以單一總排名或 RankScore 作為 Experiment A 的主軸。HV、IGD、PF Overlap、EAF Band Width、PF Drift、Diversity、Spacing 與 Runtime "
        "衡量的是不同面向；直接將它們平均成一個 RankScore 會隱含相同權重假設，例如 Runtime 與 HV 權重相同、EAF Width 與 IGD 權重相同。"
        "因此本節改以四個分析構面解讀結果。"
    )
    add_table(
        doc,
        ["分析構面", "指標", "解讀"],
        [
            ["Solution Quality", "HV, IGD", "衡量 Pareto front 解品質與收斂能力。HV 越高越好，IGD 越低越好。"],
            ["PF Stability", "PF Overlap, EAF Band Width, PF Drift", "衡量 repeated-run Pareto front 穩定性與漂移。Overlap 越高越好；EAF width 與 drift 越低越好。"],
            ["Search Behavior", "Diversity, Spacing", "衡量 final/archive PF 的分散程度與點間均勻性。"],
            ["Computational Cost", "Runtime", "衡量單次 optimization run 的運算成本，越低越好。"],
        ],
        [1900, 2600, 4860],
    )

    doc.add_heading("5.1 Solution Quality", level=2)
    doc.add_paragraph(
        "正式 OR-Library portfolio instances 是 Experiment A 中最重要的 portfolio-specific 測試場域。"
        "在 OR-Library 上，ECMADE-MOO 取得六種方法中最高 HV 與最低 IGD，代表其在受限制多目標投資組合問題上具有最佳 Pareto front 解品質與收斂能力。"
    )
    add_table(
        doc,
        ["Dataset", "Method", "HV", "IGD"],
        quality_rows("Synthetic", syn_overall) + quality_rows("OR-Library", or_overall),
        [1500, 1800, 1200, 4860],
    )

    doc.add_heading("5.2 PF Stability", level=2)
    doc.add_paragraph(
        "PF stability 與 solution quality 分開呈現。此設計避免將高 HV/低 IGD 與 repeated-run stability 混為同一個分數，"
        "也呼應後續 Experiment C 將 performance-only objective 與 stability-aware objective 分開比較的研究設計。"
    )
    add_table(
        doc,
        ["Dataset", "Method", "PF Overlap", "EAF Width", "PF Drift"],
        stability_rows("Synthetic", syn_overall) + stability_rows("OR-Library", or_overall),
        [1500, 1800, 1500, 1500, 3060],
    )

    doc.add_heading("5.3 Search Behavior", level=2)
    add_table(
        doc,
        ["Dataset", "Method", "Diversity", "Spacing"],
        behavior_rows("Synthetic", syn_overall) + behavior_rows("OR-Library", or_overall),
        [1500, 1800, 1500, 4560],
    )

    doc.add_heading("5.4 Computational Cost", level=2)
    add_table(
        doc,
        ["Dataset", "Method", "Runtime"],
        cost_rows("Synthetic", syn_overall) + cost_rows("OR-Library", or_overall),
        [1500, 1800, 6060],
    )

    doc.add_heading("6. 統計檢定摘要", level=1)
    doc.add_paragraph(
        "統計檢定放在描述統計與分構面結果之後。Friedman test 只用來回答六種演算法在同一指標上是否存在整體顯著差異；"
        "若 Friedman test 顯著，再進行 post-hoc pairwise comparison。正文只保留 ECMADE-MOO 與該指標最佳 baseline 的比較，"
        "並回報 Holm-adjusted p-value 與 Vargha-Delaney A12 effect size。完整 ECMADE-MOO vs all baselines 結果另存 supplementary CSV。"
    )
    add_table(
        doc,
        ["Dataset", "Metric", "Friedman p", "Sig.", "ECMADE-MOO vs Best Baseline", "Holm adj. p", "A12"],
        compact_stats_rows(compact_stats),
        [1100, 1400, 1300, 700, 2600, 1200, 1060],
    )
    doc.add_paragraph(
        "A12 > 0.5 表示 ECMADE-MOO 在該指標上相對 baseline 具有較高機率取得較佳表現；A12 < 0.5 則表示 baseline 較佔優勢。"
    )
    add_path_paragraph(doc, "Supplementary full pairwise table", POSTHOC_CSV)
    add_path_paragraph(doc, "Compact statistical table CSV", COMPACT_STATS_CSV)

    doc.add_heading("7. 分資料集解讀與研究動機連結", level=1)
    doc.add_heading("7.1 OR-Library portfolio instances", level=2)
    doc.add_paragraph(
        "在正式 OR-Library portfolio instances 上，ECMADE-MOO 於六種比較方法中取得最高 HV（0.9909）與最低 IGD（0.0288），"
        "顯示其在受限制多目標投資組合問題中具有最佳的 Pareto front 解品質與收斂能力。"
        "然而，其 PF Overlap 與 EAF Band Width 並未同步優於 SPEA2 與 NSGA-II，顯示高解品質並不必然伴隨高 repeated-run stability。"
        "此結果支持本文將 Pareto front stability 獨立納入 automated evolutionary algorithm design objective 的研究動機。"
    )
    doc.add_heading("7.2 Synthetic constrained portfolio instances", level=2)
    doc.add_paragraph(
        "Synthetic instances 呈現較明顯的 heterogeneous problem behavior。ECMADE-MOO 在 synthetic 上仍有高 HV（0.9609），"
        "但 IGD（0.0723）、PF Overlap（0.3465）與 EAF Width（0.1287）顯示其泛化後的穩定性與部分 baseline 仍有差距。"
        "這表示 fixed hand-crafted ECMADE-MOO configuration 不能保證在所有 instance distribution 上同時取得最佳解品質與穩定性。"
    )
    doc.add_heading("7.3 對 Experiment B/C 的銜接", level=2)
    doc.add_paragraph(
        "Experiment A 因此形成清楚的研究動機：OR-Library 與 synthetic 顯示不同搜尋行為，固定手工設計組態無法適應 heterogeneous problem instances；"
        "因此後續 Experiment B 可進一步檢驗 meta-designed ECMADE-MOO 是否能根據 instance meta-features 自動選擇更合適的 theta，"
        "而 Experiment C 則可比較 performance-only objective 與 stability-aware objective 的差異。"
    )

    doc.add_heading("8. 主要輸出檔案位置", level=1)
    file_rows = []
    for dataset, path in [("Synthetic", SYN_DIR), ("OR-Library", OR_DIR)]:
        for name in [
            "run_metrics.csv",
            "instance_method_metrics.csv",
            "overall_method_summary.csv",
            "statistical_tests.csv",
            "generation_metrics_manifest.csv",
            "metric_availability.csv",
            "missing_outputs.csv",
        ]:
            p = os.path.join(path, name)
            file_rows.append([dataset, name, f"{file_size_mb(p):.2f} MB", p])
    add_table(doc, ["Dataset", "File", "Size", "Path"], file_rows, [1100, 2200, 900, 5160])

    doc.add_heading("9. 圖表位置", level=1)
    fig_rows = []
    for dataset, path in [("Synthetic", SYN_DIR), ("OR-Library", OR_DIR)]:
        fig_dir = os.path.join(path, "figures")
        for fig in sorted(os.listdir(fig_dir)):
            if fig.lower().endswith(".png"):
                fig_rows.append([dataset, fig, os.path.join(fig_dir, fig)])
    add_table(doc, ["Dataset", "Figure file", "Path"], fig_rows, [1100, 2800, 5460])

    doc.add_heading("10. 結論摘要", level=1)
    doc.add_paragraph(
        "Experiment A 的 synthetic 與 OR-Library 數據均已補齊並通過完整性檢查。Synthetic 共 34,560 runs，OR-Library 共 3,600 runs，"
        "兩者 missing_outputs.csv 皆為 0 筆缺漏。"
    )
    doc.add_paragraph(
        "本文件不使用總排名作為主結論，而是將結果分成 solution quality、PF stability、search behavior 與 computational cost 四個構面。"
        "OR-Library 是 Experiment A 的主要亮點：ECMADE-MOO 在正式 portfolio instances 上取得最高 HV 與最低 IGD。"
        "Synthetic 則揭示 fixed hand-crafted configuration 在 heterogeneous instances 上仍存在穩定性與泛化差距。"
    )
    doc.add_paragraph(
        "所有方法差異的 Friedman tests 在 synthetic 與 OR-Library 的主要指標上皆達顯著，可支撐正文中對方法差異的統計描述。"
    )

    doc.save(OUT_DOCX)
    print(f"DOCX={OUT_DOCX}")


if __name__ == "__main__":
    build_doc()

from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESULTS_CSV = ROOT / "ecmade_paper_compliant_best_f1_f13_30runs.csv"
OUTPUT = ROOT / "docx_outputs" / "ECMADE_目前最佳參數整理.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "18324A"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
GRID = "B7C3D0"
INK = "202428"


def set_run_font(run, size=None, bold=None, color=INK, italic=None, mono=False):
    latin = "Consolas" if mono else "Calibri"
    east_asia = "Microsoft JhengHei"
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, latin="Calibri", east_asia="Microsoft JhengHei", size=11, color=INK):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), latin)
    rpr.rFonts.set(qn("w:hAnsi"), latin)
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths, indent=120):
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def format_cell(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=INK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    repeat_table_header(header)
    for idx, label in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT_BLUE)
        format_cell(header.cells[idx], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY, size=9.5)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        if row_idx % 2:
            for cell in row.cells:
                set_cell_shading(cell, "F8FAFC")
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(row.cells[idx], value, align=align)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return table


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("頁碼 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def sci(value):
    value = float(value)
    if value == 0:
        return "0.000E+00"
    return f"{value:.3E}"


def result_note(name, mean, paper):
    if name in {"f4", "f9", "f10", "f12"}:
        return "吻合或達機器精度"
    if name in {"f1", "f8", "f13"}:
        return "接近 0，量級仍有差異"
    if name == "f2":
        return "仍有少數 run 未到 0"
    if name == "f3":
        return "目前最大結構性差距"
    if name == "f5":
        return "平均值非常接近，波動較大"
    if name == "f6":
        return "比論文更接近理論最小值"
    if name == "f7":
        return "平均值與標準差接近"
    if name == "f11":
        return "同一量級，平均與波動偏高"
    return ""


def build_document():
    results = pd.read_csv(RESULTS_CSV)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        set_style_font(style, size=size, color=color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    title_style = doc.styles["Title"]
    set_style_font(title_style, size=24, color=NAVY)
    title_style.font.bold = True
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.line_spacing = 1.0

    subtitle_style = doc.styles["Subtitle"]
    set_style_font(subtitle_style, size=12, color=MUTED)
    subtitle_style.paragraph_format.space_before = Pt(0)
    subtitle_style.paragraph_format.space_after = Pt(14)
    subtitle_style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run("ECMADE Benchmark Reproduction")
    set_run_font(hr, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    add_page_field(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    kr = kicker.add_run("技術實驗摘要")
    set_run_font(kr, size=10, bold=True, color=BLUE)
    doc.add_paragraph("ECMADE 目前最佳參數整理", style="Title")
    doc.add_paragraph("依 Song et al. (2023) Table 1、Algorithm 1 與 Table 3 重現結果整理", style="Subtitle")

    metadata = [
        ("文件日期", date.today().isoformat()),
        ("目前狀態", "論文合規基準仍為整體最佳候選"),
        ("驗證範圍", "13 個 30 維 benchmark；完整結果為 30 independent runs"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(f"{label}：")
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run("結論：")
    set_run_font(r, bold=True, color=BLUE)
    r = p.add_run("目前沒有任何已測替代參數在 5-run 配對比較中勝過論文合規基準，因此建議繼續使用下列設定作為正式 30-run 實驗基線。")
    set_run_font(r)

    doc.add_heading("1. 目前最佳核心設定", level=1)
    core_rows = [
        ("維度 D", "30", "依 Table 1"),
        ("族群大小 NP", "60", "論文明確設定"),
        ("最大世代 MG", "3000", "目前正式重現採用"),
        ("獨立運行", "30 runs", "以固定 seed block 配對"),
        ("子群結構", "3 個子群，各 20 個體", "隨機且平均分群"),
        ("理論 fitness evaluations", "180,060", "初始 60 + 3000×60；論文同時寫 FEs=3000，存在矛盾"),
    ]
    add_table(doc, ["項目", "最佳值", "說明"], core_rows, [2200, 2100, 5060], numeric_cols=(1,))

    doc.add_heading("2. ECMADE 自適應與搜尋參數", level=1)
    adaptive_rows = [
        ("Recent archive size", "H", "20", "論文明確"),
        ("修正係數", "θ", "1/13 ≈ 0.076923", "論文明確"),
        ("停滯門檻", "C", "50", "連續未改善超過 50 代交換"),
        ("Exploitation 權重", "α", "0.8", "Equation (16)"),
        ("Balance 動態權重", "ω", "G/MG", "由 0 線性增加到 1"),
        ("F 分布", "Fᵢ", "Cauchy(μFᵢ, 0.1)", "F≤0 重抽；F>1 截為 1"),
        ("CR 分布", "CRᵢ", "Normal(μCRᵢ, 0.1)", "截至 [0,1]"),
    ]
    add_table(doc, ["參數", "符號", "最佳值", "實作說明"], adaptive_rows, [2250, 1050, 2300, 3760], numeric_cols=(1, 2))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("三個子群的初始參數")
    set_run_font(r, bold=True, color=NAVY)
    initial_rows = [
        ("p1", "Exploration", "DE/rand/2", "0.9", "0.9"),
        ("p2", "Exploitation", "α·best + F·diff₁ + F·diff₂", "0.8", "0.5"),
        ("p3", "Balance", "DE/rand/1 → current-to-best/1", "0.8", "0.5"),
    ]
    add_table(doc, ["子群", "角色", "Mutation", "初始 μF", "初始 μCR"], initial_rows, [850, 1550, 3860, 1550, 1550], numeric_cols=(0, 3, 4))

    doc.add_heading("3. 目前最佳實作選擇", level=1)
    num_id = add_bullet_numbering(doc)
    choices = [
        "Boundary handling：clip 到各函數在 Table 1 的上下界。",
        "Random-vector sampling：只從目標個體所屬子群抽樣，索引互異並排除 target。",
        "Adaptive state：每個個體各自保存 μF、μCR、RSF 與 RSCR；archive 最多保留最近 20 筆成功參數。",
        "Archive weighting：使用 wₖ=exp(gₖ/G−1) 的時間權重，聚合前正規化；F 用 weighted Lehmer mean，CR 用 weighted arithmetic mean。",
        "Generation update：同步 mutation/crossover/selection；trial fitness 不大於 parent 時取代。",
        "Best scope：Equation (16)/(17) 使用全族群 global best。",
        "Information exchange：global best 連續 51 代未改善時重新平均分群，將全族群 top 5% 複製至各子群並取代最差個體。",
        "Elite state：複製 decision vector、fitness、μF/μCR 與個體 archive。",
        "f11 noise：每次 fitness evaluation 都重新產生 random[0,1]；父代不在每一代重新評估。",
        "Seed 與統計：seed block 從 202305 開始；標準差採 sample standard deviation（ddof=1）。",
    ]
    for choice in choices:
        add_bullet(doc, num_id, choice)

    doc.add_heading("4. 目前 30-run 與論文 Table 3 比較", level=1)
    p = doc.add_paragraph("下表列出目前最佳合規實作的平均值與標準差。科學記號統一保留三位小數。")
    p.paragraph_format.keep_with_next = True
    result_rows = []
    for row in results.itertuples(index=False):
        ours = f"{sci(row.mean)} ± {sci(row.std)}"
        paper = f"{sci(row.reported_mean)} ± {sci(row.reported_std)}"
        result_rows.append((row.function, ours, paper, result_note(row.function, row.mean, row.reported_mean)))
    add_table(doc, ["函數", "目前 Mean ± Std", "論文 Mean ± Std", "判讀"], result_rows, [780, 2450, 2450, 3680], numeric_cols=(0, 1, 2))

    rejected_heading = doc.add_heading("5. 已測但不採用的參數", level=1)
    rejected_heading.paragraph_format.page_break_before = True
    rejected = [
        "H=25 或 30：部分函數改善，但 f5 平均值與標準差明顯惡化。",
        "θ=0.20：5-run normalized MAE=0.070319，高於基準 0.064080。",
        "θ=0.05：雖改善 f3/f5，但 f2 退到 0.198992，normalized MAE=0.093079。",
        "C=100：f3 可降至 0.243753，但 f5 與 f11 同時惡化。",
        "θ=0.05、C=75：5-run f3=0.233529，但 f5=16.6865、std=7.29，整體不如基準。",
        "高 μ 初值改配 p2/p3、subpopulation best、重新分群後重設參數、父代重評等方案，皆未通過配對篩選。",
    ]
    for item in rejected:
        add_bullet(doc, num_id, item)

    doc.add_heading("6. 建議使用方式", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("正式重現：")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run("使用本文件列出的合規基準完成 f1–f13、30 runs。不要將單函數較佳的非合規參數混入正式結果。")
    set_run_font(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("後續診斷：")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run("若要繼續縮小 f2/f3 差距，優先檢查作者未公開的亂數來源、函數實作與 MG/FEs 定義，而非繼續微調 H、θ 或 C。")
    set_run_font(r)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("資料來源：")
    set_run_font(r, size=9, bold=True, color=MUTED)
    r = p.add_run("Song et al. (2023), An enhanced distributed differential evolution algorithm for portfolio optimization problems；本地 30-run 結果 ecmade_paper_compliant_best_f1_f13_30runs.csv。")
    set_run_font(r, size=9, color=MUTED)

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.widow_control = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "ECMADE 目前最佳參數整理"
    doc.core_properties.subject = "Song et al. (2023) Table 3 benchmark reproduction"
    doc.core_properties.author = "NCHU Lab"
    doc.core_properties.keywords = "ECMADE, differential evolution, benchmark, Table 3"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())

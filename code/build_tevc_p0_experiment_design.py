# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("output/docx/TEVC_P0_最低實驗設計.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 32, 32)
MUTED = RGBColor(90, 90, 90)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_east_asia_font(run, font_name="Microsoft JhengHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_run(paragraph, text, bold=False, italic=False, color=INK):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    set_east_asia_font(run)
    return run


def set_para(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para(p)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        add_run(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, after=4, line=1.167)
        add_run(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para(p, after=4, line=1.167)
        add_run(p, item)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, after=0, line=1.10)
        run = add_run(p, h, bold=True, color=INK)
        run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row()
        for idx, val in enumerate(row_data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            set_para(p, after=0, line=1.10)
            if idx == 0 and len(headers) <= 4:
                run = add_run(p, str(val), bold=True, color=INK)
            else:
                run = add_run(p, str(val), color=INK)
            run.font.size = Pt(10)
    set_table_geometry(table, widths_dxa)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_para(p, after=2, line=1.10)
    add_run(p, title, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    set_para(p2, after=0, line=1.10)
    add_run(p2, body)
    set_table_geometry(table, [9360])


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p, after=0)
    r = p.add_run("TEVC P0 Experiment Design")
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED
    set_east_asia_font(r)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = add_run(title, "TEVC P0 最低版本實驗設計", bold=True, color=BLUE)
    run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    set_para(subtitle, after=12)
    r = add_run(subtitle, "Stability-Aware Meta-learning for Automated Design of Multi-subpopulation MOEAs", color=MUTED)
    r.font.size = Pt(11)

    add_callout(
        doc,
        "核心設計原則",
        "P0 不拆成互相獨立的八條實驗線，而是用同一套 instances、seed blocks、FE budget 與 raw outputs，派生 baseline comparison、label generation、meta-designed final test、objective validation、ablation、constraint handling、statistics 與 runtime/cost analysis。",
    )

    doc.add_heading("1. P0 實驗總目標", level=1)
    add_bullets(
        doc,
        [
            "證明 ECMADE-MOO 本身比一般 MOEA 更能產生穩定 Pareto Front。",
            "證明 Meta-designed ECMADE-MOO 優於 hand-crafted ECMADE-MOO、random configuration 與 Bayesian configuration。",
            "證明 stability-aware meta-objective 是必要的，不能只依賴 HV/IGD。",
            "證明改善不是單純來自 repair operator、elite injection、固定子群數或額外計算成本。",
        ],
    )

    doc.add_heading("2. 共同輸入與公平比較設定", level=1)
    add_table(
        doc,
        ["類別", "P0 設定"],
        [
            ["資料集", "OR-Library portfolio instances；synthetic constrained portfolio instances"],
            ["K 值", "K = 5, 10, 20, 30；另保留至少一組 unseen K 或 unseen n 作泛化測試"],
            ["目標", "maximize return；minimize risk。實作可轉成 obj1 = -return、obj2 = risk"],
            ["限制", "sum(w)=1、w>=0、cardinality <= K、weight lower/upper bound"],
            ["正式 runs", "每個 method x instance x K 做 30 independent runs"],
            ["Label generation", "每個 candidate theta 先用 10 runs 產生 training labels"],
            ["公平性", "相同 FE budget、population size、seed block、objective scaling 與 constraint handling protocol"],
            ["每次 run 保存", "final PF、archive、weights、objective values、HV/IGD log、runtime、feasible rate、constraint violation"],
        ],
        [2200, 7160],
    )

    doc.add_heading("3. Training / Validation / Test 切分", level=1)
    add_body(doc, "切分必須先於所有 meta-learning 與配置選擇完成，否則後續結果容易被質疑 data leakage。")
    add_table(
        doc,
        ["Split", "用途", "禁止事項"],
        [
            ["Training", "產生 candidate theta labels，訓練 RF/XGBoost", "不可包含 final test instances"],
            ["Validation", "調整 J 權重與 meta-learner hyperparameters", "不可回流 training labels 或最終 test"],
            ["Test", "最終比較 meta-designed ECMADE-MOO", "不可用來選 theta、調 J、挑 feature 或縮減搜尋空間"],
            ["OOD Test", "測試 unseen K 或 unseen n", "若測試 unseen K，該 K 不可出現在 training labels 中"],
        ],
        [1900, 3730, 3730],
    )

    doc.add_heading("4. 實驗 1：Baseline Main Comparison", level=1)
    add_body(doc, "目的：證明 hand-crafted ECMADE-MOO 本身有效，且比一般 MOEA 更能維持穩定 Pareto Front。")
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["輸入", "OR-Library + synthetic portfolio instances；K = 5/10/20/30；相同 FE 與 population size"],
            ["方法", "NSGA-II、SPEA2、MOEA/D、GDE3 或 DEMO、hand-crafted ECMADE-MOO"],
            ["輸出指標", "HV、IGD、PF Overlap、EAF Band Width、Diversity、Runtime、Feasible Rate"],
            ["圖形", "PF Overlay、PF Heatmap、HV/IGD boxplot、runtime boxplot"],
            ["要證明", "多子群 ECMADE-MOO 比單族群或一般 MOEA 更能維持穩定 Pareto Front"],
            ["P0-lite 備註", "若目前只有 ECMADE 與 NSGA-II，可先做 P0-lite；完整 P0 仍需補 SPEA2、MOEA/D、GDE3/DEMO"],
        ],
        [2200, 7160],
    )

    doc.add_heading("5. 實驗 2：Meta-learning Label Generation", level=1)
    add_body(doc, "目的：建立「problem features -> best theta」訓練資料，使最佳配置來自 repeated experimental evidence，而不是人工指定。")
    add_numbered(
        doc,
        [
            "建立 candidate theta set，建議 24 到 48 組，避免計算成本失控。",
            "對每個 training instance 與每個 theta，在相同 FE 下跑 10 independent runs。",
            "計算 nHV、nIGD、PF Overlap、EAF Width、Diversity、Runtime、Feasible Rate。",
            "以 stability-aware score J 產生每個 instance 的 theta ranking 與 best theta。",
            "輸出 label_table.csv，包含 instance id、theta id、metrics、J、rank、best theta 與 cost。",
        ],
    )
    add_table(
        doc,
        ["J 版本", "定義", "用途"],
        [
            ["J_equal", "+nHV - nIGD + PFOverlap - EAFWidth + Diversity - RuntimePenalty", "避免主觀調權"],
            ["J_validation", "用 validation set 調整權重", "取得較佳實務表現"],
            ["J_rank", "使用 Pareto ranking 或 Borda count", "避免單一 scalarization 偏差"],
        ],
        [1800, 4860, 2700],
    )

    doc.add_heading("6. 實驗 3：Meta-designed ECMADE-MOO Final Test", level=1)
    add_body(doc, "目的：證明 meta-designed ECMADE-MOO 不是隨機選配置，也不是普通自動調參，而是能根據問題特徵選出較穩定架構。")
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["輸入", "Unseen test instances、meta-features、候選 theta、trained Random Forest 或 XGBoost"],
            ["比較方法", "hand-crafted ECMADE-MOO、random configuration、Bayesian configuration、meta-designed ECMADE-MOO"],
            ["輸出", "selected theta、HV、IGD、PF Overlap、EAF Width、Runtime、Configuration Cost"],
            ["關鍵要求", "所有主結果必須在 unseen test set 上報告，training instance 結果只能作為輔助分析"],
            ["要證明", "meta-learner 能依據 problem features 選出比 hand-crafted/random/BO 更穩定的 ECMADE-MOO 架構"],
        ],
        [2200, 7160],
    )

    doc.add_heading("7. 實驗 4：Stability-aware Objective 與 Weight Sensitivity", level=1)
    add_body(doc, "此實驗合併原本的 stability objective 檢查與 weight sensitivity analysis，集中回答：穩定性目標是否必要，以及 J 權重是否過於主觀。")
    add_table(
        doc,
        ["版本", "說明", "主要輸出"],
        [
            ["Performance-only J", "只使用 HV/IGD", "HV、IGD、PF Overlap、EAF Width、PF Drift"],
            ["Stability-aware J_equal", "加入 PF Overlap、EAF Width、Diversity，等權重", "selected theta、stability 指標、runtime"],
            ["Stability-aware J_validation", "用 validation set 調權重", "selected theta、test performance、cost"],
            ["Stability-aware J_rank", "以 ranking / Borda count 合成", "selected theta 一致性、robustness"],
        ],
        [2450, 4210, 2700],
    )
    add_body(doc, "要證明：只看 HV/IGD 可能讓 PF 不穩；加入 PF Overlap 與 EAF Width 後 repeated-run reliability 提升，且結果不依賴單一任意權重。")

    doc.add_heading("8. 實驗 5：Ablation Study", level=1)
    add_body(doc, "目的：直接服務審稿防禦，拆解改善來源。P0 至少完成 without meta-learning、without stability objective、without adaptive exchange、single-population version 與 feature group ablation。")
    add_table(
        doc,
        ["消融版本", "比較", "主要輸出", "要回答的問題"],
        [
            ["Without meta-learning", "fixed theta vs meta-designed theta", "J、HV、IGD、PF Overlap", "Meta-learning 是否真的有用"],
            ["Without stability objective", "performance-only J vs stability-aware J", "EAF Width、PF Overlap、Diversity", "穩定性目標是否必要"],
            ["Without adaptive exchange", "adaptive vs fixed migration vs no exchange", "PF Drift、EAF Width", "穩定性是否來自資訊交換"],
            ["Without elite injection", "with vs without elite injection", "HV、Archive Diversity", "elite injection 是否造成改善"],
            ["Single-population", "multi-subpopulation vs single-population", "HV、IGD、PF Overlap", "改善是否來自多子群"],
            ["Feature group ablation", "all features vs remove early/stability/problem features", "Test J、Feature Importance", "哪些 meta-features 有用"],
        ],
        [2140, 2820, 2140, 2260],
    )

    doc.add_heading("9. 實驗 6：Constraint Handling Ablation", level=1)
    add_body(doc, "目的：確認結果不是單靠 portfolio repair operator 造成，並支撐 constraint handling 作為 design vector 的獨立元素。")
    add_table(
        doc,
        ["版本", "說明", "輸出"],
        [
            ["Repair + feasible-first", "主版本；修復後以 feasible-first selection 排序", "Feasible Rate、Violation、HV、IGD、PF Overlap"],
            ["Penalty-based", "不可行解保留，但加入 constraint violation penalty", "Violation、HV/IGD、runtime"],
            ["Discard infeasible", "不可行解直接丟棄或最低排序", "Feasible Rate、archive size、PF stability"],
        ],
        [2400, 4260, 2700],
    )
    add_body(doc, "要證明：ECMADE-MOO 的表現來自架構與 meta-design，不只是 repair operator 比較強。")

    doc.add_heading("10. 實驗 7：Statistics 與 Runtime / Cost Analysis", level=1)
    add_body(doc, "這不是額外演算法實驗，而是所有實驗共用的分析層。")
    add_table(
        doc,
        ["分析類型", "內容"],
        [
            ["兩方法比較", "Wilcoxon signed-rank test"],
            ["多方法多資料集比較", "Friedman test"],
            ["Post-hoc", "Nemenyi test"],
            ["Effect size", "Cliff's delta 或 Vargha-Delaney A12"],
            ["表格呈現", "mean、std、CV、rank、p-value、effect size、win/tie/loss"],
            ["成本分析", "Meta-training cost、preliminary cost、configuration prediction cost、final optimization runtime、net benefit ratio"],
        ],
        [2600, 6760],
    )

    doc.add_heading("11. 最終交付表格與圖", level=1)
    add_table(
        doc,
        ["編號", "名稱", "用途"],
        [
            ["Table 1", "Methods and configurations", "列出演算法、population、operator、archive 與 configuration 差異"],
            ["Table 2", "Theta search space", "列出候選配置與範圍"],
            ["Table 3", "Meta-features", "列出特徵名稱、計算方式與對應設計元素"],
            ["Table 4", "Baseline main results", "ECMADE-MOO vs 主要 baselines"],
            ["Table 5", "Meta-designed final test", "meta-designed vs hand-crafted/random/BO"],
            ["Table 6", "Objective and weight sensitivity", "performance-only vs stability-aware；J_equal/J_validation/J_rank"],
            ["Table 7", "Ablation results", "拆解 meta-learning、exchange、elite、多子群與 features 的效果"],
            ["Table 8", "Constraint handling ablation", "確認結果不是 repair 策略單獨造成"],
            ["Table 9", "Runtime / cost analysis", "回應 automated design 是否成本過高"],
            ["Figure 1", "Framework", "整體流程圖"],
            ["Figure 2", "PF Overlay", "多 runs 的 Pareto Front 疊圖"],
            ["Figure 3", "PF Heatmap", "搜尋熱區與穩定 frontier"],
            ["Figure 4", "EAF Band", "attainment uncertainty band"],
            ["Figure 5", "Feature importance", "RF/XGBoost importance 或 SHAP"],
            ["Figure 6", "Stability-diversity plot", "證明 stability 不是 diversity collapse"],
        ],
        [1450, 3050, 4860],
    )

    doc.add_heading("12. 建議執行順序", level=1)
    add_numbered(
        doc,
        [
            "先完成 data loader、objective scaling、constraint repair 與 unified run logger。",
            "用 ECMADE 與 NSGA-II 做 P0-lite smoke test，確認 PF、HV、runtime 與 feasible rate 都能正確輸出。",
            "補齊 SPEA2、MOEA/D、GDE3/DEMO 後完成 Baseline Main Comparison。",
            "建立 candidate theta set，執行 label generation，輸出 label_table.csv。",
            "訓練 Random Forest 或 XGBoost，並在 unseen test instances 上做 final test。",
            "用同一批 raw outputs 派生 stability objective、weight sensitivity、ablation、constraint handling 與 cost analysis。",
            "最後統一產生 statistics tables、PF overlay、PF heatmap、EAF band 與 feature importance。",
        ],
    )

    doc.add_heading("13. P0 最終判準", level=1)
    add_callout(
        doc,
        "一句話判準",
        "所有圖表都要回到同一個主張：meta-learning 自動配置多子群 ECMADE-MOO 架構，可以在限制條件變動與 repeated stochastic runs 下，提高 Pareto Front 的可重現性，同時維持足夠 diversity 並控制額外配置成本。",
    )

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()

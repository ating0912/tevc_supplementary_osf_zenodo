from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "TEVC_消融實驗數據報告_依審稿問題分類_含Elite0_20260720.docx"


def pct(value: float) -> str:
    if pd.isna(value):
        return "-"
    return f"{value * 100:.1f}%"


def num(value: float, digits: int = 4) -> str:
    if pd.isna(value):
        return "-"
    return f"{value:.{digits}f}"


def set_font(run, size=11, bold=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str, size=11, bold=False, color=None, style=None, after=6) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    repeat_header(table.rows[0])
    for c_idx, header in enumerate(headers):
        cell = table.cell(0, c_idx)
        shade(cell, "F2F4F7")
        write_cell(cell, header, bold=True, size=9)
    for r_idx, row_values in enumerate(rows, start=1):
        for c_idx, value in enumerate(row_values):
            write_cell(table.cell(r_idx, c_idx), value, size=9)
    doc.add_paragraph()


def write_cell(cell, text: str, bold=False, size=9) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if not p.runs:
        run = p.add_run()
    else:
        run = p.runs[0]
    run.text = text
    set_font(run, size=size, bold=bold)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_section_intro(doc: Document, comparison: str, question: str, metrics: str) -> None:
    add_table(
        doc,
        ["比較版本", "要回答的審稿問題", "主要指標"],
        [[comparison, question, metrics]],
        [2.15, 2.55, 1.80],
    )


def load_data():
    expb = pd.read_csv(ROOT / "p0_lite_outputs/experiment_b_configuration_summary_20260713/overall_configuration_comparison.csv")
    stats = pd.read_csv(ROOT / "p0_lite_outputs/experiment_b_configuration_summary_20260713/statistical_tests_meta_vs_baselines.csv")
    c_eval = pd.read_csv(ROOT / "outputs/tevc_ablation_4_5_20260717/label_objective_cross_evaluation_on_C_summary.csv")
    theta = pd.read_csv(ROOT / "outputs/tevc_ablation_6_20260717/theta_factor_main_effect_summary.csv")
    best_counts = pd.read_csv(ROOT / "outputs/tevc_ablation_6_20260717/theta_factor_best_level_counts.csv")
    feature = pd.read_csv(ROOT / "outputs/tevc_ablation_4_5_20260717/feature_group_ablation_summary.csv")
    elite = pd.read_csv(ROOT / "outputs/tevc_without_elite_injection_20260720/without_elite_injection_overall.csv")
    elite_win_loss = pd.read_csv(ROOT / "outputs/tevc_without_elite_injection_20260720/without_elite_0pct_vs_elite_ratios_win_loss.csv")
    return expb, stats, c_eval, theta, best_counts, feature, elite, elite_win_loss


def main() -> None:
    expb, stats, c_eval, theta, best_counts, feature, elite, elite_win_loss = load_data()
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("TEVC 消融實驗數據報告")
    set_font(run, size=20, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("依審稿問題與消融項目重新分類")
    set_font(r, size=11, color="666666")

    add_para(
        doc,
        "本版依照使用者提供的消融實驗順序重排內容：每一個消融項目作為一個大標，並在各節開頭列出比較版本、審稿問題與主要指標。數據主要來自 Experiment B configuration summary、TEVC ablation 4/5 與 TEVC ablation 6。",
    )

    add_para(doc, "消融設計總覽", style="Heading 1")
    add_table(
        doc,
        ["消融", "比較版本", "要回答的審稿問題", "主要指標"],
        [
            ["Without meta-learning", "meta-designed vs fixed hand-crafted", "meta-learning 是否真的有額外價值？", "J、HV、IGD、PF Overlap"],
            ["Without stability objective", "performance-only J vs stability-aware J", "穩定性目標是否必要？", "PF width / PF Overlap / Diversity"],
            ["Without adaptive exchange", "adaptive exchange vs fixed migration vs no exchange", "穩定性是否來自資訊交換？", "PF Drift / PF width"],
            ["Without elite injection", "with vs without elite injection", "elite injection 是否造成改善？", "HV / archive diversity"],
            ["Subpopulation number", "S=2、3、5", "固定子群是否只是人工假設？", "J、runtime、PF Overlap"],
            ["Feature group ablation", "all features vs no preliminary features; features vs no stability features", "哪些 meta-features 有用？", "test J、feature importance"],
        ],
        [1.45, 1.95, 1.95, 1.15],
    )

    add_para(doc, "1. Without Meta-learning", style="Heading 1")
    add_section_intro(doc, "MetaDesigned_ECMADE_MOO vs HandCrafted_ECMADE_MOO", "meta-learning 是否真的有額外價值？", "RankScore/J、HV、IGD、PF Overlap")
    rows = []
    for _, row in expb.iterrows():
        rows.append([
            str(row["method"]).replace("_ECMADE_MOO", ""),
            num(row["mean_HV"]),
            num(row["mean_IGD"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_PF_Drift"]),
            num(row["overall_RankScore"], 3),
            str(int(row["first_place_instances"])),
        ])
    add_table(doc, ["方法", "HV", "IGD", "PF Overlap", "PF Drift", "RankScore", "First-place"], rows, [1.35, .75, .75, .95, .85, .85, 1.00])
    rank_stats = stats[stats["metric"] == "RankScore"].copy()
    stat_rows = []
    for _, row in rank_stats.iterrows():
        stat_rows.append([
            str(row["baseline"]).replace("_ECMADE_MOO", ""),
            f'{int(row["wins"])}/{int(row["ties"])}/{int(row["losses"])}',
            num(row["median_improvement"], 3),
            num(row["holm_p_value"], 4),
            "Yes" if bool(row["significant_0_05"]) else "No",
        ])
    add_table(doc, ["Baseline", "Wins/Ties/Losses", "Median J gain", "Holm p", "Significant"], stat_rows, [1.85, 1.25, 1.15, 1.05, 1.20])
    add_para(doc, "解讀：Meta-designed 在 32 個 unseen test instances 上取得最低 overall RankScore=1.50，優於 hand-crafted 的 3.17。相對於 hand-crafted，RankScore wins/ties/losses=23/0/9，Holm-adjusted p=0.031，顯示 meta-learning 不只是換一組人工設定，而是在泛化測試上帶來顯著額外價值。")

    add_para(doc, "2. Without Stability Objective", style="Heading 1")
    add_section_intro(doc, "performance-only J vs stability-aware J", "穩定性目標是否必要？", "PF Overlap、PF Drift、Diversity、HV、IGD")
    ce_rows = []
    for sel in ["stability_label_top1", "performance_only_top1", "pf_stability_only_top1", "standard_label_top1", "all_theta_mean"]:
        row = c_eval[c_eval["selector"] == sel].iloc[0]
        ce_rows.append([
            sel,
            pct(row["c_top1_hit_rate"]),
            pct(row["c_top3_hit_rate"]),
            num(row["mean_C_ThetaRank"], 2),
            num(row["mean_C_regret_loss"], 2),
            num(row["mean_selected_HV"]),
            num(row["mean_selected_IGD"]),
            num(row["mean_selected_PF_Overlap"]),
            num(row["mean_selected_PF_Drift"]),
            num(row["mean_selected_Diversity"]),
        ])
    add_table(doc, ["Selector", "C top1", "C top3", "C rank", "C regret", "HV", "IGD", "PF Overlap", "PF Drift", "Diversity"], ce_rows, [1.55, .55, .55, .55, .65, .55, .55, .75, .65, .70])
    add_para(doc, "解讀：performance-only 可得到較高 HV=0.9932 與較低 IGD=0.0611，但 PF Overlap 降到 0.5647、PF Drift 升到 0.0727；stability-aware selector 的 PF Overlap=0.5831、PF Drift=0.0652，且 C regret=0。這支持穩定性目標是必要的，因為它避免只追求收斂品質時造成 PF 穩定性退化。")

    add_para(doc, "3. Without Adaptive Exchange", style="Heading 1")
    add_section_intro(doc, "adaptive exchange vs fixed migration vs no exchange", "穩定性是否來自資訊交換？", "PF Drift、PF Overlap、C loss")
    val = theta[(theta["source"] == "Validation") & (theta["objective"] == "stability_label")]
    mig = val[val["factor"] == "migration"].sort_values("level_rank_within_factor")
    add_table(
        doc,
        ["Migration", "C loss", "Mean rank", "Top3 share", "HV", "IGD", "PF Overlap", "PF Drift", "Runtime"],
        [[str(r["level"]), num(r["mean_objective_loss"], 2), num(r["mean_objective_rank"], 2), pct(r["mean_top3_share"]), num(r["mean_HV"]), num(r["mean_IGD"]), num(r["mean_PF_Overlap"]), num(r["mean_PF_Drift"]), num(r["mean_Runtime"], 2)] for _, r in mig.iterrows()],
        [.90, .70, .80, .80, .65, .65, .85, .75, .80],
    )
    add_para(doc, "解讀：在 validation set 的 stability objective 下，no exchange / migration=none 的 C loss=10.82，優於 adaptive=13.09 與 fixed=13.59，PF Overlap 也最高。這表示目前資料並不支持「交換越多越穩定」；穩定性改善更可能來自較少干擾的子群演化，而不是 adaptive exchange 本身。")

    add_para(doc, "4. Without Elite Injection", style="Heading 1")
    add_section_intro(doc, "with vs without elite injection", "elite injection 是否造成改善？", "HV、archive diversity / Diversity")
    elite = elite.copy()
    elite["level_order"] = elite["ablation_level"].map({"0%": 0, "1%": 1, "5%": 5, "10%": 10})
    elite = elite.sort_values("level_order")
    add_table(
        doc,
        ["Elite ratio", "Family rank", "First-place", "HV", "IGD", "PF Overlap", "PF Drift", "Diversity", "Runtime", "Overall rank"],
        [[
            str(r["ablation_level"]),
            num(r["mean_FamilyInstanceRank"], 2),
            str(int(r["first_place_instances"])),
            num(r["mean_HV"]),
            num(r["mean_IGD"]),
            num(r["mean_PF_Overlap"]),
            num(r["mean_PF_Drift"]),
            num(r["mean_Diversity"]),
            num(r["mean_Runtime"], 2),
            num(r["overall_RankScore"], 2),
        ] for _, r in elite.iterrows()],
        [.70, .75, .75, .60, .60, .75, .70, .70, .70, .75],
    )
    runtime_pairs = elite_win_loss[
        (elite_win_loss["method_a"] == "PDF_Abl_Elite_0pct")
        & (elite_win_loss["metric"] == "Runtime")
    ].copy()
    runtime_text = "; ".join(
        f'0% vs {row["method_b"].replace("PDF_Abl_Elite_", "").replace("pct", "%")}: {int(row["wins"])}/{int(row["ties"])}/{int(row["losses"])}'
        for _, row in runtime_pairs.iterrows()
    )
    add_para(doc, f"對照設定：本節使用 PDF-aligned direct ablation package，固定 subpops=3、operator=DE/rand、migration=fixed、stagnation threshold=10、30 runs、32 個 test instances；eliteRatio=0% 代表真正 without elite injection。Runtime 的 win/tie/loss 為 {runtime_text}。")
    add_para(doc, "解讀：eliteRatio=0% 的 family rank=1.00，first-place=32/32，overall rank=1.25，優於 1%、5%、10%。HV、IGD、PF Overlap、PF Drift 與 Diversity 在四組之間幾乎相同，主要差異來自 runtime：0% 的 mean runtime=2.48 秒，低於 1% 的 3.40 秒、5% 的 3.66 秒與 10% 的 3.93 秒。因此，這組直接消融顯示 elite injection 沒有帶來解品質或 archive diversity 的明顯收益，反而增加計算成本。")

    add_para(doc, "5. Subpopulation Number", style="Heading 1")
    add_section_intro(doc, "S=2、3、5", "固定子群是否只是人工假設？", "J/C loss、runtime、PF Overlap")
    s_rows = val[val["factor"] == "S"].sort_values("level_rank_within_factor")
    add_table(
        doc,
        ["S", "C loss", "Mean rank", "Top3 share", "HV", "IGD", "PF Overlap", "PF Drift", "Runtime"],
        [[str(int(r["level"])), num(r["mean_objective_loss"], 2), num(r["mean_objective_rank"], 2), pct(r["mean_top3_share"]), num(r["mean_HV"]), num(r["mean_IGD"]), num(r["mean_PF_Overlap"]), num(r["mean_PF_Drift"]), num(r["mean_Runtime"], 2)] for _, r in s_rows.iterrows()],
        [.45, .75, .80, .80, .65, .65, .85, .75, .85],
    )
    add_para(doc, "解讀：S=5 的 C loss=11.46 最低且 runtime=3.68 最短；S=3 的 top3 share 較高、PF Drift 較低，但整體 C loss 略高；S=2 明顯落後。這表示固定子群數不是任意人工假設，S 太少會降低 selector 可取得的穩定性與品質，而 S=5 是目前 validation data 下最穩的折衷。")

    add_para(doc, "6. Feature Group Ablation", style="Heading 1")
    add_section_intro(doc, "all features vs no preliminary features; features vs no stability features", "哪些 meta-features 有用？", "test J/C regret、feature group importance")
    f = feature[feature["objective"] == "stability_label"].sort_values("mean_regret")
    add_table(
        doc,
        ["Selector", "Top1", "Top3", "Mean rank", "C regret", "HV", "IGD", "PF Overlap", "PF Drift", "RMSE", "MAE"],
        [[str(r["selector"]), pct(r["top1_hit_rate"]), pct(r["top3_hit_rate"]), num(r["mean_selected_rank"], 2), num(r["mean_regret"], 2), num(r["mean_selected_HV"]), num(r["mean_selected_IGD"]), num(r["mean_selected_PF_Overlap"]), num(r["mean_selected_PF_Drift"]), num(r["rmse"], 2), num(r["mae"], 2)] for _, r in f.iterrows()],
        [1.45, .45, .45, .65, .65, .55, .55, .70, .65, .45, .45],
    )
    add_para(doc, "解讀：在 stability-aware objective 下，categorical-only 的 mean regret=3.16，低於 all_features=3.58；no_theta_categorical 的 top3 最高但 regret 略高。這表示 algorithm-design categorical features 是目前 selector 最主要的訊號，完整特徵不必然比精簡後的結構特徵更穩。若要完全對應圖片中的「no preliminary features / no stability features」，建議在下一輪把特徵群組命名固定為 preliminary、stability、problem、theta 四類，並輸出同名消融表。")

    add_para(doc, "整體結論", style="Heading 1")
    add_para(doc, "依圖片順序重排後，最清楚的證據鏈是：meta-learning 對 hand-crafted baseline 有顯著增益；stability-aware objective 能避免 performance-only 對 PF stability 的犧牲；migration=none 與 S=5 是 validation set 中較穩的 theta 水準；without elite injection 的直接對照顯示 eliteRatio=0% 在解品質近乎不變的情況下降低 runtime；特徵消融則顯示 categorical algorithm-design features 是 selector 的核心訊號。")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("TEVC 消融實驗數據報告")
    set_font(r, size=9, color="666666")

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()

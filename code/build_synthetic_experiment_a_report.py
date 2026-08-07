import csv
import math
import os
import statistics
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = os.environ.get("EXPERIMENT_A_ROOT", r".")
OUT_ROOT = os.environ.get(
    "EXPERIMENT_A_OUT_ROOT",
    os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio"),
)
MANIFEST = os.environ.get(
    "EXPERIMENT_A_MANIFEST",
    os.path.join(ROOT, "data", "synthetic_constrained_portfolio", "manifest.csv"),
)
METHODS = [m.strip() for m in os.environ.get(
    "EXPERIMENT_A_METHODS",
    "NSGAII,SPEA2,MOEAD,GDE3,A_MPMO,ECMADE_MOO",
).split(",") if m.strip()]
RUN_COUNT = int(os.environ.get("EXPERIMENT_A_RUNS", "30"))
RUNS = range(1, RUN_COUNT + 1)
REPORT_DIR_PREFIX = os.environ.get("EXPERIMENT_A_REPORT_PREFIX", "experiment_A_report")
REPORT_DOCX_NAME = os.environ.get("EXPERIMENT_A_DOCX_NAME", "Experiment_A_synthetic_results_report.docx")
REPORT_TITLE = os.environ.get("EXPERIMENT_A_REPORT_TITLE", "Experiment A Results Report")
REPORT_SUBTITLE = os.environ.get(
    "EXPERIMENT_A_REPORT_SUBTITLE",
    "ECMADE-MOO and Baseline Comparison on Synthetic Constrained Portfolio Instances",
)
DATASET_LABEL = os.environ.get("EXPERIMENT_A_DATASET_LABEL", "synthetic")
PF_OVERLAP_TOL = 0.01
PF_OVERLAP_SENSITIVITY_TOLS = [0.005, 0.02]
EAF_GRID_POINTS = 201
REQUIRED_FILES = [
    "pf_obj.csv",
    "pf_dec.csv",
    "final_archive_obj.csv",
    "final_archive_dec.csv",
    "pf_points.csv",
    "population_obj.csv",
    "population_dec.csv",
    "runtime.csv",
    "feasible_rate.csv",
    "generation_pf_points.csv",
    "generation_population_log.csv",
    "instance_metadata.csv",
]


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def safe_read_matrix(path):
    try:
        arr = np.loadtxt(path, delimiter=",")
        if arr.ndim == 1:
            arr = arr.reshape(1, -1)
        return arr
    except Exception:
        return np.empty((0, 2))


def read_runtime(path):
    try:
        with open(path, newline="", encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))
        return float(rows[0]["runtime_sec"]) if rows else math.nan
    except Exception:
        return math.nan


def read_feasible(path):
    try:
        with open(path, newline="", encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))
        if not rows:
            return math.nan, math.nan
        return float(rows[0]["PF_Feasible_Rate"]), float(rows[0]["Population_Feasible_Rate"])
    except Exception:
        return math.nan, math.nan


def constraint_violation(dec, k):
    dec = np.asarray(dec, dtype=float)
    if dec.size == 0:
        return np.asarray([math.nan])
    if dec.ndim == 1:
        dec = dec.reshape(1, -1)
    card_violation = np.maximum((dec > 1e-8).sum(axis=1) - k, 0)
    sum_violation = np.abs(dec.sum(axis=1) - 1.0)
    lower_violation = np.maximum(-dec, 0.0).sum(axis=1)
    upper_violation = np.maximum(dec - 1.0, 0.0).sum(axis=1)
    return card_violation + sum_violation + lower_violation + upper_violation


def nanmean(values):
    values = np.asarray(values, dtype=float)
    values = values[~np.isnan(values)]
    return float(values.mean()) if len(values) else math.nan


def nanmax(values):
    values = np.asarray(values, dtype=float)
    values = values[~np.isnan(values)]
    return float(values.max()) if len(values) else math.nan


def nondominated(points):
    points = np.asarray(points, dtype=float)
    if points.size == 0:
        return points.reshape(0, 2)
    order = np.lexsort((points[:, 1], points[:, 0]))
    pts = points[order]
    keep = []
    best_y = math.inf
    for p in pts:
        if p[1] < best_y - 1e-12:
            keep.append(p)
            best_y = p[1]
    return np.asarray(keep)


def thin_front(points, max_points=300):
    points = np.asarray(points, dtype=float)
    if len(points) <= max_points:
        return points
    order = np.argsort(points[:, 0])
    pts = points[order]
    idx = np.linspace(0, len(pts) - 1, max_points).round().astype(int)
    return pts[idx]


def load_one_run(task):
    split, instance, assets, k, k_ratio, corr, ret_dist, risk, method, run = task
    run_dir = os.path.join(OUT_ROOT, split, instance, f"K_{k:02d}", method, f"run_{run:03d}")
    missing = [f for f in REQUIRED_FILES if not os.path.exists(os.path.join(run_dir, f))]
    if missing:
        return {
            "ok": False,
            "missing": {
                "split": split,
                "instance": instance,
                "assets": assets,
                "K": k,
                "method": method,
                "run": run,
                "missing_files": ";".join(missing),
                "run_dir": run_dir,
            },
        }
    pf = safe_read_matrix(os.path.join(run_dir, "pf_obj.csv"))
    rt = read_runtime(os.path.join(run_dir, "runtime.csv"))
    pf_feas, pop_feas = read_feasible(os.path.join(run_dir, "feasible_rate.csv"))
    if os.path.exists(os.path.join(run_dir, "constraint_metrics.csv")):
        pf_dec = np.empty((0, 2))
        pop_dec = np.empty((0, 2))
    else:
        pf_dec = safe_read_matrix(os.path.join(run_dir, "pf_dec.csv"))
        pop_dec = safe_read_matrix(os.path.join(run_dir, "population_dec.csv"))
    constraint = write_constraint_metrics(run_dir, pf_dec, pop_dec, k, pf_feas, pop_feas)
    rec = {
        "split": split,
        "instance": instance,
        "assets": assets,
        "K": k,
        "k_ratio": k_ratio,
        "corr_structure": corr,
        "return_distribution": ret_dist,
        "risk_structure": risk,
        "method": method,
        "run": run,
        "PF_Size": len(pf),
        "Runtime": rt,
        "PF_Feasible_Rate": pf_feas,
        "Population_Feasible_Rate": pop_feas,
        "PF_Mean_Violation": constraint["PF_Mean_Violation"],
        "PF_Max_Violation": constraint["PF_Max_Violation"],
        "Population_Mean_Violation": constraint["Population_Mean_Violation"],
        "Population_Max_Violation": constraint["Population_Max_Violation"],
        "run_dir": run_dir,
    }
    return {"ok": True, "key": (instance, method, run), "pf": pf, "record": rec}


def normalize(points, ideal, nadir):
    points = np.asarray(points, dtype=float)
    denom = np.maximum(nadir - ideal, 1e-12)
    return np.clip((points - ideal) / denom, 0.0, 1.5)


def hv2d(points, ref=(1.1, 1.1)):
    pts = nondominated(points)
    if pts.size == 0:
        return 0.0
    pts = pts[(pts[:, 0] <= ref[0]) & (pts[:, 1] <= ref[1])]
    if pts.size == 0:
        return 0.0
    pts = pts[np.argsort(pts[:, 0])]
    area = 0.0
    prev_y = ref[1]
    for x, y in pts:
        if y < prev_y:
            area += max(ref[0] - x, 0.0) * (prev_y - y)
            prev_y = y
    return float(max(area, 0.0))


def mean_min_distance(source, target):
    if len(source) == 0 or len(target) == 0:
        return math.nan
    source = np.asarray(source)
    target = np.asarray(target)
    d = np.sqrt(((source[:, None, :] - target[None, :, :]) ** 2).sum(axis=2))
    return float(d.min(axis=1).mean())


def igd(run_pf, ref_pf):
    return mean_min_distance(ref_pf, run_pf)


def pf_overlap(run_pf, ref_pf, tol=PF_OVERLAP_TOL):
    if len(run_pf) == 0 or len(ref_pf) == 0:
        return math.nan
    run_pf = np.asarray(run_pf)
    ref_pf = np.asarray(ref_pf)
    d = np.sqrt(((ref_pf[:, None, :] - run_pf[None, :, :]) ** 2).sum(axis=2))
    return float((d.min(axis=1) <= tol).mean())


def diversity_spread(points):
    if len(points) <= 1:
        return 0.0
    pts = np.asarray(points)
    return float(np.sqrt(((pts.max(axis=0) - pts.min(axis=0)) ** 2).sum()))


def spacing(points):
    if len(points) <= 2:
        return 0.0
    pts = np.asarray(points)
    d = np.sqrt(((pts[:, None, :] - pts[None, :, :]) ** 2).sum(axis=2))
    d[d == 0] = np.nan
    nearest = np.nanmin(d, axis=1)
    return float(np.nanstd(nearest))


def write_constraint_metrics(run_dir, pf_dec, pop_dec, k, pf_feas, pop_feas):
    out_path = os.path.join(run_dir, "constraint_metrics.csv")
    if os.path.exists(out_path):
        try:
            row = pd.read_csv(out_path).iloc[0].to_dict()
            return {k: float(v) for k, v in row.items()}
        except Exception:
            pass
    pf_v = constraint_violation(pf_dec, k)
    pop_v = constraint_violation(pop_dec, k)
    row = {
        "PF_Mean_Violation": nanmean(pf_v),
        "PF_Max_Violation": nanmax(pf_v),
        "Population_Mean_Violation": nanmean(pop_v),
        "Population_Max_Violation": nanmax(pop_v),
        "PF_Feasible_Rate": pf_feas,
        "Population_Feasible_Rate": pop_feas,
    }
    pd.DataFrame([row]).to_csv(out_path, index=False, encoding="utf-8-sig")
    return row


def write_archive_metrics(run_dir, norm_pf):
    out_path = os.path.join(run_dir, "archive_metrics.csv")
    if os.path.exists(out_path):
        try:
            row = pd.read_csv(out_path).iloc[0].to_dict()
            return {
                "Archive_Size": int(row["Archive_Size"]),
                "Archive_Diversity": float(row["Archive_Diversity"]),
                "Archive_Spacing": float(row["Archive_Spacing"]),
            }
        except Exception:
            pass
    row = {
        "Archive_Size": int(len(norm_pf)),
        "Archive_Diversity": diversity_spread(norm_pf),
        "Archive_Spacing": spacing(norm_pf),
    }
    pd.DataFrame([row]).to_csv(out_path, index=False, encoding="utf-8-sig")
    return row


def write_generation_metrics(run_dir, info):
    points_path = os.path.join(run_dir, "generation_pf_points.csv")
    pop_path = os.path.join(run_dir, "generation_population_log.csv")
    out_path = os.path.join(run_dir, "generation_metrics.csv")
    if os.path.exists(out_path):
        try:
            return out_path, len(pd.read_csv(out_path, usecols=["generation"]))
        except Exception:
            pass
    if not os.path.exists(points_path):
        return out_path, 0
    try:
        gen_points = pd.read_csv(points_path)
    except Exception:
        gen_points = pd.DataFrame()
    try:
        pop_log = pd.read_csv(pop_path) if os.path.exists(pop_path) else pd.DataFrame()
    except Exception:
        pop_log = pd.DataFrame()

    pop_by_gen = {}
    if not pop_log.empty:
        for _, row in pop_log.iterrows():
            pop_by_gen[int(row["generation"])] = row.to_dict()

    generations = sorted(set(pop_by_gen.keys()) | set(gen_points["generation"].astype(int).tolist() if not gen_points.empty else []))
    rows = []
    for generation in generations:
        pts = np.empty((0, 2))
        if not gen_points.empty:
            g = gen_points[gen_points["generation"].astype(int) == generation]
            if not g.empty:
                pts = g[["risk", "minus_return"]].to_numpy(dtype=float)
        norm_pf = thin_front(nondominated(normalize(pts, info["ideal"], info["nadir"])), 120)
        pop = pop_by_gen.get(generation, {})
        rows.append(
            {
                "generation": generation,
                "evaluations": int(pop.get("evaluations", generation)),
                "HV": hv2d(norm_pf),
                "IGD": igd(norm_pf, info["ref"]),
                "PF_Size": int(len(norm_pf)),
                "Feasible_Rate": float(pop.get("feasible_rate", math.nan)),
                "Diversity": diversity_spread(norm_pf),
                "Spacing": spacing(norm_pf),
            }
        )
    pd.DataFrame(rows).to_csv(out_path, index=False, encoding="utf-8-sig")
    return out_path, len(rows)


def compute_run_metric_record(args):
    rec, pf, info = args
    norm_pf = thin_front(nondominated(normalize(pf, info["ideal"], info["nadir"])), 120)
    archive = write_archive_metrics(rec["run_dir"], norm_pf)
    gen_path, gen_count = write_generation_metrics(rec["run_dir"], info)
    metric = {
        **rec,
        "Configuration_Cost_Runtime": 0.0,
        "Configuration_Cost_FE": 0,
        "Configuration_Type": "fixed_baseline_no_search",
        "Archive_Size": archive["Archive_Size"],
        "Archive_Diversity": archive["Archive_Diversity"],
        "Archive_Spacing": archive["Archive_Spacing"],
        "HV": hv2d(norm_pf),
        "IGD": igd(norm_pf, info["ref"]),
        "PF_Overlap": pf_overlap(norm_pf, info["ref"]),
        "PF_Overlap_tol_0.005": pf_overlap(norm_pf, info["ref"], 0.005),
        "PF_Overlap_tol_0.02": pf_overlap(norm_pf, info["ref"], 0.02),
        "Diversity": diversity_spread(norm_pf),
        "Spacing": spacing(norm_pf),
    }
    generation_index = {
        "split": rec["split"],
        "instance": rec["instance"],
        "K": rec["K"],
        "method": rec["method"],
        "run": rec["run"],
        "generation_rows": gen_count,
        "generation_metrics_path": gen_path,
    }
    return metric, generation_index


def attainment_curve(points, grid):
    pts = nondominated(points)
    y = np.full(len(grid), np.nan)
    if len(pts) == 0:
        return y
    for i, gx in enumerate(grid):
        ok = pts[:, 0] <= gx
        if np.any(ok):
            y[i] = np.min(pts[ok, 1])
    return y


def eaf_band_width(norm_fronts):
    grid = np.linspace(0, 1, EAF_GRID_POINTS)
    curves = np.vstack([attainment_curve(p, grid) for p in norm_fronts if len(p) > 0])
    if curves.size == 0:
        return math.nan
    widths = []
    for col in range(curves.shape[1]):
        vals = curves[:, col]
        vals = vals[~np.isnan(vals)]
        if len(vals) >= 3:
            widths.append(np.quantile(vals, 0.9) - np.quantile(vals, 0.1))
    return float(np.mean(widths)) if widths else math.nan


def centroid(points):
    if len(points) == 0:
        return np.array([math.nan, math.nan])
    return np.asarray(points).mean(axis=0)


def normal_cdf(x):
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def gammap_ser(a, x):
    gln = math.lgamma(a)
    if x <= 0:
        return 0.0
    ap = a
    summ = 1.0 / a
    delt = summ
    for _ in range(1000):
        ap += 1
        delt *= x / ap
        summ += delt
        if abs(delt) < abs(summ) * 3e-14:
            break
    return summ * math.exp(-x + a * math.log(x) - gln)


def gammaq_cf(a, x):
    gln = math.lgamma(a)
    b = x + 1.0 - a
    c = 1.0 / 1e-300
    d = 1.0 / b
    h = d
    for i in range(1, 1000):
        an = -i * (i - a)
        b += 2.0
        d = an * d + b
        if abs(d) < 1e-300:
            d = 1e-300
        c = b + an / c
        if abs(c) < 1e-300:
            c = 1e-300
        d = 1.0 / d
        delt = d * c
        h *= delt
        if abs(delt - 1.0) < 3e-14:
            break
    return math.exp(-x + a * math.log(x) - gln) * h


def chi2_sf(x, df):
    if x < 0:
        return 1.0
    a = df / 2.0
    xx = x / 2.0
    if xx < a + 1.0:
        return max(0.0, 1.0 - gammap_ser(a, xx))
    return min(1.0, gammaq_cf(a, xx))


def friedman_test(wide, direction):
    data = wide.dropna()
    if data.empty:
        return math.nan, math.nan
    values = data[METHODS].copy()
    if direction == "max":
        ranks = values.rank(axis=1, ascending=False, method="average")
    else:
        ranks = values.rank(axis=1, ascending=True, method="average")
    n, k = ranks.shape
    rank_sums = ranks.sum(axis=0)
    stat = (12.0 / (n * k * (k + 1))) * float((rank_sums ** 2).sum()) - 3 * n * (k + 1)
    p = chi2_sf(stat, k - 1)
    return stat, p


def wilcoxon_approx(x, y, direction):
    diffs = np.asarray(x) - np.asarray(y)
    diffs = diffs[~np.isnan(diffs)]
    diffs = diffs[np.abs(diffs) > 1e-12]
    n = len(diffs)
    if n == 0:
        return math.nan, math.nan
    if direction == "max":
        signed = diffs
    else:
        signed = -diffs
    absdiff = np.abs(signed)
    order = np.argsort(absdiff)
    ranks = np.empty(n)
    i = 0
    while i < n:
        j = i
        while j + 1 < n and absdiff[order[j + 1]] == absdiff[order[i]]:
            j += 1
        rank = (i + 1 + j + 1) / 2.0
        ranks[order[i : j + 1]] = rank
        i = j + 1
    wpos = ranks[signed > 0].sum()
    mean = n * (n + 1) / 4.0
    var = n * (n + 1) * (2 * n + 1) / 24.0
    if var <= 0:
        return wpos, math.nan
    z = (wpos - mean - 0.5 * (1 if wpos > mean else -1)) / math.sqrt(var)
    p = 2 * min(normal_cdf(z), 1 - normal_cdf(z))
    return float(wpos), float(max(min(p, 1.0), 0.0))


def font(size=24, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msjhbd.ttc" if bold else r"C:\Windows\Fonts\msjh.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if path and os.path.exists(path):
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_title(draw, title, subtitle, width):
    draw.text((45, 28), title, fill=(20, 44, 78), font=font(34, True))
    if subtitle:
        draw.text((45, 72), subtitle, fill=(90, 90, 90), font=font(20))
    draw.line((45, 108, width - 45, 108), fill=(190, 205, 220), width=2)


def color_for_method(method):
    colors = {
        "NSGAII": (31, 119, 180),
        "SPEA2": (255, 127, 14),
        "MOEAD": (44, 160, 44),
        "GDE3": (214, 39, 40),
        "A_MPMO": (23, 190, 207),
        "ECMADE_MOO": (148, 103, 189),
    }
    return colors.get(method, (80, 80, 80))


def bar_chart(path, title, metrics, ylabel="", higher_better=True):
    w, h = 1200, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    draw_title(d, title, ylabel, w)
    left, top, right, bottom = 105, 150, 1120, 595
    vals = [metrics.get(m, math.nan) for m in METHODS]
    clean = [v for v in vals if not math.isnan(v)]
    if not clean:
        img.save(path)
        return
    ymin = min(0.0, min(clean))
    ymax = max(clean) * 1.15 if max(clean) > 0 else 1.0
    d.rectangle((left, top, right, bottom), outline=(210, 210, 210))
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        d.line((left, y, right, y), fill=(235, 238, 241))
        label = ymin + (ymax - ymin) * i / 5
        d.text((25, y - 10), f"{label:.3g}", fill=(80, 80, 80), font=font(16))
    gap = 28
    bw = (right - left - gap * (len(METHODS) + 1)) / len(METHODS)
    for i, m in enumerate(METHODS):
        v = vals[i]
        x0 = left + gap + i * (bw + gap)
        x1 = x0 + bw
        y0 = bottom if math.isnan(v) else bottom - (v - ymin) / (ymax - ymin) * (bottom - top)
        d.rectangle((x0, y0, x1, bottom), fill=color_for_method(m))
        d.text((x0, bottom + 15), m.replace("_", "-"), fill=(40, 40, 40), font=font(16))
        if not math.isnan(v):
            d.text((x0, y0 - 26), f"{v:.3g}", fill=(35, 35, 35), font=font(16, True))
    img.save(path)


def multi_metric_bars(path, overall):
    metrics = [
        ("HV", "mean_HV", "max"),
        ("IGD", "mean_IGD", "min"),
        ("PF overlap", "mean_PF_Overlap", "max"),
        ("EAF width", "mean_EAF_Band_Width", "min"),
        ("Diversity", "mean_Diversity", "context"),
        ("Runtime", "mean_Runtime", "min"),
        ("Feasible", "mean_Feasible_Rate", "max"),
    ]
    w, h = 1500, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    draw_title(d, "Figure 1. Overall Metric Dashboard", f"Mean values across {DATASET_LABEL} constrained instances", w)
    start_x, start_y = 70, 150
    cell_w, cell_h = 430, 210
    for idx, (label, col, direction) in enumerate(metrics):
        x = start_x + (idx % 3) * (cell_w + 35)
        y = start_y + (idx // 3) * (cell_h + 45)
        d.text((x, y), label, fill=(20, 44, 78), font=font(22, True))
        vals = {m: float(overall.loc[m, col]) if m in overall.index else math.nan for m in METHODS}
        clean = [v for v in vals.values() if not math.isnan(v)]
        if not clean:
            continue
        vmin, vmax = min(clean), max(clean)
        denom = max(vmax - vmin, 1e-12)
        for j, m in enumerate(METHODS):
            v = vals[m]
            yy = y + 38 + j * 30
            d.text((x, yy), m.replace("_", "-"), fill=(50, 50, 50), font=font(15))
            bar_x = x + 125
            bar_w = 210 if math.isnan(v) else 40 + 170 * (v - vmin) / denom
            d.rectangle((bar_x, yy + 3, bar_x + bar_w, yy + 21), fill=color_for_method(m))
            d.text((bar_x + 220, yy), "NA" if math.isnan(v) else f"{v:.3g}", fill=(45, 45, 45), font=font(15))
    img.save(path)


def scatter_chart(path, inst_metrics):
    w, h = 1050, 760
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    draw_title(d, "Figure 6. Stability-Diversity Plot", "Each point is a method x instance aggregate", w)
    left, top, right, bottom = 110, 145, 970, 620
    xs = inst_metrics["EAF_Band_Width"].astype(float).to_numpy()
    ys = inst_metrics["Diversity"].astype(float).to_numpy()
    mask = ~np.isnan(xs) & ~np.isnan(ys)
    xs, ys = xs[mask], ys[mask]
    if len(xs) == 0:
        img.save(path)
        return
    xmin, xmax = np.quantile(xs, [0.01, 0.99])
    ymin, ymax = np.quantile(ys, [0.01, 0.99])
    xmin, xmax = float(xmin), float(max(xmax, xmin + 1e-9))
    ymin, ymax = float(ymin), float(max(ymax, ymin + 1e-9))
    d.rectangle((left, top, right, bottom), outline=(210, 210, 210))
    for i in range(6):
        x = left + (right - left) * i / 5
        y = bottom - (bottom - top) * i / 5
        d.line((x, top, x, bottom), fill=(236, 238, 241))
        d.line((left, y, right, y), fill=(236, 238, 241))
    for _, r in inst_metrics.iterrows():
        xval, yval = float(r["EAF_Band_Width"]), float(r["Diversity"])
        if math.isnan(xval) or math.isnan(yval):
            continue
        x = left + (min(max(xval, xmin), xmax) - xmin) / (xmax - xmin) * (right - left)
        y = bottom - (min(max(yval, ymin), ymax) - ymin) / (ymax - ymin) * (bottom - top)
        c = color_for_method(r["method"])
        d.ellipse((x - 4, y - 4, x + 4, y + 4), fill=c)
    d.text((left + 300, bottom + 45), "EAF band width (lower is more stable)", fill=(55, 55, 55), font=font(18))
    d.text((18, top + 190), "Diversity", fill=(55, 55, 55), font=font(18))
    lx = 755
    for i, m in enumerate(METHODS):
        yy = 155 + i * 28
        d.rectangle((lx, yy, lx + 18, yy + 18), fill=color_for_method(m))
        d.text((lx + 25, yy - 2), m.replace("_", "-"), fill=(45, 45, 45), font=font(16))
    img.save(path)


def pf_overlay(path, title, instance_rows, ref_data):
    w, h = 1500, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    draw_title(d, "Figure 2. PF Overlay", title, w)
    left, top, right, bottom = 105, 145, 1380, 745
    d.rectangle((left, top, right, bottom), outline=(210, 210, 210))
    for i in range(6):
        x = left + (right - left) * i / 5
        y = bottom - (bottom - top) * i / 5
        d.line((x, top, x, bottom), fill=(238, 240, 243))
        d.line((left, y, right, y), fill=(238, 240, 243))
    for method in METHODS:
        pts = np.vstack(instance_rows.get(method, [])) if instance_rows.get(method) else np.empty((0, 2))
        if len(pts) == 0:
            continue
        ideal, nadir = ref_data
        pts = normalize(pts, ideal, nadir)
        pts = pts[np.linspace(0, len(pts) - 1, min(len(pts), 1000)).astype(int)]
        c = color_for_method(method)
        for xval, yval in pts:
            x = left + min(max(xval, 0), 1) * (right - left)
            y = bottom - min(max(yval, 0), 1) * (bottom - top)
            d.ellipse((x - 2, y - 2, x + 2, y + 2), fill=c)
    d.text((left + 460, bottom + 45), "Normalized risk", fill=(55, 55, 55), font=font(18))
    d.text((18, top + 245), "Normalized -return", fill=(55, 55, 55), font=font(18))
    lx = 1130
    for i, m in enumerate(METHODS):
        yy = 155 + i * 30
        d.rectangle((lx, yy, lx + 20, yy + 20), fill=color_for_method(m))
        d.text((lx + 30, yy - 2), m.replace("_", "-"), fill=(45, 45, 45), font=font(17))
    img.save(path)


def heatmap_chart(path, title, instance_rows, ref_data):
    w, h = 1500, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    draw_title(d, "Figure 3. PF Heatmap", title, w)
    panel_w, panel_h = 250, 250
    x0, y0 = 70, 160
    for idx, method in enumerate(METHODS):
        px = x0 + (idx % 3) * 455
        py = y0 + (idx // 3) * 320
        d.text((px, py - 32), method.replace("_", "-"), fill=(20, 44, 78), font=font(20, True))
        d.rectangle((px, py, px + panel_w, py + panel_h), outline=(200, 200, 200))
        pts = np.vstack(instance_rows.get(method, [])) if instance_rows.get(method) else np.empty((0, 2))
        if len(pts) == 0:
            continue
        ideal, nadir = ref_data
        pts = normalize(pts, ideal, nadir)
        bins = np.zeros((40, 40), dtype=int)
        for xval, yval in pts:
            xi = int(min(max(xval, 0), 0.9999) * 40)
            yi = int(min(max(yval, 0), 0.9999) * 40)
            bins[39 - yi, xi] += 1
        maxv = bins.max() or 1
        for yy in range(40):
            for xx in range(40):
                v = bins[yy, xx] / maxv
                if v <= 0:
                    continue
                col = (255 - int(205 * v), 245 - int(120 * v), 235 - int(200 * v))
                d.rectangle(
                    (
                        px + xx * panel_w / 40,
                        py + yy * panel_h / 40,
                        px + (xx + 1) * panel_w / 40 + 1,
                        py + (yy + 1) * panel_h / 40 + 1,
                    ),
                    fill=col,
                )
        d.rectangle((px, py, px + panel_w, py + panel_h), outline=(120, 120, 120))
    img.save(path)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8)
    if color:
        r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_df_table(doc, df, columns, title=None, max_rows=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(31, 77, 120)
    rows = df if max_rows is None else df.head(max_rows)
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, col in enumerate(columns):
        set_cell_text(table.cell(0, j), col, bold=True)
        set_cell_fill(table.cell(0, j), "E8EEF5")
    for i, (_, row) in enumerate(rows.iterrows(), start=1):
        for j, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.4g}"
            else:
                text = val
            set_cell_text(table.cell(i, j), text)
    doc.add_paragraph()


def apply_doc_styles(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def make_report(report_dir, completeness, availability, overall, inst_metrics, stat_rows, figure_paths, warnings, output_docx):
    doc = Document()
    apply_doc_styles(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{REPORT_TITLE}\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    sub = title.add_run(REPORT_SUBTITLE)
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Data Completeness Check", level=1)
    if completeness["missing_count"] == 0:
        doc.add_paragraph(
            f"PASS. All expected method x instance x run outputs are present: "
            f"{completeness['complete_runs']} / {completeness['expected_runs']} complete runs."
        )
    else:
        doc.add_paragraph(
            f"WARNING. Missing or incomplete outputs: {completeness['missing_count']} run/file records. "
            f"See {completeness['missing_csv']}."
        )
    add_df_table(doc, availability, ["Category", "Item", "Status", "Source_or_definition"], "Metric availability")

    if warnings:
        doc.add_heading("2. Warnings and Interpretation Notes", level=1)
        for w in warnings:
            doc.add_paragraph(w, style=None)

    doc.add_heading("3. Overall Results", level=1)
    overall_out = overall.reset_index().rename(columns={"index": "method"})
    columns = [
        "method",
        "mean_HV",
        "std_HV",
        "cv_HV",
        "mean_IGD",
        "std_IGD",
        "cv_IGD",
        "mean_PF_Overlap",
        "mean_EAF_Band_Width",
        "mean_Diversity",
        "mean_Archive_Diversity",
        "mean_Runtime",
        "mean_Configuration_Cost_Runtime",
        "mean_Feasible_Rate",
        "mean_PF_Mean_Violation",
        "RankScore",
    ]
    add_df_table(doc, overall_out[columns], columns, "Table 1. Main results by method")

    doc.add_heading("4. Statistical Tests", level=1)
    stat_df = pd.DataFrame(stat_rows)
    add_df_table(doc, stat_df, list(stat_df.columns), "Table 2. Friedman and Wilcoxon tests")

    doc.add_heading("5. Figures", level=1)
    captions = {
        "dashboard": "Figure 1. Overall metric dashboard.",
        "overlay": "Figure 2. PF overlay on a representative high-dimensional test instance.",
        "heatmap": "Figure 3. PF heatmap on the same representative instance.",
        "eaf": "Figure 4. EAF band width, lower values indicate more stable repeated-run fronts.",
        "runtime": "Figure 5. Runtime comparison, lower is better.",
        "scatter": "Figure 6. Stability-diversity plot.",
    }
    for key in ["dashboard", "overlay", "heatmap", "eaf", "runtime", "scatter"]:
        if key in figure_paths and os.path.exists(figure_paths[key]):
            doc.add_picture(figure_paths[key], width=Inches(9.5))
            cap = doc.add_paragraph(captions[key])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("6. Method Ranking Summary", level=1)
    rank_cols = ["method", "RankScore", "rank_HV", "rank_IGD", "rank_PF_Overlap", "rank_EAF_Band_Width", "rank_Runtime"]
    add_df_table(doc, overall_out.sort_values("RankScore")[rank_cols], rank_cols, "Table 3. Composite rank summary")

    doc.add_heading("7. Reproducibility", level=1)
    doc.add_paragraph(f"Output root: {OUT_ROOT}")
    doc.add_paragraph(f"Report artifact directory: {report_dir}")
    doc.add_paragraph(f"Manifest: {MANIFEST}")
    doc.add_paragraph(f"All metrics were recomputed from final PF files, generation snapshots, feasible-rate files, runtime files, and the {DATASET_LABEL} manifest.")
    doc.save(output_docx)


def main():
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_dir = os.path.join(OUT_ROOT, f"{REPORT_DIR_PREFIX}_{stamp}")
    figures_dir = os.path.join(report_dir, "figures")
    ensure_dir(figures_dir)

    manifest = pd.read_csv(MANIFEST)
    expected = len(manifest) * len(METHODS) * len(RUNS)
    run_pfs = {}
    run_records = []
    missing_rows = []
    by_instance_points = defaultdict(list)
    by_instance_method_points = defaultdict(list)
    tasks = []
    for _, row in manifest.iterrows():
        for method in METHODS:
            for run in RUNS:
                tasks.append(
                    (
                        row["split"],
                        row["instance"],
                        int(row["assets"]),
                        int(row["K"]),
                        row["k_ratio"],
                        row["corr_structure"],
                        row["return_distribution"],
                        row["risk_structure"],
                        method,
                        run,
                    )
                )

    print(f"Scanning {len(tasks)} expected runs...", flush=True)
    done = 0
    with ThreadPoolExecutor(max_workers=32) as pool:
        futures = [pool.submit(load_one_run, task) for task in tasks]
        for fut in as_completed(futures):
            res = fut.result()
            done += 1
            if done % 1000 == 0:
                print(f"  scanned {done}/{len(tasks)}", flush=True)
            if not res["ok"]:
                missing_rows.append(res["missing"])
                continue
            key = res["key"]
            pf = res["pf"]
            rec = res["record"]
            run_pfs[key] = pf
            by_instance_points[key[0]].append(pf)
            by_instance_method_points[(key[0], key[1])].append(pf)
            run_records.append(rec)

    complete_runs = expected - len({(r["instance"], r["method"], r["run"]) for r in missing_rows})
    missing_csv = os.path.join(report_dir, "missing_outputs.csv")
    if missing_rows:
        pd.DataFrame(missing_rows).to_csv(missing_csv, index=False, encoding="utf-8-sig")
    else:
        pd.DataFrame(columns=["split", "instance", "K", "method", "run", "missing_files", "run_dir"]).to_csv(
            missing_csv, index=False, encoding="utf-8-sig"
        )

    if missing_rows:
        print(f"MISSING_OUTPUTS={len(missing_rows)}")
        print(f"MISSING_CSV={missing_csv}")
        return

    print("Building empirical reference fronts...", flush=True)
    ref_info = {}
    for key, fronts in by_instance_points.items():
        union = np.vstack([f for f in fronts if len(f) > 0])
        ideal = union.min(axis=0)
        nadir = union.max(axis=0)
        norm_union = normalize(union, ideal, nadir)
        ref = thin_front(nondominated(norm_union), 120)
        ref_info[key] = {"ideal": ideal, "nadir": nadir, "ref": ref}

    print("Computing run-level metrics...", flush=True)
    metric_rows = []
    generation_index_rows = []
    records = pd.DataFrame(run_records)
    metric_tasks = []
    for rec in run_records:
        key, method, run = rec["instance"], rec["method"], rec["run"]
        metric_tasks.append((rec, run_pfs[(key, method, run)], ref_info[key]))
    done = 0
    with ThreadPoolExecutor(max_workers=32) as pool:
        futures = [pool.submit(compute_run_metric_record, task) for task in metric_tasks]
        for fut in as_completed(futures):
            metric, generation_index = fut.result()
            metric_rows.append(metric)
            generation_index_rows.append(generation_index)
            done += 1
            if done % 1000 == 0:
                print(f"  run metrics {done}/{len(metric_tasks)}", flush=True)
    run_metrics = pd.DataFrame(metric_rows)

    print("Computing method-instance stability metrics...", flush=True)
    inst_rows = []
    for (instance, method), fronts in by_instance_method_points.items():
        info = ref_info[instance]
        norm_fronts = [thin_front(nondominated(normalize(f, info["ideal"], info["nadir"])), 120) for f in fronts]
        centroids = np.vstack([centroid(f) for f in norm_fronts if len(f) > 0])
        mean_centroid = np.nanmean(centroids, axis=0)
        drifts = [float(np.sqrt(((centroid(f) - mean_centroid) ** 2).sum())) for f in norm_fronts if len(f) > 0]
        base = run_metrics[(run_metrics["instance"] == instance) & (run_metrics["method"] == method)]
        meta = base.iloc[0].to_dict()
        inst_rows.append(
            {
                "split": meta["split"],
                "instance": instance,
                "assets": int(meta["assets"]),
                "K": int(meta["K"]),
                "k_ratio": meta["k_ratio"],
                "method": method,
                "HV": base["HV"].mean(),
                "IGD": base["IGD"].mean(),
                "PF_Overlap": base["PF_Overlap"].mean(),
                "PF_Overlap_tol_0.005": base["PF_Overlap_tol_0.005"].mean(),
                "PF_Overlap_tol_0.02": base["PF_Overlap_tol_0.02"].mean(),
                "Diversity": base["Diversity"].mean(),
                "Spacing": base["Spacing"].mean(),
                "Archive_Size": base["Archive_Size"].mean(),
                "Archive_Diversity": base["Archive_Diversity"].mean(),
                "Archive_Spacing": base["Archive_Spacing"].mean(),
                "Runtime": base["Runtime"].mean(),
                "Configuration_Cost_Runtime": base["Configuration_Cost_Runtime"].mean(),
                "Configuration_Cost_FE": base["Configuration_Cost_FE"].mean(),
                "Feasible_Rate": base["PF_Feasible_Rate"].mean(),
                "PF_Mean_Violation": base["PF_Mean_Violation"].mean(),
                "PF_Max_Violation": base["PF_Max_Violation"].mean(),
                "Population_Mean_Violation": base["Population_Mean_Violation"].mean(),
                "Population_Max_Violation": base["Population_Max_Violation"].mean(),
                "EAF_Band_Width": eaf_band_width(norm_fronts),
                "PF_Drift": float(np.nanmean(drifts)) if drifts else math.nan,
            }
        )
    inst_metrics = pd.DataFrame(inst_rows)

    print("Aggregating overall summaries and statistical tests...", flush=True)
    overall = inst_metrics.groupby("method").agg(
        mean_HV=("HV", "mean"),
        std_HV=("HV", "std"),
        mean_IGD=("IGD", "mean"),
        std_IGD=("IGD", "std"),
        mean_PF_Overlap=("PF_Overlap", "mean"),
        std_PF_Overlap=("PF_Overlap", "std"),
        mean_PF_Overlap_tol_0_005=("PF_Overlap_tol_0.005", "mean"),
        mean_PF_Overlap_tol_0_02=("PF_Overlap_tol_0.02", "mean"),
        mean_EAF_Band_Width=("EAF_Band_Width", "mean"),
        std_EAF_Band_Width=("EAF_Band_Width", "std"),
        mean_PF_Drift=("PF_Drift", "mean"),
        std_PF_Drift=("PF_Drift", "std"),
        mean_Diversity=("Diversity", "mean"),
        std_Diversity=("Diversity", "std"),
        mean_Spacing=("Spacing", "mean"),
        std_Spacing=("Spacing", "std"),
        mean_Archive_Size=("Archive_Size", "mean"),
        std_Archive_Size=("Archive_Size", "std"),
        mean_Archive_Diversity=("Archive_Diversity", "mean"),
        std_Archive_Diversity=("Archive_Diversity", "std"),
        mean_Archive_Spacing=("Archive_Spacing", "mean"),
        std_Archive_Spacing=("Archive_Spacing", "std"),
        mean_Runtime=("Runtime", "mean"),
        std_Runtime=("Runtime", "std"),
        mean_Configuration_Cost_Runtime=("Configuration_Cost_Runtime", "mean"),
        mean_Configuration_Cost_FE=("Configuration_Cost_FE", "mean"),
        mean_Feasible_Rate=("Feasible_Rate", "mean"),
        std_Feasible_Rate=("Feasible_Rate", "std"),
        mean_PF_Mean_Violation=("PF_Mean_Violation", "mean"),
        std_PF_Mean_Violation=("PF_Mean_Violation", "std"),
        mean_PF_Max_Violation=("PF_Max_Violation", "mean"),
        std_PF_Max_Violation=("PF_Max_Violation", "std"),
        mean_Population_Mean_Violation=("Population_Mean_Violation", "mean"),
        std_Population_Mean_Violation=("Population_Mean_Violation", "std"),
        mean_Population_Max_Violation=("Population_Max_Violation", "mean"),
        std_Population_Max_Violation=("Population_Max_Violation", "std"),
    )
    for base in ["HV", "IGD"]:
        overall[f"cv_{base}"] = overall[f"std_{base}"] / overall[f"mean_{base}"].replace(0, np.nan)
    rank_specs = [
        ("HV", "max"),
        ("IGD", "min"),
        ("PF_Overlap", "max"),
        ("EAF_Band_Width", "min"),
        ("PF_Drift", "min"),
        ("Runtime", "min"),
        ("Feasible_Rate", "max"),
    ]
    for metric, direction in rank_specs:
        vals = overall[f"mean_{metric}"]
        overall[f"rank_{metric}"] = vals.rank(ascending=(direction == "min"), method="average")
    overall["RankScore"] = overall[[f"rank_{m}" for m, _ in rank_specs]].mean(axis=1)
    overall = overall.sort_values("RankScore")

    stat_rows = []
    for metric, direction in rank_specs[:6]:
        wide = inst_metrics.pivot_table(index="instance", columns="method", values=metric, aggfunc="mean")
        wide = wide.reindex(columns=METHODS)
        stat, p = friedman_test(wide, direction)
        stat_rows.append(
            {
                "metric": metric,
                "test": "Friedman",
                "comparison": "all methods",
                "statistic": stat,
                "p_value": p,
            }
        )
        if "ECMADE_MOO" in wide.columns:
            for m in METHODS:
                if m == "ECMADE_MOO":
                    continue
                w, p2 = wilcoxon_approx(wide["ECMADE_MOO"].to_numpy(), wide[m].to_numpy(), direction)
                stat_rows.append(
                    {
                        "metric": metric,
                        "test": "Wilcoxon approx.",
                        "comparison": f"ECMADE_MOO vs {m}",
                        "statistic": w,
                        "p_value": p2,
                    }
                )

    availability = pd.DataFrame(
        [
            ["Performance", "HV mean / std / CV", "Available", "Derived from final PF using empirical reference front."],
            ["Performance", "IGD mean / std / CV", "Available", "Derived from final PF against empirical reference front."],
            ["Stability", "PF overlap", "Available", "Main table uses normalized tolerance 0.01; sensitivity columns also report 0.005 and 0.02."],
            ["Stability", "EAF band width", "Available", "Derived from 30-run attainment curves per method and instance using a 201-point grid."],
            ["Stability", "PF drift distance", "Available", "Derived from normalized PF centroid drift across runs."],
            ["Diversity", "Spread / spacing / archive diversity", "Available", "Diversity/Spacing are in run_metrics.csv; Archive_Size/Archive_Diversity/Archive_Spacing are also written to archive_metrics.csv per run."],
            ["Cost", "Runtime", "Available", "Direct from runtime.csv."],
            ["Cost", "Configuration cost", "Available as zero fixed-baseline cost", "Experiment A has no meta-configuration search; Configuration_Cost_Runtime=0 and Configuration_Cost_FE=0 are reported explicitly."],
            ["Feasibility", "Feasible rate / violation degree", "Available", "Feasible rate is read from feasible_rate.csv; violation degree is written to constraint_metrics.csv per run and summarized in report tables."],
            ["Performance log", "Per-generation HV/IGD", "Available", "generation_metrics.csv is written per run from generation_pf_points.csv using the same empirical reference PF as final HV/IGD."],
            ["Visualization", "PF overlay / PF heatmap / EAF plots", "Available", "Generated in this report from final PF outputs."],
        ],
        columns=["Category", "Item", "Status", "Source_or_definition"],
    )

    warnings = [
        "HV and IGD use an empirical reference front constructed from the union of all final PF points across methods and runs for each instance.",
        "Per-generation HV/IGD is generated in post-processing from generation_pf_points.csv and written back to each run as generation_metrics.csv using the common empirical reference front.",
        "Violation degree is generated from saved decision vectors and written back to each run as constraint_metrics.csv. PortfolioORLIB repair usually makes final violations zero, but the explicit file is kept for audit and later constraint-handling ablation.",
        "All methods use the same PortfolioORLIB repair: clip to [0,1], retain top-K weights, zero the rest, and normalize to sum one. This makes constraint handling fair but can make PFs more similar, especially on small-K instances.",
        "Configuration cost is reported as zero for Experiment A because all methods are fixed baselines, not meta-learning configuration searches.",
    ]

    run_metrics_path = os.path.join(report_dir, "run_metrics.csv")
    inst_metrics_path = os.path.join(report_dir, "instance_method_metrics.csv")
    overall_path = os.path.join(report_dir, "overall_method_summary.csv")
    stat_path = os.path.join(report_dir, "statistical_tests.csv")
    availability_path = os.path.join(report_dir, "metric_availability.csv")
    generation_manifest_path = os.path.join(report_dir, "generation_metrics_manifest.csv")
    run_metrics.to_csv(run_metrics_path, index=False, encoding="utf-8-sig")
    inst_metrics.to_csv(inst_metrics_path, index=False, encoding="utf-8-sig")
    overall.to_csv(overall_path, encoding="utf-8-sig")
    pd.DataFrame(stat_rows).to_csv(stat_path, index=False, encoding="utf-8-sig")
    availability.to_csv(availability_path, index=False, encoding="utf-8-sig")
    pd.DataFrame(generation_index_rows).to_csv(generation_manifest_path, index=False, encoding="utf-8-sig")

    print("Drawing figures...", flush=True)
    fig_paths = {}
    fig_paths["dashboard"] = os.path.join(figures_dir, "figure_1_metric_dashboard.png")
    multi_metric_bars(fig_paths["dashboard"], overall)
    eaf_vals = overall["mean_EAF_Band_Width"].to_dict()
    fig_paths["eaf"] = os.path.join(figures_dir, "figure_4_eaf_band_width.png")
    bar_chart(fig_paths["eaf"], "Figure 4. EAF Band Width", eaf_vals, "Lower is more stable", higher_better=False)
    runtime_vals = overall["mean_Runtime"].to_dict()
    fig_paths["runtime"] = os.path.join(figures_dir, "figure_5_runtime.png")
    bar_chart(fig_paths["runtime"], "Figure 5. Runtime", runtime_vals, "Seconds per final optimization run", higher_better=False)
    fig_paths["scatter"] = os.path.join(figures_dir, "figure_6_stability_diversity.png")
    scatter_chart(fig_paths["scatter"], inst_metrics)

    candidate_pool = manifest[manifest["split"].eq("test")]
    if candidate_pool.empty:
        candidate_pool = manifest
    candidate = candidate_pool.sort_values(["assets", "K"]).iloc[-1]
    rep_instance = candidate["instance"]
    rep_key = rep_instance
    rep_title = f"{candidate['split']} | {rep_instance} | assets={candidate['assets']} | K={candidate['K']}"
    rep_rows = {m: by_instance_method_points[(rep_key, m)] for m in METHODS}
    ref_pair = (ref_info[rep_key]["ideal"], ref_info[rep_key]["nadir"])
    fig_paths["overlay"] = os.path.join(figures_dir, "figure_2_pf_overlay.png")
    pf_overlay(fig_paths["overlay"], rep_title, rep_rows, ref_pair)
    fig_paths["heatmap"] = os.path.join(figures_dir, "figure_3_pf_heatmap.png")
    heatmap_chart(fig_paths["heatmap"], rep_title, rep_rows, ref_pair)

    completeness = {
        "expected_runs": expected,
        "complete_runs": expected - len(missing_rows),
        "missing_count": len(missing_rows),
        "missing_csv": missing_csv,
    }
    output_docx = os.path.join(report_dir, REPORT_DOCX_NAME)
    print("Writing DOCX report...", flush=True)
    make_report(report_dir, completeness, availability, overall, inst_metrics, stat_rows, fig_paths, warnings, output_docx)

    print(f"REPORT_DIR={report_dir}")
    print(f"DOCX={output_docx}")
    print(f"RUN_METRICS={run_metrics_path}")
    print(f"GENERATION_METRICS_MANIFEST={generation_manifest_path}")
    print(f"OVERALL={overall_path}")
    print(f"MISSING_OUTPUTS={len(missing_rows)}")
    print(f"EXPECTED_RUNS={expected}")
    print(f"COMPLETE_RUNS={expected - len(missing_rows)}")


if __name__ == "__main__":
    main()

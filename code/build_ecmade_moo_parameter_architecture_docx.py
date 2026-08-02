from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "ECMADE-MOO_參數設定與演算法架構說明.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
HEADER_FILL = "E8EEF5"
BORDER = "B7C9DD"


def set_run_font(run, size=11, bold=False, color=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def set_table_widths(table, widths_in: list[float]) -> None:
    dxa = [int(w * 1440) for w in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, width in enumerate(dxa):
            cell = row.cells[i]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_in: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        shade_cell(hdr[i], HEADER_FILL)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        row = table.add_row()
        for i, text in enumerate(values):
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    set_table_widths(table, widths_in)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("ECMADE-MOO Parameters")
    set_run_font(r, size=9, color=MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ECMADE-MOO 參數設定與演算法架構說明")
    set_run_font(r, size=22, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    r = p.add_run("本文件整理目前 ECMADE-MOO 程式採用之非簡化版 ECMADE 設定，以及多目標 NSGA-II 環境選擇機制。")
    set_run_font(r, color=MUTED)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


def build() -> Path:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    doc.add_heading("1. 演算法定位", level=1)
    doc.add_paragraph(
        "ECMADE-MOO 是將 ECMADE 的多子群協同演化、自適應 F/CR 控制與停滯觸發式資訊交換，結合 NSGA-II 多目標環境選擇後形成的多目標搜尋架構。"
    )
    add_bullets(
        doc,
        [
            "保留 ECMADE 的三子群 DE 搜尋策略，不使用簡化版共用 F/CR 狀態。",
            "以 Pareto dominance、Fast Non-dominated Sorting、Crowding Distance 與 Elitism 建構 Pareto Front。",
            "投資組合問題採兩目標：最小化風險與最小化負報酬。",
        ],
    )

    doc.add_heading("2. 使用者執行參數", level=1)
    add_table(
        doc,
        ["參數", "預設值", "說明"],
        [
            ["--problem", "ZDT1", "單一測試問題。支援 ZDT1-4/ZDT6、DTLZ1-7、UF1-10、ORLIB、PORTFOLIO。"],
            ["--problems", "無", "批次執行多個 benchmark；若設定此參數，會覆蓋 --problem。"],
            ["--pop-size", "100", "族群大小。因同子群抽樣需要至少 6 個體，建議 pop-size >= 6 × subpops。"],
            ["--max-fe", "10000", "最大函數評估次數。初始族群評估也會計入。"],
            ["--subpops", "3", "子群數量。預設為 p1、p2、p3 三種 ECMADE 搜尋角色。"],
            ["--seed", "2026", "隨機種子，用於重現實驗結果。"],
            ["--dimension", "依問題預設", "決策變數維度。ZDT/UF 通常為 30；DTLZ 可依需求指定。"],
            ["--objectives", "3 for DTLZ", "DTLZ 目標數；ZDT/UF 由問題定義決定。"],
            ["--out-dir", "ecmade_moo_outputs", "輸出 population、Pareto front、history 與 metrics 的資料夾。"],
            ["--pf-samples", "10000", "reference Pareto Front 取樣數，用於 benchmark IGD 計算。"],
        ],
        [1.25, 1.05, 4.2],
    )

    doc.add_heading("3. ECMADE 內部參數", level=1)
    add_table(
        doc,
        ["參數", "值", "論文對應與用途"],
        [
            ["archive_size H", "20", "每個個體各自保存最近 H 筆成功 F/CR，形成 RSF_i 與 RSCR_i。"],
            ["theta", "1/13", "自適應修正係數；用於更新每個個體的 μF_i 與 μCR_i。"],
            ["stagnation_threshold C", "50", "搜尋停滯門檻；Pareto archive 連續未新增非支配解超過 C 時觸發資訊交換。"],
            ["exploitation_alpha α", "0.8", "公式 (16) 的 best-guided 權重，用於降低 p2 過度依賴 best 的風險。"],
            ["initial_mu_f", "(0.9, 0.8, 0.8)", "p1、p2、p3 的初始 μF。"],
            ["initial_mu_cr", "(0.9, 0.5, 0.5)", "p1、p2、p3 的初始 μCR。"],
        ],
        [1.65, 1.15, 3.7],
    )

    doc.add_heading("4. F/CR 自適應控制", level=1)
    add_table(
        doc,
        ["項目", "目前設定"],
        [
            ["F 產生方式", "F_i,G = randc(μF_i,G, 0.1)；若 F <= 0 則重抽，若 F > 1 則截為 1。"],
            ["CR 產生方式", "CR_i,G = randn(μCR_i,G, 0.1)，並限制在 [0, 1]。"],
            ["成功參數判定", "若 offspring 經 NSGA-II 環境選擇後存活，該次 F/CR 視為成功參數。"],
            ["時間權重", "w_k = exp(g_k / G - 1)，越接近目前世代的成功參數權重越大。"],
            ["μF 更新", "μF_i ← (1 - θ) μF_i + θ × weighted Lehmer mean(RSF_i)。"],
            ["μCR 更新", "μCR_i ← (1 - θ) μCR_i + θ × weighted arithmetic mean(RSCR_i)。"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("5. 三子群搜尋策略", level=1)
    add_table(
        doc,
        ["子群", "角色", "Mutation 策略", "公式"],
        [
            ["p1", "Exploration", "DE/rand/2", "v = x_r1 + F(x_r2 - x_r3) + F(x_r4 - x_r5)"],
            ["p2", "Exploitation", "best-guided DE/rand-best/2 變形", "v = α x_best + F(x_r1 - x_r2) + F(x_r3 - x_r4)"],
            ["p3", "Balance", "DE/rand/1 + DE/current-to-best/1", "v = (1 - ω)(x_r1 + F(x_r2 - x_r3)) + ω(x_i + F(x_best - x_i) + F(x_r4 - x_r5))"],
        ],
        [0.55, 1.05, 1.9, 3.0],
    )
    add_bullets(
        doc,
        [
            "所有 r1、r2、r3、r4、r5 皆從 target 所屬子群內抽樣，索引互異且排除 target。",
            "ω = G / MG；目前程式以 max-fe / pop-size 近似 MG，使 p3 由探索逐步轉向開發。",
            "多目標情境沒有唯一 global best，因此 x_best 由 NSGA-II rank 最佳且 crowding 較佳的代表性精英個體擔任。",
        ],
    )

    doc.add_heading("6. 多目標環境選擇", level=1)
    add_table(
        doc,
        ["步驟", "說明"],
        [
            ["合併族群", "將 parent population 與 offspring population 合併。"],
            ["Pareto dominance", "以多目標支配關係建立個體優劣。"],
            ["Fast Non-dominated Sorting", "將合併族群分成 F1、F2、... 多層 Pareto fronts。"],
            ["Crowding Distance", "當最後一層 front 無法全部保留時，優先保留分布較疏的個體。"],
            ["Elitism", "保留 N 個個體進入下一世代，形成 NSGA-II 式環境選擇。"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("7. Pareto Archive 與資訊交換", level=1)
    add_table(
        doc,
        ["機制", "目前設定"],
        [
            ["Pareto archive", "保存歷史非支配解；輸出 Pareto front 時以 archive 為準。"],
            ["停滯判斷", "若 archive 連續多代沒有新增非支配解，視為搜尋停滯。"],
            ["觸發條件", "stagnation > C，預設 C = 50。"],
            ["資訊交換", "重新隨機平均分群，並將 Pareto/NSGA-II 精英複製到各子群取代較差個體。"],
            ["狀態複製", "複製 decision vector、objective values、μF_i、μCR_i、RSF_i 與 RSCR_i。"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("8. 輸出檔案", level=1)
    add_table(
        doc,
        ["檔案", "內容"],
        [
            ["*_population_variables.csv", "最後族群的決策變數。"],
            ["*_population_objectives.csv", "最後族群的目標值。"],
            ["*_pareto_variables.csv", "Pareto archive 中非支配解的決策變數。"],
            ["*_pareto_objectives.csv", "Pareto archive 中非支配解的目標值。"],
            ["*_history.csv", "generation、evaluations、nondominated、archive_size、stagnation、exchanges。"],
            ["*_metrics.txt", "evaluations、pareto_size；benchmark 另包含 IGD。"],
        ],
        [2.3, 4.2],
    )

    doc.add_heading("9. 建議正式實驗設定", level=1)
    add_table(
        doc,
        ["用途", "建議設定"],
        [
            ["快速確認程式可跑", "pop-size=60, max-fe=1000，至少確保每個子群有 6 個體。"],
            ["一般 benchmark", "pop-size=100, max-fe=10000，可與 NSGA-II baseline 進行初步比較。"],
            ["較困難 UF/DTLZ", "pop-size=100 或以上，max-fe=30000 或以上。"],
            ["正式統計", "固定參數，多 seed 獨立執行 20-30 runs，統計 IGD mean/std。"],
            ["投資組合最佳化", "依資產數調整 pop-size；資產數高時提高 max-fe。"],
        ],
        [1.65, 4.85],
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "ECMADE-MOO 參數設定與演算法架構說明"
    doc.core_properties.subject = "ECMADE-MOO parameters and architecture"
    doc.core_properties.author = "NCHU Lab"
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())

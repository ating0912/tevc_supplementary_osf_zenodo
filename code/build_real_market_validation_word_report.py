from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
BASE = ROOT / "p0_lite_outputs" / "p1_rolling_window_market_validation_20260719"
SUMMARY = BASE / "configured_ecmade_comparison_summary"
OUT = ROOT / "Real-market_validation_experiment_report_20260727.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(95, 95, 95)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_east_asian_font(run, font_name: str = "Microsoft JhengHei") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE7")


def style_table(table, widths: list[float]) -> None:
    table.autofit = False
    set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_east_asian_font(run)
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = INK


def fixed(num: float, digits: int = 4) -> str:
    return f"{num:.{digits}f}"


def pct(num: float) -> str:
    return f"{num * 100:.2f}%"


def method_label(name: str) -> str:
    return {
        "HandCrafted_ECMADE_MOO": "Hand-crafted",
        "BayesianConfig_ECMADE_MOO": "Bayesian-config",
        "MetaDesigned_ECMADE_MOO": "Meta-designed",
        "ExperimentC_StabilityAware_ECMADE_MOO": "Stability-aware",
    }.get(name, name)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asian_font(run)
        run.font.color.rgb = BLUE if level < 3 else DARK_BLUE


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_east_asian_font(r)
        r.bold = True
        r.font.color.rgb = INK
        r2 = p.add_run(text[len(bold_prefix) :])
        set_east_asian_font(r2)
        r2.font.color.rgb = INK
    else:
        r = p.add_run(text)
        set_east_asian_font(r)
        r.font.color.rgb = INK


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.add_run(text),:
        set_east_asian_font(run)
        run.font.color.rgb = INK


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_east_asian_font(r)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_east_asian_font(r2)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_dataframe_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    style_table(table, widths)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def main() -> None:
    overall = pd.read_csv(SUMMARY / "configured_overall_summary.csv")
    cost = pd.read_csv(SUMMARY / "configured_transaction_cost_overall.csv")
    theta = pd.read_csv(SUMMARY / "configured_theta_usage_by_method.csv")
    ranked = pd.read_csv(SUMMARY / "configured_window_method_ranked.csv")
    completeness = pd.read_csv(SUMMARY / "configured_run_completeness.csv")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Real-market Validation 實驗報告")
    set_east_asian_font(r)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("Hand-crafted、Bayesian、Meta-designed 與 Stability-aware ECMADE-MOO 之 rolling-window external validation")
    set_east_asian_font(sr)
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED

    add_callout(
        doc,
        "摘要結論",
        "四種 configuration protocols 共完成 1320 runs。Hand-crafted ECMADE-MOO 在整體 RankScore 上最佳；Bayesian-config ECMADE-MOO 則取得最高 after-cost annual net return，並在 PF Overlap 與 PF Drift 上呈現最佳 Pareto-front stability。Stability-aware 版本相較 Meta-designed 改善 downside-risk 與 PF Drift，但尚未全面超越固定配置。",
    )

    add_heading(doc, "1. 實驗目的", 1)
    add_paragraph(
        doc,
        "本實驗作為 Experiment A/B/C 之外的 external validation / robustness check，目的在於檢驗 ECMADE-MOO 的不同 configuration protocols 是否能在真實市場資料、rolling-window out-of-sample 測試與交易成本條件下維持穩定表現。此實驗不重新產生 training labels，也不讓測試期間資料參與 theta selection，而是將已建立的四種配置版本直接移植到 real-market setting 中比較。",
    )

    add_heading(doc, "2. 實驗設定", 1)
    add_paragraph(
        doc,
        "實驗採用 S&P 100、NASDAQ 100 與 Taiwan 50 三個真實市場股票池。每個 universe 建立 11 個 rolling windows，共 33 個 universe-window。每個 window 使用 3 年 training window 建立 expected return、covariance matrix 與 downside-risk scenarios，後續 6 個月 testing window 僅用於 out-of-sample evaluation。每個 method-window 執行 10 independent runs，並從 final Pareto front 中選擇 training Sharpe ratio 最高的 portfolio 進入測試期。",
    )
    add_paragraph(
        doc,
        "主要交易成本設定為 10 bps，並額外測試 20 bps 與 50 bps 的 sensitivity analysis。PF stability 指標使用與 Experiment B/C 一致的 common reference front、normalization 與 post-processing pipeline。",
    )

    comp_rows = [
        [method_label(r.method), str(int(r.universe_windows)), str(int(r.runs))]
        for r in completeness.itertuples(index=False)
    ]
    add_dataframe_table(doc, ["Method", "Universe-windows", "Runs"], comp_rows, [3.2, 1.6, 1.2])

    add_heading(doc, "3. Theta configuration 使用情形", 1)
    add_paragraph(
        doc,
        "Hand-crafted 使用固定人工設定；Bayesian configuration 在所有 real-market windows 中使用 theta_02。Meta-designed selector 主要選擇 theta_21，顯示 Random Forest 在多數 real-market windows 上偏好同一類型配置。Stability-aware selector 則較分散，最常選擇 theta_144，其次為 theta_034 與 theta_089，代表 stability-aware label 對不同市場 window 的配置差異較敏感。",
    )
    theta_rows = []
    for method, group in theta.groupby("method", sort=False):
        usage = ", ".join(f"{row.theta_id}: {int(row.windows)}" for row in group.itertuples(index=False))
        theta_rows.append([method_label(method), usage])
    add_dataframe_table(doc, ["Method", "Theta usage"], theta_rows, [2.0, 4.3])

    add_heading(doc, "4. 整體結果", 1)
    add_paragraph(
        doc,
        "整體而言，Hand-crafted ECMADE-MOO 在綜合 RankScore 上表現最佳，於 33 個 universe-window 中取得 19 次第一名，顯示原始人工設計配置在真實市場 rolling validation 下仍具高度競爭力。Bayesian-config ECMADE-MOO 則取得最高 mean annual net return，且在 PF Overlap 與 PF Drift 上表現最佳，代表其 Pareto front 較接近 common reference front 且跨 runs 的 front 位置較穩定。",
    )
    add_paragraph(
        doc,
        "Meta-designed 與 Stability-aware 版本雖然沒有在整體 RankScore 上超越 Hand-crafted 或 Bayesian-config，但 Stability-aware 在 maximum drawdown、CVaR 與 PF Drift 上優於 Meta-designed，顯示 stability-aware label protocol 對風險與 front drift 的控制有一定效果。",
    )
    order = [
        "HandCrafted_ECMADE_MOO",
        "BayesianConfig_ECMADE_MOO",
        "ExperimentC_StabilityAware_ECMADE_MOO",
        "MetaDesigned_ECMADE_MOO",
    ]
    overall = overall.set_index("method").loc[order].reset_index()
    overall_rows = []
    for row in overall.itertuples(index=False):
        overall_rows.append(
            [
                method_label(row.method),
                fixed(row.mean_annual_net_return),
                fixed(row.mean_sharpe),
                fixed(row.mean_sortino),
                fixed(row.mean_max_drawdown),
                fixed(row.mean_cvar95_loss),
                fixed(row.mean_rebalance_turnover),
                fixed(row.mean_PF_Overlap),
                fixed(row.mean_PF_Drift),
                fixed(row.mean_RankScore),
                str(int(row.first_place_windows)),
            ]
        )
    add_dataframe_table(
        doc,
        ["Method", "Ann. net", "Sharpe", "Sortino", "MDD", "CVaR", "Turnover", "PF Overlap", "PF Drift", "RankScore", "1st"],
        overall_rows,
        [1.15, 0.58, 0.58, 0.58, 0.55, 0.55, 0.62, 0.65, 0.58, 0.62, 0.35],
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    nr = note.add_run("註：MDD 為負值，越接近 0 越好；CVaR(95%) loss、turnover、PF Drift 與 RankScore 越低越好；annual net return、Sharpe、Sortino 與 PF Overlap 越高越好。")
    set_east_asian_font(nr)
    nr.font.size = Pt(9)
    nr.font.color.rgb = MUTED

    add_heading(doc, "5. 交易成本敏感度", 1)
    add_paragraph(
        doc,
        "在 10 bps、20 bps 與 50 bps 三種交易成本設定下，annual net return 的排序保持一致：Bayesian-config 最高，其次為 Hand-crafted、Stability-aware，最後為 Meta-designed。此結果顯示主要結論對交易成本設定具有穩定性，並非由單一 cost assumption 驅動。",
    )
    cost_rows = []
    for scenario in ["10bps", "20bps", "50bps"]:
        group = cost[cost["cost_scenario"] == scenario].copy()
        group = group.sort_values("rank_annual_net_return")
        ranking = " > ".join(f"{method_label(r.method)} {fixed(r.mean_annual_net_return)}" for r in group.itertuples(index=False))
        cost_rows.append([scenario, method_label(group.iloc[0]["method"]), ranking])
    add_dataframe_table(doc, ["Cost", "Best method", "Annual net return ranking"], cost_rows, [0.7, 1.4, 4.2])

    add_heading(doc, "6. 各市場股票池結果", 1)
    by = (
        ranked.groupby(["universe", "method"])
        .agg(
            mean_rankscore=("RankScore", "mean"),
            mean_window_rank=("WindowRank", "mean"),
            first_windows=("WindowRank", lambda s: int((s == 1).sum())),
            annual_net_return=("annual_net_return_mean", "mean"),
            sharpe=("sharpe_mean", "mean"),
            cvar95_loss=("cvar95_loss_mean", "mean"),
            pf_overlap=("PF_Overlap_mean", "mean"),
            pf_drift=("PF_Drift_mean", "mean"),
        )
        .reset_index()
    )
    add_paragraph(
        doc,
        "在 NASDAQ100 與 SP100 中，Hand-crafted 的 mean WindowRank 均為 1.5455，分別取得 8 與 5 個 first-place windows，顯示其風險調整後表現在美股大型股 universe 中較穩定。TAIWAN50 中 Hand-crafted 仍取得 6 個 first-place windows，但 Bayesian-config 的 annual net return 最高，達 0.4255。",
    )
    universe_rows = []
    for universe in ["NASDAQ100", "SP100", "TAIWAN50"]:
        subset = by[by["universe"] == universe].sort_values("mean_window_rank")
        best = subset.iloc[0]
        top_return = subset.sort_values("annual_net_return", ascending=False).iloc[0]
        universe_rows.append(
            [
                universe,
                method_label(best["method"]),
                str(int(best["first_windows"])),
                f"{method_label(top_return['method'])} ({fixed(top_return['annual_net_return'])})",
                fixed(best["mean_window_rank"]),
            ]
        )
    add_dataframe_table(doc, ["Universe", "Best by WindowRank", "1st windows", "Highest annual net return", "Mean WindowRank"], universe_rows, [1.0, 1.5, 0.85, 2.2, 0.85])

    add_heading(doc, "7. Pareto-front stability 分析", 1)
    add_paragraph(
        doc,
        "PF stability 指標使用與 Experiment B/C 一致的 common reference front、normalization 與 post-processing pipeline。Bayesian-config 在 PF Overlap 與 PF Drift 上皆為最佳，平均 PF Overlap 為 0.9222，PF Drift 為 0.0367，表示其 front 與 common reference front 較一致，且跨 independent runs 的 front centroid 變動較小。",
    )
    add_paragraph(
        doc,
        "Stability-aware 的 PF Drift 為 0.0446，優於 Hand-crafted 的 0.0538 與 Meta-designed 的 0.0526，支持 stability-aware label 對 front drift 的控制具有效果。然而，其 PF Overlap 為 0.9033，低於其他方法，表示 stability-aware 配置可能產生較穩定但覆蓋較不足的 Pareto front。此現象可作為後續調整 J-score 權重的依據。",
    )

    add_heading(doc, "8. 討論與結論", 1)
    add_paragraph(
        doc,
        "本 real-market validation 顯示，原始 Hand-crafted ECMADE-MOO 在真實市場資料上仍是強 baseline，尤其在 Sharpe、Sortino、MDD、CVaR、turnover 與 runtime 上具有優勢。Bayesian-config 則在 annual net return 與 Pareto-front stability 上表現突出，且在不同交易成本設定下皆維持最高 after-cost annual net return，顯示 Bayesian-selected theta_02 具備良好的 external robustness。",
    )
    add_paragraph(
        doc,
        "Meta-designed selector 並未在本實驗中帶來整體優勢，可能原因是 OR-Library benchmark 的 meta-feature distribution 與真實市場 rolling windows 存在 domain shift，使 selector 在 real-market setting 中偏向選擇少數 theta，未能充分反映市場狀態差異。Stability-aware selector 相較 Meta-designed 改善了 CVaR、MDD 與 PF Drift，但仍未轉化為最高 overall RankScore，表示 stability-aware label engineering 的方向有效，但仍需進一步調整 J-score 權重或納入 real-market instances 進行 calibration。",
    )
    add_paragraph(
        doc,
        "綜合而言，Hand-crafted ECMADE-MOO 是最穩健的 overall performer；Bayesian-config ECMADE-MOO 在 after-cost annual return 與 Pareto-front stability 上最佳；Stability-aware ECMADE-MOO 能改善相對於 Meta-designed 的 downside risk 與 PF Drift，但尚未全面超越 fixed configurations。此結果可在論文中定位為 robustness check：所提出的 automated/stability-aware design 在 benchmark 外仍具合理表現，但 real-market deployment 需要額外的 domain calibration 與 J-score sensitivity tuning。",
    )

    add_heading(doc, "9. 結果位置", 1)
    add_paragraph(
        doc,
        rf"完整彙整結果位於：{SUMMARY}",
    )
    add_dataframe_table(
        doc,
        ["File", "Content"],
        [
            ["configured_overall_summary.csv", "四種方法整體彙整與總排名"],
            ["configured_window_method_summary.csv", "每個 universe-window-method 的平均與標準差"],
            ["configured_window_method_ranked.csv", "每個 universe-window 的 RankScore 與 WindowRank"],
            ["configured_run_metrics_with_pf_stability.csv", "1320 runs 的 run-level metrics 與 PF stability 指標"],
            ["configured_transaction_cost_overall.csv", "10/20/50 bps 交易成本敏感度整體表"],
            ["configured_theta_usage_by_method.csv", "四種 protocol 的 theta 使用次數"],
        ],
        [2.6, 3.7],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Real-market validation report")
    set_east_asian_font(fr)
    fr.font.size = Pt(9)
    fr.font.color.rgb = MUTED

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

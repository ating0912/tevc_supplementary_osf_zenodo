import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = r"C:\Users\yiting\Documents\Playground"
OUT_ROOT = os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio")
REPORT_DIR = os.path.join(OUT_ROOT, "experiment_A_report_20260702_134204")
FIG_DIR = os.path.join(REPORT_DIR, "figures")
DOCX_OUT = os.path.join(REPORT_DIR, "Experiment_A_synthetic_results_report_zh_final.docx")


def set_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, rgb, before, after in [
        ("Heading 1", 16, (46, 116, 181), 12, 6),
        ("Heading 2", 13, (46, 116, 181), 8, 4),
        ("Heading 3", 11.5, (31, 77, 120), 6, 3),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(*rgb)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("" if pd.isna(text) else str(text))
    set_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, headers, widths=None, title=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_font(r, size=10.5, bold=True, color=(31, 77, 120))
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, j), "E8EEF5")
    for i, row in enumerate(rows, start=1):
        for j, h in enumerate(headers):
            val = row.get(h, "")
            set_cell_text(table.cell(i, j), val, size=8.2)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths):
                row.cells[j].width = Inches(width)
    doc.add_paragraph()
    return table


def add_note(doc, label, body):
    p = doc.add_paragraph()
    r = p.add_run(label)
    set_font(r, size=10.2, bold=True, color=(31, 77, 120))
    r2 = p.add_run(" " + body)
    set_font(r2, size=10.2)


def fmt(x, digits=4):
    if pd.isna(x):
        return ""
    if abs(float(x)) >= 100:
        return f"{float(x):.2f}"
    return f"{float(x):.{digits}g}"


def metric_winner(overall, metric, direction):
    col = f"mean_{metric}"
    if direction == "max":
        idx = overall[col].idxmax()
    else:
        idx = overall[col].idxmin()
    return idx, overall.loc[idx, col]


def add_picture_with_caption(doc, file_name, caption, explanation, interpretation):
    path = os.path.join(FIG_DIR, file_name)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(9.4))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, size=9.2, bold=True, color=(31, 77, 120))
    add_note(doc, "圖表說明：", explanation)
    add_note(doc, "結果解讀：", interpretation)


def main():
    overall = pd.read_csv(os.path.join(REPORT_DIR, "overall_method_summary.csv"), index_col="method")
    stats = pd.read_csv(os.path.join(REPORT_DIR, "statistical_tests.csv"))
    availability = pd.read_csv(os.path.join(REPORT_DIR, "metric_availability.csv"))
    inst = pd.read_csv(os.path.join(REPORT_DIR, "instance_method_metrics.csv"))
    missing = pd.read_csv(os.path.join(REPORT_DIR, "missing_outputs.csv"))

    rank_order = overall.sort_values("RankScore")
    best_rank = rank_order.index[0]
    hv_best, hv_val = metric_winner(overall, "HV", "max")
    igd_best, igd_val = metric_winner(overall, "IGD", "min")
    overlap_best, overlap_val = metric_winner(overall, "PF_Overlap", "max")
    eaf_best, eaf_val = metric_winner(overall, "EAF_Band_Width", "min")
    drift_best, drift_val = metric_winner(overall, "PF_Drift", "min")
    runtime_best, runtime_val = metric_winner(overall, "Runtime", "min")
    diversity_best, diversity_val = metric_winner(overall, "Diversity", "max")

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("實驗 A 新版結果報告書\n")
    set_font(r, size=22, bold=True, color=(11, 37, 69))
    r = title.add_run("ECMADE-MOO 與主要 Baseline 於 Synthetic Constrained Portfolio Instances 之比較")
    set_font(r, size=11, color=(80, 80, 80))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}｜資料來源：{REPORT_DIR}")
    set_font(r, size=9.5, color=(90, 90, 90))

    doc.add_heading("1. 執行狀態與資料完整性", level=1)
    expected_runs = 6 * 192 * 30
    complete_runs = expected_runs - len(missing)
    doc.add_paragraph(
        f"本版報告已納入新版 ECMADE_MOO 與新增 A_MPMO，並與 NSGAII、SPEA2、MOEAD、GDE3 共同重算指標。"
        f"預期輸出為 6 methods × 192 instances × 30 runs = {expected_runs:,} runs；缺漏數為 {len(missing)}，完整率為 {complete_runs / expected_runs:.1%}。"
    )
    add_table(
        doc,
        [
            {"項目": "資料集", "內容": "Synthetic constrained portfolio instances"},
            {"項目": "splits", "內容": "train / validation / test"},
            {"項目": "資產數 n", "內容": "50、100、200、500"},
            {"項目": "K 定義", "內容": "實際 cardinality。synthetic instance 名稱中的 k05/k10/k20/k30 表示 K/n 比率；例如 n=500,k30 對應 K=150。"},
            {"項目": "methods", "內容": "NSGAII、SPEA2、MOEAD、GDE3、A_MPMO、ECMADE_MOO"},
            {"項目": "runs", "內容": "每 method × instance 30 independent runs"},
            {"項目": "共同設定", "內容": "population size N=100、maxFE=10000、seed=run index、rng=mcg16807、相同 PortfolioORLIB constraint handling"},
        ],
        ["項目", "內容"],
        widths=[1.4, 8.2],
        title="Table 1. 實驗輸入與共同設定",
    )

    doc.add_heading("2. 指標與輸出檢查", level=1)
    avail_rows = [
        {
            "類別": r["Category"],
            "指標": r["Item"],
            "狀態": r["Status"],
            "來源或定義": r["Source_or_definition"],
        }
        for _, r in availability.iterrows()
    ]
    add_table(doc, avail_rows, ["類別", "指標", "狀態", "來源或定義"], widths=[1.1, 2.0, 1.5, 5.0], title="Table 2. 指標可用性")
    add_note(
        doc,
        "注意：",
        "HV、IGD、PF overlap、EAF band width、PF drift、diversity、runtime、feasible rate 均已由最終 PF、generation snapshots、runtime.csv 與 feasible_rate.csv 重算。"
        "目前 per-generation HV/IGD 未以原始 log 直接保存，但每個 run 都有 generation_pf_points.csv，因此可在後處理階段補算。",
    )

    doc.add_heading("3. 主要結論", level=1)
    doc.add_paragraph(
        f"綜合 RankScore 排名第一為 {best_rank}。若分開看單一指標，HV 最佳為 {hv_best} ({fmt(hv_val)})，"
        f"IGD 最佳為 {igd_best} ({fmt(igd_val)})，PF overlap 最佳為 {overlap_best} ({fmt(overlap_val)})。"
        f"穩定性方面，EAF band width 最低為 {eaf_best} ({fmt(eaf_val)})，PF drift 最低為 {drift_best} ({fmt(drift_val)})。"
        f"Runtime 最低為 {runtime_best} ({fmt(runtime_val)} 秒/run)，diversity 最大為 {diversity_best} ({fmt(diversity_val)})。"
    )
    doc.add_paragraph(
        "整體而言，SPEA2 在此批 synthetic constrained instances 上呈現最均衡的 performance/stability 表現；"
        "NSGAII 速度最快且綜合排名第二；ECMADE_MOO 的 HV、IGD 與 diversity 具競爭力，但 EAF band width、PF drift 與 runtime 變異較大，顯示新版結果中穩定性仍需進一步改善。"
        "A_MPMO 已成功納入 baseline，但目前綜合排名略低於 ECMADE_MOO。"
    )

    doc.add_heading("4. Overall Method Ranking", level=1)
    rows = []
    for method, row in rank_order.iterrows():
        rows.append(
            {
                "Method": method,
                "RankScore": fmt(row["RankScore"], 3),
                "HV mean": fmt(row["mean_HV"]),
                "IGD mean": fmt(row["mean_IGD"]),
                "PF overlap": fmt(row["mean_PF_Overlap"]),
                "EAF width": fmt(row["mean_EAF_Band_Width"]),
                "PF drift": fmt(row["mean_PF_Drift"]),
                "Diversity": fmt(row["mean_Diversity"]),
                "Runtime": fmt(row["mean_Runtime"]),
            }
        )
    add_table(
        doc,
        rows,
        ["Method", "RankScore", "HV mean", "IGD mean", "PF overlap", "EAF width", "PF drift", "Diversity", "Runtime"],
        widths=[1.15, 0.8, 0.9, 0.9, 0.95, 0.9, 0.9, 0.9, 0.9],
        title="Table 3. 新版 overall results 與綜合排名",
    )

    doc.add_heading("5. 圖表結果與解讀", level=1)
    add_picture_with_caption(
        doc,
        "figure_1_metric_dashboard.png",
        "Figure 1. Overall metric dashboard.",
        "此圖將 performance、stability、diversity、cost 與 feasibility 指標並列，方便快速比較各方法的整體形狀。",
        "SPEA2 在 HV、IGD、PF overlap 與穩定性之間最平衡；ECMADE_MOO 的 HV 與 IGD 位居第二且 diversity 最高，但穩定性與 runtime 指標拖累綜合排名。",
    )
    add_picture_with_caption(
        doc,
        "figure_2_pf_overlay.png",
        "Figure 2. PF overlay on representative test instance.",
        "此圖疊合代表性 test instance 中各方法 30 runs 的 final Pareto front points，用來觀察 front 覆蓋範圍與重複執行的一致性。",
        "PF overlay 可補足單看 HV/IGD 的不足；若 front 離散或偏移，即使部分解品質高，也可能代表 repeated-run reliability 不足。",
    )
    add_picture_with_caption(
        doc,
        "figure_3_pf_heatmap.png",
        "Figure 3. PF heatmap on representative test instance.",
        "此圖將不同方法的 PF points 轉為密度熱圖，顏色越集中代表 repeated runs 較常落在相近的 Pareto 區域。",
        "熱圖可用來判斷方法是否只是偶爾找到好解，或能穩定重現相似的 PF 區域；這是 stability-aware 評估的重要視覺證據。",
    )
    add_picture_with_caption(
        doc,
        "figure_4_eaf_band_width.png",
        "Figure 4. EAF band width.",
        "EAF band width 衡量 repeated runs 的 attainment surface 不確定帶寬，數值越小代表 front 重複性越穩定。",
        f"本批資料中 {eaf_best} 的 EAF band width 最低；ECMADE_MOO 的 EAF width 較高，代表其 front 在不同 runs 間仍有較明顯的不確定性。",
    )
    add_picture_with_caption(
        doc,
        "figure_5_runtime.png",
        "Figure 5. Runtime comparison.",
        "Runtime 代表 final optimization cost，所有方法使用相同 maxFE 與 population size，因此可直接比較每 run 平均耗時。",
        f"{runtime_best} 是最快方法；MOEAD 最慢。ECMADE_MOO 與 A_MPMO 的 runtime 受演算法內部結構影響，需在後續 cost/runtime analysis 中一併討論。",
    )
    add_picture_with_caption(
        doc,
        "figure_6_stability_diversity.png",
        "Figure 6. Stability-diversity plot.",
        "此圖同時呈現 EAF band width 與 diversity，用來檢查方法是否以犧牲多樣性換取穩定，或以高多樣性造成 front 漂移。",
        "ECMADE_MOO diversity 最高，但 EAF width 也偏高，顯示多樣性並未完全轉化為穩定 front；SPEA2 則呈現較佳的整體折衷。",
    )

    doc.add_heading("6. 統計檢定摘要", level=1)
    fried = stats[stats["test"].eq("Friedman")].copy()
    stat_rows = []
    for _, row in fried.iterrows():
        stat_rows.append(
            {
                "Metric": row["metric"],
                "Test": row["test"],
                "Statistic": fmt(row["statistic"], 4),
                "p-value": f"{row['p_value']:.3e}",
                "Interpretation": "methods 間有顯著差異" if row["p_value"] < 0.05 else "未達顯著差異",
            }
        )
    add_table(doc, stat_rows, ["Metric", "Test", "Statistic", "p-value", "Interpretation"], widths=[1.4, 1.4, 1.2, 1.4, 3.8], title="Table 4. Friedman test 摘要")
    add_note(
        doc,
        "解讀：",
        "所有主要指標的 Friedman test 皆顯示 methods 間存在顯著差異。Wilcoxon approx. 已完整輸出於 statistical_tests.csv，可用於後續撰寫逐對比較段落。",
    )

    doc.add_heading("7. 分組補充觀察", level=1)
    split_summary = (
        inst.groupby(["split", "method"])
        .agg(HV=("HV", "mean"), IGD=("IGD", "mean"), EAF=("EAF_Band_Width", "mean"), Runtime=("Runtime", "mean"))
        .reset_index()
    )
    split_rows = []
    for split in ["train", "validation", "test"]:
        part = split_summary[split_summary["split"].eq(split)].copy()
        if part.empty:
            continue
        hv_m = part.loc[part["HV"].idxmax(), "method"]
        igd_m = part.loc[part["IGD"].idxmin(), "method"]
        eaf_m = part.loc[part["EAF"].idxmin(), "method"]
        split_rows.append({"Split": split, "Best HV": hv_m, "Best IGD": igd_m, "Lowest EAF": eaf_m})
    add_table(doc, split_rows, ["Split", "Best HV", "Best IGD", "Lowest EAF"], widths=[1.2, 2.1, 2.1, 2.1], title="Table 5. Train/validation/test 分組最佳方法")

    doc.add_heading("8. 檔案與再現性", level=1)
    doc.add_paragraph(f"原始輸出根目錄：{OUT_ROOT}")
    doc.add_paragraph(f"新版報告資料夾：{REPORT_DIR}")
    doc.add_paragraph(f"Run-level metrics：{os.path.join(REPORT_DIR, 'run_metrics.csv')}")
    doc.add_paragraph(f"Instance-method metrics：{os.path.join(REPORT_DIR, 'instance_method_metrics.csv')}")
    doc.add_paragraph(f"Overall summary：{os.path.join(REPORT_DIR, 'overall_method_summary.csv')}")
    doc.add_paragraph(f"Statistical tests：{os.path.join(REPORT_DIR, 'statistical_tests.csv')}")

    doc.add_heading("9. 後續建議", level=1)
    doc.add_paragraph(
        "下一步建議以本版結果作為 Experiment A baseline 報告，接著進行："
        "（1）檢查 ECMADE_MOO 穩定性來源，針對 adaptive exchange、elite injection、single-population 做 ablation；"
        "（2）針對 A_MPMO 補方法參數表與演算法差異說明；"
        "（3）若論文需要 generation-level HV/IGD 曲線，可由 generation_pf_points.csv 補算並追加 convergence plot。"
    )

    doc.save(DOCX_OUT)
    print(f"DOCX={DOCX_OUT}")


if __name__ == "__main__":
    main()

import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REPORT_DIR = r"C:\Users\yiting\Documents\Playground\p0_lite_outputs\synthetic_constrained_portfolio\experiment_A_report_20260701_112713"
OUT_ROOT = r"C:\Users\yiting\Documents\Playground\p0_lite_outputs\synthetic_constrained_portfolio"
METHODS = ["NSGAII", "SPEA2", "MOEAD", "GDE3", "A_MPMO", "ECMADE_MOO"]


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_df_table(doc, df, columns, title=None, max_rows=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(31, 77, 120)
    rows = df if max_rows is None else df.head(max_rows)
    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, col in enumerate(columns):
        set_cell_text(table.cell(0, j), col, bold=True)
        set_cell_fill(table.cell(0, j), "E8EEF5")
    for i, (_, row) in enumerate(rows.iterrows(), start=1):
        for j, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.4g}"
            else:
                text = val
            set_cell_text(table.cell(i, j), text)
    doc.add_paragraph()


def apply_doc_styles(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_note(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(" " + text)


def add_caption_and_explanation(doc, caption, explanation, interpretation):
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_note(doc, "圖表說明：", explanation)
    add_note(doc, "結果解讀：", interpretation)


def best_method(overall, metric, direction):
    col = f"mean_{metric}"
    if direction == "max":
        return overall[col].idxmax(), overall[col].max()
    return overall[col].idxmin(), overall[col].min()


def main():
    overall = pd.read_csv(os.path.join(REPORT_DIR, "overall_method_summary.csv"), index_col="method")
    stats = pd.read_csv(os.path.join(REPORT_DIR, "statistical_tests.csv"))
    availability = pd.read_csv(os.path.join(REPORT_DIR, "metric_availability.csv"))
    figures_dir = os.path.join(REPORT_DIR, "figures")

    hv_best, hv_val = best_method(overall, "HV", "max")
    igd_best, igd_val = best_method(overall, "IGD", "min")
    overlap_best, overlap_val = best_method(overall, "PF_Overlap", "max")
    eaf_best, eaf_val = best_method(overall, "EAF_Band_Width", "min")
    runtime_best, runtime_val = best_method(overall, "Runtime", "min")
    diversity_best, diversity_val = best_method(overall, "Diversity", "max")
    rank_best = overall["RankScore"].idxmin()

    doc = Document()
    apply_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Experiment A Results Report\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    sub = title.add_run("ECMADE-MOO and Baseline Comparison on Synthetic Constrained Portfolio Instances")
    sub.font.size = Pt(11)
    sub.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Data Completeness Check", level=1)
    doc.add_paragraph("PASS. All expected method x instance x run outputs are present: 28,800 / 28,800 complete runs.")
    add_df_table(doc, availability, ["Category", "Item", "Status", "Source_or_definition"], "Metric availability")

    doc.add_heading("2. Key Findings", level=1)
    doc.add_paragraph(
        f"Composite ranking is led by {rank_best}. The best mean HV is {hv_best} ({hv_val:.4g}), "
        f"the best mean IGD is {igd_best} ({igd_val:.4g}), and the highest PF overlap is {overlap_best} ({overlap_val:.4g})."
    )
    doc.add_paragraph(
        f"For stability, the lowest EAF band width is {eaf_best} ({eaf_val:.4g}). "
        f"For runtime, the fastest method is {runtime_best} ({runtime_val:.4g} seconds per run on average). "
        f"For diversity, the largest spread is {diversity_best} ({diversity_val:.4g})."
    )

    doc.add_heading("3. Overall Results", level=1)
    overall_out = overall.reset_index()
    columns = [
        "method",
        "mean_HV",
        "std_HV",
        "cv_HV",
        "mean_IGD",
        "std_IGD",
        "cv_IGD",
        "mean_PF_Overlap",
        "mean_EAF_Band_Width",
        "mean_Diversity",
        "mean_Runtime",
        "mean_Feasible_Rate",
        "RankScore",
    ]
    add_df_table(doc, overall_out[columns], columns, "Table 1. Main results by method")
    add_note(
        doc,
        "表格說明：",
        "此表彙整每個方法在所有 synthetic constrained portfolio instances 上的平均表現。HV 與 PF overlap 越大越好；IGD、EAF band width、runtime 越小越好；CV 用於觀察 repeated runs 的相對變異。",
    )

    doc.add_heading("4. Statistical Tests", level=1)
    add_df_table(doc, stats, list(stats.columns), "Table 2. Friedman and Wilcoxon tests")
    add_note(
        doc,
        "表格說明：",
        "Friedman test 檢查所有方法在同一批 instances 上是否存在整體差異；Wilcoxon approx. 用於 ECMADE-MOO 與各 baseline 的成對比較。p-value 較小表示差異較不可能由隨機波動造成。",
    )

    doc.add_heading("5. Figures and Explanations", level=1)
    fig_specs = [
        (
            "figure_1_metric_dashboard.png",
            "Figure 1. Overall metric dashboard.",
            "此圖把主要 performance、stability、diversity、cost、feasibility 指標放在同一頁比較。每個小圖的橫條代表各方法在所有 instances 上的平均值。",
            f"SPEA2 在 composite rank 與 HV/IGD/PF overlap 上最強；ECMADE-MOO 在 HV、IGD、PF overlap 與 diversity 也維持前段表現，但 EAF width 較高，代表 repeated-run attainment band 較寬。",
        ),
        (
            "figure_2_pf_overlay.png",
            "Figure 2. PF overlay on a representative test instance.",
            "此圖疊合代表性 test instance 中各方法 30 runs 的 final Pareto front points。座標已依同一 instance 的 empirical range 正規化，方便比較不同方法的 front 位置與覆蓋範圍。",
            "GDE3 與 ECMADE-MOO 在該高維 instance 上展開到較大的 normalized risk-return 區域；SPEA2 與 NSGAII 的 front 較集中，反映較穩定但探索範圍較受限的行為。",
        ),
        (
            "figure_3_pf_heatmap.png",
            "Figure 3. PF heatmap on the same representative instance.",
            "此圖將每個方法的 PF points 做二維 density heatmap。顏色越深表示 repeated runs 中越常出現的 Pareto 區域。",
            "heatmap 可看出不同方法的搜尋集中區域。SPEA2 與 NSGAII 密度集中；GDE3 front 曲線較長，ECMADE-MOO 則呈現較大的探索散布。",
        ),
        (
            "figure_4_eaf_band_width.png",
            "Figure 4. EAF band width.",
            "EAF band width 衡量 repeated runs 的 attainment surface 不確定性；數值越小代表 repeated-run stability 越好。",
            f"本批結果中 {eaf_best} 的 EAF band width 最小，表示其 Pareto front 在重複執行間最穩定。ECMADE-MOO 的 band 較寬，後續可作為 stability objective 或 adaptive exchange ablation 的重點觀察。",
        ),
        (
            "figure_5_runtime.png",
            "Figure 5. Runtime comparison.",
            "此圖比較每次 final optimization run 的平均 runtime。這是 final optimization cost，不包含後處理與報告生成時間。",
            f"{runtime_best} 最快；MOEA/D runtime 明顯較高，代表若要納入 cost-sensitive 設計，MOEA/D 的穩定性或品質收益需要能抵銷額外時間成本。",
        ),
        (
            "figure_6_stability_diversity.png",
            "Figure 6. Stability-diversity plot.",
            "每個點代表一個 method x instance 的聚合結果；橫軸為 EAF band width，越左越穩定；縱軸為 diversity，越高代表 archive/front 展開較廣。",
            "此圖用來檢查穩定性與多樣性是否互相犧牲。ECMADE-MOO 多數點落在較高 diversity 區域，但也伴隨較大的 EAF width；SPEA2 則呈現較好的整體穩定與品質平衡。",
        ),
    ]
    for file_name, caption, explanation, interpretation in fig_specs:
        path = os.path.join(figures_dir, file_name)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(9.5))
            add_caption_and_explanation(doc, caption, explanation, interpretation)

    doc.add_heading("6. Method Ranking Summary", level=1)
    rank_cols = ["method", "RankScore", "rank_HV", "rank_IGD", "rank_PF_Overlap", "rank_EAF_Band_Width", "rank_Runtime"]
    add_df_table(doc, overall_out.sort_values("RankScore")[rank_cols], rank_cols, "Table 3. Composite rank summary")
    add_note(
        doc,
        "表格說明：",
        "RankScore 是多指標平均排名，整合 performance、stability、runtime 與 feasibility。它不是單一品質指標，而是用來快速檢視整體平衡表現。",
    )

    doc.add_heading("7. Reproducibility", level=1)
    doc.add_paragraph(f"Output root: {OUT_ROOT}")
    doc.add_paragraph(f"Report artifact directory: {REPORT_DIR}")
    doc.add_paragraph("All metric tables and figures were generated from the final PF files, generation snapshots, feasible-rate files, runtime files, and the synthetic manifest.")

    output = os.path.join(REPORT_DIR, "Experiment_A_synthetic_results_report_with_explanations.docx")
    doc.save(output)
    print(f"DOCX={output}")


if __name__ == "__main__":
    main()

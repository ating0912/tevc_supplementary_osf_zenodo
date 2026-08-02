from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

from build_experiment_c_stability_report_docx import (
    BASELINES,
    C_FINAL_DIR,
    C_METHOD,
    COMPARISON_DIR,
    DOCX_OUT,
    ROOT,
    SELECTOR_DIR,
    TRAIN_REPORT,
    VAL_REPORT,
    add_bullets,
    add_heading,
    add_kv_table,
    add_table,
    build_confirmatory_tests,
    build_sign_tests,
    data_completeness_rows,
    fmt,
    fmt_pct,
    load_data,
    setup_styles,
)


REPORT_STEM = "Experiment_C_穩定性感知_theta_selector_中文報告_20260717"


def set_east_asia_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        if style.element.rPr is not None:
            style.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft JhengHei")


def add_title_zh(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("實驗 C：穩定性感知 Theta Selector 報告")
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Validation label 後處理、穩定性感知 selector 訓練，以及 unseen test 最終比較")
    r.font.name = "Microsoft JhengHei"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)

    add_kv_table(
        doc,
        [
            ("生成日期", date.today().isoformat()),
            ("資料根目錄", str(ROOT)),
            ("核心方法", C_METHOD),
            ("比較基準", "Hand-crafted、Random、Bayesian、Meta-designed"),
        ],
    )


def add_report_content_zh(
    doc: Document,
    data: dict[str, pd.DataFrame],
    sign_tests: pd.DataFrame,
    confirmatory_tests: pd.DataFrame,
) -> None:
    overall = data["overall"]
    selector_summary = data["selector_summary"]
    assignment = data["assignment"]

    add_heading(doc, "1. 摘要")
    doc.add_paragraph(
        "實驗 C 的目的，是將原本以整體效能為主的 theta label，改寫成更強調 Pareto front 穩定性的標籤目標。"
        "本報告整理 validation label 後處理、C 專用 selector 訓練與驗證、unseen test final run，"
        "以及和實驗 B 四種 baseline 的 common-reference 比較結果。"
    )
    c = overall[overall["method"] == C_METHOD].iloc[0]
    meta = overall[overall["method"] == "MetaDesigned_ECMADE_MOO"].iloc[0]
    add_bullets(
        doc,
        [
            f"實驗 C 在 final common-reference comparison 中排名第一，overall_RankScore={fmt(c['overall_RankScore'], 4)}。",
            f"實驗 C 在 32 個 unseen test instances 中取得 {int(c['first_place_instances'])} 個 first-place instance；實驗 B Meta-designed 為 {int(meta['first_place_instances'])} 個。",
            "正式統計主端點預先指定為 per-instance RankScore，也就是 OverallRankScore 在配對檢定單位上的版本；數值越低越好。",
            "Validation label generation 已補齊，theta_089 不再缺 run；每個 theta x validation instance 都有 30 runs。",
            "目前主實驗 raw data 已足夠，不需要再跑主要 C final test；後續重點應放在統計檢定、圖表與論文文字整理。",
        ],
    )

    add_heading(doc, "2. 資料完整性與輸出檔案")
    add_table(
        doc,
        ["資料項目", "instance groups", "theta/method count", "完整性摘要"],
        data_completeness_rows(data),
        widths=[2500, 1700, 2100, 3060],
    )
    doc.add_paragraph(
        "Validation raw runs 的完整性檢查結果顯示，每個 theta-by-validation-instance group 的 run 數皆為 30，沒有缺失 group。"
        "Experiment C final test 也完成 32 groups x 30 runs，共 960 runs。"
    )

    add_heading(doc, "3. C 專用 Stability-aware Label")
    doc.add_paragraph(
        "實驗 C 從 raw outputs 後處理出的 HV、IGD、PF_Overlap、PF_Drift 指標重新建立 label。"
        "每個 theta 在同一個 instance 與 K 下先轉成 rank，再用以下公式計算 C_LabelScore，且 C_LabelScore 越大越好："
    )
    p = doc.add_paragraph()
    r = p.add_run("C_LabelScore = -0.2*rank_HV - 0.2*rank_IGD - 0.3*rank_PF_Overlap - 0.3*rank_PF_Drift")
    r.bold = True
    r.font.name = "Calibri"
    doc.add_paragraph(
        "這個設計讓 label 不只追求收斂品質，也把 Pareto front 的 overlap 與 drift 納入選擇標準。"
        "因此，C selector 的目標不是單純模仿原本 top1 theta，而是選出在穩定性與 front 品質上更平衡的 theta。"
    )

    add_heading(doc, "4. Selector 訓練與 Validation 結果")
    doc.add_paragraph(
        "C selector 使用 Training labels 訓練 Random Forest regression model，輸入包含 instance meta-features 與 theta encoding，"
        "目標為預測 C_LabelScore。對每個待測 instance，selector 選擇 predicted_C_LabelScore 最高的 theta。"
    )
    val_rows = []
    for _, row in selector_summary.iterrows():
        val_rows.append(
            [
                row["selector"],
                int(row["groups"]),
                fmt_pct(row["top1_hit_rate"]),
                fmt_pct(row["top3_hit_rate"]),
                fmt(row["mean_C_ThetaRank"], 2),
                fmt(row["mean_C_regret"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_PF_Drift"], 4),
            ]
        )
    add_table(
        doc,
        ["selector", "groups", "top1", "top3", "mean C rank", "C regret", "PF overlap", "PF drift"],
        val_rows,
        widths=[2600, 900, 900, 900, 1200, 1200, 1300, 1360],
    )
    doc.add_paragraph(
        "Validation 結果顯示，C selector 並不是 C-oracle 的完美近似；top1 hit rate 為 24.14%，top3 hit rate 為 48.28%。"
        "但相較於 all-theta mean，selector 選出的 theta 具有更高 PF_Overlap 與更低 PF_Drift，表示它確實往穩定性感知目標移動。"
    )

    add_heading(doc, "5. Unseen Test Final Comparison")
    doc.add_paragraph(
        "Final test 在 32 個 unseen test instances 上執行，每個 instance 使用 C selector 選出的 theta，並跑 30 次獨立 runs。"
        "比較方式採用 common-reference 後處理，使 Experiment C 與實驗 B 的四種 baseline 在同一個 metric 空間下排名。"
    )
    rows = []
    for _, row in overall.iterrows():
        rows.append(
            [
                row["method"],
                fmt(row["mean_HV"], 4),
                fmt(row["mean_IGD"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_PF_Drift"], 4),
                fmt(row["mean_Runtime"], 3),
                fmt(row["overall_RankScore"], 4),
                int(row["first_place_instances"]),
            ]
        )
    add_table(
        doc,
        ["method", "HV", "IGD", "PF overlap", "PF drift", "runtime", "overall rank", "1st-place inst."],
        rows,
        widths=[3100, 850, 850, 1100, 1000, 900, 1100, 1460],
    )

    add_heading(doc, "6. Confirmatory C vs B 配對檢定")
    doc.add_paragraph(
        "本報告預先指定 per-instance RankScore 為 primary endpoint，作為 OverallRankScore 的 test-instance 配對檢定版本。"
        "Experiment C 與 Experiment B Meta-designed 之比較採 one-sided paired Wilcoxon signed-rank test；directional improvement 為正代表 C 較佳。"
    )
    primary_rows = []
    for _, row in confirmatory_tests[confirmatory_tests["endpoint_role"] == "primary"].iterrows():
        primary_rows.append(
            [
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["median_directional_improvement"], 5),
                fmt(row["raw_p_value_one_sided"], 6),
                fmt(row["holm_p_value"], 6),
                "yes" if row["significant_raw_0_05"] else "no",
            ]
        )
    add_table(
        doc,
        ["primary endpoint", "W/T/L", "median improvement", "one-sided p", "Holm p", "raw p < .05"],
        primary_rows,
        widths=[1900, 1050, 1700, 1400, 1200, 2110],
    )
    doc.add_paragraph("Secondary endpoints 以 Holm correction 控制 C vs B 多端點檢定的 family-wise error rate。")
    secondary_rows = []
    for _, row in confirmatory_tests[confirmatory_tests["endpoint_role"] != "primary"].iterrows():
        secondary_rows.append(
            [
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["median_directional_improvement"], 5),
                fmt(row["raw_p_value_one_sided"], 6),
                fmt(row["holm_p_value"], 6),
                "yes" if row["significant_after_holm_0_05"] else "no",
            ]
        )
    add_table(
        doc,
        ["secondary endpoint", "W/T/L", "median improvement", "one-sided p", "Holm p", "Holm < .05"],
        secondary_rows,
        widths=[1900, 1050, 1700, 1400, 1200, 2110],
    )

    add_heading(doc, "7. Exploratory sign-test evidence")
    doc.add_paragraph(
        "下表列出 Experiment C 對各 baseline 的 paired per-instance 比較。W/T/L 代表 C 在該 metric 上勝/平/負的 instance 數；"
        "p-value 使用 exact two-sided sign test，排除 ties 後計算。這些 sign-test p-values 未做多重校正，因此僅作 exploratory evidence。"
    )
    selected_tests = sign_tests[sign_tests["metric"].isin(["HV", "IGD", "PF_Overlap", "PF_Drift"])]
    pair_rows = []
    for _, row in selected_tests.iterrows():
        pair_rows.append(
            [
                row["baseline"].replace("_ECMADE_MOO", ""),
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["sign_test_p_two_sided"], 4),
                fmt(row["mean_improvement_directional"], 5),
            ]
        )
    add_table(
        doc,
        ["baseline", "metric", "W/T/L", "sign-test p", "平均方向性改善"],
        pair_rows,
        widths=[2800, 1400, 1300, 1400, 2460],
    )

    add_heading(doc, "8. Theta 使用情形")
    usage = assignment["theta_id"].value_counts().rename_axis("theta_id").reset_index(name="test_instances")
    usage_rows = [[row["theta_id"], int(row["test_instances"])] for _, row in usage.iterrows()]
    add_table(doc, ["theta_id", "被選中的 test instances 數"], usage_rows, widths=[2500, 6860])
    doc.add_paragraph(
        "C selector 並不是單一全域 theta 策略，而是會依照 unseen instance 的 meta-features 選擇不同 theta。"
        "這點符合 meta-learning selector 的設計目的，也讓 C 和 Bayesian global selection baseline 有清楚區別。"
    )

    add_heading(doc, "9. 結果解讀")
    add_bullets(
        doc,
        [
            "Experiment C 的主要貢獻不是單純提升 HV，而是用 label engineering 讓 selector 更重視 PF_Overlap 與 PF_Drift。",
            "相較於 Experiment B Meta-designed，C 的 HV 幾乎持平，IGD 與 PF_Overlap 更好，且 first-place instances 明顯增加。",
            "PF_Drift 與 Meta-designed 幾乎接近，表示 C 在強化 front overlap 與 IGD 時，沒有明顯犧牲穩定性。",
            "Validation top1 hit rate 不高，因此論文中應避免宣稱 C selector 已精準學到 oracle；比較穩健的說法是：C 的 label 目標在 final test 上帶來更好的穩定性感知選擇效果。",
        ],
    )

    add_heading(doc, "10. 可寫入論文的重點")
    add_bullets(
        doc,
        [
            "Stability-aware labels 可以直接由既有 multi-run raw outputs 後處理而成，不需要額外增加 label-generation runs。",
            "在 32 個 unseen test instances 上，Experiment C 在 Hand-crafted、Random、Bayesian、Meta-designed 與 C 五組方法中取得最佳 overall rank；正式 C vs B 主張應連結到預先指定 per-instance RankScore 的 paired Wilcoxon 檢定。",
            "Experiment C 同時取得最高 PF_Overlap 與最低 IGD，表示它不只改善穩定性，也保有良好的 front quality。",
            "結果顯示，在 meta-designed parameter selection 中，label objective 的設計和模型架構同樣重要。",
        ],
    )

    add_heading(doc, "11. 主要輸出路徑")
    add_kv_table(
        doc,
        [
            ("C Training labels", str(TRAIN_REPORT / "experiment_c_stability_regression_labels.csv")),
            ("C Validation labels", str(VAL_REPORT / "experiment_c_stability_regression_labels.csv")),
            ("Selector 訓練輸出", str(SELECTOR_DIR)),
            ("C final test raw outputs", str(C_FINAL_DIR)),
            ("B/C comparison report", str(COMPARISON_DIR)),
            ("C vs B Wilcoxon/Holm table", str(COMPARISON_DIR / "experiment_c_vs_b_paired_wilcoxon_holm.csv")),
            ("Pairwise sign-test table", str(COMPARISON_DIR / "experiment_c_pairwise_sign_tests.csv")),
        ],
    )


def md_table(frame: pd.DataFrame) -> str:
    text = frame.copy()
    for col in text.columns:
        text[col] = text[col].map(lambda value: fmt(value, 4) if isinstance(value, (int, float)) else str(value))
    headers = [str(col) for col in text.columns]
    rows = text.astype(str).values.tolist()
    out = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    out.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(out)


def build_markdown_zh(
    data: dict[str, pd.DataFrame],
    sign_tests: pd.DataFrame,
    confirmatory_tests: pd.DataFrame,
    path: Path,
) -> None:
    overall = data["overall"]
    selector_summary = data["selector_summary"]
    c = overall[overall["method"] == C_METHOD].iloc[0]
    lines = [
        "# 實驗 C：穩定性感知 Theta Selector 報告",
        "",
        "## 摘要",
        "",
        f"- Experiment C 在 final common-reference comparison 中排名第一：overall_RankScore={fmt(c['overall_RankScore'], 4)}。",
        f"- First-place test instances：{int(c['first_place_instances'])}/{int(c['instances'])}。",
        "- Primary endpoint：per-instance RankScore，亦即 OverallRankScore 的配對檢定單位版本；數值越低越好。",
        "- 主結論：Experiment C 透過 stability-aware label engineering 改善 theta selection，不需要新增主要 raw runs。",
        "",
        "## Confirmatory C vs B Paired Test",
        "",
        "以 one-sided paired Wilcoxon signed-rank test 比較 Experiment C 與 Experiment B Meta-designed；secondary endpoints 採 Holm correction。未校正 sign-test p-values 僅列為 exploratory evidence。",
        "",
        md_table(confirmatory_tests),
        "",
        "## Stability-aware Label",
        "",
        "`C_LabelScore = -0.2*rank_HV - 0.2*rank_IGD - 0.3*rank_PF_Overlap - 0.3*rank_PF_Drift`",
        "",
        "## Validation Selector Summary",
        "",
        md_table(selector_summary),
        "",
        "## Final Comparison",
        "",
        md_table(
            overall[
                [
                    "method",
                    "instances",
                    "mean_HV",
                    "mean_IGD",
                    "mean_PF_Overlap",
                    "mean_PF_Drift",
                    "overall_RankScore",
                    "first_place_instances",
                ]
            ]
        ),
        "",
        "## Exploratory Pairwise Sign Tests",
        "",
        md_table(sign_tests),
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    DOCX_OUT.mkdir(exist_ok=True)
    data = load_data()
    sign_tests = build_sign_tests(data["ranked"])
    confirmatory_tests = build_confirmatory_tests(data["ranked"])

    md_path = DOCX_OUT / f"{REPORT_STEM}.md"
    docx_path = DOCX_OUT / f"{REPORT_STEM}.docx"

    build_markdown_zh(data, sign_tests, confirmatory_tests, md_path)

    doc = Document()
    setup_styles(doc)
    set_east_asia_font(doc)
    add_title_zh(doc)
    add_report_content_zh(doc, data, sign_tests, confirmatory_tests)
    doc.save(docx_path)

    print(f"MARKDOWN={md_path}")
    print(f"DOCX={docx_path}")


if __name__ == "__main__":
    main()

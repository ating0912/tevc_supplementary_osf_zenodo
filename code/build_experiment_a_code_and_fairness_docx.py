from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\yiting\Documents\Playground")
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "Experiment_A_程式檔案與公平性固定設定_修正版.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths_in: list[float]) -> None:
    table.style = "Table Grid"
    set_table_width(table, widths_in)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        p.paragraph_format.space_after = Pt(0)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_table(doc, ["項目", "內容"], [[k, v] for k, v in rows], [1.9, 4.6])


def add_code_call_flow(doc: Document) -> None:
    steps = [
        ("1", "批次或單方法啟動", "run_p0_lite_synthetic_all_methods.bat 或各 run_p0_lite_synthetic_*.m"),
        ("2", "方法入口", "SyntheticRunner.runAlgorithm(@AlgorithmHandle, 'METHOD')"),
        ("3", "讀取設定與資料", "SyntheticRunner.defaultConfig → applyWorkspaceOverrides → loadManifest → filterManifest"),
        ("4", "建立問題", "PortfolioORLIB(dataPath, K)；使用共同 portfolio repair / feasibility handling"),
        ("5", "呼叫演算法", "NSGAII / SPEA2 / MOEAD / GDE3 / ECMADE_MOO"),
        ("6", "每 run 輸出", "pf_obj.csv、runtime.csv、feasible_rate.csv、generation_pf_points.csv、generation_population_log.csv"),
        ("7", "彙整與後處理", "build_synthetic_experiment_a_report.py 掃描所有 run outputs"),
        ("8", "計算指標", "HV、IGD、PF Overlap、EAF Band Width、PF Drift、Diversity、Spacing、Runtime、Feasible Rate"),
        ("9", "輸出報告", "overall_method_summary.csv、run_metrics.csv、figures、Experiment_A_synthetic_results_report.docx"),
    ]
    add_table(doc, ["步驟", "呼叫階段", "主要程式 / 輸出"], steps, [0.55, 1.75, 4.2])


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Experiment A code and fairness settings")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("實驗 A 程式檔案、呼叫順序與公平性固定設定")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.add_run("ECMADE_MOO 設定基準：experiment_A_report_20260701_112713；目前實驗 A 程式方法：NSGAII、SPEA2、MOEAD、GDE3、ECMADE_MOO、A_MPMO。").italic = True

    doc.add_heading("1. 實驗版本與輸出位置", level=1)
    add_kv_table(
        doc,
        [
            ("工作目錄", str(ROOT)),
            ("資料 manifest", str(ROOT / "data" / "synthetic_constrained_portfolio" / "manifest.csv")),
            ("Synthetic instances", str(ROOT / "data" / "synthetic_constrained_portfolio" / "instances")),
            ("實驗輸出根目錄", str(ROOT / "p0_lite_outputs" / "synthetic_constrained_portfolio")),
            ("本版報告目錄", str(ROOT / "p0_lite_outputs" / "synthetic_constrained_portfolio" / "experiment_A_report_20260701_112713")),
            ("備註", "目前實驗 A 程式已納入 A_MPMO；ECMADE_MOO 主版本採用 experiment_A_report_20260701_112713 對應設定。"),
        ],
    )

    doc.add_heading("2. 程式檔案清單", level=1)
    add_table(
        doc,
        ["用途", "檔名", "位置"],
        [
            ["主控 runner", "SyntheticRunner.m", str(ROOT / "SyntheticRunner.m")],
            ["Portfolio problem / repair", "PortfolioORLIB.m", str(ROOT / "PortfolioORLIB.m")],
            ["工具函式", "P0LiteUtils.m", str(ROOT / "P0LiteUtils.m")],
            ["Synthetic instance 生成", "generate_synthetic_portfolio_instances.py", str(ROOT / "generate_synthetic_portfolio_instances.py")],
            ["報告 / 指標統整", "build_synthetic_experiment_a_report.py", str(ROOT / "build_synthetic_experiment_a_report.py")],
            ["參數鎖定檢查", "check_paper_locked_parameters.py", str(ROOT / "check_paper_locked_parameters.py")],
            ["參數鎖定說明", "ALGORITHM_PARAMETER_LOCK.md", str(ROOT / "ALGORITHM_PARAMETER_LOCK.md")],
            ["一鍵跑全部方法", "run_p0_lite_synthetic_all_methods.bat", str(ROOT / "run_p0_lite_synthetic_all_methods.bat")],
            ["NSGA-II 啟動", "run_p0_lite_synthetic_nsga2.m", str(ROOT / "run_p0_lite_synthetic_nsga2.m")],
            ["SPEA2 啟動", "run_p0_lite_synthetic_spea2.m", str(ROOT / "run_p0_lite_synthetic_spea2.m")],
            ["MOEA/D 啟動", "run_p0_lite_synthetic_moead.m", str(ROOT / "run_p0_lite_synthetic_moead.m")],
            ["GDE3 啟動", "run_p0_lite_synthetic_gde3.m", str(ROOT / "run_p0_lite_synthetic_gde3.m")],
            ["ECMADE-MOO 啟動", "run_p0_lite_synthetic_ecmade_moo.m", str(ROOT / "run_p0_lite_synthetic_ecmade_moo.m")],
            ["ECMADE-MOO 演算法", "ECMADE_MOO.m", str(ROOT / "ECMADE_MOO.m")],
            ["ECMADE-MOO Python 參考版", "ecmade_moo.py", str(ROOT / "ecmade_moo.py")],
            ["A-MPMO baseline", "A_MPMO_NSGAII_v290.m", str(ROOT / "A_MPMO_NSGAII_v290.m")],
            ["A-MPMO 啟動", "run_p0_lite_synthetic_ampmo.m", str(ROOT / "run_p0_lite_synthetic_ampmo.m")],
        ],
        [1.45, 2.05, 3.0],
    )

    doc.add_heading("3. PlatEMO Baseline 位置", level=1)
    add_table(
        doc,
        ["Method", "PlatEMO 檔案"],
        [
            ["NSGA-II", str(ROOT / "PlatEMO" / "PlatEMO" / "Algorithms" / "Multi-objective optimization" / "NSGA-II" / "NSGAII.m")],
            ["SPEA2", str(ROOT / "PlatEMO" / "PlatEMO" / "Algorithms" / "Multi-objective optimization" / "SPEA2" / "SPEA2.m")],
            ["MOEA/D", str(ROOT / "PlatEMO" / "PlatEMO" / "Algorithms" / "Multi-objective optimization" / "MOEA-D" / "MOEAD.m")],
            ["GDE3", str(ROOT / "PlatEMO" / "PlatEMO" / "Algorithms" / "Multi-objective optimization" / "GDE3" / "GDE3.m")],
        ],
        [1.35, 5.15],
    )

    doc.add_heading("4. 程式碼間的呼叫順序", level=1)
    add_code_call_flow(doc)

    doc.add_heading("5. 公平性固定設定", level=1)
    add_table(
        doc,
        ["設定", "固定值 / 說明"],
        [
            ["平台", "MATLAB + PlatEMO"],
            ["Problem class", "PortfolioORLIB"],
            ["Instance set", "Synthetic constrained portfolio"],
            ["Manifest rows", "192"],
            ["Split", "train 112、validation 48、test 32"],
            ["Assets n", "50、100、200、500"],
            ["K/n", "0.05、0.10、0.20、0.30"],
            ["K 實際值", "n=50: 3/5/10/15；n=100: 5/10/20/30；n=200: 10/20/40/60；n=500: 25/50/100/150"],
            ["Runs", "每個 method × instance 30 independent runs"],
            ["Population size", "N = 100"],
            ["Function evaluations", "maxFE = 10000"],
            ["RNG", "mcg16807"],
            ["Seed block", "seed = run index；所有方法共用同一 run index"],
            ["Constraint handling", "共用 PortfolioORLIB repair / feasible handling"],
            ["Reference front", "後處理由所有方法同一 instance 的 PF union 建 empirical reference front"],
        ],
        [1.85, 4.65],
    )

    doc.add_heading("6. 每 Run 輸出與後處理指標", level=1)
    add_table(
        doc,
        ["類別", "內容"],
        [
            ["每 run 輸出", "pf_obj.csv、runtime.csv、feasible_rate.csv、generation_pf_points.csv、generation_population_log.csv"],
            ["Performance", "HV、IGD"],
            ["Stability", "PF Overlap、EAF Band Width、PF Drift"],
            ["Diversity", "Diversity / Spread、Spacing"],
            ["Cost", "Runtime"],
            ["Feasibility", "Feasible Rate"],
            ["主要彙整表", "overall_method_summary.csv、run_metrics.csv、instance_method_metrics.csv"],
        ],
        [1.55, 4.95],
    )

    doc.add_heading("7. ECMADE-MOO 本版固定參數", level=1)
    add_table(
        doc,
        ["參數", "值"],
        [
            ["subpops", "3"],
            ["archiveSize / H", "20"],
            ["theta", "1/13"],
            ["stagnationThreshold / C", "50"],
            ["exploitationAlpha", "0.8"],
            ["initMuF", "[0.9 0.8 0.8]"],
            ["initMuCR", "[0.9 0.5 0.5]"],
            ["fScale", "0.1"],
            ["crScale", "0.1"],
            ["fMax", "1.0"],
            ["exchangeMode", "paper"],
            ["consensusArchive", "false"],
            ["archiveConsWeight", "0.0"],
            ["bestGuide", "rank"],
            ["minSubpopSize", "1"],
        ],
        [2.3, 4.2],
    )

    doc.add_heading("8. 方法與啟動對應", level=1)
    add_table(
        doc,
        ["Method", "啟動檔", "Algorithm handle"],
        [
            ["NSGAII", "run_p0_lite_synthetic_nsga2.m", "@NSGAII"],
            ["SPEA2", "run_p0_lite_synthetic_spea2.m", "@SPEA2"],
            ["MOEAD", "run_p0_lite_synthetic_moead.m", "@MOEAD"],
            ["GDE3", "run_p0_lite_synthetic_gde3.m", "@GDE3"],
            ["ECMADE_MOO", "run_p0_lite_synthetic_ecmade_moo.m", "@ECMADE_MOO"],
            ["A_MPMO", "run_p0_lite_synthetic_ampmo.m", "@A_MPMO_NSGAII_v290"],
        ],
        [1.4, 2.7, 2.4],
    )

    doc.add_heading("9. 本版結果定位", level=1)
    add_kv_table(
        doc,
        [
            ("ECMADE-MOO 優勢", "HV 第 2、IGD 第 2、PF Overlap 第 2、Diversity 最高。"),
            ("ECMADE-MOO 弱點", "EAF Band Width 與 PF Drift 較差，不宜描述為所有穩定性指標全面勝過 GDE3。"),
            ("建議表述", "ECMADE-MOO 在 PF Overlap 與整體 RankScore 優於 GDE3，但 EAF / PF Drift 仍需後續穩定化設計支撐。"),
        ],
    )

    add_footer(doc)
    doc.save(OUT_PATH)
    print(f"DOCX={OUT_PATH}")


if __name__ == "__main__":
    main()

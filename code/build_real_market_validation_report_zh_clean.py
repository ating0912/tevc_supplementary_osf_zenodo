from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA_DIR = (
    ROOT
    / "p0_lite_outputs"
    / "p1_rolling_window_market_validation_20260719"
    / "p0_real_market_validation"
)
OUT_DIR = ROOT / "outputs" / "real_market_validation_report_20260725"
DOCX_PATH = OUT_DIR / "TEVC_real_market_validation_report_zh_20260725.docx"
MD_PATH = OUT_DIR / "TEVC_real_market_validation_report_zh_20260725.md"


FONT_EN = "Calibri"
FONT_ZH = "Microsoft JhengHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"


def set_run_font(run, size: float = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=3, line_spacing=1.05)
    r = p.add_run("TEVC Real-Market Rolling-Window Validation 實驗報告")
    set_run_font(r, size=20, bold=True, color=RGBColor(0, 0, 0))

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=12)
    r = p.add_run("最小可行版本：3 年 training window、6 個月 testing window、10/20/50 bps 交易成本敏感度")
    set_run_font(r, size=10.5, color=RGBColor(85, 85, 85))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix) :])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=4)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.2) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line_spacing=1.0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def fmt_pct(value: float, decimals: int = 2) -> str:
    return f"{value * 100:.{decimals}f}%"


def fmt_num(value: float, decimals: int = 3) -> str:
    return f"{value:.{decimals}f}"


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4)
    r = p.add_run(title)
    set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], GRAY_FILL)
        set_cell_text(table.rows[0].cells[i], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)

    doc.add_paragraph()


def load_data() -> tuple[pd.DataFrame, pd.DataFrame]:
    overall = pd.read_csv(DATA_DIR / "p0_real_market_overall_summary.csv", encoding="utf-8-sig")
    cost = pd.read_csv(DATA_DIR / "p0_real_market_transaction_cost_overall.csv", encoding="utf-8-sig")
    return overall, cost


def build_markdown(overall: pd.DataFrame, cost: pd.DataFrame) -> str:
    overall_md = overall[
        [
            "method",
            "windows",
            "mean_annual_net_return",
            "mean_sharpe",
            "mean_sortino",
            "mean_max_drawdown",
            "mean_cvar95_loss",
            "mean_rebalance_turnover",
            "stability_rank_score",
        ]
    ].copy()
    for col in [
        "mean_annual_net_return",
        "mean_max_drawdown",
        "mean_cvar95_loss",
        "mean_rebalance_turnover",
    ]:
        overall_md[col] = overall_md[col].astype(float).map(lambda x: f"{x:.4f}")
    for col in ["mean_sharpe", "mean_sortino", "stability_rank_score"]:
        overall_md[col] = overall_md[col].astype(float).map(lambda x: f"{x:.3f}")

    cost_md = cost[
        [
            "cost_scenario",
            "method",
            "mean_annual_net_return",
            "mean_turnover",
            "mean_cvar95_loss",
            "mean_max_drawdown",
            "rank_annual_net_return",
        ]
    ].copy()
    for col in ["mean_annual_net_return", "mean_turnover", "mean_cvar95_loss", "mean_max_drawdown"]:
        cost_md[col] = cost_md[col].astype(float).map(lambda x: f"{x:.4f}")
    cost_md["rank_annual_net_return"] = cost_md["rank_annual_net_return"].astype(float).map(lambda x: f"{x:.0f}")

    return f"""# TEVC Real-Market Rolling-Window Validation 實驗報告

## 實驗目的

本實驗提供最小可行版本的 real-market validation，用於檢查在合成資料與 benchmark instances 之外，所提出的 ECMADE-MOO 是否能在真實市場資料上維持合理的 downside-risk 與交易後績效。此實驗定位為外部驗證與限制分析，而不是用來宣稱 ECMADE-MOO 在真實市場報酬上全面優於所有 baselines。

## 資料與 rolling 設定

- Universe：SP100、NASDAQ100、TAIWAN50。
- Rolling windows：每個 universe 11 個 rolling windows，合計 33 個 universe-window。
- Training window：3 年，約 755 至 756 個交易日。
- Testing window：6 個月，約 123 至 126 個交易日。
- 每個 window-method 使用 10 runs。
- Portfolio selection：從每次 run 的 final Pareto front 中選擇 training Sharpe 最高的 portfolio，接著在 out-of-sample testing window 上計算績效。
- 交易成本：主要 wealth curve 使用 10 bps；另做 10 bps、20 bps、50 bps sensitivity。

## 評估指標

- Annual net return：扣除交易成本後的年化報酬。
- CVaR(95%) loss：testing window 中最差 5% 日報酬的平均損失，數值越低代表 downside risk 越小。
- Maximum drawdown, MDD：testing wealth curve 的最大回撤，絕對值越小越穩定。
- Turnover：相鄰 rebalancing portfolio 的權重變動總量。
- After-cost wealth curve：以扣除交易成本後的 daily return 串接累積財富。

## 主要結果

{overall_md.to_markdown(index=False)}

## 交易成本敏感度

{cost_md.to_markdown(index=False)}

## 結論

在 33 個 real-market rolling windows 上，ECMADE-MOO 並未在 after-cost annual return 上勝過 return-oriented baselines；MOEAD 在 10/20/50 bps 三種交易成本下皆取得最高 mean annual net return。然而，ECMADE-MOO 取得最低的 mean CVaR(95%) loss，表示其在真實市場 out-of-sample data 上呈現較保守的 downside-risk profile。交易成本敏感度結果顯示，方法在 annual net return 的排序於 10 bps、20 bps 與 50 bps 下大致穩定，表示主要結論不是由單一交易成本設定造成。

因此，此 real-market validation 應作為外部驗證與限制分析呈現：ECMADE-MOO 的優勢主要反映在 downside-risk control，而不是在真實市場 after-cost return 上的全面 dominance。
"""


def build_docx(overall: pd.DataFrame, cost: pd.DataFrame) -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    add_heading(doc, "1. 實驗目的", 1)
    add_para(
        doc,
        "本實驗提供最小可行版本的 real-market validation，用於檢查在合成資料與 benchmark instances 之外，所提出的 ECMADE-MOO 是否能在真實市場資料上維持合理的 downside-risk 與交易後績效。此實驗定位為外部驗證與限制分析，而不是用來宣稱 ECMADE-MOO 在真實市場報酬上全面優於所有 baselines。",
    )

    add_heading(doc, "2. 資料與 rolling evaluation 設定", 1)
    add_bullets(
        doc,
        [
            "Universe：SP100、NASDAQ100、TAIWAN50。",
            "Rolling windows：每個 universe 11 個 rolling windows，合計 33 個 universe-window。",
            "Training window：3 年，約 755 至 756 個交易日。",
            "Testing window：6 個月，約 123 至 126 個交易日。",
            "每個 window-method 使用 10 runs。",
            "Portfolio selection：從每次 run 的 final Pareto front 中選擇 training Sharpe 最高的 portfolio，再於 out-of-sample testing window 上計算績效。",
            "交易成本：主要 wealth curve 使用 10 bps；另做 10 bps、20 bps、50 bps sensitivity。",
        ],
    )

    add_heading(doc, "3. 評估指標", 1)
    add_bullets(
        doc,
        [
            "Annual net return：扣除交易成本後的年化報酬，數值越高越好。",
            "CVaR(95%) loss：testing window 中最差 5% 日報酬的平均損失，數值越低代表 downside risk 越小。",
            "Maximum drawdown, MDD：testing wealth curve 的最大回撤，絕對值越小越穩定。",
            "Turnover：相鄰 rebalancing portfolio 的權重變動總量，數值越低代表交易更穩定。",
            "After-cost wealth curve：以扣除交易成本後的 daily return 串接累積財富。",
        ],
    )

    add_heading(doc, "4. 主要結果", 1)
    headers = ["Method", "Windows", "Annual net return", "Sharpe", "Sortino", "MDD", "CVaR95 loss", "Turnover", "Stability rank"]
    rows = []
    for _, row in overall.iterrows():
        rows.append(
            [
                row["method"],
                str(int(row["windows"])),
                fmt_pct(float(row["mean_annual_net_return"])),
                fmt_num(float(row["mean_sharpe"])),
                fmt_num(float(row["mean_sortino"])),
                fmt_pct(float(row["mean_max_drawdown"])),
                fmt_pct(float(row["mean_cvar95_loss"])),
                fmt_pct(float(row["mean_rebalance_turnover"])),
                fmt_num(float(row["stability_rank_score"])),
            ]
        )
    add_table(doc, "Table RM-1. Real-market rolling-window validation summary", headers, rows)

    add_para(
        doc,
        "整體結果顯示，MOEAD 取得最高 mean annual net return，而 ECMADE-MOO 取得最低 mean CVaR(95%) loss。這表示 ECMADE-MOO 的主要優勢不在於追求最高 out-of-sample return，而是在 downside-risk control 上較保守。",
    )

    add_heading(doc, "5. 交易成本敏感度", 1)
    headers = ["Cost", "Method", "Annual net return", "Turnover", "CVaR95 loss", "MDD", "Return rank"]
    rows = []
    for _, row in cost.iterrows():
        rows.append(
            [
                row["cost_scenario"],
                row["method"],
                fmt_pct(float(row["mean_annual_net_return"])),
                fmt_pct(float(row["mean_turnover"])),
                fmt_pct(float(row["mean_cvar95_loss"])),
                fmt_pct(float(row["mean_max_drawdown"])),
                str(int(float(row["rank_annual_net_return"]))),
            ]
        )
    add_table(doc, "Table RM-2. Transaction-cost sensitivity under 10/20/50 bps", headers, rows)

    add_para(
        doc,
        "交易成本由 10 bps 提高到 50 bps 後，所有方法的 annual net return 皆下降，但排序大致穩定。MOEAD 在三種成本設定下皆維持 return rank 1；ECMADE-MOO 在 return rank 上維持第 5，但 CVaR(95%) loss 仍為六個方法中最低。",
    )

    add_heading(doc, "6. After-cost wealth curve", 1)
    plot_path = DATA_DIR / "p0_real_market_wealth_curve_10bps.png"
    if plot_path.exists():
        doc.add_picture(str(plot_path), width=Inches(6.8))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(cap, before=3, after=8)
        r = cap.add_run("Figure RM-1. After-cost wealth curve under 10 bps transaction cost.")
        set_run_font(r, size=9, color=RGBColor(85, 85, 85))

    add_heading(doc, "7. 結論與論文呈現方式", 1)
    add_para(
        doc,
        "在 33 個 real-market rolling windows 上，ECMADE-MOO 並未在 after-cost annual return 上勝過 return-oriented baselines；因此不應將此實驗寫成 return dominance 的證據。較合理的呈現方式是：real-market validation 支持 ECMADE-MOO 在 downside-risk control 上具有外部資料一致性，但也揭示其在真實市場報酬最大化上仍有改善空間。",
    )
    add_para(
        doc,
        "建議在論文中將本節放在 Experiment C 之後或 supplementary validation 中，標題可用 Real-market rolling-window validation。正文需明確說明此實驗使用真實市場資料、rolling out-of-sample evaluation、交易成本敏感度，以及 after-cost wealth curve。",
    )

    add_heading(doc, "8. 數據位置", 1)
    add_bullets(
        doc,
        [
            str(DATA_DIR / "p0_real_market_overall_summary.csv"),
            str(DATA_DIR / "p0_real_market_transaction_cost_overall.csv"),
            str(DATA_DIR / "p0_real_market_window_method_summary.csv"),
            str(DATA_DIR / "p0_real_market_wealth_curve_mean.csv"),
            str(DATA_DIR / "p0_real_market_wealth_curve_run_level.csv"),
            str(DATA_DIR / "p0_real_market_wealth_curve_10bps.png"),
        ],
    )

    doc.save(DOCX_PATH)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    overall, cost = load_data()
    build_docx(overall, cost)
    MD_PATH.write_text(build_markdown(overall, cost), encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()

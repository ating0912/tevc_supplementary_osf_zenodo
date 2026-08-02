from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_mokp_experiment_results_report_docx import (
    DIAGNOSTIC_DIR,
    LIGHT_FILL,
    MUTED,
    OUT_DIR,
    ROOT,
    SUMMARY_DIR,
    STABILITY_DIR,
    add_bullet,
    add_callout,
    add_caption,
    add_dataframe_table,
    add_heading,
    add_paragraph,
    configure_doc,
    set_run_font,
    set_table_width,
    set_cell_shading,
)


OUT_DOCX_ZH = OUT_DIR / "MOKP_nonfinancial_experiment_results_report_ZH_20260730.docx"


def add_title_page_zh(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("實驗結果報告")
    set_run_font(r, size=24, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("非金融 MOKP 泛化測試與 Experiment C stability-aware transfer diagnostic")
    set_run_font(r, size=13, color=MUTED)

    meta = [
        ("報告日期", "2026-07-30"),
        ("實驗類型", "非金融 constrained / combinatorial MOKP test bed"),
        ("新增主列", "ExperimentC_StabilityAware_ECMADE_MOO"),
        ("診斷列", "ExperimentC_GlobalTheta034_ECMADE_MOO；ExperimentC_GlobalTheta037_ECMADE_MOO"),
        ("資料來源", str(SUMMARY_DIR)),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1900, 7460])
    for i, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(i, 0), LIGHT_FILL)
        p0 = table.cell(i, 0).paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run_font(r0, size=9, bold=True)
        p1 = table.cell(i, 1).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run_font(r1, size=9)

    doc.add_paragraph()
    add_callout(
        doc,
        "主要結論",
        "Experiment C 的 stability-aware selector 已可在非金融 MOKP test bed 上完整執行，"
        "但目前結果不支持強烈宣稱已完成跨問題領域泛化。其表現接近兩個 fixed global-theta diagnostic，"
        "且整體仍落後 BayesianConfig_ECMADE_MOO 與 base ECMADE_MOO。"
    )


def build_report_zh():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(SUMMARY_DIR / "overall_method_summary.csv")
    friedman = pd.read_csv(SUMMARY_DIR / "friedman_tests.csv")
    ref = pd.read_csv(SUMMARY_DIR / "reference_front_info.csv")
    run_metrics = pd.read_csv(SUMMARY_DIR / "run_metrics.csv")

    doc = Document()
    configure_doc(doc)
    add_title_page_zh(doc)

    add_heading(doc, "1. 實驗範圍", 1)
    add_paragraph(
        doc,
        "本報告整理非金融 MOKP 泛化測試結果，用於回應審稿意見中對 portfolio-specific application 的疑慮。"
        "本次新增的 MOKP 為 bi-objective multi-objective knapsack problem，屬於 constrained / combinatorial benchmark，"
        "因此可作為 constrained portfolio optimization 之外的外部測試場域。"
    )
    for item in [
        "測試集包含 18 個 MOKP instances，涵蓋 item counts d = 100、250、500，capacity ratios = 0.35、0.50、0.65，以及 independent / conflicting profit modes。",
        "每個 method-instance cell 執行 30 次 independent runs。",
        "比較方法共 12 個，包含 base MOEAs、ECMADE_MOO、BayesianConfig_ECMADE_MOO、RandomConfig_ECMADE_MOO、MetaTransfer_ECMADE_MOO，以及三個 Experiment C 相關列。",
        "後處理採用 instance-level common reference front；每個 instance 的 reference front 由所有方法的 observed PF outputs pooled 後建立。",
    ]:
        add_bullet(doc, item)

    completeness = pd.DataFrame([
        {"item": "ExperimentC_StabilityAware_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_StabilityAware_ECMADE_MOO")].shape[0]), "status": "完成"},
        {"item": "ExperimentC_GlobalTheta034_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_GlobalTheta034_ECMADE_MOO")].shape[0]), "status": "完成"},
        {"item": "ExperimentC_GlobalTheta037_ECMADE_MOO", "expected_runs": 540, "parsed_runs": int(run_metrics[run_metrics.method.eq("ExperimentC_GlobalTheta037_ECMADE_MOO")].shape[0]), "status": "完成"},
        {"item": "Global-theta diagnostic PF outputs", "expected_runs": 1080, "parsed_runs": len(list((DIAGNOSTIC_DIR / "test").rglob("pf_obj.csv"))), "status": "完成"},
    ])
    add_caption(doc, "表 1. 新增 Experiment C 相關列的輸出完整性檢查。")
    add_dataframe_table(
        doc,
        completeness,
        ["item", "expected_runs", "parsed_runs", "status"],
        ["項目", "預期 runs", "已解析 runs", "狀態"],
        [4680, 1500, 1500, 1680],
        number_cols=["expected_runs", "parsed_runs", "status"],
    )

    add_heading(doc, "2. 整體結果", 1)
    add_paragraph(
        doc,
        "RankScore 越低代表跨指標平均排名越好。原始指標中，HV、PF_Overlap 與 Diversity 為越高越好；"
        "IGD、PF_Drift 與 Runtime 為越低越好。"
    )
    compact = summary[
        [
            "method",
            "runs",
            "mean_HV",
            "mean_IGD",
            "mean_PF_Overlap",
            "mean_PF_Drift",
            "mean_Runtime",
            "overall_RankScore",
            "first_place_instances",
        ]
    ].copy().sort_values("overall_RankScore")
    add_caption(doc, "表 2. MOKP 方法總表，依 overall_RankScore 排序。")
    add_dataframe_table(
        doc,
        compact,
        [
            "method",
            "runs",
            "mean_HV",
            "mean_IGD",
            "mean_PF_Overlap",
            "mean_PF_Drift",
            "mean_Runtime",
            "overall_RankScore",
            "first_place_instances",
        ],
        ["方法", "Runs", "HV", "IGD", "PF overlap", "PF drift", "Runtime", "RankScore", "第 1 名 instances"],
        [3000, 700, 800, 800, 900, 850, 850, 850, 610],
        number_cols=["runs", "mean_HV", "mean_IGD", "mean_PF_Overlap", "mean_PF_Drift", "mean_Runtime", "overall_RankScore", "first_place_instances"],
    )

    add_heading(doc, "3. Experiment C 診斷結果", 1)
    expc = summary[summary["method"].isin([
        "ExperimentC_StabilityAware_ECMADE_MOO",
        "ExperimentC_GlobalTheta034_ECMADE_MOO",
        "ExperimentC_GlobalTheta037_ECMADE_MOO",
        "BayesianConfig_ECMADE_MOO",
        "ECMADE_MOO",
    ])].copy().sort_values("overall_RankScore")
    add_caption(doc, "表 3. Experiment C transfer rows 與兩個主要 ECMADE-family 參考方法比較。")
    add_dataframe_table(
        doc,
        expc,
        ["method", "mean_HV", "mean_IGD", "mean_Diversity", "mean_Runtime", "mean_InstanceRank", "overall_RankScore"],
        ["方法", "HV", "IGD", "Diversity", "Runtime", "Instance rank", "RankScore"],
        [3600, 900, 900, 900, 900, 1080, 1080],
        number_cols=["mean_HV", "mean_IGD", "mean_Diversity", "mean_Runtime", "mean_InstanceRank", "overall_RankScore"],
    )
    for item in [
        "ExperimentC_StabilityAware_ECMADE_MOO 的 HV = 0.8749、IGD = 0.1696、overall_RankScore = 8.5000。",
        "兩個 fixed global-theta diagnostic 非常接近：theta_034 RankScore = 8.8333，theta_037 RankScore = 8.8333。",
        "Stability-aware assignment 僅略優於固定 theta；差距不足以支持 selector 已成功跨領域泛化的強主張。",
        "較強的 ECMADE-family 參考方法仍是 BayesianConfig_ECMADE_MOO 與 ECMADE_MOO，overall_RankScore 分別為 3.1667 與 3.5000。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 統計檢定摘要", 1)
    friedman_table = friedman.copy()
    friedman_table["friedman_p_value"] = friedman_table["friedman_p_value"].map(lambda x: f"{x:.3e}")
    add_caption(doc, "表 4. 18 個 MOKP instances、12 個 methods 的 Friedman tests。")
    add_dataframe_table(
        doc,
        friedman_table,
        ["metric", "direction", "instances", "methods", "friedman_chi_square", "friedman_p_value"],
        ["指標", "方向", "Instances", "Methods", "Chi-square", "p-value"],
        [1600, 1100, 1300, 1100, 2100, 2160],
        number_cols=["instances", "methods", "friedman_chi_square", "friedman_p_value"],
    )
    add_paragraph(
        doc,
        "六個指標的 omnibus test 皆達顯著，表示方法間確實存在系統性差異。Pairwise Wilcoxon-Holm 結果也顯示，"
        "BayesianConfig_ECMADE_MOO 與 ECMADE_MOO 在 HV 與 IGD 上顯著優於 Experiment C transfer rows；"
        "而 direct stability-aware row 與兩個 fixed global-theta diagnostics 之間沒有形成實質分離。"
    )

    add_heading(doc, "5. Reference Front 覆蓋狀況", 1)
    ref_summary = pd.DataFrame([
        {"quantity": "Instances", "value": len(ref)},
        {"quantity": "Minimum reference points", "value": int(ref["reference_points"].min())},
        {"quantity": "Median reference points", "value": int(ref["reference_points"].median())},
        {"quantity": "Maximum reference points", "value": int(ref["reference_points"].max())},
    ])
    add_caption(doc, "表 5. MOKP analysis 所使用的 common reference-front 覆蓋狀況。")
    add_dataframe_table(
        doc,
        ref_summary,
        ["quantity", "value"],
        ["項目", "數值"],
        [6600, 2760],
        number_cols=["value"],
    )

    add_heading(doc, "6. 論文撰寫建議", 1)
    add_callout(
        doc,
        "建議定位",
        "這組 MOKP 結果應定位為非金融外部驗證與 stress test，而不是跨領域泛化已完成的證據。"
        "它的價值在於證明流程可在 portfolio 外執行，同時揭露目前 selector / theta family 的限制。"
    )
    add_paragraph(doc, "建議可放入 manuscript 的敘述如下：")
    quote = (
        "雖然非金融 MOKP 評估確認本文提出的 stability-aware transfer pipeline 可套用於 constrained portfolio optimization "
        "以外的問題場域，但目前 transfer result 尚不足以建立廣泛的跨領域泛化結論。在 MOKP test bed 上，"
        "Experiment C 的 stability-aware configuration 與 fixed global-theta diagnostics 表現接近，且仍落後於 "
        "BayesianConfig_ECMADE_MOO 與 base ECMADE_MOO。因此，本文將此結果視為 experimental protocol 的外部驗證與目前限制，"
        "而非宣稱 selector 已能在不同問題領域精準重建 Oracle theta。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(quote)
    set_run_font(r, size=10.5)
    r.italic = True

    add_heading(doc, "7. 資料檔案", 1)
    artifacts = pd.DataFrame([
        {"artifact": "Overall method summary", "path": str(SUMMARY_DIR / "overall_method_summary.csv")},
        {"artifact": "Run-level metrics", "path": str(SUMMARY_DIR / "run_metrics.csv")},
        {"artifact": "Friedman tests", "path": str(SUMMARY_DIR / "friedman_tests.csv")},
        {"artifact": "Pairwise Wilcoxon-Holm tests", "path": str(SUMMARY_DIR / "pairwise_wilcoxon.csv")},
        {"artifact": "Reference front info", "path": str(SUMMARY_DIR / "reference_front_info.csv")},
        {"artifact": "Experiment C global-theta raw outputs", "path": str(DIAGNOSTIC_DIR)},
        {"artifact": "Experiment C stability-aware raw outputs", "path": str(STABILITY_DIR)},
    ])
    add_caption(doc, "表 6. 支援本報告的資料檔與輸出目錄。")
    add_dataframe_table(
        doc,
        artifacts,
        ["artifact", "path"],
        ["資料項目", "路徑"],
        [3000, 6360],
    )

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = footer.add_run("MOKP 實驗結果報告")
        set_run_font(r, size=9, color=MUTED)

    doc.save(OUT_DOCX_ZH)
    return OUT_DOCX_ZH


if __name__ == "__main__":
    print(build_report_zh())

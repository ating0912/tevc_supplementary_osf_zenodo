import math
import os
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = r"."
SYN_DIR = os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio", "experiment_A_report_20260705_094925")
OR_DIR = os.path.join(ROOT, "p0_lite_outputs", "orlib_constrained_portfolio", "experiment_A_orlib_report_20260705_235837")
OUT_DIR = os.path.join(ROOT, "docx_outputs")
OUT_DOCX = os.path.join(OUT_DIR, "Experiment_A_數據來源與計算流程說明_統計檢定修正版.docx")
POSTHOC_CSV = os.path.join(OUT_DIR, "Experiment_A_posthoc_ECMADE_vs_baselines_Holm_A12.csv")
COMPACT_CSV = os.path.join(OUT_DIR, "Experiment_A_compact_statistical_tests.csv")

STAT_METRICS = [
    ("HV", "max"),
    ("IGD", "min"),
    ("PF_Overlap", "max"),
    ("EAF_Band_Width", "min"),
    ("PF_Drift", "min"),
    ("Diversity", "max"),
    ("Runtime", "min"),
]


def qname(name):
    return qn(name)


def set_ea(run):
    run._element.get_or_add_rPr()
    if run._element.rPr.rFonts is None:
        run._element.rPr.append(OxmlElement("w:rFonts"))
    run._element.rPr.rFonts.set(qname("w:eastAsia"), "Microsoft JhengHei")


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qname("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qname("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qname("w:w"), str(width))
    tc_w.set(qname("w:type"), "dxa")
    cell.width = Inches(width / 1440)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell_width(cell, widths[ci])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)
                    set_ea(run)
            if ri == 0:
                set_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    doc.add_paragraph()
    return table


def add_path(doc, label, path):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label + ": ")
    r.bold = True
    set_ea(r)
    r2 = p.add_run(path)
    set_ea(r2)


def fmt(x, d=4):
    try:
        return f"{float(x):.{d}f}"
    except Exception:
        return str(x)


def fmt_p(x):
    x = float(x)
    if x == 0:
        return "<1e-300"
    if x < 0.001:
        return f"{x:.2e}"
    return f"{x:.4f}"


def normal_cdf(x):
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def gammap_ser(a, x):
    gln = math.lgamma(a)
    if x <= 0:
        return 0.0
    ap = a
    total = 1.0 / a
    delta = total
    for _ in range(1000):
        ap += 1
        delta *= x / ap
        total += delta
        if abs(delta) < abs(total) * 3e-14:
            break
    return total * math.exp(-x + a * math.log(x) - gln)


def gammaq_cf(a, x):
    gln = math.lgamma(a)
    b = x + 1.0 - a
    c = 1.0 / 1e-300
    d = 1.0 / b
    h = d
    for i in range(1, 1000):
        an = -i * (i - a)
        b += 2.0
        d = an * d + b
        if abs(d) < 1e-300:
            d = 1e-300
        c = b + an / c
        if abs(c) < 1e-300:
            c = 1e-300
        d = 1.0 / d
        delta = d * c
        h *= delta
        if abs(delta - 1.0) < 3e-14:
            break
    return math.exp(-x + a * math.log(x) - gln) * h


def chi2_sf(x, df):
    if x < 0:
        return 1.0
    a = df / 2.0
    xx = x / 2.0
    if xx < a + 1.0:
        return max(0.0, 1.0 - gammap_ser(a, xx))
    return min(1.0, gammaq_cf(a, xx))


def friedman_from_wide(wide, direction):
    data = wide.dropna()
    if data.empty:
        return math.nan, math.nan
    methods = list(data.columns)
    ranks = data.rank(axis=1, ascending=(direction == "min"), method="average")
    n, k = ranks.shape
    rank_sums = ranks.sum(axis=0)
    stat = (12.0 / (n * k * (k + 1))) * float((rank_sums ** 2).sum()) - 3 * n * (k + 1)
    return stat, chi2_sf(stat, k - 1)


def wilcoxon_approx(x, y, direction):
    diffs = np.asarray(x, dtype=float) - np.asarray(y, dtype=float)
    diffs = diffs[~np.isnan(diffs)]
    diffs = diffs[np.abs(diffs) > 1e-12]
    n = len(diffs)
    if n == 0:
        return math.nan, math.nan
    signed = diffs if direction == "max" else -diffs
    absdiff = np.abs(signed)
    order = np.argsort(absdiff)
    ranks = np.empty(n)
    i = 0
    while i < n:
        j = i
        while j + 1 < n and absdiff[order[j + 1]] == absdiff[order[i]]:
            j += 1
        ranks[order[i : j + 1]] = (i + 1 + j + 1) / 2.0
        i = j + 1
    wpos = ranks[signed > 0].sum()
    mean = n * (n + 1) / 4.0
    var = n * (n + 1) * (2 * n + 1) / 24.0
    z = (wpos - mean - 0.5 * (1 if wpos > mean else -1)) / math.sqrt(var)
    p = 2 * min(normal_cdf(z), 1 - normal_cdf(z))
    return float(wpos), float(max(min(p, 1.0), 0.0))


def holm_adjust(p_values):
    indexed = sorted(enumerate(p_values), key=lambda item: item[1])
    out = [math.nan] * len(p_values)
    running = 0.0
    m = len(p_values)
    for rank, (idx, p) in enumerate(indexed, start=1):
        adj = min(1.0, p * (m - rank + 1))
        running = max(running, adj)
        out[idx] = running
    return out


def a12(x, y, direction):
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    if direction == "min":
        x = -x
        y = -y
    x = x[~np.isnan(x)]
    y = y[~np.isnan(y)]
    greater = 0.0
    ties = 0.0
    for xv in x:
        greater += np.sum(xv > y)
        ties += np.sum(xv == y)
    return float((greater + 0.5 * ties) / (len(x) * len(y)))


def load_overall(path):
    df = pd.read_csv(os.path.join(path, "overall_method_summary.csv"))
    if "method" not in df.columns:
        df = df.rename(columns={df.columns[0]: "method"})
    return df


def build_stats():
    all_posthoc = []
    compact = []
    for dataset, path in [("Synthetic", SYN_DIR), ("OR-Library", OR_DIR)]:
        inst = pd.read_csv(os.path.join(path, "instance_method_metrics.csv"))
        overall = load_overall(path)
        friedman = pd.read_csv(os.path.join(path, "statistical_tests.csv"))
        friedman = friedman[friedman["test"].eq("Friedman")].set_index("metric")
        for metric, direction in STAT_METRICS:
            wide = inst.pivot_table(index="instance", columns="method", values=metric, aggfunc="mean")
            baselines = [m for m in wide.columns if m != "ECMADE_MOO"]
            means = overall[overall["method"].isin(baselines)].set_index("method")[f"mean_{metric}"]
            best = means.idxmax() if direction == "max" else means.idxmin()
            rows = []
            for baseline in baselines:
                paired = wide[["ECMADE_MOO", baseline]].dropna()
                stat, raw_p = wilcoxon_approx(paired["ECMADE_MOO"], paired[baseline], direction)
                effect = a12(paired["ECMADE_MOO"], paired[baseline], direction)
                e_mean = overall.loc[overall["method"].eq("ECMADE_MOO"), f"mean_{metric}"].iloc[0]
                b_mean = overall.loc[overall["method"].eq(baseline), f"mean_{metric}"].iloc[0]
                rows.append({
                    "Dataset": dataset,
                    "Metric": metric,
                    "Direction": "higher is better" if direction == "max" else "lower is better",
                    "Comparison": f"ECMADE_MOO vs {baseline}",
                    "Baseline": baseline,
                    "Best baseline for metric": "Yes" if baseline == best else "No",
                    "ECMADE mean": e_mean,
                    "Baseline mean": b_mean,
                    "Wilcoxon statistic": stat,
                    "raw p-value": raw_p,
                    "A12": effect,
                })
            adjusted = holm_adjust([r["raw p-value"] for r in rows])
            for row, adj in zip(rows, adjusted):
                row["Holm adjusted p-value"] = adj
            all_posthoc.extend(rows)
            best_row = next(r for r in rows if r["Baseline"] == best)
            if metric in friedman.index:
                f_p = friedman.loc[metric, "p_value"]
            else:
                _, f_p = friedman_from_wide(wide, direction)
            compact.append({
                "Dataset": dataset,
                "Metric": metric,
                "Friedman p-value": f_p,
                "Significant": "Yes" if float(f_p) < 0.05 else "No",
                "ECMADE-MOO vs Best Baseline": best,
                "Holm adjusted p-value": best_row["Holm adjusted p-value"],
                "Effect Size A12": best_row["A12"],
            })
    posthoc = pd.DataFrame(all_posthoc)
    compact_df = pd.DataFrame(compact)
    posthoc.to_csv(POSTHOC_CSV, index=False, encoding="utf-8-sig")
    compact_df.to_csv(COMPACT_CSV, index=False, encoding="utf-8-sig")
    return compact_df


def rows_solution(dataset, overall):
    rows = []
    for _, r in overall.sort_values(["mean_HV", "mean_IGD"], ascending=[False, True]).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_HV"]), fmt(r["mean_IGD"])])
    return rows


def rows_stability(dataset, overall):
    rows = []
    for _, r in overall.sort_values("mean_PF_Overlap", ascending=False).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_PF_Overlap"]), fmt(r["mean_EAF_Band_Width"]), fmt(r["mean_PF_Drift"])])
    return rows


def rows_behavior(dataset, overall):
    rows = []
    for _, r in overall.sort_values("mean_Diversity", ascending=False).iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_Diversity"]), fmt(r["mean_Spacing"])])
    return rows


def rows_cost(dataset, overall):
    rows = []
    for _, r in overall.sort_values("mean_Runtime").iterrows():
        rows.append([dataset, r["method"], fmt(r["mean_Runtime"], 3)])
    return rows


def compact_rows(compact):
    rows = []
    for _, r in compact.iterrows():
        rows.append([
            r["Dataset"],
            r["Metric"],
            fmt_p(r["Friedman p-value"]),
            r["Significant"],
            r["ECMADE-MOO vs Best Baseline"],
            fmt_p(r["Holm adjusted p-value"]),
            fmt(r["Effect Size A12"], 3),
        ])
    return rows


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qname("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st._element.rPr.rFonts.set(qname("w:eastAsia"), "Microsoft JhengHei")
    return doc


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    syn_overall = load_overall(SYN_DIR)
    or_overall = load_overall(OR_DIR)
    compact = build_stats()
    syn_runs = len(pd.read_csv(os.path.join(SYN_DIR, "run_metrics.csv"), usecols=["method"]))
    or_runs = len(pd.read_csv(os.path.join(OR_DIR, "run_metrics.csv"), usecols=["method"]))
    syn_missing = len(pd.read_csv(os.path.join(SYN_DIR, "missing_outputs.csv")))
    or_missing = len(pd.read_csv(os.path.join(OR_DIR, "missing_outputs.csv")))

    doc = setup_doc()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Experiment A 數據來源與計算流程說明")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    set_ea(r)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Synthetic constrained portfolio + OR-Library portfolio instances")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor.from_string("555555")
    set_ea(sr)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. 文件目的與完成狀態", level=1)
    doc.add_paragraph("本文件整理 Experiment A 的資料來源、實驗數據產生流程、指標計算方式、統計檢定放置方式，以及主要檔案位置。")
    add_table(doc, ["Dataset", "Complete runs", "Missing outputs", "Status", "Report folder"], [
        ["Synthetic", syn_runs, syn_missing, "Complete", SYN_DIR],
        ["OR-Library", or_runs, or_missing, "Complete", OR_DIR],
    ], [1200, 1200, 1200, 1000, 4760])

    doc.add_heading("2. 數據來源", level=1)
    doc.add_paragraph("Synthetic 資料由 synthetic constrained portfolio manifest 定義，runner 依 manifest 逐一執行六種演算法。")
    add_path(doc, "Synthetic manifest", os.path.join(ROOT, "data", "synthetic_constrained_portfolio", "manifest.csv"))
    add_path(doc, "Synthetic report folder", SYN_DIR)
    doc.add_paragraph("OR-Library 資料只使用 port1.txt 到 port5.txt，並以固定 K = 5, 10, 20, 30 展開為 20 個 portfolio test instances。")
    add_path(doc, "OR-Library source folder", os.environ.get("TEVC_ORLIB_SOURCE", os.path.join(ROOT, "data", "orlib_constrained_portfolio")))
    add_path(doc, "Workspace OR-Library data folder", os.path.join(ROOT, "data", "orlib_constrained_portfolio"))
    add_path(doc, "OR-Library manifest", os.path.join(ROOT, "data", "orlib_constrained_portfolio", "manifest.csv"))

    doc.add_heading("3. 數據如何產生", level=1)
    add_table(doc, ["Step", "Description"], [
        ["Manifest definition", "manifest.csv defines instance, assets, K, split, and input data path."],
        ["Runner execution", "SyntheticRunner and ORLibraryRunner call NSGAII, SPEA2, MOEAD, GDE3, ECMADE_MOO, and A_MPMO."],
        ["Fair setting", "All methods use the same function evaluations, population size, seed block, and PortfolioORLIB repair-based constraint handling."],
        ["Run outputs", "Each run saves PF, population, archive, runtime, feasible rate, constraint metrics, and generation snapshots."],
        ["Post-processing", "Report scripts recompute HV, IGD, PF overlap, EAF band width, PF drift, diversity, spacing, runtime, and feasibility."],
    ], [1900, 7460])

    doc.add_heading("4. 指標構面", level=1)
    add_table(doc, ["Analysis aspect", "Metrics", "Interpretation"], [
        ["Solution Quality", "HV, IGD", "Pareto front solution quality and convergence. Higher HV and lower IGD are better."],
        ["PF Stability", "PF Overlap, EAF Band Width, PF Drift", "Repeated-run front stability. Higher overlap, lower EAF width, and lower drift are better."],
        ["Search Behavior", "Diversity, Spacing", "Spread and distribution behavior of final/archive fronts."],
        ["Computational Cost", "Runtime", "Optimization runtime per run. Lower is better."],
    ], [1900, 2600, 4860])

    doc.add_heading("5. Overall Performance Comparison", level=1)
    doc.add_paragraph("本報告不以單一總排名或 RankScore 作為主軸，因為不同指標衡量不同面向，直接平均 rank 會隱含相同權重假設。")
    doc.add_heading("5.1 Solution Quality", level=2)
    doc.add_paragraph("OR-Library 是 Experiment A 的主要亮點：ECMADE-MOO 取得最高 HV（0.9909）與最低 IGD（0.0288）。")
    add_table(doc, ["Dataset", "Method", "HV", "IGD"], rows_solution("Synthetic", syn_overall) + rows_solution("OR-Library", or_overall), [1500, 1800, 1200, 4860])
    doc.add_heading("5.2 PF Stability", level=2)
    add_table(doc, ["Dataset", "Method", "PF Overlap", "EAF Width", "PF Drift"], rows_stability("Synthetic", syn_overall) + rows_stability("OR-Library", or_overall), [1500, 1800, 1500, 1500, 3060])
    doc.add_heading("5.3 Search Behavior", level=2)
    add_table(doc, ["Dataset", "Method", "Diversity", "Spacing"], rows_behavior("Synthetic", syn_overall) + rows_behavior("OR-Library", or_overall), [1500, 1800, 1500, 4560])
    doc.add_heading("5.4 Computational Cost", level=2)
    add_table(doc, ["Dataset", "Method", "Runtime"], rows_cost("Synthetic", syn_overall) + rows_cost("OR-Library", or_overall), [1500, 1800, 6060])

    doc.add_heading("6. 統計檢定建議放法", level=1)
    doc.add_paragraph("正文先報描述統計與指標結果，再報 Friedman test。若 Friedman 顯著，再做 ECMADE-MOO 對各 baseline 的 Wilcoxon signed-rank post-hoc test，並加入 Holm correction 與 A12 effect size。")
    doc.add_paragraph("正文不放所有 pairwise p-value，只放下表的濃縮統計結果；完整 pairwise 結果放 supplementary CSV。")
    add_table(doc, ["Dataset", "Metric", "Friedman p", "Sig.", "ECMADE-MOO vs Best Baseline", "Holm adj. p", "A12"], compact_rows(compact), [1100, 1400, 1300, 700, 2600, 1200, 1060])
    doc.add_paragraph("A12 > 0.5 表示 ECMADE-MOO 較可能優於該 baseline；A12 < 0.5 表示該 baseline 較可能優於 ECMADE-MOO。")
    add_path(doc, "Supplementary full pairwise table", POSTHOC_CSV)
    add_path(doc, "Compact statistical table CSV", COMPACT_CSV)

    doc.add_heading("7. 分資料集解讀與後續實驗銜接", level=1)
    doc.add_paragraph("在正式 OR-Library portfolio instances 上，ECMADE-MOO 取得最高 HV 與最低 IGD，顯示其在受限制多目標投資組合問題中具有最佳解品質與收斂能力。")
    doc.add_paragraph("但 ECMADE-MOO 的 PF Overlap 與 EAF Band Width 並未同步優於 SPEA2 與 NSGA-II，表示高解品質不必然伴隨高 repeated-run stability。")
    doc.add_paragraph("Synthetic instances 顯示 fixed hand-crafted configuration 在 heterogeneous problem instances 上仍存在泛化與穩定性差距，這可自然銜接到 Experiment B 的 meta-designed ECMADE-MOO，以及 Experiment C 的 stability-aware objective。")

    doc.add_heading("8. 主要檔案位置", level=1)
    for dataset, path in [("Synthetic", SYN_DIR), ("OR-Library", OR_DIR)]:
        doc.add_heading(dataset, level=2)
        for name in ["run_metrics.csv", "instance_method_metrics.csv", "overall_method_summary.csv", "statistical_tests.csv", "generation_metrics_manifest.csv", "missing_outputs.csv"]:
            add_path(doc, name, os.path.join(path, name))
        add_path(doc, "figures", os.path.join(path, "figures"))

    doc.save(OUT_DOCX)
    print(f"DOCX={OUT_DOCX}")
    print(f"POSTHOC_CSV={POSTHOC_CSV}")
    print(f"COMPACT_CSV={COMPACT_CSV}")


if __name__ == "__main__":
    main()

# -*- coding: utf-8 -*-
"""Build the TEVC P0-lite research code location and workflow DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r".")
OUT = ROOT / "TEVC_P0_lite_研究程式碼位置與流程.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(90, 99, 110)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C2CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_cell_text(cell, bold=False, color=INK, size=9.5) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.12
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
        style_cell_text(cell, bold=True, color=DARK_BLUE, size=9.5)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = text
            style_cell_text(cell, size=9.2)
    doc.add_paragraph()


def add_command(doc: Document, command: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.18)
    run = p.add_run(command)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.8)
    run.font.color.rgb = RGBColor(45, 55, 72)


def add_note_box(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10)
    p.add_run("  " + body).font.size = Pt(10)
    style_cell_text(cell, size=10)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("TEVC P0-lite Research Code Reference")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("TEVC P0-lite 研究程式碼位置與流程")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("更新日期：2026-06-27｜工作目錄：.")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    add_note_box(
        doc,
        "使用目的",
        "這份文件整理目前 P0-lite 研究程式的位置、資料來源、演算法啟動方式、輸出結構與後續實驗流程，方便之後接續跑 OR-Library 與 synthetic constrained portfolio instances。",
    )

    doc.add_heading("1. 專案根目錄", level=1)
    add_table(
        doc,
        ["項目", "位置 / 說明"],
        [
            ["研究程式碼根目錄", r"."],
            ["PlatEMO 版本", r"PlatEMO_v2.9.0\PlatEMO"],
            ["主要輸出根目錄", r"p0_lite_outputs"],
        ],
        [2300, 7060],
    )

    doc.add_heading("2. 資料與 Instance", level=1)
    add_table(
        doc,
        ["資料類型", "位置", "用途"],
        [
            ["OR-Library port1", r"data\orlib\port1.txt", "P0-lite baseline 的真實資料測試。"],
            ["Synthetic instances", r"data\synthetic_constrained_portfolio\instances", "訓練 / 驗證 / 測試切分後的 synthetic constrained portfolio instances。"],
            ["Synthetic manifest", r"data\synthetic_constrained_portfolio\manifest.csv", "記錄 split、n、K、K/n、correlation、return distribution、risk structure 與檔案路徑。"],
            ["Synthetic summary", r"data\synthetic_constrained_portfolio\summary.json", "快速確認 synthetic instance 總數與各因子層級。"],
        ],
        [1900, 4200, 3260],
    )

    doc.add_heading("3. 核心程式", level=1)
    add_table(
        doc,
        ["檔案", "角色"],
        [
            ["PortfolioORLIB.m", "PlatEMO problem class；讀取 OR-Library-like txt，處理 cardinality、sum(w)=1、w>=0，目標為 minimize variance 與 maximize return。"],
            ["P0LiteUtils.m", "共用工具；負責讀資料、first front、輸出 archive、runtime、feasible rate、generation logs 與 port1 summary。"],
            ["SyntheticRunner.m", "Synthetic 共用 runner；讀 manifest，逐 instance 使用對應 K 與資料路徑，並統一輸出格式。"],
            ["generate_synthetic_portfolio_instances.py", "Synthetic instance 產生器；依 n、K/n、correlation、return distribution、risk structure 產生 OR-Library-like txt。"],
            ["ECMADE_MOO.m", "已加入 PlatEMO 的 ECMADE-MOO 演算法版本。"],
        ],
        [3000, 6360],
    )

    doc.add_heading("4. OR-Library port1 啟動檔", level=1)
    add_table(
        doc,
        ["Method", "啟動檔"],
        [
            ["NSGA-II", "run_p0_lite_port1_nsga2.m"],
            ["SPEA2", "run_p0_lite_port1_spea2.m"],
            ["MOEA/D", "run_p0_lite_port1_moead.m"],
            ["GDE3", "run_p0_lite_port1_gde3.m"],
            ["ECMADE-MOO", "run_p0_lite_port1_ecmade_moo.m"],
        ],
        [2100, 7260],
    )

    doc.add_heading("5. Synthetic 啟動檔", level=1)
    add_table(
        doc,
        ["Method", "啟動檔"],
        [
            ["NSGA-II", "run_p0_lite_synthetic_nsga2.m"],
            ["SPEA2", "run_p0_lite_synthetic_spea2.m"],
            ["MOEA/D", "run_p0_lite_synthetic_moead.m"],
            ["GDE3", "run_p0_lite_synthetic_gde3.m"],
            ["ECMADE-MOO", "run_p0_lite_synthetic_ecmade_moo.m"],
        ],
        [2100, 7260],
    )

    doc.add_heading("6. 標準流程", level=1)
    doc.add_heading("6.1 產生 synthetic constrained portfolio instances", level=2)
    add_command(doc, r"python .\generate_synthetic_portfolio_instances.py")
    doc.add_paragraph("目前預設會建立 192 個 instances，包含 train、validation、test，且涵蓋圖片要求的 n、K/n、correlation、return distribution 與 risk structure。")

    doc.add_heading("6.2 執行 OR-Library port1 baseline", level=2)
    for cmd in [
        "matlab -batch \"cd('.'); run_p0_lite_port1_nsga2;\"",
        "matlab -batch \"cd('.'); run_p0_lite_port1_spea2;\"",
        "matlab -batch \"cd('.'); run_p0_lite_port1_moead;\"",
        "matlab -batch \"cd('.'); run_p0_lite_port1_gde3;\"",
        "matlab -batch \"cd('.'); run_p0_lite_port1_ecmade_moo;\"",
    ]:
        add_command(doc, cmd)

    doc.add_heading("6.3 執行 synthetic baseline", level=2)
    for cmd in [
        "matlab -batch \"cd('.'); run_p0_lite_synthetic_nsga2;\"",
        "matlab -batch \"cd('.'); run_p0_lite_synthetic_spea2;\"",
        "matlab -batch \"cd('.'); run_p0_lite_synthetic_moead;\"",
        "matlab -batch \"cd('.'); run_p0_lite_synthetic_gde3;\"",
        "matlab -batch \"cd('.'); run_p0_lite_synthetic_ecmade_moo;\"",
    ]:
        add_command(doc, cmd)

    doc.add_heading("6.4 後處理與繪圖", level=2)
    add_command(doc, "matlab -batch \"cd('.'); postprocess_p0_lite_port1_metrics_figures;\"")
    doc.add_paragraph("目前完整後處理主要對 port1 結果使用；synthetic 結果已用相同輸出檔命名，後續可擴充共用後處理。")

    doc.add_heading("7. 輸出位置與檔案", level=1)
    add_table(
        doc,
        ["類型", "位置"],
        [
            ["port1 結果", r"p0_lite_outputs\port1_nsga2_spea2_logged"],
            ["synthetic 結果", r"p0_lite_outputs\synthetic_constrained_portfolio"],
            ["port1 圖與 summary", r"p0_lite_outputs\port1_nsga2_spea2_logged\figures 與 summary CSV"],
        ],
        [2100, 7260],
    )
    add_table(
        doc,
        ["每個 run 輸出檔", "內容"],
        [
            ["final_archive_obj.csv / final_archive_dec.csv", "最終 non-dominated archive 的 objective 與 decision。"],
            ["pf_obj.csv / pf_dec.csv / pf_points.csv", "最終 Pareto front points。"],
            ["generation_pf_points.csv", "每代 PF points，可用於 PF overlay、PF heatmap、EAF 與 drift 類分析。"],
            ["generation_population_log.csv", "每代 feasible count、feasible rate、PF size 等 log。"],
            ["feasible_rate.csv", "final PF 與 population 的 feasible rate。"],
            ["runtime.csv", "單次 run 的 runtime seconds。"],
            ["instance_metadata.csv", "synthetic runner 會記錄 instance、split、n、K、K/n 與資料結構。"],
        ],
        [3300, 6060],
    )

    doc.add_heading("8. Synthetic Runner 常用控制變數", level=1)
    add_table(
        doc,
        ["變數", "用途", "範例"],
        [
            ["SYNTHETIC_SMOKE", "只跑 1 個 train instance、1 run，用於測試。", "SYNTHETIC_SMOKE=true"],
            ["SYNTHETIC_SPLITS", "限制 split，例如只跑 test。", "SYNTHETIC_SPLITS={'test'}"],
            ["SYNTHETIC_MAX_INSTANCES", "限制 manifest 前 N 個 instances。", "SYNTHETIC_MAX_INSTANCES=4"],
            ["SYNTHETIC_RUNS", "覆蓋每個 instance 的 independent runs。", "SYNTHETIC_RUNS=30"],
            ["SYNTHETIC_N", "覆蓋 population size。", "SYNTHETIC_N=100"],
            ["SYNTHETIC_MAXFE", "覆蓋 function evaluations。", "SYNTHETIC_MAXFE=10000"],
        ],
        [2300, 3860, 3200],
    )
    doc.add_paragraph("例：只跑 test split 前 4 題 NSGA-II：")
    add_command(doc, "matlab -batch \"cd('.'); SYNTHETIC_SPLITS={'test'}; SYNTHETIC_MAX_INSTANCES=4; run_p0_lite_synthetic_nsga2;\"")

    doc.add_heading("9. 目前研究流程狀態", level=1)
    add_table(
        doc,
        ["項目", "目前狀態"],
        [
            ["OR-Library port1", "已完成多演算法 runner 與後處理/繪圖流程。"],
            ["Synthetic instances", "已建立 192 題 P0-lite synthetic constrained portfolio dataset。"],
            ["Synthetic runner", "已新增並通過 NSGA-II smoke test；可用相同格式跑各 baseline。"],
            ["ECMADE-MOO", "已加入 MATLAB/PlatEMO 版本，可與 baseline 一起跑 port1 與 synthetic。"],
            ["下一步", "建議先用 synthetic smoke / 小批次確認每個方法，再擴大到 validation/test 與完整 30 runs。"],
        ],
        [2500, 6860],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

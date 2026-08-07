from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ANALYSIS_DIR = ROOT / "outputs" / "selector_level_ablation_20260728" / "final_test_analysis"
OUT_DOCX = ANALYSIS_DIR / "Selector_Level_Ablation_正式消融數據報告_中文版_20260730.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "C9D1D9"

METHOD_LABELS = {
    "SelectorAblation_FullSelector_ECMADE_MOO": "FullSelector",
    "SelectorAblation_NoInstanceFeatures_ECMADE_MOO": "NoInstanceFeatures",
    "SelectorAblation_NoThetaFeatures_ECMADE_MOO": "NoThetaFeatures",
    "SelectorAblation_RandomizedLabels_ECMADE_MOO": "RandomizedLabels",
}


def fmt(value, digits: int = 4) -> str:
    try:
        f = float(value)
    except Exception:
        return str(value)
    return f"{f:.{digits}f}"


def set_run_font(run, *, name="Calibri", east_asia="Microsoft JhengHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_widths(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_dxa):
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_cell_text(cell, text, *, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "TEVC Selector-Level Ablation"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Generated 2026-07-30"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("正式 Selector-Level 消融實驗數據報告")
    set_run_font(r, size=22, bold=True, color=RGBColor(0, 0, 0))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("Experiment C / ECMADE-MOO selector 特徵與標籤層級 final-test ablation")
    set_run_font(r2, size=11.5, color=MUTED)


def add_para(doc, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    set_table_borders(table, color="D7DBE2", size="4")
    set_cell_margins(table, top=130, bottom=130, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_dxa, *, font_size=8.7, left_cols=None):
    left_cols = set(left_cols or [])
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths_dxa)
    set_table_borders(table)
    set_cell_margins(table)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_shading(hdr.cells[idx], LIGHT_FILL)
        set_cell_text(hdr.cells[idx], header, bold=True, size=8.7)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if idx in left_cols else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, size=font_size, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def method_label(name: str) -> str:
    return METHOD_LABELS.get(str(name), str(name))


def load_tables():
    return {
        "completeness": pd.read_csv(ANALYSIS_DIR / "selector_ablation_run_completeness.csv"),
        "overall": pd.read_csv(ANALYSIS_DIR / "selector_ablation_overall_summary.csv"),
        "friedman": pd.read_csv(ANALYSIS_DIR / "selector_ablation_friedman_tests.csv"),
        "pairwise": pd.read_csv(ANALYSIS_DIR / "selector_ablation_pairwise_wilcoxon_holm.csv"),
    }


def main():
    data = load_tables()
    doc = Document()
    setup_document(doc)
    add_title(doc)

    add_callout(
        doc,
        "核心結論",
        "正式 selector-level final-test 消融已完成並產出統計表；但是結果不支持 FullSelector 在 RankScore 上顯著優於移除特徵或 randomized-label variants。此結果應作為 limitation 與 selector 設計診斷，而不是作為 FullSelector 有效性的正向證據。",
    )

    add_heading(doc, "1. 實驗完成狀態與資料來源")
    add_para(
        doc,
        "本報告整理 outputs/selector_level_ablation_20260728/final_test_analysis 下的正式 final-test 分析結果。"
        "每個變體使用 common-reference post-processing，配對單位為 split-instance-K；RankScore 為 HV、IGD、PF_Overlap、PF_Drift、Diversity 與 Runtime 的平均排名，數值越低越好。"
    )
    add_para(
        doc,
        "RNG 與 seed policy：Selector-level final-test ablation 使用 MATLAB/PlatEMO 的 mcg16807 random stream，並以 run index 作為 optimizer seed；各 selector variants 在相同 test instance 使用一致的 seed assignment。"
    )
    comp = data["completeness"].copy()
    comp_rows = [[method_label(r["method"]), int(r["instances"]), int(r["runs"])] for _, r in comp.iterrows()]
    add_table(
        doc,
        ["Selector 變體", "instances", "runs"],
        comp_rows,
        [5000, 1600, 1600],
        left_cols={0},
    )
    add_para(
        doc,
        "注意：RandomizedLabels 只完成 29 個 instances、284 runs；其他三個變體皆為 32 instances、320 runs。"
        "因此涉及 RandomizedLabels 的配對統計只使用共同可配對的 29 個 units。"
    )

    add_heading(doc, "2. Overall 描述統計")
    overall = data["overall"].copy()
    overall["label"] = overall["method"].map(method_label)
    rows = []
    for _, r in overall.iterrows():
        rows.append(
            [
                r["label"],
                int(r["instances"]),
                fmt(r["mean_RankScore"]),
                fmt(r["mean_InstanceRank"]),
                int(r["first_place_instances"]),
                fmt(r["mean_HV"]),
                fmt(r["mean_IGD"]),
                fmt(r["mean_PF_Overlap"]),
                fmt(r["mean_PF_Drift"]),
                fmt(r["mean_Runtime"], 3),
                fmt(r["overall_RankScore"]),
            ]
        )
    add_table(
        doc,
        ["變體", "N", "mean RankScore", "mean InstRank", "1st", "HV", "IGD", "Overlap", "Drift", "Runtime", "overall"],
        rows,
        [2100, 550, 1050, 1050, 550, 800, 800, 850, 800, 800, 810],
        font_size=7.8,
        left_cols={0},
    )
    add_para(
        doc,
        "描述統計上，RandomizedLabels 的 overall_RankScore 最低，但其 instances 不完整；NoThetaFeatures 在 32-instance 完整變體中有最低 mean_RankScore 與最多 first-place instances。FullSelector 的 overall_RankScore 為 3.1667，並非描述統計最佳。"
    )

    add_heading(doc, "3. Friedman 整體差異檢定")
    friedman = data["friedman"].copy()
    f_rows = []
    for _, r in friedman.iterrows():
        f_rows.append(
            [
                r["metric"],
                r["direction"],
                int(r["n_paired_units"]),
                fmt(r["friedman_chi_square"]),
                fmt(r["p_value"], 6),
                "顯著" if bool(r["significant"]) else "不顯著",
            ]
        )
    add_table(
        doc,
        ["Metric", "方向", "paired N", "chi-square", "p-value", "結論"],
        f_rows,
        [1800, 900, 1000, 1500, 1400, 1400],
    )
    add_para(
        doc,
        "Friedman test 顯示 RankScore、HV、PF_Overlap、PF_Drift、Diversity 與 Runtime 在 selector variants 間存在整體差異；IGD 則未達顯著。"
    )

    add_heading(doc, "4. RankScore Pairwise Wilcoxon + Holm")
    pairwise = data["pairwise"].copy()
    rank_pairs = pairwise[pairwise["metric"].eq("RankScore")].copy()
    p_rows = []
    for _, r in rank_pairs.iterrows():
        p_rows.append(
            [
                f"Full vs {method_label(r['baseline'])}",
                int(r["n_paired_units"]),
                f"{int(r['wins'])}/{int(r['ties'])}/{int(r['losses'])}",
                fmt(r["median_signed_improvement"]),
                fmt(r["raw_p_value"], 6),
                fmt(r["holm_p_value"], 6),
                "是" if bool(r["significant_after_holm"]) else "否",
            ]
        )
    add_table(
        doc,
        ["Comparison", "paired N", "W/T/L", "median improvement", "raw p", "Holm p", "顯著"],
        p_rows,
        [2650, 900, 950, 1500, 1100, 1100, 820],
        left_cols={0},
    )
    add_para(
        doc,
        "以 FullSelector 為 primary method 時，RankScore 的三個 pairwise comparisons 經 Holm correction 後皆不顯著。FullSelector 對 NoInstanceFeatures 的 raw p=0.0491，但 Holm p=0.1473，不能作 confirmatory evidence。"
    )

    add_heading(doc, "5. 其他指標的輔助觀察")
    selected = pairwise[
        pairwise["metric"].isin(["Diversity_mean", "Runtime_mean"])
        & pairwise["baseline"].isin([
            "SelectorAblation_NoInstanceFeatures_ECMADE_MOO",
            "SelectorAblation_NoThetaFeatures_ECMADE_MOO",
        ])
    ].copy()
    aux_rows = []
    for _, r in selected.iterrows():
        aux_rows.append(
            [
                r["metric"],
                f"Full vs {method_label(r['baseline'])}",
                f"{int(r['wins'])}/{int(r['ties'])}/{int(r['losses'])}",
                fmt(r["median_signed_improvement"]),
                fmt(r["holm_p_value"], 6),
                "是" if bool(r["significant_after_holm"]) else "否",
            ]
        )
    add_table(
        doc,
        ["Metric", "Comparison", "W/T/L", "median improvement", "Holm p", "顯著"],
        aux_rows,
        [1550, 2800, 950, 1500, 1250, 850],
        left_cols={0, 1},
    )
    add_para(
        doc,
        "輔助指標顯示 FullSelector 在 Runtime 上顯著快於 NoInstanceFeatures 與 NoThetaFeatures；在 Diversity 上只顯著優於 NoThetaFeatures。這些結果不足以推翻 RankScore 主結論，但可作為診斷訊息。"
    )

    add_heading(doc, "6. 可寫入論文的保守結論")
    add_bullet(doc, "正式 selector-level 消融已完成；完整輸出包含 run-level metrics、instance summary、ranked table、Friedman tests 與 pairwise Wilcoxon-Holm。")
    add_bullet(doc, "FullSelector 未在 RankScore 上顯著優於 NoInstanceFeatures、NoThetaFeatures 或 RandomizedLabels，因此不能宣稱完整 feature selector 已被消融實驗證明為必要。")
    add_bullet(doc, "NoThetaFeatures 在完整 32-instance variants 中表現較佳，暗示目前 theta-feature encoding 或 full-feature model 可能存在噪音、過擬合或 selector 訓練不穩定。")
    add_bullet(doc, "此結果應放在 limitation / diagnostic ablation 段落，建議後續做 feature pruning、regularization、重新抽樣訓練或更嚴格的 nested validation。")

    add_heading(doc, "附錄：主要輸出檔案")
    source_rows = [
        ["Run metrics", str(ANALYSIS_DIR / "selector_ablation_run_metrics.csv")],
        ["Overall summary", str(ANALYSIS_DIR / "selector_ablation_overall_summary.csv")],
        ["Friedman tests", str(ANALYSIS_DIR / "selector_ablation_friedman_tests.csv")],
        ["Wilcoxon + Holm", str(ANALYSIS_DIR / "selector_ablation_pairwise_wilcoxon_holm.csv")],
        ["Completeness", str(ANALYSIS_DIR / "selector_ablation_run_completeness.csv")],
    ]
    add_table(doc, ["項目", "路徑"], source_rows, [2100, 7260], font_size=7.8, left_cols={0, 1})

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()

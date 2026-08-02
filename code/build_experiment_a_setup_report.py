import os
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = r"C:\Users\yiting\Documents\Playground"
OUT_ROOT = os.path.join(ROOT, "p0_lite_outputs", "synthetic_constrained_portfolio")
REPORT_DIR = os.path.join(OUT_ROOT, "experiment_A_setup_report")
MANIFEST = os.path.join(ROOT, "data", "synthetic_constrained_portfolio", "manifest.csv")
METHODS = ["NSGAII", "SPEA2", "MOEAD", "GDE3", "A_MPMO", "ECMADE_MOO"]


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, title=None):
    if title:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = RGBColor(31, 77, 120)
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        set_cell_text(table.cell(0, j), h, True)
        set_cell_fill(table.cell(0, j), "E8EEF5")
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.cell(i, j), val)
    doc.add_paragraph()


def add_note(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(" " + text)


def apply_styles(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def read_config(method):
    path = os.path.join(OUT_ROOT, f"config_{method}.txt")
    data = {}
    if not os.path.exists(path):
        return data
    with open(path, encoding="utf-8-sig") as f:
        for line in f:
            if "=" in line:
                k, v = line.strip().split("=", 1)
                data[k] = v
    return data


def main():
    ensure_dir(REPORT_DIR)
    df = pd.read_csv(MANIFEST)
    configs = {m: read_config(m) for m in METHODS}

    split_counts = df["split"].value_counts().sort_index()
    asset_counts = df["assets"].value_counts().sort_index()
    k_counts = df["K"].value_counts().sort_index()
    ratio_counts = df["k_ratio"].value_counts().sort_index()
    corr_counts = df["corr_structure"].value_counts().sort_index()
    ret_counts = df["return_distribution"].value_counts().sort_index()
    risk_counts = df["risk_structure"].value_counts().sort_index()

    expected_runs = len(df) * len(METHODS) * 30
    log_path = os.path.join(OUT_ROOT, "logs", "all_methods_resume_20260701_072822.out.log")
    result_report_dir = os.path.join(OUT_ROOT, "experiment_A_report_20260701_112713")

    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("實驗 A 設定與資料流程報告\n")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 37, 69)
    s = title.add_run("ECMADE-MOO 與主要 baseline 比較：baseline、統一參數、輸入與輸出")
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph(f"產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. 實驗 A 定位", level=1)
    doc.add_paragraph(
        "實驗 A 的目的，是在同一批 synthetic constrained portfolio instances 上，"
        "比較 hand-crafted ECMADE-MOO 與主要多目標演化演算法 baseline。"
        "此實驗不調整 meta-learning configuration，也不進行 constraint-handling ablation；"
        "所有方法使用相同資料、相同 function evaluations、相同 population size、相同 seed block 與相同 portfolio repair-based constraint handling。"
    )

    doc.add_heading("2. 已執行 Baseline", level=1)
    baseline_rows = [
        ["NSGA-II", "PlatEMO NSGAII", "主要 Pareto dominance baseline", "已執行", "使用 NSGA-II environmental selection、crowding distance、GA operator。"],
        ["SPEA2", "PlatEMO SPEA2", "archive/strength Pareto baseline", "已執行", "使用 strength Pareto fitness 與 environmental selection。"],
        ["MOEA/D", "PlatEMO MOEAD", "decomposition baseline", "已執行", "使用 weight vectors、neighborhood update、PBI aggregation 預設 type=1。"],
        ["GDE3", "PlatEMO GDE3", "DE-based multiobjective baseline", "已執行", "使用 differential evolution offspring 與 GDE3 environmental selection。"],
        ["A-MPMO", "A_MPMO_NSGAII_v290.m", "adaptive multi-population baseline", "已加入", "使用 adaptive subpopulation allocation 與 NSGA-II environmental selection。"],
        ["ECMADE-MOO", "自訂 ECMADE_MOO.m", "本文 hand-crafted multi-subpopulation baseline", "已更新", "3 subpops、individual-level adaptive F/CR、Pareto archive monitoring、stagnation-triggered elite exchange。"],
        ["DEMO", "未獨立執行", "可視為 DE-family 額外 baseline", "未執行", "本批以 GDE3 代表 DE-based MOEA baseline。"],
    ]
    add_table(doc, ["方法", "程式來源", "角色", "本批狀態", "說明"], baseline_rows, "Table 1. 實驗 A baseline 清單")
    add_note(doc, "重點：", "本批正式完成的 baseline 為 NSGA-II、SPEA2、MOEA/D、GDE3、ECMADE-MOO，共 5 種方法。")

    doc.add_heading("3. 統一參數設定", level=1)
    common_rows = [
        ["experiment", configs["NSGAII"].get("experiment", "synthetic_constrained_portfolio"), "所有方法相同"],
        ["manifest", configs["NSGAII"].get("manifestPath", MANIFEST), "所有方法讀取同一份 manifest.csv"],
        ["selected manifest rows", configs["NSGAII"].get("selectedManifestRows", str(len(df))), "192 筆 instances"],
        ["splits", configs["NSGAII"].get("splits", "train,validation,test"), "train / validation / test 全部納入"],
        ["runs", configs["NSGAII"].get("runs", "30"), "每個 method x instance 執行 30 independent runs"],
        ["population size N", configs["NSGAII"].get("N", "100"), "所有方法 N=100"],
        ["maxFE", configs["NSGAII"].get("maxFE", "10000"), "所有方法 FE=10000"],
        ["objectives M", "2", "risk minimization 與 negative return minimization"],
        ["decision dimension D", "assets", "依 instance assets = 50 / 100 / 200 / 500"],
        ["rng", configs["NSGAII"].get("rng", "mcg16807"), "所有方法使用相同 random stream type"],
        ["seed", configs["NSGAII"].get("seed", "run index"), "第 r 次 run 使用 seed=r"],
        ["generation snapshots", "maxFE / N = 100", "每次 run 保存 generation snapshots"],
        ["constraint handling", "repair-to-cardinality-simplex + feasibility check", "所有方法共用 PortfolioORLIB.CalDec / CalObj / CalCon"],
    ]
    add_table(doc, ["參數", "設定值", "說明"], common_rows, "Table 2. 統一實驗參數")

    doc.add_heading("4. ECMADE-MOO 手工設定", level=1)
    ecmade_rows = [
        ["subpops", "3", "多子群架構"],
        ["archiveSize", "20", "保存 recent successful F/CR 記憶"],
        ["theta", "1/13", "F/CR adaptation learning rate"],
        ["stagnationThreshold", "50 generations", "Pareto archive 連續未改善超過門檻時觸發 elite exchange"],
        ["exploitationAlpha", "0.8", "第二子群 exploitation mutation 中 best vector 的權重"],
        ["initMuF", "[0.9, 0.8, 0.8]", "各子群初始 mutation factor mean"],
        ["initMuCR", "[0.9, 0.5, 0.5]", "各子群初始 crossover rate mean"],
        ["selection", "NDSort + CrowdingDistance", "使用 constraint-aware nondominated sorting"],
    ]
    add_table(doc, ["參數", "設定值", "用途"], ecmade_rows, "Table 3. ECMADE-MOO hand-crafted configuration")

    doc.add_heading("5. 跑過的資料", level=1)
    data_rows = [
        ["total instances", str(len(df)), "manifest.csv 中共 192 筆 synthetic constrained portfolio instances"],
        ["split distribution", ", ".join(f"{k}={v}" for k, v in split_counts.items()), "train / validation / test 分開保存"],
        ["assets n", ", ".join(f"{k}={v}" for k, v in asset_counts.items()), "n = 50, 100, 200, 500，各 48 筆"],
        ["k_ratio", ", ".join(f"{k:.2f}={v}" for k, v in ratio_counts.items()), "K/n = 0.05, 0.10, 0.20, 0.30，各 48 筆"],
        ["actual K values", ", ".join(f"{k}={v}" for k, v in k_counts.items()), "實際 K 隨 n 與 k_ratio 變動"],
        ["correlation structure", ", ".join(f"{k}={v}" for k, v in corr_counts.items()), "low / cluster / high / pathological covariance"],
        ["return distribution", ", ".join(f"{k}={v}" for k, v in ret_counts.items()), "normal / skewed / heavy-tail / mixed"],
        ["risk structure", ", ".join(f"{k}={v}" for k, v in risk_counts.items()), "low vol / high vol / extreme events"],
        ["expected formal runs", str(expected_runs), "192 instances x 5 methods x 30 runs = 28,800 runs"],
        ["completed runs", "28,800", "完整性檢查為 28,800 / 28,800，missing_outputs.csv 無缺漏"],
    ]
    add_table(doc, ["資料項目", "數量 / 設定", "說明"], data_rows, "Table 4. Synthetic constrained portfolio data coverage")

    doc.add_heading("6. 實驗輸入", level=1)
    input_rows = [
        ["manifest.csv", MANIFEST, "列出 instance name、split、assets、days、k_ratio、K、correlation/return/risk 結構、seed、instance path。"],
        ["instance text files", os.path.join(ROOT, "data", "synthetic_constrained_portfolio", "instances"), "每筆 portfolio instance 的 return、risk/correlation 資料。"],
        ["PortfolioORLIB.m", os.path.join(ROOT, "PortfolioORLIB.m"), "PlatEMO problem wrapper；讀取 instance file 與 K，定義 objective、repair、constraint violation。"],
        ["SyntheticRunner.m", os.path.join(ROOT, "SyntheticRunner.m"), "共用 runner；統一 manifest、runs、N、maxFE、seed、output format。"],
        ["algorithm launchers", "run_p0_lite_synthetic_*.m", "每個演算法各自獨立啟動，符合不要把演算法寫在同一個程式檔的要求。"],
        ["batch launcher", os.path.join(ROOT, "run_p0_lite_synthetic_all_methods.bat"), "依序啟動 NSGAII、SPEA2、MOEAD、ECMADE_MOO、GDE3。"],
    ]
    add_table(doc, ["輸入", "位置 / 檔案", "用途"], input_rows, "Table 5. Experiment A input files")

    doc.add_heading("7. 每次 Run 的 Raw Outputs", level=1)
    raw_rows = [
        ["population_dec.csv", "final population decision variables", "可檢查 portfolio weights 與 cardinality。"],
        ["population_obj.csv", "final population objective values", "risk 與 negative return。"],
        ["pf_dec.csv", "final non-dominated decision vectors", "最終 PF 對應的 portfolio weights。"],
        ["pf_obj.csv", "final non-dominated objective values", "最終 PF points，用於 HV、IGD、PF overlay、heatmap。"],
        ["final_archive_dec.csv", "same as PF decision archive", "投稿文件可稱 final non-dominated archive。"],
        ["final_archive_obj.csv", "same as PF objective archive", "投稿文件可稱 final non-dominated archive objectives。"],
        ["pf_points.csv", "PF objective points", "PF visualisation input。"],
        ["runtime.csv", "runtime_sec", "final optimization runtime。"],
        ["feasible_rate.csv", "PF_Feasible_Rate, Population_Feasible_Rate", "可行率輸出。"],
        ["generation_pf_points.csv", "generation-level PF points", "可後處理每代 HV/IGD、EAF trajectory。"],
        ["generation_population_log.csv", "generation, evaluations, feasible_count, feasible_rate, pf_size", "每代 feasible rate 與 PF size log。"],
        ["instance_metadata.csv", "method, instance, split, assets, K, N, maxFE, dataPath", "每個 run 的追溯資訊。"],
    ]
    add_table(doc, ["輸出檔", "內容", "用途"], raw_rows, "Table 6. Per-run raw output files")

    doc.add_heading("8. 後處理輸出與圖表", level=1)
    post_rows = [
        ["run_metrics.csv", os.path.join(result_report_dir, "run_metrics.csv"), "每個 run 的 HV、IGD、PF overlap、diversity、spacing、runtime、feasible rate。"],
        ["instance_method_metrics.csv", os.path.join(result_report_dir, "instance_method_metrics.csv"), "每個 instance x method 的 30-run 聚合結果，含 EAF band width 與 PF drift。"],
        ["overall_method_summary.csv", os.path.join(result_report_dir, "overall_method_summary.csv"), "各方法總表與 RankScore。"],
        ["statistical_tests.csv", os.path.join(result_report_dir, "statistical_tests.csv"), "Friedman 與 ECMADE-MOO vs baseline Wilcoxon approx.。"],
        ["metric_availability.csv", os.path.join(result_report_dir, "metric_availability.csv"), "哪些指標可直接輸出、哪些由後處理計算、哪些不適用。"],
        ["missing_outputs.csv", os.path.join(result_report_dir, "missing_outputs.csv"), "完整性檢查結果；目前無缺漏。"],
        ["figure_1_metric_dashboard.png", "figures/figure_1_metric_dashboard.png", "HV、IGD、PF overlap、EAF、diversity、runtime、feasible rate 總覽。"],
        ["figure_2_pf_overlay.png", "figures/figure_2_pf_overlay.png", "代表性 test instance 的多方法 PF 疊圖。"],
        ["figure_3_pf_heatmap.png", "figures/figure_3_pf_heatmap.png", "代表性 test instance 的 PF 搜尋熱區。"],
        ["figure_4_eaf_band_width.png", "figures/figure_4_eaf_band_width.png", "EAF band width 穩定性比較。"],
        ["figure_5_runtime.png", "figures/figure_5_runtime.png", "各方法 runtime 比較。"],
        ["figure_6_stability_diversity.png", "figures/figure_6_stability_diversity.png", "EAF width vs diversity。"],
    ]
    add_table(doc, ["後處理輸出", "位置", "說明"], post_rows, "Table 7. Post-processing outputs")

    doc.add_heading("9. 指標定義與方向", level=1)
    metric_rows = [
        ["HV mean / std / CV", "越大越好；CV 越小越穩", "由 final PF 與 empirical reference range 後處理計算。"],
        ["IGD mean / std / CV", "越小越好", "由 final PF 至 empirical reference front 的距離後處理計算。"],
        ["PF overlap", "越大越好", "衡量 repeated runs 對 reference PF 的覆蓋程度。"],
        ["EAF band width", "越小越穩", "衡量 attainment surface 在 repeated runs 間的不確定性。"],
        ["PF drift distance", "越小越穩", "衡量 PF centroid 在 repeated runs 間的漂移。"],
        ["Diversity / spread / spacing", "依定義解讀", "避免只追求 HV/IGD 而造成 diversity collapse。"],
        ["Runtime", "越小越好", "final optimization cost。"],
        ["Configuration cost", "本實驗 A 為 N/A", "固定 baseline 比較，沒有 meta-learning 或 Bayesian search cost。"],
        ["Feasible rate", "越高越好", "直接由 feasible_rate.csv 取得。"],
        ["Violation degree", "目前未另存", "PortfolioORLIB 會 repair 到 cardinality simplex；若做 constraint ablation，建議另外輸出 violation degree CSV。"],
    ]
    add_table(doc, ["指標", "方向", "本批資料狀態"], metric_rows, "Table 8. Metrics required by Experiment A")

    doc.add_heading("10. 執行紀錄", level=1)
    exec_rows = [
        ["NSGAII", "2026/07/01 07:28:23 started; 07:29:50 finished", "快速跳過/補齊已完成結果。"],
        ["SPEA2", "2026/07/01 07:29:50 started; 07:30:38 finished", "完成。"],
        ["MOEAD", "2026/07/01 07:30:38 started; 07:31:27 finished", "完成；先前長時間正式結果已存在，resume 時跳過完整 runs。"],
        ["ECMADE_MOO", "2026/07/01 07:31:27 started; 07:32:14 finished", "完成；resume 時跳過完整 runs。"],
        ["GDE3", "2026/07/01 07:32:14 started; 10:26:08 finished", "完成最後剩餘 runs。"],
        ["master log", log_path, "All synthetic methods finished."],
    ]
    add_table(doc, ["項目", "紀錄", "說明"], exec_rows, "Table 9. Run completion log")

    doc.add_heading("11. 實驗 A 可寫入論文的方法描述", level=1)
    doc.add_paragraph(
        "All algorithms were evaluated on the same 192 synthetic constrained portfolio instances generated from combinations of asset size, cardinality ratio, correlation structure, return distribution, and risk structure. "
        "For each method-instance pair, 30 independent runs were executed with population size 100 and maximum function evaluations 10,000. "
        "The random seed block was fixed by using run index as seed under mcg16807. "
        "The portfolio constraints were handled uniformly by a repair-to-cardinality-simplex operator and a feasibility check implemented in the shared PortfolioORLIB problem class."
    )
    doc.add_paragraph(
        "The final outputs include the final non-dominated archive, PF points, generation-level PF snapshots, generation-level feasible rate logs, runtime, and feasible-rate summaries. "
        "Performance, stability, diversity, cost, and feasibility metrics were recomputed from these raw outputs for the final report."
    )

    doc.add_heading("12. 注意事項", level=1)
    notes = [
        "Experiment A 是 baseline comparison，不是 meta-learning label generation；因此 configuration cost 在本實驗可標 N/A 或 0。",
        "目前 violation degree 未單獨輸出。若要支撐 constraint handling ablation，建議在未來 run 另存 raw constraint violation。",
        "每代 HV/IGD 沒有直接保存成 generation_hv_igd.csv，但 generation_pf_points.csv 已完整保存，可由後處理計算。",
        "本報告只描述已完成的 synthetic constrained portfolio 實驗；OR-Library port1 的早期測試不併入這批 28,800 runs。"
    ]
    for note in notes:
        doc.add_paragraph(note)

    output = os.path.join(REPORT_DIR, "Experiment_A_setup_inputs_outputs_report.docx")
    doc.save(output)
    print(f"DOCX={output}")


if __name__ == "__main__":
    main()

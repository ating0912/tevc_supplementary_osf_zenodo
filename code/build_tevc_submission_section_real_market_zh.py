from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
BASE_DIR = ROOT / "p0_lite_outputs" / "p1_rolling_window_market_validation_20260719"
REAL_DIR = BASE_DIR / "p0_real_market_validation"
SUMMARY_DIR = BASE_DIR / "summary"
OUT_DIR = ROOT / "outputs" / "tevc_submission_real_market_section_20260725"
DOCX_PATH = OUT_DIR / "TEVC_submission_section_real_market_validation_ABC_consistency_20260725.docx"
MD_PATH = OUT_DIR / "TEVC_submission_section_real_market_validation_ABC_consistency_20260725.md"


FONT_EN = "Calibri"
FONT_ZH = "Microsoft JhengHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(85, 85, 85)
HEADER_FILL = "F2F4F7"


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def style_doc(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document):
    p = doc.add_paragraph()
    set_spacing(p, after=3, line=1.05)
    r = p.add_run("投稿章節草稿：Real-market validation 與 A/B/C 實驗一致性")
    set_run(r, size=19, bold=True)

    p = doc.add_paragraph()
    set_spacing(p, after=12)
    r = p.add_run("適用位置：Experiment C 後方、Robustness Check、或 Supplementary Validation")
    set_run(r, size=10.5, color=GRAY)


def para(doc: Document, text: str):
    p = doc.add_paragraph()
    set_spacing(p)
    r = p.add_run(text)
    set_run(r)


def heading(doc: Document, text: str, level=1):
    doc.add_heading(text, level=level)


def bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, after=4)
        r = p.add_run(item)
        set_run(r)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold=False, size=8.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0, line=1.0)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], font_size=7.8):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=4)
    r = p.add_run(caption)
    set_run(r, size=9.3, bold=True, color=DARK_BLUE)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(headers):
        shade_cell(t.rows[0].cells[i], HEADER_FILL)
        set_cell(t.rows[0].cells[i], h, bold=True, size=font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, size=font_size)
    doc.add_paragraph()


def pct(v, d=2):
    return f"{float(v) * 100:.{d}f}%"


def num(v, d=3):
    return f"{float(v):.{d}f}"


def load_tables():
    overall = pd.read_csv(REAL_DIR / "p0_real_market_overall_summary.csv", encoding="utf-8-sig")
    method_summary = pd.read_csv(SUMMARY_DIR / "method_overall_summary.csv", encoding="utf-8-sig")
    cost = pd.read_csv(REAL_DIR / "p0_real_market_transaction_cost_overall.csv", encoding="utf-8-sig")
    windows = pd.read_csv(BASE_DIR / "windows" / "rolling_window_manifest.csv", encoding="utf-8-sig")

    merged = overall.merge(
        method_summary[["method", "mean_annual_volatility"]],
        on="method",
        how="left",
    )
    return merged, cost, windows


def df_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(c) for c in df.columns]
    rows = [[str(v) for v in row] for row in df.to_numpy().tolist()]
    widths = []
    for idx, header in enumerate(headers):
        values = [row[idx] for row in rows]
        widths.append(max([len(header), *[len(v) for v in values]]))

    def fmt_row(values: list[str]) -> str:
        return "| " + " | ".join(value.ljust(widths[i]) for i, value in enumerate(values)) + " |"

    separator = "| " + " | ".join("-" * width for width in widths) + " |"
    return "\n".join([fmt_row(headers), separator, *[fmt_row(row) for row in rows]])


def build_docx(overall: pd.DataFrame, cost: pd.DataFrame, windows: pd.DataFrame):
    doc = Document()
    style_doc(doc)
    add_title(doc)

    heading(doc, "4.1 Real-market validation 補強設計與最小可行執行", 1)
    para(
        doc,
        "為避免本文被解讀為僅針對 OR-Library 或合成 benchmark 調整方法，本研究新增 real-market rolling-window validation 作為 external validation / robustness check。此補強實驗的目的不是取代 Experiment A/B/C 的主要證據，而是檢查所提出方法在真實市場報酬、風險與交易成本條件下是否仍維持合理的 out-of-sample 行為。",
    )
    para(
        doc,
        "最小可行版本使用 SP100、NASDAQ100 與 TAIWAN50 三個真實市場股票池。每一個 rolling window 使用前段資料估計 expected return、covariance 與 downside-risk scenarios，並以後段資料評估 out-of-sample after-cost performance。每個 universe 建立 11 個 rolling windows，合計 33 個 universe-window；training window 為 3 年，testing window 為 6 個月。每個 window-method 執行 10 runs，最後從 final Pareto front 中選取 training Sharpe 最高的 portfolio 進入 out-of-sample 測試。",
    )
    bullets(
        doc,
        [
            f"Universe：{', '.join(sorted(windows['universe'].unique()))}。",
            f"Rolling windows：每個 universe {windows.groupby('universe')['window_id'].nunique().iloc[0]} 個，合計 {len(windows)} 個 universe-window。",
            "Training/testing：3 年 training window + 6 個月 testing window。",
            "交易成本：10 bps 作為主要 after-cost wealth curve；10 bps、20 bps、50 bps 作為 sensitivity analysis。",
            "目前 MVP 比較方法：NSGA-II、SPEA2、MOEA/D、GDE3、A-MPMO、ECMADE-MOO。",
        ],
    )

    heading(doc, "4.1.1 評估指標", 2)
    para(
        doc,
        "Real-market validation 回報 after-cost annualized return、annualized volatility、Sharpe ratio、Sortino ratio、CVaR(95%) loss、maximum drawdown、turnover 與 after-cost cumulative wealth curve。其中 CVaR(95%) loss 以 testing window 中最差 5% daily returns 的平均損失衡量；MDD 由 testing wealth curve 計算；turnover 衡量相鄰 rolling rebalance portfolio 的權重變動。對於 return、Sharpe 與 Sortino，數值越高越好；對於 volatility、CVaR loss、turnover，數值越低越好；MDD 因以負值表示，數值越接近 0 代表 drawdown 越小。",
    )

    rows = []
    for _, row in overall.iterrows():
        rows.append(
            [
                row["method"],
                str(int(row["windows"])),
                pct(row["mean_annual_net_return"]),
                pct(row["mean_annual_volatility"]),
                num(row["mean_sharpe"]),
                num(row["mean_sortino"]),
                pct(row["mean_cvar95_loss"]),
                pct(row["mean_max_drawdown"]),
                pct(row["mean_rebalance_turnover"]),
                num(row["stability_rank_score"]),
            ]
        )
    table(
        doc,
        "Table 4.1. Real-market rolling-window validation summary.",
        ["Method", "Win.", "Ann. net ret.", "Vol.", "Sharpe", "Sortino", "CVaR95", "MDD", "Turnover", "Stab. rank"],
        rows,
    )

    heading(doc, "4.1.2 主要結果與解讀", 2)
    para(
        doc,
        "結果顯示，MOEA/D 在 after-cost annualized return 上取得最佳平均表現，而 ECMADE-MOO 並未在 return-oriented 指標上主導所有 baseline。另一方面，ECMADE-MOO 取得最低的 mean CVaR(95%) loss，表示其在真實市場 out-of-sample data 下呈現較保守的 downside-risk profile。此結果支持本文將 real-market validation 定位為 robustness check：它提供 ECMADE-MOO 在 downside-risk control 上的外部證據，同時也清楚揭示其並非在真實市場報酬最大化上全面勝出。",
    )

    cost_rows = []
    for _, row in cost.iterrows():
        cost_rows.append(
            [
                row["cost_scenario"],
                row["method"],
                pct(row["mean_annual_net_return"]),
                pct(row["mean_turnover"]),
                pct(row["mean_cvar95_loss"]),
                pct(row["mean_max_drawdown"]),
                str(int(float(row["rank_annual_net_return"]))),
            ]
        )
    table(
        doc,
        "Table 4.2. Transaction-cost sensitivity under 10/20/50 bps.",
        ["Cost", "Method", "Ann. net ret.", "Turnover", "CVaR95", "MDD", "Return rank"],
        cost_rows,
    )
    para(
        doc,
        "交易成本敏感度結果顯示，當交易成本由 10 bps 提高至 50 bps 時，所有方法的 annualized net return 皆下降，但 return ranking 大致穩定。此結果表示 real-market conclusion 並非由單一交易成本設定造成。建議投稿時將 after-cost wealth curve 作為圖呈現，交易成本敏感度作為補充表格。",
    )

    plot_path = REAL_DIR / "p0_real_market_wealth_curve_10bps.png"
    if plot_path.exists():
        doc.add_picture(str(plot_path), width=Inches(6.7))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=2, after=8)
        r = p.add_run("Figure 4.1. After-cost cumulative wealth curve under 10 bps transaction cost.")
        set_run(r, size=9, color=GRAY)

    heading(doc, "4.1.3 與審稿建議的對應與限制", 2)
    para(
        doc,
        "目前已完成的 MVP 滿足 real-market rolling-window validation 的核心要求：真實市場股票池、3 年 training window、6 個月 out-of-sample testing window、10/20/50 bps 交易成本敏感度，以及 CVaR(95%)、MDD、turnover 與 after-cost wealth curve。尚未納入的項目包括 Hand-crafted、Bayesian、Meta-designed 與 Stability-aware ECMADE-MOO 四種配置版本在 real-market setting 下的直接比較，以及 average holdings、PF Overlap、PF Drift 等 Pareto-front stability 指標。因此，若投稿時間有限，建議將此節明確寫為 external validation / robustness check；若要達到完整版 reviewer expectation，則需將 Experiment B/C 的四種 configuration protocol 移植到 real-market windows 後再補跑。",
    )

    heading(doc, "4.2 Experiment A/B/C 與 real-market validation 的角色分工", 1)
    abc_rows = [
        ["Experiment A", "診斷 hand-crafted ECMADE-MOO 的優勢與限制", "NSGA-II、SPEA2、MOEA/D、GDE3、A-MPMO、ECMADE-MOO", "HV、IGD、PF Overlap、EAF Width、PF Drift、Diversity、Runtime", "Baseline diagnosis，證明固定配置不足"],
        ["Experiment B", "檢驗 instance-aware meta-designed configuration", "Hand-crafted、Random、Bayesian、Meta-designed", "HV、IGD、PF Overlap、PF Drift、Diversity、Runtime、RankScore", "主實驗一：證明 automated configuration 有價值"],
        ["Experiment C", "檢驗 stability-aware label engineering", "Hand-crafted、Random、Bayesian、Meta-designed、Stability-aware", "HV、IGD、PF Overlap、PF Drift、Diversity、Runtime、OverallRankScore、J-score", "主實驗二：證明 label objective 是 automated design 的關鍵"],
        ["Real-market validation", "檢查方法是否能外推到真實市場 rolling windows", "MVP：NSGA-II、SPEA2、MOEA/D、GDE3、A-MPMO、ECMADE-MOO；完整版：四種 ECMADE-MOO configuration", "After-cost wealth、annualized return、volatility、Sharpe、Sortino、CVaR、MDD、turnover", "External validation / robustness check，不作為主要訓練來源"],
    ]
    table(
        doc,
        "Table 4.3. Experimental roles in the revised manuscript.",
        ["Experiment", "Purpose", "Compared methods", "Main metrics", "Manuscript role"],
        abc_rows,
        font_size=7.2,
    )

    heading(doc, "4.3 A/B/C 實驗一致性與 post-processing protocol", 1)
    para(
        doc,
        "為避免 A/B/C 的結論受到後處理差異影響，正式投稿版本應明確宣告三組主要實驗使用一致的 post-processing pipeline。所有 raw runs 均先進行 feasibility check 與 constraint repair 記錄，再抽取 final Pareto front、final archive、objective values、decision variables、runtime、constraint metrics 與 archive metrics。接著在相同 instance 與相同 K 的 comparison group 中建立 common reference front，並以此計算 HV、IGD、PF Overlap、PF Drift、EAF Width 與 diversity。Experiment B/C 的 ranking 與 J-score 皆基於相同的 normalized metric table 產生，避免因不同 normalization 範圍導致不可比。",
    )
    bullets(
        doc,
        [
            "Common reference front：在同一 instance、K 與 split 下，合併所有 comparison methods 的 non-dominated objective points 建立 common reference front。",
            "Normalization：benefit metrics 採 min-max normalization；cost metrics 採 reversed min-max normalization；範圍限於同一 comparison group。",
            "Seed policy：所有方法使用一致的 run index 與 seed assignment；若某方法為 stochastic selector，selector seed 與 optimizer seed 分開記錄。",
            "Budget control：同一比較表內固定 maxFE、population size N、number of runs 與 termination criterion。",
            "Constraint handling：所有方法使用相同 portfolio feasibility definition，包括 cardinality K、weight bounds、budget sum-to-one repair 與 feasible-first selection。",
            "Missing-output check：每個 method-instance-run 必須檢查 final_archive_obj、final_archive_dec、runtime、feasible_rate、constraint_metrics 與 archive_metrics 是否存在；缺漏 run 不參與 label selection，並在缺漏表中列示。",
            "Tie handling：rank-based 指標使用 average rank 處理 ties；summary ranking 使用 dense rank 顯示並列名次。",
        ],
    )

    heading(doc, "4.4 投稿用結論文字", 1)
    para(
        doc,
        "Overall, Experiments A-C establish the proposed automated design pipeline on controlled benchmark and synthetic portfolio instances, while the additional real-market rolling-window validation serves as an external robustness check. The real-market results show that ECMADE-MOO does not dominate return-oriented baselines in after-cost annualized return, but it achieves the lowest CVaR(95%) loss among the compared methods. This indicates that the proposed design is more closely associated with downside-risk control than with unconditional return maximization under real-market data. Therefore, the revised manuscript reports real-market validation as supporting evidence for robustness and limitation analysis, while retaining Experiments B and C as the primary evidence for instance-aware and stability-aware automated configuration.",
    )

    heading(doc, "4.5 數據與輸出位置", 1)
    bullets(
        doc,
        [
            str(REAL_DIR / "p0_real_market_overall_summary.csv"),
            str(REAL_DIR / "p0_real_market_transaction_cost_overall.csv"),
            str(REAL_DIR / "p0_real_market_window_method_summary.csv"),
            str(REAL_DIR / "p0_real_market_wealth_curve_mean.csv"),
            str(REAL_DIR / "p0_real_market_wealth_curve_run_level.csv"),
            str(REAL_DIR / "p0_real_market_wealth_curve_10bps.png"),
            str(SUMMARY_DIR / "friedman_tests.csv"),
            str(SUMMARY_DIR / "pairwise_wilcoxon.csv"),
        ],
    )

    doc.save(DOCX_PATH)


def make_markdown(overall: pd.DataFrame, cost: pd.DataFrame, windows: pd.DataFrame) -> str:
    overall_md = overall[
        [
            "method",
            "windows",
            "mean_annual_net_return",
            "mean_annual_volatility",
            "mean_sharpe",
            "mean_sortino",
            "mean_cvar95_loss",
            "mean_max_drawdown",
            "mean_rebalance_turnover",
            "stability_rank_score",
        ]
    ].copy()
    for col in ["mean_annual_net_return", "mean_annual_volatility", "mean_cvar95_loss", "mean_max_drawdown", "mean_rebalance_turnover"]:
        overall_md[col] = overall_md[col].astype(float).map(lambda x: f"{x:.4f}")
    for col in ["mean_sharpe", "mean_sortino", "stability_rank_score"]:
        overall_md[col] = overall_md[col].astype(float).map(lambda x: f"{x:.3f}")

    cost_md = cost[
        [
            "cost_scenario",
            "method",
            "mean_annual_net_return",
            "mean_turnover",
            "mean_cvar95_loss",
            "mean_max_drawdown",
            "rank_annual_net_return",
        ]
    ].copy()
    for col in ["mean_annual_net_return", "mean_turnover", "mean_cvar95_loss", "mean_max_drawdown"]:
        cost_md[col] = cost_md[col].astype(float).map(lambda x: f"{x:.4f}")
    cost_md["rank_annual_net_return"] = cost_md["rank_annual_net_return"].astype(float).map(lambda x: f"{x:.0f}")

    overall_table = df_to_markdown(overall_md)
    cost_table = df_to_markdown(cost_md)

    return f"""# 投稿章節草稿：Real-market validation 與 A/B/C 實驗一致性

## 4.1 Real-market validation 補強設計與最小可行執行

為避免本文被解讀為僅針對 OR-Library 或合成 benchmark 調整方法，本研究新增 real-market rolling-window validation 作為 external validation / robustness check。此補強實驗的目的不是取代 Experiment A/B/C 的主要證據，而是檢查所提出方法在真實市場報酬、風險與交易成本條件下是否仍維持合理的 out-of-sample 行為。

最小可行版本使用 SP100、NASDAQ100 與 TAIWAN50 三個真實市場股票池。每一個 rolling window 使用前段資料估計 expected return、covariance 與 downside-risk scenarios，並以後段資料評估 out-of-sample after-cost performance。每個 universe 建立 11 個 rolling windows，合計 33 個 universe-window；training window 為 3 年，testing window 為 6 個月。每個 window-method 執行 10 runs，最後從 final Pareto front 中選取 training Sharpe 最高的 portfolio 進入 out-of-sample 測試。

### 4.1.1 評估指標

Real-market validation 回報 after-cost annualized return、annualized volatility、Sharpe ratio、Sortino ratio、CVaR(95%) loss、maximum drawdown、turnover 與 after-cost cumulative wealth curve。其中 CVaR(95%) loss 以 testing window 中最差 5% daily returns 的平均損失衡量；MDD 由 testing wealth curve 計算；turnover 衡量相鄰 rolling rebalance portfolio 的權重變動。

### Table 4.1. Real-market rolling-window validation summary

{overall_table}

### 4.1.2 主要結果與解讀

結果顯示，MOEA/D 在 after-cost annualized return 上取得最佳平均表現，而 ECMADE-MOO 並未在 return-oriented 指標上主導所有 baseline。另一方面，ECMADE-MOO 取得最低的 mean CVaR(95%) loss，表示其在真實市場 out-of-sample data 下呈現較保守的 downside-risk profile。此結果支持本文將 real-market validation 定位為 robustness check：它提供 ECMADE-MOO 在 downside-risk control 上的外部證據，同時也清楚揭示其並非在真實市場報酬最大化上全面勝出。

### Table 4.2. Transaction-cost sensitivity under 10/20/50 bps

{cost_table}

交易成本敏感度結果顯示，當交易成本由 10 bps 提高至 50 bps 時，所有方法的 annualized net return 皆下降，但 return ranking 大致穩定。此結果表示 real-market conclusion 並非由單一交易成本設定造成。

### 4.1.3 與審稿建議的對應與限制

目前已完成的 MVP 滿足 real-market rolling-window validation 的核心要求：真實市場股票池、3 年 training window、6 個月 out-of-sample testing window、10/20/50 bps 交易成本敏感度，以及 CVaR(95%)、MDD、turnover 與 after-cost wealth curve。尚未納入的項目包括 Hand-crafted、Bayesian、Meta-designed 與 Stability-aware ECMADE-MOO 四種配置版本在 real-market setting 下的直接比較，以及 average holdings、PF Overlap、PF Drift 等 Pareto-front stability 指標。因此，若投稿時間有限，建議將此節明確寫為 external validation / robustness check；若要達到完整版 reviewer expectation，則需將 Experiment B/C 的四種 configuration protocol 移植到 real-market windows 後再補跑。

## 4.2 Experiment A/B/C 與 real-market validation 的角色分工

| Experiment | Purpose | Compared methods | Main metrics | Manuscript role |
|---|---|---|---|---|
| Experiment A | 診斷 hand-crafted ECMADE-MOO 的優勢與限制 | NSGA-II、SPEA2、MOEA/D、GDE3、A-MPMO、ECMADE-MOO | HV、IGD、PF Overlap、EAF Width、PF Drift、Diversity、Runtime | Baseline diagnosis，證明固定配置不足 |
| Experiment B | 檢驗 instance-aware meta-designed configuration | Hand-crafted、Random、Bayesian、Meta-designed | HV、IGD、PF Overlap、PF Drift、Diversity、Runtime、RankScore | 主實驗一：證明 automated configuration 有價值 |
| Experiment C | 檢驗 stability-aware label engineering | Hand-crafted、Random、Bayesian、Meta-designed、Stability-aware | HV、IGD、PF Overlap、PF Drift、Diversity、Runtime、OverallRankScore、J-score | 主實驗二：證明 label objective 是 automated design 的關鍵 |
| Real-market validation | 檢查方法是否能外推到真實市場 rolling windows | MVP：NSGA-II、SPEA2、MOEA/D、GDE3、A-MPMO、ECMADE-MOO；完整版：四種 ECMADE-MOO configuration | After-cost wealth、annualized return、volatility、Sharpe、Sortino、CVaR、MDD、turnover | External validation / robustness check，不作為主要訓練來源 |

## 4.3 A/B/C 實驗一致性與 post-processing protocol

為避免 A/B/C 的結論受到後處理差異影響，正式投稿版本應明確宣告三組主要實驗使用一致的 post-processing pipeline。所有 raw runs 均先進行 feasibility check 與 constraint repair 記錄，再抽取 final Pareto front、final archive、objective values、decision variables、runtime、constraint metrics 與 archive metrics。接著在相同 instance 與相同 K 的 comparison group 中建立 common reference front，並以此計算 HV、IGD、PF Overlap、PF Drift、EAF Width 與 diversity。Experiment B/C 的 ranking 與 J-score 皆基於相同的 normalized metric table 產生，避免因不同 normalization 範圍導致不可比。

- Common reference front：在同一 instance、K 與 split 下，合併所有 comparison methods 的 non-dominated objective points 建立 common reference front。
- Normalization：benefit metrics 採 min-max normalization；cost metrics 採 reversed min-max normalization；範圍限於同一 comparison group。
- Seed policy：所有方法使用一致的 run index 與 seed assignment；若某方法為 stochastic selector，selector seed 與 optimizer seed 分開記錄。
- Budget control：同一比較表內固定 maxFE、population size N、number of runs 與 termination criterion。
- Constraint handling：所有方法使用相同 portfolio feasibility definition，包括 cardinality K、weight bounds、budget sum-to-one repair 與 feasible-first selection。
- Missing-output check：每個 method-instance-run 必須檢查 final_archive_obj、final_archive_dec、runtime、feasible_rate、constraint_metrics 與 archive_metrics 是否存在；缺漏 run 不參與 label selection，並在缺漏表中列示。
- Tie handling：rank-based 指標使用 average rank 處理 ties；summary ranking 使用 dense rank 顯示並列名次。

## 4.4 投稿用結論文字

Overall, Experiments A-C establish the proposed automated design pipeline on controlled benchmark and synthetic portfolio instances, while the additional real-market rolling-window validation serves as an external robustness check. The real-market results show that ECMADE-MOO does not dominate return-oriented baselines in after-cost annualized return, but it achieves the lowest CVaR(95%) loss among the compared methods. This indicates that the proposed design is more closely associated with downside-risk control than with unconditional return maximization under real-market data. Therefore, the revised manuscript reports real-market validation as supporting evidence for robustness and limitation analysis, while retaining Experiments B and C as the primary evidence for instance-aware and stability-aware automated configuration.
"""


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    overall, cost, windows = load_tables()
    build_docx(overall, cost, windows)
    MD_PATH.write_text(make_markdown(overall, cost, windows), encoding="utf-8")
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()

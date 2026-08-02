from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SUMMARY_DIR = ROOT / "p0_lite_outputs" / "experiment_b_configuration_summary_20260713"
OUT_DIR = ROOT / "docx_outputs"
OUT = OUT_DIR / "Experiment_B_configuration_strategy_comparison_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(92, 105, 120)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


METHOD_NAMES = {
    "MetaDesigned_ECMADE_MOO": "Meta-designed ECMADE-MOO",
    "BayesianConfig_ECMADE_MOO": "Bayesian configuration ECMADE-MOO",
    "RandomConfig_ECMADE_MOO": "Random configuration ECMADE-MOO",
    "HandCrafted_ECMADE_MOO": "Hand-crafted ECMADE-MOO",
}


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_para(doc, text="", style=None, size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=4)


def fill_cell(cell, text, bold=False, color=INK, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, numeric_cols=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    numeric_cols = set(numeric_cols or [])
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_GRAY)
        fill_cell(cell, h, bold=True, color=DARK_BLUE, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            fill_cell(
                cells[idx],
                str(value),
                size=9.3,
                align=WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )
    set_table_geometry(table, widths_dxa)
    add_para(doc, "", after=4)
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    left = p.add_run("Experiment B")
    set_run_font(left, size=9, color=MUTED, bold=True)
    p.add_run("\t")
    right = p.add_run("Configuration Strategy Comparison")
    set_run_font(right, size=9, color=MUTED)
    paragraph_border_bottom(p, color="D7DBE2", size="4", space="4")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("ECMADE-MOO 實驗報告")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc):
    add_para(doc, "實驗 B 報告", size=10, color=MUTED, bold=True, after=2)
    title = add_para(
        doc,
        "ECMADE-MOO Configuration Strategy 比較內容說明",
        size=22,
        color=INK,
        bold=True,
        after=6,
    )
    subtitle = add_para(
        doc,
        "比較 hand-crafted、random、Bayesian 與 meta-designed 四種參數組合選擇策略在 unseen instances 上的品質、穩定性、多樣性與執行成本。",
        size=11.5,
        color=MUTED,
        after=12,
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="2E74B5", size="10", space="3")

    meta = [
        ("主題", "Experiment B: configuration strategy comparison"),
        ("核心問題", "不同 theta configuration selection strategy 是否會影響 ECMADE-MOO 的泛化表現"),
        ("測試設定", "32 個 unseen test instances；每個 instance 30 次 independent runs；N=100；maxFE=10000"),
        ("整體指標", "RankScore，整合 HV、IGD、PF overlap、PF drift、diversity 與 runtime"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1900, 7460])
    for idx, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(idx, 0), LIGHT_GRAY)
        fill_cell(table.cell(idx, 0), label, bold=True, color=DARK_BLUE, size=9.7)
        fill_cell(table.cell(idx, 1), value, size=9.7)
    add_para(doc, "", after=4)


def fmt(x, digits=4):
    return f"{float(x):.{digits}f}"


def main():
    OUT_DIR.mkdir(exist_ok=True)
    overall = pd.read_csv(SUMMARY_DIR / "overall_configuration_comparison.csv")
    stats = pd.read_csv(SUMMARY_DIR / "statistical_tests_meta_vs_baselines.csv")
    friedman = pd.read_csv(SUMMARY_DIR / "friedman_tests_all_methods.csv")

    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    add_title_block(doc)

    add_callout(
        doc,
        "一句話結論",
        "實驗 B 不是在比較四個完全不同的演算法，而是在同一個 ECMADE-MOO 架構下，比較不同 configuration selection strategy 如何選擇 theta 參數組合，以及這些選擇是否能改善 unseen instances 的整體表現。",
    )

    add_heading(doc, "1. 實驗 B 要比較的內容", level=1)
    add_para(
        doc,
        "本實驗的比較焦點是 configuration selection strategy。四種方法使用相同的 ECMADE-MOO 主體與相同測試設定，但 theta configuration 的來源或選擇方式不同。因此，實驗 B 主要回答：當測試問題具有異質性時，固定參數、隨機參數、自動調參與 meta-designed 選參策略，哪一種能取得較好的 Pareto-front 品質與穩定性。",
    )

    strategy_rows = [
        ["Hand-crafted ECMADE-MOO", "人工設定固定 theta configuration", "作為傳統 baseline，檢查固定手調參數是否能泛化到不同 unseen instances。"],
        ["Random configuration ECMADE-MOO", "從 L24 theta set 隨機選擇 configuration", "檢查隨機選參數是否可能靠運氣達到可接受表現，並作為最低限度的自動選擇 baseline。"],
        ["Bayesian configuration ECMADE-MOO", "使用 Bayesian tuning 選擇 configuration", "作為較強的自動調參 baseline，檢查 tuning-based strategy 的上限與限制。"],
        ["Meta-designed ECMADE-MOO", "根據 instance features 選擇 theta configuration", "檢查 meta-learning/instance-aware selection 是否能針對不同問題選出更合適的參數組合。"],
    ]
    add_table(
        doc,
        ["方法", "theta 選擇方式", "比較目的"],
        strategy_rows,
        [2500, 2800, 4060],
    )

    add_heading(doc, "2. 比較指標與判讀方向", level=1)
    add_para(
        doc,
        "為避免只看單一指標造成誤判，實驗 B 同時比較解的品質、Pareto front 穩定性、分布多樣性與執行時間。最後再以 RankScore 統整多個指標，作為整體排名依據。",
    )
    metric_rows = [
        ["HV", "越大越好", "衡量解集覆蓋目標空間的能力；越大通常代表 Pareto front 品質較佳。"],
        ["IGD", "越小越好", "衡量解集與 reference front 的距離；越小表示越接近理想前緣。"],
        ["PF overlap", "越大越好", "衡量所得 Pareto front 與 common reference front 的重疊程度。"],
        ["PF drift", "越小越好", "衡量 front 偏離 reference front 的程度；越小表示穩定性較好。"],
        ["Diversity", "越大越好", "衡量解在 front 上的分布是否夠廣、夠均勻。"],
        ["Runtime", "越小越好", "衡量執行成本；用來確認品質提升是否伴隨過高時間代價。"],
        ["RankScore", "越小越好", "綜合多個指標後的整體排名分數，是本實驗最重要的總結性指標。"],
    ]
    add_table(doc, ["指標", "方向", "代表意義"], metric_rows, [1500, 1500, 6360])

    add_heading(doc, "3. 整體比較結果", level=1)
    add_para(
        doc,
        "從 overall comparison 來看，Meta-designed ECMADE-MOO 在整體 RankScore 上排名第一，且 HV、IGD、PF overlap 與 PF drift 皆取得四種策略中最佳的 overall rank。這表示它不只是某一個指標較好，而是在品質與穩定性上呈現較一致的優勢。",
    )
    order = [
        "MetaDesigned_ECMADE_MOO",
        "BayesianConfig_ECMADE_MOO",
        "RandomConfig_ECMADE_MOO",
        "HandCrafted_ECMADE_MOO",
    ]
    overall = overall.set_index("method").loc[order].reset_index()
    overall_rows = []
    for _, row in overall.iterrows():
        overall_rows.append([
            METHOD_NAMES[row["method"]],
            fmt(row["mean_HV"], 4),
            fmt(row["mean_IGD"], 4),
            fmt(row["mean_PF_Overlap"], 4),
            fmt(row["mean_PF_Drift"], 4),
            fmt(row["mean_Runtime"], 3),
            fmt(row["overall_RankScore"], 3),
            f"{int(row['first_place_instances'])}/32",
        ])
    add_table(
        doc,
        ["方法", "HV", "IGD", "PF overlap", "PF drift", "Runtime", "Overall RankScore", "First-place"],
        overall_rows,
        [2450, 850, 850, 1150, 1050, 950, 1250, 860],
        numeric_cols={1, 2, 3, 4, 5, 6, 7},
    )

    fig_path = SUMMARY_DIR / "figures" / "overall_rank_score.png"
    if fig_path.exists():
        add_para(doc, "圖 1. 四種 configuration strategy 的 overall RankScore，比較方向為越低越好。", size=9.5, color=MUTED, italic=True, after=4)
        doc.add_picture(str(fig_path), width=Inches(5.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_para(doc, "", after=4)

    add_heading(doc, "4. 統計檢定重點", level=1)
    rank_friedman = friedman[friedman["metric"] == "RankScore"].iloc[0]
    add_para(
        doc,
        f"Friedman test 顯示四種 strategy 在 RankScore 上存在顯著差異：chi-square={rank_friedman['friedman_chi_square']:.3f}, p={rank_friedman['p_value']:.4f}。進一步以 one-sided Wilcoxon signed-rank test 比較 Meta-designed 與各 baseline，並使用 Holm correction 校正多重比較。",
    )
    stat_rows = []
    for baseline in [
        "HandCrafted_ECMADE_MOO",
        "RandomConfig_ECMADE_MOO",
        "BayesianConfig_ECMADE_MOO",
    ]:
        row = stats[(stats["metric"] == "RankScore") & (stats["baseline"] == baseline)].iloc[0]
        sig = "顯著" if bool(row["significant_0_05"]) else "未顯著"
        stat_rows.append([
            f"Meta-designed vs {METHOD_NAMES[baseline]}",
            f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
            fmt(row["median_improvement"], 3),
            f"{row['holm_p_value']:.4f}",
            sig,
        ])
    add_table(
        doc,
        ["比較組合", "Win/Tie/Loss", "RankScore median improvement", "Holm p-value", "結論"],
        stat_rows,
        [3500, 1450, 2200, 1250, 960],
        numeric_cols={1, 2, 3, 4},
    )

    add_heading(doc, "5. 可用於論文的解讀", level=1)
    add_para(
        doc,
        "實驗 B 的結果支持 instance-aware configuration selection 的必要性。Hand-crafted strategy 雖然簡單且容易重現，但固定參數難以同時適應不同 portfolio instances。Random strategy 顯示 theta set 中確實存在可用 configuration，但缺乏穩定的選擇機制。Bayesian strategy 可作為強 baseline，然而它偏向尋找整體較佳的 tuning choice，未必能針對每個 unseen instance 做出最合適的選擇。相較之下，Meta-designed strategy 將 instance features 納入 configuration selection，使 ECMADE-MOO 能在異質測試問題上取得更好的整體 RankScore。",
    )
    add_callout(
        doc,
        "建議寫法",
        "Experiment B compares four ECMADE-MOO configuration strategies under identical evaluation settings. The results show that the meta-designed strategy achieves the best overall RankScore and significantly outperforms hand-crafted, random, and Bayesian configuration baselines, indicating that instance-aware theta selection improves the generalization ability of ECMADE-MOO on unseen portfolio optimization instances.",
    )

    add_heading(doc, "6. 小結", level=1)
    summary_rows = [
        ["比較對象", "四種 ECMADE-MOO configuration strategy，而非四個完全不同的演算法核心。"],
        ["比較重點", "theta configuration 的選擇方式是否能改善 unseen instances 上的 Pareto-front 品質與穩定性。"],
        ["主要證據", "Meta-designed ECMADE-MOO 的 overall RankScore 最低，並在 RankScore paired tests 中顯著優於三個 baseline。"],
        ["論文意義", "證明 ECMADE-MOO 的參數設計不應只依賴固定手調，而應考慮 instance-aware configuration selection。"],
    ]
    add_table(doc, ["項目", "說明"], summary_rows, [1900, 7460])

    doc.save(OUT)
    print(f"OUT={OUT}")


if __name__ == "__main__":
    main()

from __future__ import annotations

import math
from datetime import date
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from scipy.stats import wilcoxon


ROOT = Path(__file__).resolve().parent
DOCX_OUT = ROOT / "docx_outputs"
REPORT_STEM = "Experiment_C_stability_aware_selector_report_20260717"

TRAIN_REPORT = ROOT / (
    "p0_lite_outputs/theta24_70_15_15_training_label_full_20260706/"
    "knowledge_base_parameter_report"
)
VAL_REPORT = ROOT / (
    "p0_lite_outputs/theta24_70_15_15_validation_label_full_20260713/"
    "knowledge_base_parameter_report"
)
SELECTOR_DIR = ROOT / "p0_lite_outputs/experiment_c_stability_selector_training"
C_FINAL_DIR = ROOT / "p0_lite_outputs/experiment_c_stability_ecmade_moo_20260717"
COMPARISON_DIR = ROOT / "p0_lite_outputs/experiment_c_stability_comparison_20260717"

C_METHOD = "ExperimentC_StabilityAware_ECMADE_MOO"
PRIMARY_BASELINE = "MetaDesigned_ECMADE_MOO"
PRIMARY_ENDPOINT = "RankScore"
BASELINES = [
    "MetaDesigned_ECMADE_MOO",
    "BayesianConfig_ECMADE_MOO",
    "RandomConfig_ECMADE_MOO",
    "HandCrafted_ECMADE_MOO",
]
METRICS = [
    ("HV", "max"),
    ("IGD", "min"),
    ("PF_Overlap", "max"),
    ("PF_Drift", "min"),
    ("Runtime", "min"),
]
CONFIRMATORY_METRICS = [(PRIMARY_ENDPOINT, "min"), *METRICS]
ALPHA = 0.05


def fmt(value: object, digits: int = 4) -> str:
    if value is None:
        return ""
    try:
        f = float(value)
    except Exception:
        return str(value)
    if math.isnan(f):
        return "-"
    return f"{f:.{digits}f}"


def fmt_pct(value: object) -> str:
    try:
        f = float(value)
    except Exception:
        return "-"
    if math.isnan(f):
        return "-"
    return f"{f * 100:.1f}%"


def exact_sign_test_p(wins: int, losses: int) -> float:
    n = wins + losses
    if n == 0:
        return math.nan
    k = min(wins, losses)
    tail = sum(math.comb(n, i) for i in range(k + 1)) / (2**n)
    return min(1.0, 2.0 * tail)


def holm_adjust(p_values: list[float]) -> list[float]:
    if not p_values:
        return []
    values = np.asarray(p_values, dtype=float)
    adjusted = np.full(len(values), np.nan, dtype=float)
    finite_idx = np.flatnonzero(np.isfinite(values))
    if len(finite_idx) == 0:
        return adjusted.tolist()
    ordered = finite_idx[np.argsort(values[finite_idx])]
    running = 0.0
    m = len(ordered)
    for rank, idx in enumerate(ordered):
        candidate = min(1.0, (m - rank) * float(values[idx]))
        running = max(running, candidate)
        adjusted[idx] = running
    return adjusted.tolist()


def paired_directional_diff(c: pd.Series, baseline: pd.Series, direction: str) -> np.ndarray:
    joined = pd.concat([c, baseline], axis=1, keys=["c", "baseline"]).dropna()
    if direction == "max":
        diff = joined["c"] - joined["baseline"]
    else:
        diff = joined["baseline"] - joined["c"]
    return diff.to_numpy(dtype=float)


def safe_one_sided_wilcoxon(diff: np.ndarray) -> tuple[float, float]:
    diff = diff[np.isfinite(diff)]
    nonzero = diff[np.abs(diff) > 1e-12]
    if len(nonzero) == 0:
        return 0.0, 1.0
    stat, p_value = wilcoxon(nonzero, alternative="greater", zero_method="wilcox")
    return float(stat), float(p_value)


def load_data() -> dict[str, pd.DataFrame]:
    return {
        "train_labels": pd.read_csv(TRAIN_REPORT / "experiment_c_stability_regression_labels.csv"),
        "val_labels": pd.read_csv(VAL_REPORT / "experiment_c_stability_regression_labels.csv"),
        "val_run_metrics": pd.read_csv(VAL_REPORT / "run_metrics.csv"),
        "selector_summary": pd.read_csv(SELECTOR_DIR / "validation_selector_summary.csv"),
        "assignment": pd.read_csv(SELECTOR_DIR / "experiment_c_stability_theta_assignment.csv"),
        "overall": pd.read_csv(COMPARISON_DIR / "overall_configuration_comparison.csv"),
        "ranked": pd.read_csv(COMPARISON_DIR / "combined_instance_method_metrics_ranked.csv"),
        "pairwise": pd.read_csv(COMPARISON_DIR / "pairwise_win_tie_loss_by_metric.csv"),
    }


def build_sign_tests(ranked: pd.DataFrame) -> pd.DataFrame:
    rows = []
    index_cols = ["split", "instance", "K"]
    for baseline in BASELINES:
        for metric, direction in METRICS:
            pivot = ranked.pivot_table(index=index_cols, columns="method", values=metric, aggfunc="mean")
            if C_METHOD not in pivot or baseline not in pivot:
                continue
            c = pivot[C_METHOD]
            b = pivot[baseline]
            if direction == "max":
                wins = int((c > b).sum())
                losses = int((c < b).sum())
                mean_diff = float((c - b).mean())
                median_diff = float((c - b).median())
            else:
                wins = int((c < b).sum())
                losses = int((c > b).sum())
                mean_diff = float((b - c).mean())
                median_diff = float((b - c).median())
            ties = int((c == b).sum())
            rows.append(
                {
                    "baseline": baseline,
                    "metric": metric,
                    "direction": direction,
                    "wins": wins,
                    "ties": ties,
                    "losses": losses,
                    "non_tied": wins + losses,
                    "sign_test_p_two_sided": exact_sign_test_p(wins, losses),
                    "mean_improvement_directional": mean_diff,
                    "median_improvement_directional": median_diff,
                }
            )
    out = pd.DataFrame(rows)
    out.to_csv(COMPARISON_DIR / "experiment_c_pairwise_sign_tests.csv", index=False, encoding="utf-8-sig")
    return out


def build_confirmatory_tests(ranked: pd.DataFrame) -> pd.DataFrame:
    rows = []
    index_cols = ["split", "instance", "K"]
    for metric, direction in CONFIRMATORY_METRICS:
        pivot = ranked.pivot_table(index=index_cols, columns="method", values=metric, aggfunc="mean")
        if C_METHOD not in pivot or PRIMARY_BASELINE not in pivot:
            continue
        diff = paired_directional_diff(pivot[C_METHOD], pivot[PRIMARY_BASELINE], direction)
        stat, p_value = safe_one_sided_wilcoxon(diff)
        rows.append(
            {
                "endpoint_role": "primary" if metric == PRIMARY_ENDPOINT else "secondary_holm_family",
                "metric": metric,
                "direction": direction,
                "paired_unit": "test_instance",
                "comparison": f"{C_METHOD} vs {PRIMARY_BASELINE}",
                "primary": C_METHOD,
                "baseline": PRIMARY_BASELINE,
                "alternative": "Experiment C better than Experiment B Meta-designed",
                "n_instances": int(len(diff)),
                "wins": int((diff > 1e-12).sum()),
                "ties": int((np.abs(diff) <= 1e-12).sum()),
                "losses": int((diff < -1e-12).sum()),
                "median_directional_improvement": float(np.nanmedian(diff)),
                "mean_directional_improvement": float(np.nanmean(diff)),
                "wilcoxon_stat": stat,
                "raw_p_value_one_sided": p_value,
                "alpha": ALPHA,
            }
        )
    out = pd.DataFrame(rows)
    if not out.empty:
        out["holm_p_value"] = holm_adjust(out["raw_p_value_one_sided"].tolist())
        out["significant_raw_0_05"] = out["raw_p_value_one_sided"] < ALPHA
        out["significant_after_holm_0_05"] = out["holm_p_value"] < ALPHA
    out.to_csv(COMPARISON_DIR / "experiment_c_vs_b_paired_wilcoxon_holm.csv", index=False, encoding="utf-8-sig")
    return out


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths: list[int] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]], widths: list[int] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = str(value)
    style_table(table, widths)
    return table


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_table(doc, ["項目", "內容"], rows, widths=[2700, 6660])


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Experiment C stability-aware selector report")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Experiment C Stability-aware Theta Selector Report")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Validation label post-processing, stability-aware selector training, and final test comparison")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)

    meta = [
        ("生成日期", date.today().isoformat()),
        ("資料根目錄", str(ROOT)),
        ("核心方法", C_METHOD),
        ("比較基準", "Hand-crafted, Random, Bayesian, Meta-designed"),
    ]
    add_kv_table(doc, meta)


def data_completeness_rows(data: dict[str, pd.DataFrame]) -> list[tuple[str, str, str, str]]:
    train = data["train_labels"]
    val = data["val_labels"]
    runs = data["val_run_metrics"]
    c_final_runs = pd.read_csv(C_FINAL_DIR / "knowledge_base_parameter_report" / "run_metrics.csv")
    validation_counts = runs.groupby(["split", "instance", "K", "method"])["run"].nunique()
    c_counts = c_final_runs.groupby(["split", "instance", "K", "method"])["run"].nunique()
    return [
        (
            "Training labels",
            str(train[["split", "instance", "K"]].drop_duplicates().shape[0]),
            str(train["method"].nunique()),
            f"{len(train)} label rows",
        ),
        (
            "Validation labels",
            str(val[["split", "instance", "K"]].drop_duplicates().shape[0]),
            str(val["method"].nunique()),
            f"{len(val)} label rows",
        ),
        (
            "Validation raw runs",
            str(validation_counts.shape[0]),
            "24 theta",
            f"min={validation_counts.min()}, max={validation_counts.max()}, bad={(validation_counts != 30).sum()}",
        ),
        (
            "Experiment C final test",
            str(c_counts.shape[0]),
            "1 selected theta per instance",
            f"{len(c_final_runs)} runs; min={c_counts.min()}, max={c_counts.max()}, bad={(c_counts != 30).sum()}",
        ),
    ]


def add_report_content(
    doc: Document,
    data: dict[str, pd.DataFrame],
    sign_tests: pd.DataFrame,
    confirmatory_tests: pd.DataFrame,
) -> None:
    overall = data["overall"]
    selector_summary = data["selector_summary"]
    assignment = data["assignment"]
    ranked = data["ranked"]

    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Experiment C introduces a stability-aware label objective for theta selection. "
        "The final selector was trained on Training labels, checked on Validation labels, "
        "and then evaluated on 32 unseen test instances with 30 independent runs per instance."
    )
    c = overall[overall["method"] == C_METHOD].iloc[0]
    meta = overall[overall["method"] == "MetaDesigned_ECMADE_MOO"].iloc[0]
    add_bullets(
        doc,
        [
            f"Experiment C ranks first in the final common-reference comparison: overall_RankScore={fmt(c['overall_RankScore'], 4)}.",
            f"Experiment C obtains {int(c['first_place_instances'])} first-place test instances out of {int(c['instances'])}, compared with {int(meta['first_place_instances'])} for Experiment B Meta-designed.",
            "The pre-specified confirmatory endpoint is per-instance RankScore, the paired test-unit analogue of OverallRankScore; lower values are better.",
            f"C improves stability-oriented metrics versus all-theta mean on validation: PF_Overlap increases and PF_Drift decreases.",
            "No additional raw experiment runs are required for the main C result; the remaining work is paper-ready reporting and optional figure generation.",
        ],
    )

    add_heading(doc, "2. Data Completeness and Artifacts")
    add_table(
        doc,
        ["資料項目", "instance groups", "theta/method count", "完整性摘要"],
        data_completeness_rows(data),
        widths=[2500, 1700, 2100, 3060],
    )
    doc.add_paragraph(
        "The validation label generation is now complete. The earlier theta_089 gap is closed: every theta-by-validation-instance group contains 30 runs."
    )

    add_heading(doc, "3. Stability-aware Label Definition")
    doc.add_paragraph(
        "Experiment C replaces the original averaged label objective with a stability-aware score. "
        "For each instance and K, theta candidates are ranked by HV, IGD, PF_Overlap, and PF_Drift. "
        "The C score is maximized:"
    )
    p = doc.add_paragraph()
    r = p.add_run("C_LabelScore = -0.2*rank_HV - 0.2*rank_IGD - 0.3*rank_PF_Overlap - 0.3*rank_PF_Drift")
    r.bold = True
    doc.add_paragraph(
        "All four ranks use the same convention: rank 1 is best and smaller ranks are better. The negative coefficients therefore reward better HV, IGD, PF overlap, and PF drift ranks when C_LabelScore is maximized. C_ThetaRank is then assigned by descending C_LabelScore."
    )

    add_heading(doc, "4. Selector Training and Validation")
    doc.add_paragraph(
        "The selector is a Random Forest regression model using instance meta-features and theta encodings. "
        "It is trained to predict C_LabelScore and selects the theta with the highest predicted C_LabelScore."
    )
    val_rows = []
    for _, row in selector_summary.iterrows():
        val_rows.append(
            [
                row["selector"],
                int(row["groups"]),
                fmt_pct(row["top1_hit_rate"]),
                fmt_pct(row["top3_hit_rate"]),
                fmt(row["mean_C_ThetaRank"], 2),
                fmt(row["mean_C_regret"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_PF_Drift"], 4),
            ]
        )
    add_table(
        doc,
        ["selector", "groups", "top1", "top3", "mean C rank", "C regret", "PF overlap", "PF drift"],
        val_rows,
        widths=[2600, 900, 900, 900, 1200, 1200, 1300, 1360],
    )
    doc.add_paragraph(
        "Validation indicates that the learned selector does not perfectly recover the C-oracle top theta, but it moves the selected theta set toward the stability objective relative to the average theta pool."
    )

    add_heading(doc, "5. Final Test Comparison")
    doc.add_paragraph(
        "All final-test methods are evaluated under a common-reference post-processing pipeline. "
        "The comparison includes Experiment C and the four Experiment B baselines."
    )
    overall_cols = [
        "method",
        "mean_HV",
        "mean_IGD",
        "mean_PF_Overlap",
        "mean_PF_Drift",
        "mean_Runtime",
        "overall_RankScore",
        "first_place_instances",
    ]
    rows = []
    for _, row in overall[overall_cols].iterrows():
        rows.append(
            [
                row["method"],
                fmt(row["mean_HV"], 4),
                fmt(row["mean_IGD"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_PF_Drift"], 4),
                fmt(row["mean_Runtime"], 3),
                fmt(row["overall_RankScore"], 4),
                int(row["first_place_instances"]),
            ]
        )
    add_table(
        doc,
        ["method", "HV", "IGD", "PF overlap", "PF drift", "runtime", "overall rank", "1st-place inst."],
        rows,
        widths=[3100, 850, 850, 1100, 1000, 900, 1100, 1460],
    )

    add_heading(doc, "6. Confirmatory C vs B Paired Test")
    doc.add_paragraph(
        "The primary endpoint is pre-specified as per-instance RankScore, which is the test-instance-level analogue of OverallRankScore. "
        "Experiment C is compared against Experiment B Meta-designed with a one-sided paired Wilcoxon signed-rank test; positive improvement means C is better."
    )
    primary = confirmatory_tests[confirmatory_tests["endpoint_role"] == "primary"]
    primary_rows = []
    for _, row in primary.iterrows():
        primary_rows.append(
            [
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["median_directional_improvement"], 5),
                fmt(row["raw_p_value_one_sided"], 6),
                fmt(row["holm_p_value"], 6),
                "yes" if row["significant_raw_0_05"] else "no",
            ]
        )
    add_table(
        doc,
        ["primary endpoint", "W/T/L", "median improvement", "one-sided p", "Holm p", "raw p < .05"],
        primary_rows,
        widths=[1900, 1050, 1700, 1400, 1200, 2110],
    )
    doc.add_paragraph(
        "Secondary endpoints are reported as a Holm-corrected family to control the family-wise error rate across the C vs B endpoint checks."
    )
    secondary_rows = []
    for _, row in confirmatory_tests[confirmatory_tests["endpoint_role"] != "primary"].iterrows():
        secondary_rows.append(
            [
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["median_directional_improvement"], 5),
                fmt(row["raw_p_value_one_sided"], 6),
                fmt(row["holm_p_value"], 6),
                "yes" if row["significant_after_holm_0_05"] else "no",
            ]
        )
    add_table(
        doc,
        ["secondary endpoint", "W/T/L", "median improvement", "one-sided p", "Holm p", "Holm < .05"],
        secondary_rows,
        widths=[1900, 1050, 1700, 1400, 1200, 2110],
    )

    add_heading(doc, "7. Exploratory Sign-test Evidence")
    doc.add_paragraph(
        "The table below reports per-instance paired comparisons between Experiment C and each baseline. "
        "Wins are counted in the metric's favorable direction. These exact two-sided sign-test p-values are uncorrected and should be read as exploratory evidence only."
    )
    pair_rows = []
    selected_tests = sign_tests[
        sign_tests["metric"].isin(["HV", "IGD", "PF_Overlap", "PF_Drift"])
    ].copy()
    for _, row in selected_tests.iterrows():
        pair_rows.append(
            [
                row["baseline"].replace("_ECMADE_MOO", ""),
                row["metric"],
                f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
                fmt(row["sign_test_p_two_sided"], 4),
                fmt(row["mean_improvement_directional"], 5),
            ]
        )
    add_table(
        doc,
        ["baseline", "metric", "W/T/L", "sign-test p", "mean directional gain"],
        pair_rows,
        widths=[2800, 1400, 1300, 1400, 2460],
    )

    add_heading(doc, "8. Theta Usage")
    usage = assignment["theta_id"].value_counts().rename_axis("theta_id").reset_index(name="test_instances")
    usage_rows = [[row["theta_id"], int(row["test_instances"])] for _, row in usage.iterrows()]
    add_table(doc, ["theta_id", "selected test instances"], usage_rows, widths=[2500, 6860])
    doc.add_paragraph(
        "The C selector is not a single-theta policy. It selects multiple theta configurations across the unseen test set, which supports the meta-learning setup rather than reducing the experiment to one global parameter choice."
    )

    add_heading(doc, "9. Interpretation")
    add_bullets(
        doc,
        [
            "The strongest C result is not raw HV dominance alone; it is the combined improvement in overlap, IGD, and per-instance first-place frequency under the common-reference ranking.",
            "Compared with Experiment B Meta-designed, C trades a very small HV difference for better IGD and PF_Overlap, and wins far more first-place instances.",
            "PF_Drift is nearly tied with Meta-designed in the final comparison, which means the C objective improved stability enough to be competitive without visibly collapsing convergence quality.",
            "The validation top1 hit rate is modest, so the report should frame C as a stability-aware selection strategy with empirical final-test benefit, not as a perfect oracle imitator.",
        ],
    )

    add_heading(doc, "10. Recommended Paper-ready Claims")
    add_bullets(
        doc,
        [
            "A stability-aware label objective can be constructed from post-processed multi-run Pareto-front metrics without requiring extra raw optimization runs.",
            "On 32 unseen test instances, the stability-aware selector achieved the best overall common-reference rank among Hand-crafted, Random, Bayesian, Meta-designed, and Experiment C strategies; the confirmatory C vs B statement should be tied to the pre-specified per-instance RankScore paired Wilcoxon test.",
            "Experiment C produced the highest PF_Overlap and the lowest IGD among the compared methods, while maintaining PF_Drift comparable to the best baseline.",
            "The result supports using label engineering, not only model architecture, as an effective lever in meta-designed parameter selection.",
        ],
    )

    add_heading(doc, "11. Files Produced")
    add_kv_table(
        doc,
        [
            ("C labels", str(TRAIN_REPORT / "experiment_c_stability_regression_labels.csv")),
            ("Validation labels", str(VAL_REPORT / "experiment_c_stability_regression_labels.csv")),
            ("Selector output", str(SELECTOR_DIR)),
            ("Final test output", str(C_FINAL_DIR)),
            ("Comparison output", str(COMPARISON_DIR)),
            ("C vs B Wilcoxon/Holm table", str(COMPARISON_DIR / "experiment_c_vs_b_paired_wilcoxon_holm.csv")),
            ("Sign-test table", str(COMPARISON_DIR / "experiment_c_pairwise_sign_tests.csv")),
        ],
    )


def build_markdown(
    data: dict[str, pd.DataFrame],
    sign_tests: pd.DataFrame,
    confirmatory_tests: pd.DataFrame,
    path: Path,
) -> None:
    overall = data["overall"]
    selector_summary = data["selector_summary"]
    c = overall[overall["method"] == C_METHOD].iloc[0]
    def md_table(frame: pd.DataFrame) -> str:
        text = frame.copy()
        for col in text.columns:
            text[col] = text[col].map(lambda value: fmt(value, 4) if isinstance(value, (int, float)) else str(value))
        headers = [str(col) for col in text.columns]
        rows = text.astype(str).values.tolist()
        out = [
            "| " + " | ".join(headers) + " |",
            "| " + " | ".join(["---"] * len(headers)) + " |",
        ]
        out.extend("| " + " | ".join(row) + " |" for row in rows)
        return "\n".join(out)

    lines = [
        "# Experiment C Stability-aware Theta Selector Report",
        "",
        "## Summary",
        "",
        f"- Experiment C ranks first in final common-reference comparison: overall_RankScore={fmt(c['overall_RankScore'], 4)}.",
        f"- First-place test instances: {int(c['first_place_instances'])}/{int(c['instances'])}.",
        "- Primary endpoint: per-instance RankScore, the paired test-unit analogue of OverallRankScore; lower is better.",
        "- Main conclusion: Experiment C improves stability-aware selection without requiring new raw label-generation runs.",
        "",
        "## Confirmatory C vs B Paired Test",
        "",
        "One-sided paired Wilcoxon signed-rank tests compare Experiment C against Experiment B Meta-designed. Secondary endpoints are Holm-corrected; uncorrected sign tests below are exploratory only.",
        "",
        md_table(confirmatory_tests),
        "",
        "## Stability-aware Label",
        "",
        "`C_LabelScore = -0.2*rank_HV - 0.2*rank_IGD - 0.3*rank_PF_Overlap - 0.3*rank_PF_Drift`",
        "",
        "## Validation Selector Summary",
        "",
        md_table(selector_summary),
        "",
        "## Final Comparison",
        "",
        md_table(overall[
            [
                "method",
                "instances",
                "mean_HV",
                "mean_IGD",
                "mean_PF_Overlap",
                "mean_PF_Drift",
                "overall_RankScore",
                "first_place_instances",
            ]
        ]),
        "",
        "## Exploratory Pairwise Sign Tests",
        "",
        md_table(sign_tests),
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    DOCX_OUT.mkdir(exist_ok=True)
    data = load_data()
    sign_tests = build_sign_tests(data["ranked"])
    confirmatory_tests = build_confirmatory_tests(data["ranked"])

    md_path = DOCX_OUT / f"{REPORT_STEM}.md"
    docx_path = DOCX_OUT / f"{REPORT_STEM}.docx"

    build_markdown(data, sign_tests, confirmatory_tests, md_path)

    doc = Document()
    setup_styles(doc)
    add_title(doc)
    add_report_content(doc, data, sign_tests, confirmatory_tests)
    doc.save(docx_path)
    print(f"MARKDOWN={md_path}")
    print(f"DOCX={docx_path}")


if __name__ == "__main__":
    main()

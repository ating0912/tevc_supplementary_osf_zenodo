from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\yiting\Documents\Playground")
REPORT_DIR = ROOT / "p0_lite_outputs" / "synthetic_constrained_portfolio" / "experiment_A_report_20260703_232819"
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "Experiment_A_數據報告_含A_MPMO_20260703_232819.docx"


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths: list[float]) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                for run in p.runs:
                    run.font.size = Pt(9)
    for cell in table.rows[0].cells:
        set_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph()


def fmt(value: float, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def configure(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color in [("Heading 1", 15, "2E74B5"), ("Heading 2", 12.5, "2E74B5"), ("Heading 3", 11.5, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("實驗 A 數據報告")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    p.paragraph_format.space_after = Pt(2)
    sub = doc.add_paragraph()
    sub.add_run("數據版本：experiment_A_report_20260703_232819；方法包含 NSGAII、SPEA2、MOEAD、GDE3、ECMADE_MOO、A_MPMO。").italic = True


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    overall = pd.read_csv(REPORT_DIR / "overall_method_summary.csv")
    inst = pd.read_csv(REPORT_DIR / "instance_method_metrics.csv")
    tests = pd.read_csv(REPORT_DIR / "statistical_tests.csv")
    avail = pd.read_csv(REPORT_DIR / "metric_availability.csv")

    overall = overall.sort_values("RankScore").reset_index(drop=True)
    overall.insert(0, "OverallRank", range(1, len(overall) + 1))

    ec = overall[overall["method"] == "ECMADE_MOO"].iloc[0]
    methods = set(overall["method"])

    doc = Document()
    configure(doc)
    add_title(doc)

    doc.add_heading("1. 資料來源與完整性", level=1)
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["報告資料夾", str(REPORT_DIR)],
            ["方法", ", ".join(overall["method"].tolist())],
            ["Instance 數", str(inst["instance"].nunique())],
            ["Runs", "每個 method × instance 30 runs"],
            ["完整性", "missing_outputs.csv 無缺漏；Feasible Rate 全方法平均為 1.0"],
            ["A_MPMO 備註", "此數據版本已納入 A_MPMO，所有表格與排名均以 6 個方法共同計算。"],
        ],
        [1.55, 5.95],
    )

    doc.add_heading("2. Overall Method Summary", level=1)
    add_table(
        doc,
        ["Rank", "Method", "HV ↑", "IGD ↓", "PF Overlap ↑", "EAF Width ↓", "PF Drift ↓", "Diversity", "Runtime ↓", "RankScore ↓"],
        [
            [
                int(row["OverallRank"]),
                row["method"],
                fmt(row["mean_HV"], 4),
                fmt(row["mean_IGD"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_EAF_Band_Width"], 4),
                fmt(row["mean_PF_Drift"], 4),
                fmt(row["mean_Diversity"], 4),
                fmt(row["mean_Runtime"], 3),
                fmt(row["RankScore"], 3),
            ]
            for _, row in overall.iterrows()
        ],
        [0.42, 1.0, 0.62, 0.62, 0.82, 0.82, 0.78, 0.72, 0.72, 0.72],
    )

    doc.add_heading("3. ECMADE_MOO 重點數據", level=1)
    add_table(
        doc,
        ["指標", "ECMADE_MOO", "名次", "解讀"],
        [
            ["HV ↑", fmt(ec["mean_HV"], 6), int(ec["rank_HV"]), "解品質排名第 2，接近 SPEA2。"],
            ["IGD ↓", fmt(ec["mean_IGD"], 6), int(ec["rank_IGD"]), "與 reference front 距離排名第 2。"],
            ["PF Overlap ↑", fmt(ec["mean_PF_Overlap"], 6), int(ec["rank_PF_Overlap"]), "穩定覆蓋度排名第 2，高於 GDE3 / MOEAD。"],
            ["EAF Width ↓", fmt(ec["mean_EAF_Band_Width"], 6), int(ec["rank_EAF_Band_Width"]), "attainment band 較寬，重複執行不確定性仍偏高。"],
            ["PF Drift ↓", fmt(ec["mean_PF_Drift"], 6), int(ec["rank_PF_Drift"]), "PF centroid 漂移較大，是後續穩定化重點。"],
            ["Diversity", fmt(ec["mean_Diversity"], 6), "-", "Diversity 最高，顯示沒有 diversity collapse。"],
        ],
        [1.3, 1.0, 0.65, 4.55],
    )

    doc.add_heading("4. ECMADE_MOO vs Baselines", level=1)
    rows = []
    for baseline in ["SPEA2", "NSGAII", "GDE3", "MOEAD"]:
        if baseline not in methods:
            continue
        b = overall[overall["method"] == baseline].iloc[0]
        rows.append(
            [
                baseline,
                fmt(ec["mean_HV"] - b["mean_HV"], 4),
                fmt(ec["mean_IGD"] - b["mean_IGD"], 4),
                fmt(ec["mean_PF_Overlap"] - b["mean_PF_Overlap"], 4),
                fmt(ec["mean_EAF_Band_Width"] - b["mean_EAF_Band_Width"], 4),
                fmt(ec["mean_PF_Drift"] - b["mean_PF_Drift"], 4),
                fmt(ec["RankScore"] - b["RankScore"], 3),
            ]
        )
    add_table(doc, ["Baseline", "ΔHV", "ΔIGD", "ΔPF Overlap", "ΔEAF", "ΔPF Drift", "ΔRankScore"], rows, [1.0, 0.82, 0.82, 1.0, 0.82, 0.88, 0.9])

    doc.add_paragraph("註：HV / PF Overlap 的差值越大越好；IGD / EAF / PF Drift / RankScore 的差值越小越好。")

    doc.add_heading("5. 依 K/n 分組的 ECMADE_MOO 表現", level=1)
    by_k = (
        inst[inst["method"] == "ECMADE_MOO"]
        .groupby("k_ratio")
        .agg(
            HV=("HV", "mean"),
            IGD=("IGD", "mean"),
            PF_Overlap=("PF_Overlap", "mean"),
            EAF_Band_Width=("EAF_Band_Width", "mean"),
            PF_Drift=("PF_Drift", "mean"),
            Diversity=("Diversity", "mean"),
            Runtime=("Runtime", "mean"),
        )
        .reset_index()
    )
    add_table(
        doc,
        ["K/n", "HV", "IGD", "PF Overlap", "EAF", "PF Drift", "Diversity", "Runtime"],
        [[fmt(r["k_ratio"], 2), fmt(r["HV"], 4), fmt(r["IGD"], 4), fmt(r["PF_Overlap"], 4), fmt(r["EAF_Band_Width"], 4), fmt(r["PF_Drift"], 4), fmt(r["Diversity"], 4), fmt(r["Runtime"], 3)] for _, r in by_k.iterrows()],
        [0.6, 0.72, 0.72, 0.92, 0.72, 0.82, 0.82, 0.82],
    )

    doc.add_heading("6. 統計檢定摘要", level=1)
    selected_tests = tests[
        tests["metric"].isin(["HV", "IGD", "PF_Overlap", "EAF_Band_Width", "PF_Drift"])
        & tests["comparison"].isin(["all methods", "ECMADE_MOO vs GDE3", "ECMADE_MOO vs NSGAII", "ECMADE_MOO vs SPEA2"])
    ].copy()
    selected_tests["p_value_fmt"] = selected_tests["p_value"].map(lambda x: f"{float(x):.3e}")
    add_table(
        doc,
        ["Metric", "Test", "Comparison", "p-value"],
        selected_tests[["metric", "test", "comparison", "p_value_fmt"]].values.tolist()[:24],
        [1.05, 1.25, 3.0, 1.0],
    )

    doc.add_heading("7. 指標可用性", level=1)
    add_table(
        doc,
        ["類別", "項目", "狀態", "來源 / 定義"],
        avail.values.tolist(),
        [1.0, 1.8, 1.0, 3.7],
    )

    doc.add_heading("8. 圖表", level=1)
    fig_dir = REPORT_DIR / "figures"
    figures = [
        ("Figure 1", "Metric dashboard", fig_dir / "figure_1_metric_dashboard.png"),
        ("Figure 2", "PF overlay", fig_dir / "figure_2_pf_overlay.png"),
        ("Figure 3", "PF heatmap", fig_dir / "figure_3_pf_heatmap.png"),
        ("Figure 4", "EAF band width", fig_dir / "figure_4_eaf_band_width.png"),
        ("Figure 5", "Runtime", fig_dir / "figure_5_runtime.png"),
        ("Figure 6", "Stability-diversity", fig_dir / "figure_6_stability_diversity.png"),
    ]
    for label, caption, path in figures:
        if path.exists():
            p = doc.add_paragraph()
            p.add_run(f"{label}. {caption}").bold = True
            doc.add_picture(str(path), width=Inches(6.9))

    doc.add_heading("9. 結論", level=1)
    ampmo = overall[overall["method"] == "A_MPMO"].iloc[0] if "A_MPMO" in set(overall["method"]) else None
    ec_rank = int(overall[overall["method"] == "ECMADE_MOO"]["OverallRank"].iloc[0])
    conclusions = [
        f"ECMADE_MOO 在此 6 方法數據版本中整體排名第 {ec_rank}，主要優勢是 HV、IGD 與 diversity；其中 HV/IGD 排名第 2，diversity 最高。",
        "ECMADE_MOO 的弱點是 PF Overlap、EAF Band Width 與 PF Drift，代表 repeated-run stability 尚未優於 SPEA2 / NSGAII。",
    ]
    if ampmo is not None:
        ampmo_rank = int(overall[overall["method"] == "A_MPMO"]["OverallRank"].iloc[0])
        conclusions.append(
            f"A_MPMO 整體排名第 {ampmo_rank}，PF Overlap、EAF Band Width、PF Drift 均優於此版 ECMADE_MOO，但 HV/IGD 低於 ECMADE_MOO。"
        )
    conclusions.append(
        "若論文要強調多子群優勢，較合適的敘事是：ECMADE_MOO 保有較好的解品質與 diversity，A_MPMO 提供較好的多子群穩定性對照；後續需用 ablation 說明多子群、exchange 與 archive 設計各自的貢獻。"
    )
    for text in conclusions:
        p = doc.add_paragraph(style=None)
        p.add_run(text)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Experiment A data report")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUT_PATH)
    print(f"DOCX={OUT_PATH}")


if __name__ == "__main__":
    main()

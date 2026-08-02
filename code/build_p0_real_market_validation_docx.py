from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "p0_lite_outputs" / "p1_rolling_window_market_validation_20260719" / "p0_real_market_validation"
OUT_PATH = ROOT / "TEVC_P0_real_market_validation_交易成本敏感度_20260724.docx"


def fmt(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        if abs(value) >= 1:
            return f"{value:.3f}"
        return f"{value:.6f}"
    return str(value)


def set_font(run, size=10):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)


def qn(tag: str):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    set_font(run, 14 if level == 1 else 12)


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_font(run, 10)


def add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = col
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, 8)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = fmt(row[col])
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    set_font(run, 8)


def main() -> None:
    overall = pd.read_csv(DATA_DIR / "p0_real_market_overall_summary.csv", encoding="utf-8-sig")
    cost = pd.read_csv(DATA_DIR / "p0_real_market_transaction_cost_overall.csv", encoding="utf-8-sig")
    plot_path = DATA_DIR / "p0_real_market_wealth_curve_10bps.png"

    overall_cols = [
        "method",
        "windows",
        "mean_annual_net_return",
        "mean_sharpe",
        "mean_sortino",
        "mean_max_drawdown",
        "mean_cvar95_loss",
        "mean_rebalance_turnover",
        "stability_rank_score",
    ]
    cost_cols = [
        "cost_scenario",
        "method",
        "windows",
        "mean_annual_net_return",
        "mean_turnover",
        "mean_cvar95_loss",
        "rank_annual_net_return",
    ]

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    add_heading(doc, "TEVC P0 Real-Market Rolling-Window Validation 與交易成本敏感度", 1)
    add_paragraph(
        doc,
        "本補充實驗使用 SP100、NASDAQ100 與 TW50 三組真實市場 universe，每組 11 個 rolling windows，共 33 個 universe-window。"
        "每個 window 採用約 3 年 training window 與約 6 個月 out-of-sample testing window；每個方法每個 window 執行 10 runs。"
        "測試投資組合由 final Pareto front 中 training Sharpe 最高者選出，並回報 after-transaction-cost return、turnover、CVaR(95%)、maximum drawdown、Sharpe、Sortino、wealth curve 與 stability rank indicators。",
    )

    add_table(doc, overall[overall_cols], "Table P0-RM1. Real-market rolling-window validation")
    add_table(doc, cost[cost_cols], "Table P0-RM2. Transaction-cost sensitivity")

    add_heading(doc, "Figure P0-RM1. After-cost wealth curve, 10 bps", 2)
    if plot_path.exists():
        doc.add_picture(str(plot_path), width=Inches(9.2))

    add_heading(doc, "建議補稿文字", 2)
    add_paragraph(
        doc,
        "Across 33 real-market rolling windows, ECMADE-MOO did not dominate the return-oriented baselines in annualized after-cost return. "
        "However, it achieved the lowest mean CVaR(95%) loss among the six compared methods, indicating a more conservative downside-risk profile under real-market data. "
        "The transaction-cost sensitivity analysis further shows that the ranking of annualized net return is stable across 10, 20, and 50 bps settings. "
        "Therefore, this real-market experiment should be reported as external validation and limitation analysis rather than as evidence of return dominance.",
    )

    add_heading(doc, "輸出資料位置", 2)
    for name in [
        "p0_real_market_overall_summary.csv",
        "p0_real_market_transaction_cost_overall.csv",
        "p0_real_market_window_method_summary.csv",
        "p0_real_market_wealth_curve_mean.csv",
        "p0_real_market_wealth_curve_10bps.png",
        "P0_real_market_validation_paper_ready.md",
    ]:
        add_paragraph(doc, str(DATA_DIR / name))

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()

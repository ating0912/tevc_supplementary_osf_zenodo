from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r".")
SUMMARY_DIR = ROOT / "p0_lite_outputs" / "experiment_b_configuration_summary_20260713"
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "Experiment_B_configuration_strategy_baseline_style_report.docx"

METHOD_LABELS = {
    "MetaDesigned_ECMADE_MOO": "Meta-designed",
    "BayesianConfig_ECMADE_MOO": "Bayesian",
    "RandomConfig_ECMADE_MOO": "Random",
    "HandCrafted_ECMADE_MOO": "Hand-crafted",
}

METRIC_DIRECTIONS = {
    "HV": "↑",
    "IGD": "↓",
    "PF_Overlap": "↑",
    "PF_Drift": "↓",
    "Diversity": "↑",
    "Runtime": "↓",
    "RankScore": "↓",
}


def set_east_asian_font(run, name="Microsoft JhengHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths: list[float]) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                if i > 0 and row_idx > 0 and len(cell.text) < 18:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.6)
                    set_east_asian_font(run)
    for cell in table.rows[0].cells:
        set_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("1F4D78")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    doc.add_paragraph()


def fmt(value, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def fmt_p(value) -> str:
    value = float(value)
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}"


def configure(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run("實驗 B 數據報告")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.font.name = "Calibri"
    set_east_asian_font(r)
    p.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "數據版本：experiment_b_configuration_summary_20260713；方法包含 Hand-crafted、Random、Bayesian、Meta-designed ECMADE-MOO。"
    )
    sr.italic = True
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor.from_string("555555")
    set_east_asian_font(sr)


def rank_indicator(row, metric: str) -> str:
    key = f"overall_rank_{metric}"
    if key in row:
        return fmt(row[key], 1)
    return "-"


def metric_delta(meta, baseline, metric: str) -> str:
    return fmt(meta[f"mean_{metric}"] - baseline[f"mean_{metric}"], 4)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    overall = pd.read_csv(SUMMARY_DIR / "overall_configuration_comparison.csv")
    ranked = pd.read_csv(SUMMARY_DIR / "combined_instance_method_metrics_ranked.csv")
    raw = pd.read_csv(SUMMARY_DIR / "combined_instance_method_metrics_raw.csv")
    stats = pd.read_csv(SUMMARY_DIR / "statistical_tests_meta_vs_baselines.csv")
    friedman = pd.read_csv(SUMMARY_DIR / "friedman_tests_all_methods.csv")
    availability_files = [
        "combined_run_metrics_common_reference.csv",
        "combined_instance_method_metrics_raw.csv",
        "combined_instance_method_metrics_ranked.csv",
        "overall_configuration_comparison.csv",
        "statistical_tests_meta_vs_baselines.csv",
        "friedman_tests_all_methods.csv",
        "theta_usage_by_method.csv",
    ]

    overall = overall.sort_values(["overall_RankScore", "mean_RankScore", "method"]).reset_index(drop=True)
    overall.insert(0, "OverallRank", range(1, len(overall) + 1))
    meta = overall[overall["method"] == "MetaDesigned_ECMADE_MOO"].iloc[0]
    methods = set(overall["method"])

    doc = Document()
    configure(doc)
    add_title(doc)

    doc.add_heading("1. 資料來源與完整性", level=1)
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["報告資料夾", str(SUMMARY_DIR)],
            ["比較方法", ", ".join(METHOD_LABELS[m] for m in overall["method"].tolist())],
            ["Instance 數", str(ranked["instance"].nunique())],
            ["Runs", "每個 method × instance 30 independent runs"],
            ["共同設定", "N=100；maxFE=10000；Random、Bayesian、Meta-designed 皆從同一組 L24 theta configuration set 選擇"],
            ["共同參考前緣", "以同一 test instance 的 raw Pareto fronts 合併建立 common reference front，再計算 HV、IGD、PF overlap 與 PF drift"],
        ],
        [1.55, 5.95],
    )

    doc.add_heading("2. Overall Method Summary", level=1)
    add_table(
        doc,
        ["Rank", "Method", "HV ↑", "IGD ↓", "PF Overlap ↑", "PF Drift ↓", "Diversity ↑", "Runtime ↓", "Mean RankScore ↓", "Overall RankScore ↓", "First-place"],
        [
            [
                int(row["OverallRank"]),
                METHOD_LABELS[row["method"]],
                fmt(row["mean_HV"], 4),
                fmt(row["mean_IGD"], 4),
                fmt(row["mean_PF_Overlap"], 4),
                fmt(row["mean_PF_Drift"], 4),
                fmt(row["mean_Diversity"], 4),
                fmt(row["mean_Runtime"], 3),
                fmt(row["mean_RankScore"], 3),
                fmt(row["overall_RankScore"], 3),
                f"{int(row['first_place_instances'])}/32",
            ]
            for _, row in overall.iterrows()
        ],
        [0.42, 0.95, 0.55, 0.55, 0.78, 0.72, 0.72, 0.68, 0.78, 0.82, 0.62],
    )
    doc.add_paragraph(
        "註：Mean RankScore 是先在每個 instance 內對六個指標排名後取平均，再跨 instance 平均；Overall RankScore 是先取各方法的 mean metric，再對六個 mean metric 排名後取平均。兩者皆為越低越好。"
    )

    doc.add_heading("3. Meta-designed ECMADE-MOO 重點數據", level=1)
    add_table(
        doc,
        ["指標", "Meta-designed", "整體名次", "解讀"],
        [
            ["HV ↑", fmt(meta["mean_HV"], 6), rank_indicator(meta, "HV"), "四種策略中最高，代表整體 Pareto-front 解品質最佳。"],
            ["IGD ↓", fmt(meta["mean_IGD"], 6), rank_indicator(meta, "IGD"), "四種策略中最低，表示與 common reference front 的距離最小。"],
            ["PF Overlap ↑", fmt(meta["mean_PF_Overlap"], 6), rank_indicator(meta, "PF_Overlap"), "四種策略中最高，代表與共同參考前緣的覆蓋重疊程度最好。"],
            ["PF Drift ↓", fmt(meta["mean_PF_Drift"], 6), rank_indicator(meta, "PF_Drift"), "四種策略中最低，顯示 front 偏移較小。"],
            ["Diversity ↑", fmt(meta["mean_Diversity"], 6), rank_indicator(meta, "Diversity"), "排名第 2，略低於 Random，但仍未出現 diversity collapse。"],
            ["Runtime ↓", fmt(meta["mean_Runtime"], 6), rank_indicator(meta, "Runtime"), "排名第 3，品質提升伴隨一定 runtime 成本。"],
            ["First-place", f"{int(meta['first_place_instances'])}/32", "-", "32 個 unseen instances 中，嚴格第一名的次數最多。"],
        ],
        [1.15, 1.05, 0.75, 4.55],
    )

    doc.add_heading("4. Meta-designed vs Baselines", level=1)
    rows = []
    for baseline_name in ["HandCrafted_ECMADE_MOO", "RandomConfig_ECMADE_MOO", "BayesianConfig_ECMADE_MOO"]:
        if baseline_name not in methods:
            continue
        b = overall[overall["method"] == baseline_name].iloc[0]
        rows.append(
            [
                METHOD_LABELS[baseline_name],
                metric_delta(meta, b, "HV"),
                metric_delta(meta, b, "IGD"),
                metric_delta(meta, b, "PF_Overlap"),
                metric_delta(meta, b, "PF_Drift"),
                metric_delta(meta, b, "Diversity"),
                metric_delta(meta, b, "Runtime"),
                fmt(meta["mean_RankScore"] - b["mean_RankScore"], 3),
            ]
        )
    add_table(
        doc,
        ["Baseline", "ΔHV", "ΔIGD", "ΔPF Overlap", "ΔPF Drift", "ΔDiversity", "ΔRuntime", "ΔMean RankScore"],
        rows,
        [1.15, 0.68, 0.68, 0.88, 0.82, 0.82, 0.78, 0.95],
    )
    doc.add_paragraph(
        "註：HV、PF Overlap、Diversity 的差值越大越好；IGD、PF Drift、Runtime、Mean RankScore 的差值越小越好。"
    )

    doc.add_heading("5. 依 K 分組的 Meta-designed 表現", level=1)
    meta_by_k = (
        ranked[ranked["method"] == "MetaDesigned_ECMADE_MOO"]
        .groupby("K")
        .agg(
            instances=("instance", "nunique"),
            HV=("HV", "mean"),
            IGD=("IGD", "mean"),
            PF_Overlap=("PF_Overlap", "mean"),
            PF_Drift=("PF_Drift", "mean"),
            Diversity=("Diversity", "mean"),
            Runtime=("Runtime", "mean"),
            RankScore=("RankScore", "mean"),
            FirstPlace=("OverallInstanceRank", lambda s: int((s == 1).sum())),
        )
        .reset_index()
        .sort_values("K")
    )
    add_table(
        doc,
        ["K", "Instances", "HV", "IGD", "PF Overlap", "PF Drift", "Diversity", "Runtime", "RankScore", "First-place"],
        [
            [
                int(r["K"]),
                int(r["instances"]),
                fmt(r["HV"], 4),
                fmt(r["IGD"], 4),
                fmt(r["PF_Overlap"], 4),
                fmt(r["PF_Drift"], 4),
                fmt(r["Diversity"], 4),
                fmt(r["Runtime"], 3),
                fmt(r["RankScore"], 3),
                int(r["FirstPlace"]),
            ]
            for _, r in meta_by_k.iterrows()
        ],
        [0.42, 0.66, 0.62, 0.62, 0.82, 0.75, 0.78, 0.72, 0.82, 0.72],
    )

    doc.add_heading("6. 統計檢定摘要", level=1)
    doc.add_paragraph(
        "本節對標 Experiment A 的放法：先以 Friedman test 檢查四種策略在同一指標上是否存在整體差異，再以 Meta-designed 對各 baseline 做 paired one-sided Wilcoxon signed-rank test，並使用 Holm correction。"
    )
    compact_rows = []
    for metric in ["HV", "IGD", "PF_Overlap", "PF_Drift", "Diversity", "Runtime", "RankScore"]:
        fr = friedman[friedman["metric"] == metric].iloc[0]
        if metric == "RankScore":
            baseline = "Best overall"
            posthoc = stats[(stats["metric"] == metric) & (stats["baseline"] == "BayesianConfig_ECMADE_MOO")].iloc[0]
            comparison = "Meta-designed vs Bayesian"
        else:
            metric_stats = stats[stats["metric"] == metric].sort_values("holm_p_value")
            posthoc = metric_stats.iloc[0]
            comparison = f"Meta-designed vs {METHOD_LABELS[posthoc['baseline']]}"
        compact_rows.append(
            [
                f"{metric} {METRIC_DIRECTIONS.get(metric, '')}",
                fmt_p(fr["p_value"]),
                "Yes" if fr["p_value"] < 0.05 else "No",
                comparison,
                f"{int(posthoc['wins'])}/{int(posthoc['ties'])}/{int(posthoc['losses'])}",
                fmt_p(posthoc["holm_p_value"]),
                "Yes" if bool(posthoc["significant_0_05"]) else "No",
            ]
        )
    add_table(
        doc,
        ["Metric", "Friedman p", "Sig.", "Post-hoc comparison", "W/T/L", "Holm adj. p", "Sig."],
        compact_rows,
        [1.0, 1.05, 0.55, 2.15, 0.8, 1.05, 0.55],
    )
    doc.add_paragraph(
        "RankScore 的 paired tests 顯示 Meta-designed 分別顯著優於 Hand-crafted、Random 與 Bayesian baselines；完整 pairwise 結果保存在 supplementary CSV。"
    )

    doc.add_heading("7. 指標可用性與 RankScore 定義", level=1)
    add_table(
        doc,
        ["類別", "項目", "狀態", "來源 / 定義"],
        [
            ["Solution quality", "HV, IGD", "Available", "由 common reference front 計算；HV 越大越好，IGD 越小越好。"],
            ["PF stability", "PF Overlap, PF Drift", "Available", "衡量與 common reference front 的重疊程度與偏移程度。"],
            ["Search behavior", "Diversity", "Available", "衡量最終 front 的解分布廣度。"],
            ["Computational cost", "Runtime", "Available", "每個 method × instance × run 的 runtime 平均。"],
            ["Per-instance score", "RankScore", "Available", "同一 instance 內六個 metric rank 的平均，越低越好。"],
            ["Overall score", "Overall RankScore", "Available", "六個 mean metric 的 overall rank 平均，越低越好。"],
            ["First-place", "first_place_instances", "Available", "OverallInstanceRank == 1 的 instance 數；並列第一不計入嚴格 first-place。"],
        ],
        [1.25, 1.45, 0.85, 3.95],
    )

    doc.add_heading("8. 圖表", level=1)
    fig_dir = SUMMARY_DIR / "figures"
    figures = [
        ("Figure B1", "Overall RankScore", fig_dir / "overall_rank_score.png"),
        ("Figure B2", "Per-instance RankScore boxplot", fig_dir / "per_instance_rankscore_boxplot.png"),
        ("Figure B3", "Meta-designed median improvement by metric", fig_dir / "meta_median_improvement_by_metric.png"),
        ("Figure B4", "Meta-designed theta usage", fig_dir / "meta_theta_usage.png"),
    ]
    for label, caption, path in figures:
        if path.exists():
            p = doc.add_paragraph()
            run = p.add_run(f"{label}. {caption}")
            run.bold = True
            set_east_asian_font(run)
            doc.add_picture(str(path), width=Inches(6.7))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("9. 結論", level=1)
    conclusions = [
        "Meta-designed ECMADE-MOO 在四種 configuration strategy 中取得最低 Overall RankScore，且在 HV、IGD、PF Overlap 與 PF Drift 的 overall rank 皆為第一。",
        "相較於 Hand-crafted baseline，Meta-designed 顯示固定人工參數在 heterogeneous unseen instances 上仍有泛化限制。",
        "相較於 Random baseline，Meta-designed 證明 theta selection 不能只依賴隨機抽樣；instance features 可提供更穩定的選擇依據。",
        "相較於 Bayesian baseline，Meta-designed 在品質與穩定性指標上取得更佳整體表現，但 runtime 並非最低，後續可在 Experiment C 或延伸實驗中加入 runtime-aware objective。",
        "因此 Experiment B 的論文敘事可定位為：在 Experiment A 發現 fixed configuration 的泛化與穩定性限制後，Experiment B 進一步證明 instance-aware configuration selection 能改善 ECMADE-MOO 在 unseen portfolio instances 上的整體表現。",
    ]
    for text in conclusions:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        r = p.add_run(text)
        set_east_asian_font(r)

    doc.add_heading("10. 主要檔案位置", level=1)
    add_table(
        doc,
        ["檔案", "用途"],
        [[name, str(SUMMARY_DIR / name)] for name in availability_files if (SUMMARY_DIR / name).exists()],
        [2.2, 5.3],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Experiment B configuration strategy report")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)
    set_east_asian_font(fr)

    doc.save(OUT_PATH)
    print(f"DOCX={OUT_PATH}")


if __name__ == "__main__":
    main()

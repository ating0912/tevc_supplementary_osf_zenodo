from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docx_outputs"
OUT_PATH = OUT_DIR / "TEVC_消融實驗數據報告_依審稿問題分類_正式修正版_含Subpopulation統計檢定_20260722.docx"


def pct(value: float) -> str:
    if pd.isna(value):
        return "-"
    return f"{value * 100:.1f}%"


def num(value: float, digits: int = 4) -> str:
    if pd.isna(value):
        return "-"
    return f"{value:.{digits}f}"


def pval(value: float) -> str:
    if pd.isna(value):
        return "-"
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}"


def set_font(run, size=11, bold=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str, size=11, bold=False, color=None, style=None, after=6) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color="2E74B5" if level < 3 else "1F4D78")


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, bold=False, size=8.3, align_center=False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_margins(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    p.text = ""
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8.3,
    center_from: int = 1,
) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    repeat_header(table.rows[0])
    for c_idx, header in enumerate(headers):
        cell = table.cell(0, c_idx)
        shade(cell, "F2F4F7")
        write_cell(cell, header, bold=True, size=font_size, align_center=c_idx >= center_from)
    for row_values in rows:
        for c_idx, value in enumerate(row_values):
            write_cell(table.cell(len(table.rows) - len(rows) + rows.index(row_values), c_idx), value, size=font_size, align_center=c_idx >= center_from)
    doc.add_paragraph()


def add_rows(table, rows, font_size, center_from):
    for r_idx, row_values in enumerate(rows, start=1):
        for c_idx, value in enumerate(row_values):
            write_cell(table.cell(r_idx, c_idx), value, size=font_size, align_center=c_idx >= center_from)


def add_table_fixed(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
    font_size: float = 8.3,
    center_from: int = 1,
) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    repeat_header(table.rows[0])
    for c_idx, header in enumerate(headers):
        cell = table.cell(0, c_idx)
        shade(cell, "F2F4F7")
        write_cell(cell, header, bold=True, size=font_size, align_center=c_idx >= center_from)
    for r_idx, row_values in enumerate(rows, start=1):
        for c_idx, value in enumerate(row_values):
            write_cell(table.cell(r_idx, c_idx), value, size=font_size, align_center=c_idx >= center_from)
    doc.add_paragraph()


def add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(f"{label}：")
    set_font(r1, size=10.5, bold=True, color="1F4D78")
    r2 = p.add_run(text)
    set_font(r2, size=10.5)


def source_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"資料來源：{text}")
    set_font(r, size=8.5, color="666666")


def read_csv(rel: str) -> pd.DataFrame:
    return pd.read_csv(ROOT / rel)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, side, Inches(1.0))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1
    return doc


def build_report() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = setup_document()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("TEVC 消融實驗數據報告")
    set_font(r, size=20, bold=True, color="0B2545")
    add_para(doc, "依審稿問題分類之補齊版：新增 J、EAF width、archive diversity、test J、feature importance，並補充各表格指標計算口徑。", size=11, color="555555", after=10)
    add_note(doc, "整體口徑", "J 在本文中對應多指標排名分數 RankScore，數值越低越好；HV、PF Overlap、Diversity 越高越好，IGD、PF Drift、Runtime 越低越好。C-label 相關指標均以 stability-aware selector 的驗證集合 C 口徑彙總。K 代表 cardinality constraint，也就是投資組合中需選取的資產數量限制。")

    add_heading(doc, "1. Without Meta-learning", 1)
    meta = read_csv("p0_lite_outputs/experiment_b_configuration_summary_20260713/overall_configuration_comparison.csv")
    order = ["MetaDesigned_ECMADE_MOO", "BayesianConfig_ECMADE_MOO", "RandomConfig_ECMADE_MOO", "HandCrafted_ECMADE_MOO"]
    meta = meta.set_index("method").loc[order].reset_index()
    rows = []
    for _, row in meta.iterrows():
        rows.append([
            row["method"].replace("_ECMADE_MOO", ""),
            num(row["mean_HV"]),
            num(row["mean_IGD"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_PF_Drift"]),
            num(row["mean_Diversity"]),
            num(row["mean_Runtime"], 3),
            num(row["overall_RankScore"], 3),
            num(row["mean_InstanceRank"], 3),
            str(int(row["first_place_instances"])),
        ])
    add_table_fixed(doc, ["設定版本", "HV", "IGD", "PF\nOverlap", "PF\nDrift", "Diversity", "Runtime", "J", "Mean\nrank", "First"], rows, [1.28, .55, .55, .7, .63, .7, .65, .48, .58, .38], 7.2)
    source_line(doc, "p0_lite_outputs/experiment_b_configuration_summary_20260713/overall_configuration_comparison.csv")
    add_note(doc, "補齊 J", "Without Meta-learning 原本缺少的 J 以 overall_RankScore 補入；此值是把各設定版本在主要指標上的排序取平均，代表整體多目標表現的彙總排名，越低越好。")
    stat = read_csv("p0_lite_outputs/experiment_b_configuration_summary_20260713/statistical_tests_meta_vs_baselines.csv")
    stat = stat[stat["metric"].eq("RankScore")]
    rows = []
    for _, row in stat.iterrows():
        rows.append([
            row["baseline"],
            f"{int(row['wins'])}/{int(row['ties'])}/{int(row['losses'])}",
            num(row["median_improvement"], 3),
            num(row["holm_p_value"], 4),
            "是" if bool(row["significant_0_05"]) else "否",
        ])
    add_table_fixed(doc, ["Meta-designed 對照", "勝/平/負", "J 中位改善", "Holm p", "顯著"], rows, [2.1, 1.0, 1.1, 1.0, .8], 8.4)
    add_note(doc, "結論", "Meta-designed 設定在 J、HV、IGD 與 PF Overlap 上皆優於固定手工與資料驅動基線，顯示 meta-learning 的功能不是只改變參數外觀，而是能穩定選出較佳 theta 組合。")

    add_heading(doc, "2. Without Stability Objective", 1)
    csum = read_csv("outputs/tevc_ablation_4_5_20260717/label_objective_cross_evaluation_on_C_summary.csv")
    c_eaf = read_csv("outputs/tevc_ablation_4_5_20260717/label_objective_cross_evaluation_eaf_summary.csv")
    c_eaf_lookup = c_eaf.set_index("selector")["mean_selected_EAF_Band_Width_IQR"].to_dict()
    rows = []
    selector_names = {
        "stability_label_top1": "stability-aware J",
        "standard_label_top1": "standard J",
        "performance_only_top1": "performance-only J",
        "pf_stability_only_top1": "PF-stability-only J",
        "all_theta_mean": "all theta mean",
    }
    for key in ["stability_label_top1", "standard_label_top1", "performance_only_top1", "pf_stability_only_top1", "all_theta_mean"]:
        row = csum[csum["selector"].eq(key)].iloc[0]
        rows.append([
            selector_names[key],
            pct(row["c_top1_hit_rate"]),
            pct(row["c_top3_hit_rate"]),
            num(row["mean_C_ThetaRank"], 3),
            num(row["mean_C_regret_loss"], 3),
            num(row["mean_selected_PF_Overlap"]),
            num(row["mean_selected_PF_Drift"]),
            num(c_eaf_lookup.get(key, float("nan"))),
            num(row["mean_selected_Diversity"]),
            num(row["mean_selected_Runtime"], 3),
        ])
    add_table_fixed(doc, ["Selector", "C top1", "C top3", "C rank", "C regret", "PF\nOverlap", "PF\nDrift", "EAF width\n(IQR)", "Diversity", "Runtime"], rows, [1.3, .58, .58, .55, .62, .65, .58, .62, .65, .62], 7.2)
    source_line(doc, "outputs/tevc_ablation_4_5_20260717/label_objective_cross_evaluation_on_C_summary.csv")
    add_note(doc, "C top1 / C top3 / C rank / C regret", "C top1 是被 selector 選到的 theta 在 C-label oracle 排名為第 1 名的比例；C top3 是排名落在前 3 名的比例；C rank 是所選 theta 的 C_ThetaRank 平均；C regret 是所選 theta 的 C_LabelScore loss 與 C oracle 最佳 loss 之差，0 代表等同 oracle，越低越好。")
    add_note(doc, "EAF width", "EAF width 由 validation theta 的 30-run pf_obj.csv 重算：每個 theta-instance 先以 union front 做 min-max normalization，再在 101-point grid 上建立 attainment curves，最後取 q75-q25 的平均帶寬作 IQR EAF width，越低代表 repeated-run attainment band 越窄。")
    add_note(doc, "解讀提醒", "EAF width 是獨立穩定性診斷；本表中 standard J 的 EAF width 略低，但 stability-aware J 在 C top1、C regret 與 PF Overlap 上最佳，因此正文應寫成 stability-aware objective 改善 C-label 一致性與 overlap，而非宣稱所有穩定性指標都單向最佳。")

    add_heading(doc, "3. Without Adaptive Exchange", 1)
    theta = read_csv("outputs/tevc_ablation_6_20260717/theta_factor_main_effect_summary.csv")
    mig_eaf = read_csv("outputs/tevc_ablation_6_20260717/theta_factor_main_effect_eaf_summary.csv")
    mig_eaf_lookup = mig_eaf.set_index("level")["mean_EAF_Band_Width_IQR"].to_dict()
    mig = theta[(theta["source"].eq("Validation")) & (theta["objective"].eq("stability_label")) & (theta["factor"].eq("migration"))]
    rows = []
    for key in ["none", "adaptive", "fixed"]:
        row = mig[mig["level"].eq(key)].iloc[0]
        rows.append([
            key,
            num(row["mean_objective_loss"], 3),
            num(row["mean_objective_rank"], 3),
            pct(row["mean_top3_share"]),
            num(row["mean_HV"]),
            num(row["mean_IGD"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_PF_Drift"]),
            num(mig_eaf_lookup.get(key, float("nan"))),
            num(row["mean_Diversity"]),
            num(row["mean_Runtime"], 3),
        ])
    add_table_fixed(doc, ["Migration", "C loss", "Mean\nrank", "Top3\nshare", "HV", "IGD", "PF\nOverlap", "PF\nDrift", "EAF width\n(IQR)", "Diversity", "Runtime"], rows, [0.75, .65, .62, .63, .5, .5, .65, .58, .62, .62, .58], 7.0)
    source_line(doc, "outputs/tevc_ablation_6_20260717/theta_factor_main_effect_summary.csv")
    add_note(doc, "C loss / Mean rank / Top3 share", "對每個驗證 instance group，先把同一 migration level 的 theta 候選彙總，再計算 C_LabelScore loss、C_ThetaRank 與是否進入 top 3；表中的 C loss、Mean rank、Top3 share 是跨驗證 groups 的平均。EAF width 使用相同 migration level 分組，先對每個 theta-instance 重算 30-run IQR band，再跨 theta 與 instance groups 平均。")
    pdf = read_csv("p0_lite_outputs/tevc_pdf_direct_ablation_full_20260717/pdf_direct_ablation_summary/pdf_direct_ablation_overall.csv")
    direct_eaf = read_csv("p0_lite_outputs/tevc_p0_requirement_tables_20260717/tevc_p0_instance_eaf_band_width.csv")
    direct_eaf = (
        direct_eaf[direct_eaf["experiment"].eq("TEVC_PDF_Direct_Ablation")]
        .groupby("method", dropna=False)["EAF_Band_Width_IQR"]
        .mean()
        .to_dict()
    )
    mig2 = pdf[pdf["ablation_family"].eq("migration")]
    rows = []
    for key in ["adaptive", "fixed", "none"]:
        row = mig2[mig2["ablation_level"].eq(key)].iloc[0]
        rows.append([
            key,
            num(row["mean_RankScore"], 3),
            num(row["overall_RankScore"], 3),
            num(row["mean_FamilyInstanceRank"], 3),
            str(int(row["first_place_instances"])),
            num(direct_eaf.get(row["method"], float("nan"))),
            num(row["mean_PF_Overlap"]),
            num(row["mean_Runtime"], 3),
        ])
    add_table_fixed(doc, ["Direct variant", "J-score", "Overall J\nrank", "Family rank\n(1 best)", "First-place", "EAF width\n(IQR)", "PF Overlap", "Runtime"], rows, [0.9, .5, .62, .72, .7, .72, .68, .58], 7.2)
    sig = read_csv("outputs/tevc_adaptive_exchange_significance_20260722/adaptive_none_vs_adaptive_wilcoxon_all.csv")
    metric_labels = {
        "objective_loss": "C loss",
        "objective_rank": "Mean rank",
        "top3_share": "Top3 share",
        "PF_Overlap": "PF Overlap",
        "PF_Drift": "PF Drift",
        "EAF_Band_Width_IQR": "EAF width",
        "Runtime": "Runtime",
        "RankScore": "J-score",
        "FamilyInstanceRank": "Family rank",
    }
    rows = []
    c_sig = sig[sig["context"].eq("Validation C-label main effect")]
    for metric in ["objective_loss", "objective_rank", "top3_share", "PF_Overlap", "PF_Drift", "EAF_Band_Width_IQR", "Runtime"]:
        row = c_sig[c_sig["metric"].eq(metric)].iloc[0]
        rows.append([
            metric_labels[metric],
            num(row["none_mean"], 4),
            num(row["adaptive_mean"], 4),
            num(row["median_none_minus_adaptive"], 4),
            f"{int(row['none_better'])}/{int(row['ties'])}/{int(row['adaptive_better'])}",
            pval(row["holm_p"]),
            "是" if bool(row["significant_0_05_holm"]) else "否",
        ])
    add_table_fixed(doc, ["C-label metric", "none\nmean", "adaptive\nmean", "Median difference\n(none - adaptive)", "none/ tie/\nadapt wins", "Holm p", "顯著"], rows, [1.0, .62, .7, 1.0, .88, .62, .43], 6.9)
    direct_sig = sig[sig["context"].eq("Direct test ablation")]
    rows = []
    for metric in ["RankScore", "FamilyInstanceRank", "EAF_Band_Width_IQR", "Runtime"]:
        row = direct_sig[direct_sig["metric"].eq(metric)].iloc[0]
        rows.append([
            metric_labels[metric],
            num(row["none_mean"], 4),
            num(row["adaptive_mean"], 4),
            num(row["median_none_minus_adaptive"], 4),
            f"{int(row['none_better'])}/{int(row['ties'])}/{int(row['adaptive_better'])}",
            pval(row["holm_p"]),
            "是" if bool(row["significant_0_05_holm"]) else "否",
        ])
    add_table_fixed(doc, ["Direct metric", "none\nmean", "adaptive\nmean", "Median difference\n(none - adaptive)", "none/ tie/\nadapt wins", "Holm p", "顯著"], rows, [1.0, .62, .7, 1.0, .88, .62, .43], 6.9)
    source_line(doc, "adaptive_none_vs_adaptive_wilcoxon_all.csv")
    add_table_fixed(
        doc,
        ["檢定表", "資料範圍", "比較組數", "彙整層級", "實驗控制與用途"],
        [
            [
                "C-label main effect",
                "Validation set；使用既有 theta-label validation data。",
                "29 組 validation instance-K 配對。",
                "跨所有 theta 配置按 migration level 平均；其他 theta 因子沒有固定。",
                "使用 validation theta runs 的相同資料生成流程與 30 runs；用來檢查既有 theta-label 資料中的外部一致性與主效應趨勢。",
            ],
            [
                "Direct test ablation",
                "Test set；使用 TEVC_PDF_Direct_Ablation 重新執行結果。",
                "32 組 test instance-K 配對。",
                "固定其他 theta 因子，只改變 migration = adaptive / fixed / none。",
                "各 variant 使用同一 direct ablation 實驗預算與 30 runs；用來判斷 migration 設計的因果效果，應作為主要證據。",
            ],
        ],
        [1.05, 1.25, .8, 1.55, 1.85],
        6.8,
        center_from=2,
    )
    add_note(doc, "統計檢定口徑", "none 與 adaptive 以相同 instance-K 配對比較，使用雙尾 paired Wilcoxon signed-rank test，並採 Holm correction 控制多重比較問題。Median difference 定義為 none - adaptive；對 C loss、Mean rank、PF Drift、EAF width、Runtime、J-score 與 Family rank 而言，負值代表 none 較佳；對 Top-3 share 與 PF Overlap 而言，正值代表 none 較佳。J-score 數值越低越好，Family rank 以 rank 1 為最佳。")
    add_note(doc, "表間差異說明", "兩張檢定表在 EAF width 與 Runtime 上的結論不完全一致，主要原因在於資料範圍與彙整層級不同。C-label main effect 以 validation label data 進行因子水準平均，可能同時受到其他 theta 配置分布的影響；Direct test ablation 則在 test set 上固定其他 theta 因子，只改變 migration 設定。因此，涉及 migration 因果效果的結論以 Direct test ablation 為主要依據，C-label main effect 則作為不同資料範圍下的外部一致性或效應趨勢分析。")
    add_note(doc, "因果解讀", "在 C-label main-effect 分析中，migration=none 在 C loss、Mean rank、Top-3 share、PF Overlap、PF Drift 及 EAF width 等指標上均顯著優於 adaptive，顯示在既有 theta-label 資料中，不進行資訊交換與較佳的解品質及穩定性表現具有一致關聯；然而，adaptive 在此分析中的 Runtime 顯著較低，表示其具有運算時間上的局部優勢。為進一步排除其他 theta 因子的干擾，本研究在 test set 上固定 subpopulation number、operator、elite ratio 與 stagnation threshold 等條件，只改變 migration 設定進行直接配對消融。結果顯示，none 在 J-score 與 Family rank 上均取得 31 勝、0 平、1 負，Holm 校正後的 p-value 皆為 8.61 x 10^-7，且 Runtime 在 32 組比較中全部優於 adaptive，差異同樣達統計顯著；相較之下，兩者的 EAF width 完全相同，未呈現顯著差異。綜合而言，在目前 test set 與固定 theta 條件下，adaptive exchange 未能改善 run-to-run 穩定性，反而增加運算時間並降低整體配置績效，因此 migration=none 為目前較適合的設定。此因果結論限於本研究所使用的資料集、theta 設定及運算預算，不宜延伸為所有問題情境下的一般性結論。")

    add_heading(doc, "4. Without Elite Injection", 1)
    elite = read_csv("outputs/tevc_without_elite_injection_20260720/without_elite_injection_overall.csv")
    archive_path = ROOT / "outputs/tevc_without_elite_injection_20260720/archive_metrics_by_method.csv"
    archive_lookup = {}
    if archive_path.exists():
        arch = pd.read_csv(archive_path)
        if {"method", "mean_Archive_Diversity"}.issubset(arch.columns):
            archive_lookup = dict(zip(arch["method"], arch["mean_Archive_Diversity"]))
    rows = []
    for _, row in elite.iterrows():
        ratio = row["ablation_level"]
        method = row["method"]
        rows.append([
            ratio,
            num(row["mean_RankScore"], 3),
            num(row["overall_RankScore"], 3),
            num(row["mean_FamilyInstanceRank"], 3),
            str(int(row["first_place_instances"])),
            num(archive_lookup.get(method, float("nan")), 4) if archive_lookup else "待補算",
            num(row["mean_HV"]),
            num(row["mean_IGD"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_Diversity"]),
            num(row["mean_Runtime"], 3),
        ])
    add_table_fixed(doc, ["Elite ratio", "J", "Overall J", "Family\nrank", "First-place", "Archive\ndiversity", "HV", "IGD", "PF\nOverlap", "Diversity", "Runtime"], rows, [.7, .48, .65, .62, .62, .78, .48, .48, .62, .62, .55], 7.0)
    source_line(doc, "outputs/tevc_without_elite_injection_20260720/without_elite_injection_overall.csv")
    add_note(doc, "Family rank / First-place", "Family rank 是在同一消融家族內、同一 instance 上只比較 eliteRatio 變體後得到的 family instance rank，再跨 32 個 instances 取平均；First-place 是該變體在 family rank 中取得第 1 名的 instances 數。")
    add_note(doc, "Archive diversity", "Archive diversity 與一般 objective-space Diversity 不同；本表已由每個 run 的 archive_metrics.csv 彙總至 archive_metrics_by_method.csv 後填入。")
    add_note(doc, "結論", "eliteRatio = 0% 補齊後顯示其 J、Family rank、First-place 與 Runtime 均最佳；因此本段可回答 elite injection 是否造成改善：在目前 direct ablation 表中，加入 elite injection 反而增加成本且未帶來更佳整體排名。")

    add_heading(doc, "5. Subpopulation Number", 1)
    sub = pdf[pdf["ablation_family"].eq("subpopulation_number")]
    rows = []
    for key in ["S=2", "S=3", "S=5"]:
        row = sub[sub["ablation_level"].eq(key)].iloc[0]
        rows.append([
            key,
            num(row["mean_RankScore"], 3),
            num(row["overall_RankScore"], 3),
            num(row["mean_FamilyInstanceRank"], 3),
            str(int(row["first_place_instances"])),
            num(direct_eaf.get(row["method"], float("nan"))),
            num(row["mean_HV"]),
            num(row["mean_IGD"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_Diversity"]),
            num(row["mean_Runtime"], 3),
        ])
    add_table_fixed(doc, ["S", "J-score", "Overall J\nrank", "Family\nrank", "First-place", "EAF width\n(IQR)", "HV", "IGD", "PF\nOverlap", "Diversity", "Runtime"], rows, [.42, .48, .62, .6, .6, .68, .45, .45, .62, .6, .55], 6.9)
    rows = []
    sf = theta[(theta["source"].eq("Validation")) & (theta["objective"].eq("stability_label")) & (theta["factor"].eq("S"))]
    for key in ["2", "3", "5"]:
        row = sf[sf["level"].astype(str).eq(key)].iloc[0]
        rows.append([
            key,
            num(row["mean_objective_loss"], 3),
            num(row["mean_objective_rank"], 3),
            pct(row["mean_top3_share"]),
            num(row["mean_PF_Overlap"]),
            num(row["mean_Diversity"]),
            num(row["mean_Runtime"], 3),
        ])
    add_table_fixed(doc, ["S", "C loss", "Mean rank", "Top3 share", "PF Overlap", "Diversity", "Runtime"], rows, [.55, .85, .85, .85, .85, .85, .8], 8.2)
    add_note(doc, "C loss / Mean rank / Top3 share", "計算方式同 adaptive exchange：在 validation C 口徑下，按 S 的 factor level 彙總 theta 候選，C loss 越低、Mean rank 越低、Top3 share 越高代表越容易選到穩定且表現好的 theta。")
    sub_sig = read_csv("outputs/tevc_subpopulation_significance_20260722/subpopulation_s3_vs_s2_s5_wilcoxon_all.csv")
    sub_metric_labels = {
        "objective_loss": "C loss",
        "objective_rank": "Mean rank",
        "top3_share": "Top3 share",
        "PF_Overlap": "PF Overlap",
        "PF_Drift": "PF Drift",
        "EAF_Band_Width_IQR": "EAF width",
        "Runtime": "Runtime",
        "RankScore": "J-score",
        "FamilyInstanceRank": "Family rank",
        "HV": "HV",
        "IGD": "IGD",
    }
    rows = []
    c_sub_sig = sub_sig[sub_sig["context"].eq("Validation C-label S main effect")]
    for comp in ["3 vs 2", "3 vs 5"]:
        for metric in ["objective_loss", "objective_rank", "top3_share", "PF_Overlap", "EAF_Band_Width_IQR", "Runtime"]:
            row = c_sub_sig[(c_sub_sig["comparison"].eq(comp)) & (c_sub_sig["metric"].eq(metric))].iloc[0]
            rows.append([
                comp,
                sub_metric_labels[metric],
                num(row["reference_mean"], 4),
                num(row["comparator_mean"], 4),
                num(row["median_reference_minus_comparator"], 4),
                f"{int(row['reference_better'])}/{int(row['ties'])}/{int(row['comparator_better'])}",
                pval(row["holm_p"]),
                "是" if bool(row["significant_0_05_holm"]) else "否",
            ])
    add_table_fixed(doc, ["Comparison", "C-label metric", "S=3\nmean", "Other\nmean", "Median difference\n(S=3 - other)", "S=3/tie/\nother wins", "Holm p", "顯著"], rows, [.82, .88, .58, .62, 1.0, .92, .58, .42], 6.2)
    rows = []
    direct_sub_sig = sub_sig[sub_sig["context"].eq("Direct test S ablation")]
    for comp in ["S=3 vs S=2", "S=3 vs S=5"]:
        for metric in ["RankScore", "FamilyInstanceRank", "HV", "IGD", "PF_Overlap", "EAF_Band_Width_IQR", "Runtime"]:
            row = direct_sub_sig[(direct_sub_sig["comparison"].eq(comp)) & (direct_sub_sig["metric"].eq(metric))].iloc[0]
            rows.append([
                comp,
                sub_metric_labels[metric],
                num(row["reference_mean"], 4),
                num(row["comparator_mean"], 4),
                num(row["median_reference_minus_comparator"], 4),
                f"{int(row['reference_better'])}/{int(row['ties'])}/{int(row['comparator_better'])}",
                pval(row["holm_p"]),
                "是" if bool(row["significant_0_05_holm"]) else "否",
            ])
    add_table_fixed(doc, ["Comparison", "Direct metric", "S=3\nmean", "Other\nmean", "Median difference\n(S=3 - other)", "S=3/tie/\nother wins", "Holm p", "顯著"], rows, [.82, .88, .58, .62, 1.0, .92, .58, .42], 6.2)
    source_line(doc, "subpopulation_s3_vs_s2_s5_wilcoxon_all.csv")
    add_note(doc, "統計檢定口徑", "以 S=3 作為 reference，分別與 S=2、S=5 以相同 instance-K 配對比較，使用雙尾 paired Wilcoxon signed-rank test，並採 Holm correction 控制多重比較。Median difference 定義為 S=3 - other；對 C loss、Mean rank、PF Drift、EAF width、Runtime、J-score 與 Family rank 而言，負值代表 S=3 較佳；對 Top-3 share、PF Overlap 與 HV 而言，正值代表 S=3 較佳。J-score 數值越低越好，Family rank 以 rank 1 為最佳。")
    add_note(doc, "結論", "Validation C-label main effect 顯示 S=3 相對 S=2 在 C loss、Mean rank、Top-3 share、PF Overlap 與 Runtime 上達顯著改善，但 PF Drift 與 EAF width 未達顯著；S=3 相對 S=5 則只有 Runtime 顯著較慢，其他 C-label 指標經 Holm correction 後未達顯著。Direct test ablation 作為固定其他 theta 因子的主要證據，顯示 S=3 相對 S=2 在 J-score、Family rank、HV、IGD 與 PF Overlap 上顯著較佳，但 Runtime 顯著較慢，EAF width 未顯著不同；相對 S=5，S=3 在 Family rank 與 PF Overlap 上顯著較佳，但 J-score、HV、IGD、EAF width 與 Runtime 未達顯著。因此報告應寫成：S=3 是解品質與整體排名的較佳折衷設定，但其優勢不是所有指標全面顯著，且相對 S=2 需承擔較高 runtime。")

    add_heading(doc, "6. Feature Group Ablation", 1)
    fg = read_csv("outputs/tevc_ablation_4_5_20260717/feature_group_ablation_summary.csv")
    fg = fg[fg["objective"].eq("stability_label")]
    order = ["categorical_only", "all_features", "no_theta_categorical", "theta_only_no_instance", "numeric_only", "no_problem_categorical", "instance_only_no_theta"]
    rows = []
    for key in order:
        row = fg[fg["selector"].eq(key)].iloc[0]
        rows.append([
            key,
            pct(row["top1_hit_rate"]),
            pct(row["top3_hit_rate"]),
            num(row["mean_selected_rank"], 3),
            num(row["mean_regret"], 3),
            num(-row["mean_selected_score"], 3),
            num(row["mean_selected_HV"]),
            num(row["mean_selected_IGD"]),
            num(row["mean_selected_PF_Overlap"]),
            num(row["mean_selected_Diversity"]),
            num(row["mean_selected_Runtime"], 3),
            num(row["rmse"], 3),
            num(row["mae"], 3),
        ])
    add_table_fixed(doc, ["Feature set", "Top1", "Top3", "Mean\nrank", "C regret", "test J", "HV", "IGD", "PF\nOverlap", "Diversity", "Runtime", "RMSE", "MAE"], rows, [1.05, .38, .38, .45, .50, .48, .40, .40, .52, .52, .48, .42, .42], 6.2)
    source_line(doc, "outputs/tevc_ablation_4_5_20260717/feature_group_ablation_summary.csv")
    add_note(doc, "test J", "本表的 mean_selected_score 是 selector 所選 theta 的實際 C score；因 stability objective 以 loss 表示，test J 以 -mean_selected_score 換回 selected C loss，越低越好。")
    add_note(doc, "其他表格指標", "Top1/Top3 是 selector 選中 oracle 第 1 名或前 3 名的比例；Mean rank 是 selected theta 的實際 C rank 平均；C regret 是 selected C loss 減 oracle C loss；HV、IGD、PF Overlap、Diversity、Runtime 是所選 theta 在測試/驗證 group 上的實際表現平均；RMSE 與 MAE 是模型預測分數與實際 score 的誤差。")
    imp = read_csv("p0_lite_outputs/experiment_c_stability_selector_training/feature_importance.csv").head(12)
    rows = [[row["feature"], num(row["importance"], 4)] for _, row in imp.iterrows()]
    add_table_fixed(doc, ["Feature", "Importance"], rows, [4.7, 1.1], 8.8)
    source_line(doc, "p0_lite_outputs/experiment_c_stability_selector_training/feature_importance.csv")
    add_note(doc, "Feature importance", "feature importance 來自 stability-aware selector 訓練後的特徵重要度彙總，可用來回答哪些 meta-features 對選擇 theta 有用；目前 K、subpops、k_ratio、replicate、assets 與 pathological covariance 相關特徵權重最高。")

    add_heading(doc, "建議合併寫法", 1)
    add_para(doc, "Feature Group Ablation 建議和 Without Stability Objective 放在同一個大章中的相鄰小節，因為兩者都在回答「stability-aware selector 是否真的需要 C-label 與 meta-features」。但在報告版面上仍保留獨立大標，可讓審稿問題逐項對應。", size=11)
    add_para(doc, "Without Adaptive Exchange、Without Elite Injection 與 Subpopulation Number 則建議放在「搜尋機制與族群設計消融」章節下，因為它們直接改變演算法機制，而不是 selector 的特徵或 objective 定義。", size=11)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_report()

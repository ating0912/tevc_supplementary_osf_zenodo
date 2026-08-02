from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "docx_outputs"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "NSGAII_baseline_reproduction_parameters.docx"

FINAL_DIR = ROOT / "nsga2_outputs" / "final_all22_r2020b_v290_mcg_randomtie"
SEED_DIR = ROOT / "nsga2_outputs" / "seed_blocks_final_params_all22"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


comparison = read_csv(FINAL_DIR / "comparison_table.csv")
relative_summary = read_csv(FINAL_DIR / "relative_diff_summary.csv")
seed_ranking = read_csv(SEED_DIR / "seed_block_paper_closeness_ranking.csv")
seed_effect = read_csv(SEED_DIR / "seed_effect_by_problem.csv")


def fmt(value: str | float, digits: int = 4) -> str:
    x = float(value)
    if abs(x) >= 100:
        return f"{x:.2f}"
    if abs(x) >= 1:
        return f"{x:.4f}"
    return f"{x:.4e}"


def pct(value: str | float) -> str:
    return f"{float(value):.2f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def rgb_from_hex(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_east_asia_font(target, font_name: str = "Microsoft JhengHei") -> None:
    element = target._element
    if hasattr(element, "get_or_add_rPr"):
        r_pr = element.get_or_add_rPr()
    else:
        r_pr = element.rPr
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            element.insert(0, r_pr)

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_width(table, widths: list[int], indent: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths: list[int]) -> None:
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, widths)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_table_width(table, [9360])
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(" " + body)
    p.paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
set_east_asia_font(styles["Normal"])
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    s = styles[style_name]
    s.font.name = "Calibri"
    set_east_asia_font(s)
    s.font.size = Pt(size)
    s.font.color.rgb = rgb_from_hex(color)
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "NSGA-II baseline reproduction parameter memo"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
r = title.add_run("NSGA-II Baseline 復現參數與差異說明")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(11, 37, 69)
set_east_asia_font(r)

subtitle = doc.add_paragraph()
subtitle.add_run("依據原文設定、PlatEMO v2.9 原生流程與補充 seed 測試整理").italic = True

doc.add_heading("1. 結論摘要", level=1)
add_note(
    doc,
    "採用設定。",
    "目前最接近論文 Table 3 NSGA-II baseline 的可重現設定為 MATLAB R2020b、PlatEMO v2.9、NSGA-II、mcg16807、seed 1:30、random tie、N=100、maxFE=10000、30 runs、PlatEMO 原生 PF(10000) 與 raw IGD。"
)
add_bullet(doc, "22 題整體 Mean relative diff 為 41.44%，Median relative diff 為 31.43%。")
add_bullet(doc, "不同 seed block 的整體差異有限，mean relative diff 約落在 41.44% 到 43.33%。")
add_bullet(doc, "因此，無法重現論文 baseline 並非單純由 random seed 選取造成。")
add_bullet(doc, "差異與資料集特性有關：ZDT 差異最大，UF 差異最小，DTLZ 居中但部分退化或困難問題差異明顯。")

doc.add_heading("2. 參數採用原則", level=1)
principle_rows = [
    ["1", "論文有明確提及", "直接採用論文設定，例如 R2020b、N=100、MaxIt=10000、30 runs、SBX、Polynomial Mutation、proC/etaC/proM/etaM。"],
    ["2", "論文未明講但 PlatEMO 有原生流程", "採用 PlatEMO v2.9 的 NSGA-II、GA、NDSort、CrowdingDistance、PF、IGD 流程。"],
    ["3", "論文與平台皆未指定", "採用測試後最接近且可重現的設定，例如 mcg16807、seed 1:30、random tie。"],
]
add_table(doc, ["優先順序", "依據", "使用方式"], principle_rows, [1000, 2300, 6060])

doc.add_heading("3. 詳細復現參數", level=1)
param_rows = [
    ["MATLAB", "R2020b", "論文指定"],
    ["平台", "PlatEMO", "論文提及"],
    ["PlatEMO 版本", "v2.9", "論文未指定；採用 R2020b 可穩定執行版本"],
    ["演算法", "NSGA-II", "論文 baseline"],
    ["測試資料集", "DTLZ1-DTLZ7、ZDT1-ZDT4、ZDT6、UF1-UF10，共 22 題", "論文 Table 3"],
    ["Population size", "N = 100", "論文指定"],
    ["最大評估次數", "maxFE = 10000", "由論文 MaxIt = 10000 對應到 PlatEMO evaluation"],
    ["獨立執行次數", "30 runs", "論文指定"],
    ["交叉算子", "Simulated Binary Crossover (SBX)", "論文指定 / PlatEMO GA"],
    ["交叉機率", "proC = 1", "論文 Table 2"],
    ["交叉分佈指數", "etaC = 20", "論文 Table 2"],
    ["變異算子", "Polynomial Mutation", "論文指定 / PlatEMO GA"],
    ["變異機率", "proM = 1", "論文 Table 2"],
    ["逐變數變異率", "proM / D", "PlatEMO GA.m 原生解讀"],
    ["變異分佈指數", "etaM = 20", "論文 Table 2"],
    ["選擇機制", "Tournament Selection", "PlatEMO NSGA-II 原生流程"],
    ["非支配排序", "PlatEMO v2.9 NDSort", "PlatEMO 原生流程"],
    ["擁擠距離", "PlatEMO v2.9 CrowdingDistance", "PlatEMO 原生流程"],
    ["同值處理", "random tie", "論文未指定；測試後較接近且可重現"],
    ["亂數產生器", "mcg16807", "論文未指定；測試後較接近且可重現"],
    ["亂數種子", "每個 run 固定 seed，seed = run index (1:30)", "論文未提供；為可重現性補充"],
    ["輸出解集", "final population 中可行且非支配解", "PlatEMO Metric 流程"],
    ["IGD 計算", "raw IGD，不做 objective normalization", "PlatEMO 原生 IGD.m"],
    ["PF 來源", "PlatEMO v2.9 problem 內建 PF(10000)", "PlatEMO 原生流程"],
]
add_table(doc, ["參數類別", "設定", "依據"], param_rows, [2200, 3700, 3460])

doc.add_heading("4. 與論文 Table 3 的比較", level=1)
doc.add_paragraph("下表列出 22 題 NSGA-II baseline 的 IGD mean 對照。Relative diff 以 abs(reproduced mean - paper mean) / abs(paper mean) x 100% 計算。")
comp_rows = []
for row in comparison:
    comp_rows.append([
        row["problem"],
        fmt(row["paper_mean_igd"]),
        fmt(row["reproduced_mean_igd"]),
        fmt(row["reproduced_std_igd"]),
        pct(row["mean_relative_diff_percent"]),
    ])
add_table(doc, ["Problem", "Paper IGD", "Reproduced mean", "Reproduced std", "Relative diff"], comp_rows, [1200, 1900, 2200, 2000, 2060])

doc.add_heading("5. Relative Diff Summary", level=1)
summary_rows = [
    [
        row["family"],
        row["problem_count"],
        pct(row["mean_relative_diff_percent"]),
        pct(row["median_relative_diff_percent"]),
    ]
    for row in relative_summary
]
add_table(doc, ["Family", "Problems", "Mean relative diff", "Median relative diff"], summary_rows, [1800, 1500, 3000, 3060])

doc.add_heading("6. Seed 測試結果", level=1)
doc.add_paragraph("為確認差異是否由 seed 造成，本研究在最終復現設定下測試四組 seed blocks。每組皆為 22 題 x 30 runs。")
seed_rows = [
    [
        row["seed_block"],
        row["GroupCount"],
        pct(row["mean_relative_diff_percent"]),
        pct(row["median_relative_diff_percent"]),
    ]
    for row in seed_ranking
]
add_table(doc, ["Seed block", "Problems", "Mean relative diff", "Median relative diff"], seed_rows, [2400, 1400, 2800, 2760])

top_sensitive = seed_effect[:8]
sensitive_rows = [
    [
        row["problem"],
        pct(row["block_mean_range_percent"]),
        row["closest_seed_block"],
        pct(row["closest_relative_diff_percent"]),
    ]
    for row in top_sensitive
]
doc.add_paragraph("Seed 敏感度最高的問題如下；block mean range 表示四組 seed block 的 mean IGD 差距相對於 grand mean 的比例。")
add_table(doc, ["Problem", "Block mean range", "Closest seed block", "Closest relative diff"], sensitive_rows, [1500, 2500, 3000, 2360])

doc.add_heading("7. 為何仍會與論文有差異", level=1)
diff_rows = [
    ["PlatEMO 版本未指定", "論文提及使用 PlatEMO，但未明確指定版本。不同版本的 NSGA-II、GA operator、NDSort、CrowdingDistance、PF sampling 可能不同。"],
    ["MaxIt 解讀差異", "論文寫 MaxIt = 10000；本研究依 PlatEMO 對應為 evaluation/maxFE = 10000。若作者以 generation 或其他停止條件解讀，結果會改變。"],
    ["random seed 未提供", "原文未提供 30 runs 的 seed。雖然本研究測試不同 seed blocks，但仍無法使整體 baseline 完全對齊。"],
    ["RNG 未指定", "RNG 類型會影響初始化、交叉、變異與 tournament selection。測試顯示 mcg16807 + random tie 較接近，但仍不能完全重現。"],
    ["IGD reference set 未完整說明", "本研究使用 PlatEMO v2.9 內建 PF(10000)。若作者使用不同點數或不同 PF 來源，IGD 會不同。"],
    ["Baseline 可能非原生 NSGA-II", "論文中的 NSGA-II baseline 可能來自作者改寫版本、不同平台版本或不同資料彙整流程。"],
]
add_table(doc, ["差異來源", "說明"], diff_rows, [2300, 7060])

doc.add_heading("8. 是否與資料集特性相關", level=1)
add_bullet(doc, "ZDT 差異最大，Mean relative diff = 71.00%，Median relative diff = 90.92%。ZDT1-ZDT3 在標準 PlatEMO NSGA-II 下 IGD 明顯低於論文，顯示論文 baseline 可能使用了不同停止條件、版本或資料整理方式。")
add_bullet(doc, "DTLZ 差異中等偏大，Mean relative diff = 49.46%，Median relative diff = 43.93%。DTLZ5、DTLZ6 等退化或特殊幾何問題對 PF sampling、解集分布與 seed 敏感。")
add_bullet(doc, "UF 差異最小，Mean relative diff = 21.05%，Median relative diff = 13.80%。UF 問題本身搜尋地形較複雜，不同實作間的結果差異可能被問題難度部分壓平。")
add_bullet(doc, "因此 baseline 差異不是單一固定偏差，而是與問題特性、Pareto front 幾何、搜尋難度、seed 敏感度與 IGD 評估流程皆有關。")

doc.add_heading("9. 建議論文敘述", level=1)
quote = (
    "本研究依據三層原則決定復現設定：首先採用原文已明確給定之設定；其次，原文未明確說明但 PlatEMO 平台具有原生實作者，採用 PlatEMO v2.9 原生流程；最後，原文與平台皆未指定者，採用經測試後最接近且可重現之設定。"
    "在上述設定下，部分 NSGA-II baseline 數值仍與原文 Table 3 存在差異。進一步 seed block、RNG 與 sequential stream 測試顯示，差異無法單純由 random seed 解釋。"
    "由於原文未提供完整程式碼、PlatEMO 版本與 random seeds，本文後續採用上述最接近且可重現之設定作為比較基準。"
)
add_note(doc, "可用文字。", quote)

footer = section.footer.paragraphs[0]
footer.text = "Generated reproduction memo - NSGA-II baseline"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)

doc.save(OUTPUT)
print(OUTPUT)
